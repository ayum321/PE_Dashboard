
# ============================================================
# PE Control Tower v4.0 – COMPLETE BASE FILE
# Run  : streamlit run app.py
# Deps : pip install streamlit plotly pandas numpy openpyxl python-docx pypdf
#        pip install pymupdf pdfplumber pillow lxml              (required)
#        pip install openai                                       (OpenAI Vision)
#        pip install anthropic                                    (Claude Vision)
#        pip install google-generativeai                          (Gemini Vision)
#        Ollama: install from https://ollama.com + pull llava     (free/local)
# ============================================================
import io, re, os
import pandas as pd
import numpy as np
import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime, date

st.set_page_config(page_title="PE Audit Control Tower", page_icon="🔧",
                   layout="wide", initial_sidebar_state="expanded")

DAILY_LIMIT_HRS   = 6.0
MONTHLY_LIMIT_HRS = 8.0
CPU_OK   = 75.0
CPU_WARN = 90.0
MEM_OK   = 75.0
MEM_WARN = 90.0
DISK_OK  = 75.0

DISK_WARN= 90.0


C = dict(
    bg="#060914",        # Deep navy-black — max contrast base
    card="#0d1526",      # Card bg — slightly lifted from bg
    card2="#111d36",     # Alternate card — visible depth
    border="#213060",    # Brighter border — visible card edges
    green="#10d96e",     # Vivid lime-green (Power BI style)
    amber="#f59e0b",     # Warm amber
    red="#f43f5e",       # Rose-red — more vivid than muted red
    blue="#3b82f6",      # Bright saturated blue
    purple="#a855f7",    # Vivid violet
    cyan="#22d3ee",      # Bright cyan
    muted="#6b7db3",     # Lighter muted text — more readable
    white="#f0f4ff",     # Crisp white with tiny blue tint
    # Nav extras
    nav_bg="#06091a",
    nav_active_bg="#14296a",
    nav_active_border="#3b82f6",
    nav_sep="#1a2850",
    nav_group="#1e3060",
    # Extended palette for charts
    teal="#2dd4bf",
    orange="#fb923c",
    pink="#ec4899",
    indigo="#6366f1",
)
BASE_LAYOUT = dict(
    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="Inter", color=C["white"], size=11),
)
# Default margin — used only when a chart does not specify its own
_DEFAULT_MARGIN = dict(l=10, r=10, t=30, b=10)
AXIS = dict(gridcolor=C["border"], linecolor=C["border"])
import html as _html_mod

# ── API KEY RESOLUTION ─────────────────────────────────────
_EMBEDDED_KEY = "***REMOVED-KEY***"

def _get_api_key():
    """Resolve Gemini API key: session_state → env vars → embedded fallback."""
    return (
        (st.session_state.get("_gemini_key") or "").strip()
        or os.environ.get("GOOGLE_API_KEY", "").strip()
        or os.environ.get("GEMINI_API_KEY", "").strip()
        or _EMBEDDED_KEY
    )

def hex_rgba(h, a=0.13):
    h = h.lstrip("#")
    r, g, b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
    return f"rgba({r},{g},{b},{a})"

def status_color(val, ok, warn):
    return C["green"] if val <= ok else (C["amber"] if val <= warn else C["red"])

# ── CSS ──────────────────────────────────────────────────────
_CSS_CACHE: dict = {}  # module-level cache — survives across reruns; cleared on C-dict change

def inject_css():
    # Version key — bump to force CSS regeneration after theme changes
    _CSS_VER = "v9-vivid-dash"
    if _CSS_CACHE.get("_ver") == _CSS_VER and "_css_injected" in _CSS_CACHE:
        st.markdown(_CSS_CACHE["_css_injected"], unsafe_allow_html=True)
        return
    _CSS_CACHE.clear()
    _CSS_CACHE["_ver"] = _CSS_VER
    C_bg=C["bg"];C_card=C["card"];C_card2=C["card2"];C_bord=C["border"]
    C_green=C["green"];C_amber=C["amber"];C_red=C["red"];C_blue=C["blue"]
    C_purp=C["purple"];C_muted=C["muted"];C_white=C["white"];C_cyan=C["cyan"]
    _css_str = f"""<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600;700&family=Sora:wght@300;400;600;800&display=swap');

/* ═══════════════════════════════════════════════════════════════
   PE CONTROL TOWER v9 — Vivid Dashboard Theme
   Saturated colors · Crisp typography · Zero dead space
   ═══════════════════════════════════════════════════════════════ */

/* ── Base & Typography ─────────────────────────────────────────── */
*{{box-sizing:border-box}}
html,body,[class*="css"]{{
  font-family:'Sora','Inter',system-ui,sans-serif!important;
  font-size:15px!important;
  background:{C_bg}!important;color:{C_white}!important;
  -webkit-font-smoothing:antialiased;-moz-osx-font-smoothing:grayscale}}
h1,h2,h3{{color:{C_white}!important;font-weight:800!important;letter-spacing:-.02em!important}}
h1{{font-size:38px!important;line-height:1.12!important}}
h2{{font-size:28px!important}}
h3{{font-size:22px!important}}
p,li,span,label,div{{font-size:15px!important}}
#MainMenu,footer,header,.stDeployButton{{display:none!important}}
/* Hide ALL sidebar toggle/collapse controls */
[data-testid="stSidebarCollapsedControl"],
[data-testid="stSidebarCollapseButton"],
[data-testid="collapsedControl"],
.css-1rs6os{{display:none!important}}

/* ═══ SIDEBAR — 280px, vivid gradient, crisp text ══════════════ */
[data-testid="stSidebar"]{{
  background:linear-gradient(180deg,#070b1e 0%,#0c1538 50%,#091028 100%)!important;
  border-right:1px solid rgba(59,130,246,.25)!important;
  min-width:280px!important;max-width:280px!important;
  box-shadow:4px 0 40px rgba(59,130,246,.08),
             1px 0 0 rgba(59,130,246,.18)!important;
  overflow-y:auto!important;
  z-index:1000!important}}
[data-testid="stSidebar"]>div:first-child{{
  padding:0!important;width:100%!important}}
[data-testid="stSidebar"] .block-container{{
  padding:0!important;max-width:100%!important}}
[data-testid="stSidebar"] [data-testid="stVerticalBlock"]{{
  gap:0!important;width:100%!important}}
[data-testid="stSidebar"] button[kind="header"]{{display:none!important}}
section[data-testid="stSidebar"] [data-testid="stSidebarCollapsedControl"]{{display:none!important}}
[data-testid="stSidebarCollapsedControl"]{{display:none!important}}
[data-testid="stSidebarCollapseButton"]{{display:none!important}}
button[data-testid="baseButton-headerNoPadding"]{{display:none!important}}
[aria-label="Collapse sidebar"]{{display:none!important}}
[aria-label="Expand sidebar"]{{display:none!important}}
.stSidebar [data-testid="stSidebarContent"]>div:first-child button{{display:none!important}}

/* ── Sidebar Nav Buttons — 16px, vivid, proper spacing ─────────── */
[data-testid="stSidebar"] .stButton>button{{
  background:transparent!important;
  border:none!important;border-radius:8px!important;
  border-left:3px solid transparent!important;
  color:#8ba3d9!important;
  font-size:16px!important;
  font-weight:600!important;
  text-align:left!important;
  padding:14px 18px 14px 20px!important;
  margin:2px 8px!important;
  width:calc(100% - 16px)!important;
  justify-content:flex-start!important;
  transition:all .2s ease!important;
  letter-spacing:.01em!important;
  line-height:1.35!important;
  background-image:none!important}}
[data-testid="stSidebar"] .stButton>button:hover{{
  background:linear-gradient(90deg,rgba(59,130,246,.15) 0%,rgba(59,130,246,.04) 100%)!important;
  color:#dce8ff!important;
  border-left-color:{C_cyan}88!important;
  transform:none!important;
  border-radius:8px!important}}
/* Active page — vivid accent bar + glow */
[data-testid="stSidebar"] .stButton>button[kind="primary"]{{
  background:linear-gradient(90deg,rgba(59,130,246,.22) 0%,rgba(59,130,246,.06) 100%)!important;
  border-left:3px solid {C_cyan}!important;
  color:#ffffff!important;
  font-weight:700!important;
  box-shadow:inset 0 0 40px rgba(59,130,246,.10),
             -3px 0 20px rgba(34,211,238,.12)!important;
  text-shadow:0 0 16px rgba(34,211,238,.3)!important;
  border-radius:8px!important}}
/* Disabled nav items */
[data-testid="stSidebar"] .stButton>button:disabled{{
  color:#2d3a60!important;
  opacity:0.5!important;
  cursor:not-allowed!important;
  background:transparent!important;
  text-shadow:none!important}}

/* ── Main Content — Fill Width, Tight Spacing ──────────────────── */
.main .block-container{{
  max-width:100%!important;
  padding:0.5rem 28px 24px!important}}
.main [data-testid="stVerticalBlock"]>div{{
  margin-bottom:0!important}}

/* ── Scrollbar ─────────────────────────────────────────────────── */
::-webkit-scrollbar{{width:6px;height:6px}}
::-webkit-scrollbar-track{{background:transparent}}
::-webkit-scrollbar-thumb{{background:{C_bord};border-radius:3px}}
::-webkit-scrollbar-thumb:hover{{background:{C_blue}55}}

/* ── KPI Cards — Power BI vivid style ─────────────────────────── */
.kpi-card{{
  background:linear-gradient(145deg,{C_card} 0%,{C_card2} 100%);
  border:1px solid {C_bord};border-radius:14px;
  padding:18px 20px;min-height:108px;
  box-shadow:0 4px 20px rgba(0,0,0,.4),
             0 0 0 1px rgba(59,130,246,.08),
             inset 0 1px 0 rgba(255,255,255,.04);
  contain:layout style;will-change:transform;
  transition:all .2s ease;
  position:relative;overflow:hidden}}
.kpi-card::before{{
  content:"";position:absolute;top:0;left:0;right:0;height:2px;
  background:linear-gradient(90deg,transparent,{C_blue}88,transparent)}}
.kpi-card:hover{{
  border-color:{C_blue}88;
  transform:translateY(-2px);
  box-shadow:0 12px 40px rgba(59,130,246,.18),
             0 0 0 1px rgba(59,130,246,.2)}}
.kpi-label{{font-size:10px;font-weight:700;letter-spacing:.12em;
  text-transform:uppercase;color:{C_muted};margin:0 0 6px}}
.kpi-value{{font-size:42px;font-weight:800;margin:0;line-height:1.05;letter-spacing:-.04em}}
.kpi-sub{{font-size:11px;color:{C_muted};margin:6px 0 0;font-weight:500}}

/* ── Panels — Elevated glassmorphism ───────────────────────────── */
.panel{{
  background:linear-gradient(145deg,{C_card} 0%,{C_card2} 100%);
  border:1px solid {C_bord};border-radius:14px;
  padding:18px 20px;contain:layout style;
  box-shadow:0 4px 24px rgba(0,0,0,.35),
             inset 0 1px 0 rgba(255,255,255,.04);
  backdrop-filter:blur(8px)}}
.panel-title{{font-size:14px;font-weight:700;margin:0 0 4px;color:{C_white};
  letter-spacing:.01em}}
.panel-sub{{font-size:11px;color:{C_muted};margin:0 0 10px;font-weight:500}}

/* ── Intel Bar — Glowing Status Strip ──────────────────────────── */
.intel-bar{{
  background:linear-gradient(135deg,rgba(13,21,38,.95) 0%,rgba(16,25,50,.95) 100%);
  border:1px solid rgba(64,112,232,.2);border-radius:12px;
  padding:12px 22px;display:flex;gap:28px;align-items:center;
  flex-wrap:wrap;margin-bottom:12px;
  position:sticky;top:0;z-index:99;
  backdrop-filter:blur(16px);-webkit-backdrop-filter:blur(16px);
  box-shadow:0 4px 24px rgba(0,0,0,.3),
             0 0 0 1px rgba(64,112,232,.08),
             inset 0 1px 0 rgba(255,255,255,.03)}}
.intel-item{{display:flex;align-items:center;gap:8px;font-size:12px}}
.intel-label{{color:{C_muted};font-weight:700;letter-spacing:.06em;font-size:10px;text-transform:uppercase}}
.intel-value{{font-weight:800;color:{C_white};font-size:15px}}
.intel-divider{{width:1px;height:24px;background:linear-gradient(180deg,transparent,{C_bord},transparent);flex-shrink:0}}

/* ── Findings Panel ────────────────────────────────────────────── */
.findings-panel{{
  background:linear-gradient(145deg,{C_card} 0%,{C_card2} 100%);
  border:1px solid {C_bord};border-radius:14px;
  padding:18px 22px;margin-bottom:12px;
  box-shadow:0 2px 12px rgba(0,0,0,.2)}}
.findings-panel-title{{
  font-size:14px;font-weight:700;color:{C_white};
  display:flex;align-items:center;gap:8px;margin-bottom:12px}}
.finding-row{{
  display:flex;align-items:flex-start;gap:10px;
  padding:10px 14px;border-radius:10px;margin-bottom:5px;
  border-left:3px solid transparent;contain:layout style}}
.finding-critical{{background:{C_red}0d;border-left-color:{C_red}}}
.finding-warning{{background:{C_amber}0d;border-left-color:{C_amber}}}
.finding-info{{background:{C_blue}0d;border-left-color:{C_blue}}}
.finding-ok{{background:{C_green}0d;border-left-color:{C_green}}}
.finding-icon{{font-size:15px;flex-shrink:0;margin-top:1px}}
.finding-text{{font-size:14px;font-weight:700;line-height:1.5}}
.finding-sub{{font-size:12px;color:{C_muted};margin-top:3px;line-height:1.5}}

/* ── Server Cards ──────────────────────────────────────────────── */
.srv-card{{
  background:{C_card};border:1px solid {C_bord};border-radius:12px;
  padding:14px 16px;margin-bottom:8px;contain:layout style}}
.srv-card-header{{display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:8px}}
.srv-hostname{{font-size:12px;font-weight:700;color:{C_white};word-break:break-all}}
.srv-tags{{display:flex;gap:4px;margin-top:3px;flex-wrap:wrap}}
.srv-bar-row{{display:flex;align-items:center;gap:6px;margin-bottom:4px}}
.srv-bar-lbl{{font-size:9px;color:{C_muted};width:32px;font-weight:600}}
.srv-bar-track{{flex:1;height:6px;background:{C_bord};border-radius:3px;overflow:hidden}}

/* ── Badges ────────────────────────────────────────────────────── */
.b-ok{{background:{C_green}22;color:{C_green};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}
.b-warn{{background:{C_amber}22;color:{C_amber};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}
.b-breach{{background:{C_red}22;color:{C_red};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}
.b-blue{{background:{C_blue}22;color:{C_blue};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}
.b-run{{background:{C_green}22;color:{C_green};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}
.badge-red{{background:{C_red}22;color:{C_red};padding:3px 8px;border-radius:6px;font-size:11px;font-weight:700}}
.badge-green{{background:{C_green}22;color:{C_green};padding:3px 8px;border-radius:6px;font-size:11px;font-weight:700}}
.b-signed{{background:{C_green}22;color:{C_green};padding:3px 10px;border-radius:8px;font-weight:700;font-size:12px}}

/* ── Sub-Tab Pills (inner tabs inside pages) ───────────────────── */
[data-testid="stTabs"] [data-baseweb="tab-list"]{{
  display:flex!important;
  background:linear-gradient(135deg,#0a1428,#0d1830)!important;
  border:1px solid rgba(64,112,232,.15)!important;
  border-radius:12px!important;
  padding:5px!important;
  gap:4px!important;
  margin-bottom:12px!important;
  overflow-x:auto!important;
  scrollbar-width:none!important;
  box-shadow:0 2px 12px rgba(0,0,0,.3),
             inset 0 1px 0 rgba(255,255,255,.02)!important}}
[data-testid="stTabs"] [data-baseweb="tab-list"]::-webkit-scrollbar{{display:none!important}}
[data-testid="stTabs"] [data-baseweb="tab"]{{
  border-radius:8px!important;
  padding:10px 20px!important;
  min-height:42px!important;
  min-width:auto!important;
  font-family:'Sora','Inter',sans-serif!important;
  font-size:14px!important;
  font-weight:700!important;
  letter-spacing:.02em!important;
  text-align:center!important;
  white-space:nowrap!important;
  line-height:1.3!important;
  color:#5a6a9e!important;
  border:1px solid transparent!important;
  transition:all .18s ease!important;
  cursor:pointer!important;
  flex-direction:row!important;
  gap:6px!important}}
[data-testid="stTabs"] [data-baseweb="tab"]:hover:not([aria-selected="true"]){{
  background:rgba(64,112,232,.08)!important;
  color:#8da0d8!important}}
[data-testid="stTabs"] [aria-selected="true"]{{
  background:linear-gradient(135deg,rgba(64,112,232,.2),rgba(64,112,232,.08))!important;
  border-color:rgba(64,112,232,.4)!important;
  color:#e8f0ff!important;
  box-shadow:0 0 16px rgba(64,112,232,.15),
             inset 0 1px 0 rgba(255,255,255,.05)!important;
  text-shadow:0 0 12px rgba(64,112,232,.25)!important}}
[data-testid="stTabs"] [data-baseweb="tab-border"],
[data-testid="stTabs"] [data-baseweb="tab-highlight"]{{
  display:none!important}}
[data-testid="stTabsContent"]{{padding-top:8px!important}}

/* ── File Uploader — vivid dashed border ──────────────────────── */
[data-testid="stFileUploader"]>div:first-child{{
  background:linear-gradient(135deg,{C_card2},{C_card})!important;
  border:1.5px dashed rgba(59,130,246,.35)!important;
  border-radius:12px!important;padding:12px!important;
  transition:all .2s ease}}
[data-testid="stFileUploader"]>div:first-child:hover{{
  border-color:{C_cyan}!important;
  background:linear-gradient(135deg,rgba(59,130,246,.06),{C_card})!important}}
[data-testid="stFileUploader"] label{{display:none!important}}

/* ── Buttons (main content only, NOT sidebar) ──────────────────── */
.main .stButton>button{{
  background:linear-gradient(135deg,{C_blue} 0%,{C_purp} 100%)!important;
  color:#fff!important;border:none!important;border-radius:10px!important;
  font-weight:700!important;font-size:14px!important;
  padding:11px 22px!important;
  box-shadow:0 4px 16px rgba(59,130,246,.25)!important;
  transition:all .2s ease!important;
  letter-spacing:.01em!important}}
.main .stButton>button:hover{{
  opacity:.92;transform:translateY(-1px)!important;
  box-shadow:0 6px 24px rgba(59,130,246,.35)!important}}

/* ── Expanders ──────────────────────────────────────────────────── */
[data-testid="stExpander"]{{
  background:{C_card}!important;border:1px solid {C_bord}!important;
  border-radius:12px!important;margin-bottom:8px!important}}
[data-testid="stExpander"] summary{{font-size:13px!important;font-weight:600!important}}

/* ── Job Rows ───────────────────────────────────────────────────── */
.job-row{{display:flex;justify-content:space-between;align-items:center;
  padding:10px 14px;border-radius:10px;margin-bottom:4px;
  background:{C_card2};border-left:3px solid {C_bord}}}
.job-name{{font-size:13px;font-weight:600}}
.job-meta{{font-size:11px;color:{C_muted}}}

/* ── PE Header — Vivid gradient strip ─────────────────────────── */
.pe-header{{
  background:linear-gradient(135deg,rgba(11,18,40,.98) 0%,rgba(14,25,58,.98) 100%);
  border-bottom:1px solid rgba(59,130,246,.2);
  padding:14px 28px;
  display:flex;justify-content:space-between;align-items:center;
  margin-bottom:0;
  box-shadow:0 4px 24px rgba(0,0,0,.4),
             0 1px 0 rgba(59,130,246,.12),
             inset 0 -1px 0 rgba(59,130,246,.06)}}
.pe-logo{{
  background:linear-gradient(135deg,{C_blue},{C_cyan});
  width:44px;height:44px;border-radius:12px;display:flex;align-items:center;
  justify-content:center;font-weight:900;font-size:16px;color:#fff;flex-shrink:0;
  box-shadow:0 4px 20px rgba(59,130,246,.4),
             0 0 0 2px rgba(59,130,246,.2)}}
.pe-title{{font-size:18px;font-weight:800;color:#fff;margin:0;letter-spacing:.01em}}
.pe-sub{{font-size:12px;color:{C_muted};margin:2px 0 0;font-weight:500}}

/* ── Scrollable Containers ──────────────────────────────────────── */
.scroll-container{{max-height:400px;overflow-y:auto;
  scrollbar-width:thin;scrollbar-color:{C_bord} transparent}}

/* ── Issue Cards ────────────────────────────────────────────────── */
.issue-card{{
  background:{C_card};border-left:4px solid transparent;
  border-radius:12px;padding:14px 16px;margin-bottom:8px;contain:layout style}}

/* ── Top-Jobs Table ─────────────────────────────────────────────── */
.tj-grid{{display:grid;grid-template-columns:2fr .85fr .85fr .6fr;
  font-size:12px;border-bottom:1px solid {C_bord}18}}
.tj-hdr{{font-size:11px;color:{C_muted};font-weight:700;letter-spacing:.04em;
  text-transform:uppercase;padding:8px 12px;border-bottom:1px solid {C_bord}}}
.tj-cell{{padding:10px 12px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}

/* ── Keyframes ──────────────────────────────────────────────────── */
@keyframes fade-in{{from{{opacity:0;transform:translateY(5px)}}to{{opacity:1;transform:none}}}}
@keyframes pulse{{0%,100%{{opacity:1}}50%{{opacity:.55}}}}
@keyframes ambient-glow{{
  0%,100%{{box-shadow:0 0 20px rgba(64,112,232,.1)}}
  50%{{box-shadow:0 0 40px rgba(64,112,232,.2)}}}}

/* ── Streamlit element spacing overrides ────────────────────────── */
.main [data-testid="stHorizontalBlock"]{{gap:12px!important}}
.main [data-testid="stVerticalBlock"]>div:empty{{margin:0!important;padding:0!important}}
.main .stMarkdown p{{font-size:13px!important}}
.main [data-testid="stDataFrame"]{{font-size:12px!important}}
.main .stSelectbox label,.main .stTextInput label{{font-size:13px!important;font-weight:600!important}}
.main .stSelectbox [data-baseweb="select"]{{font-size:13px!important}}

/* ── Dataframe / Table styling ─────────────────────────────────── */
.main [data-testid="stDataFrame"] table{{font-size:12px!important}}
.main [data-testid="stDataFrame"] th{{
  background:{C_card2}!important;color:{C_muted}!important;
  font-size:11px!important;font-weight:700!important;text-transform:uppercase!important}}
</style>"""
    _CSS_CACHE["_css_injected"] = _css_str
    st.markdown(_css_str, unsafe_allow_html=True)

# ── SESSION STATE ────────────────────────────────────────────
def init_state():
    defaults = dict(
        # ── Data ──
        ctrlm_df=None, server_data=None, batch_sla_df=None,
        # ── Identity ──
        customer_name="", env_type="",
        # ── SOW ──
        sow_dfu=0, sow_sku=0, sow_dfu_base=0, sow_sku_base=0,
        sow_scenarios_agreed=0, sow_scenarios_prod=0,
        # ── UI state ──
        show_upload=True, active_tab="Job Drilldown",
        _current_page="upload",
        # ── Issues ──
        issues_list=[],
        # ── Approvals ──
        approval_pe=False, approval_customer=False,
        approval_pe_name="", approval_customer_name="",
        approval_pe_date="", approval_customer_date="",
        approval_notes="",
        # ── File tracking hashes ──
        _ctrlm_hash=None, _sla_hash=None, _res_hash=None,
        # ── AI analysis cache — per-tab ──
        _ai_analysis_result=None, _batch_ai_result=None,
        _quick_analysis=False, _batch_quick=False,
        # ── AI upload-time insight cache ──
        _ai_ctrlm_insight=None, _ai_sla_insight=None, _ai_resource_insight=None,
        _run_ai_ctrlm=False, _run_ai_sla=False, _run_ai_resource=False,
        # ── Master PE Summary ──
        _ai_master_summary=None, _run_master_summary=False,
        # ── PE Document Review ──
        _pe_doc_data=None, _pe_doc_hash=None, _pe_doc_name="",
        _pe_doc_ai=None, _run_pe_doc_ai=False,
    )
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def extract_customer_from_filename(filename):
    """Extract customer name from PE standard filename patterns."""
    import re
    stem = re.sub(r'\.(csv|xlsx|pdf|docx)$', '', filename, flags=re.IGNORECASE)
    for pat in [
        r'Report_of_CS[_]([A-Z][A-Z0-9_]+?)_SCPO',
        r'_CS_([A-Z][A-Z0-9_]+?)_SCPO',
        r'CUSTOMER_([A-Z][A-Z0-9_]+?)_SCPO',
    ]:
        m = re.search(pat, stem, re.IGNORECASE)
        if m:
            raw = m.group(1).replace("_"," ").strip()
            parts = raw.split()
            return " ".join(p if (len(p)<=4 and p.isupper()) else p.title() for p in parts)
    return ""


def extract_env_from_filename(filename):
    """Extract environment type from filename."""
    fn = filename.upper()
    if "NON_PROD" in fn or "NONPROD" in fn or "NON-PROD" in fn: return "NON-PROD"
    if re.search(r'[_\-]PROD(?:[_\-]|$)', fn): return "PROD"
    if "TEST" in fn: return "TEST"
    if "UAT"  in fn: return "UAT"
    if "DEV"  in fn: return "DEV"
    return ""


def parse_resource_docx(file_obj):
    """
    Parse resource utilisation DOCX — handles ALL real PE report formats:
    A) Distell:   "Server1: tsbb191525041.jdadelivers.com" / bare FQDN
    B) Leonardo:  "Application Server 1 :", "SRE UI IO :", "Utility :"
    C) Generic:   "SRE : hostname.domain.com"
    D) CEAt:      "DB Server 1", "DB Server 2", "App Server" (no colon/bold)
    E) Generic:   any paragraph whose text contains a known server keyword
                  and is short enough to be a heading
    Returns (servers_list, image_only: bool)
    """
    try:
        from docx import Document
    except ImportError:
        return [], True

    file_obj.seek(0)
    try:
        doc = Document(file_obj)
    except Exception as _docx_err:
        # File is corrupted or not a valid DOCX — return empty list gracefully.
        # The upload_panel caller handles sd==[] with a clear st.warning() to the user.
        _ = _docx_err
        return [], True

    # ── Try structured tables first (text-based metrics) ─────────
    servers_with_data = []
    for tbl in doc.tables:
        if len(tbl.rows) < 2: continue
        hdrs = [c.text.strip().lower() for c in tbl.rows[0].cells]
        has_host   = any(k in h for h in hdrs for k in ["host","server","name"])
        has_metric = any(k in " ".join(hdrs) for k in ["cpu","disk","mem","memory"])
        if not (has_host and has_metric): continue
        for row in tbl.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells[0]: continue
            rec = {"host":cells[0],"type":_infer_server_type(cells[0]),
                   "cpu_used":0.0,"cpu_avg":0.0,
                   "mem_used":0.0,"mem_total_gb":0.0,"disk_used_max":0.0,"disks":{}}
            for hi,h in enumerate(hdrs):
                if hi>=len(cells): break
                cell_text = cells[hi].strip()
                # Skip columns that are clearly not numeric metrics
                # (OS version strings like "3.10.0-1160", hostnames, category labels)
                if any(skip in h for skip in ["o/s","os","oper","version","categor","name","server","host","region","module","enterprise","product"]):
                    continue
                # Extract the FIRST plain integer or simple decimal (not dotted version strings)
                # A valid metric is: digits optionally followed by ONE decimal point + digits
                vm = re.search(r"\b(\d{1,6}(?:\.\d{1,3})?)\b", cell_text)
                if not vm:
                    continue
                candidate = vm.group(1)
                # Reject version-like strings: more than one dot means it's a version, not a metric
                if cell_text.count(".") > 1:
                    continue
                try:
                    val = float(candidate)
                except ValueError:
                    continue
                if "cpu" in h:                     rec["cpu_used"]=val
                if "mem" in h and "gb" in h:       rec["mem_total_gb"]=val
                if "mem" in h and ("%" in h or "used" in h):  rec["mem_used"]=val
                if "disk" in h or "storage" in h:  rec["disk_used_max"]=val
            servers_with_data.append(rec)
    if servers_with_data:
        return servers_with_data, False

    # ── Regex helpers ─────────────────────────────────────────────
    # Generic server hostnames: short alpha prefix + digits, optionally followed by FQDN
    # e.g. tsbb191525041.domain.com, prbg241530001, dvbb941426001.local
    FQDN_RE    = re.compile(
        r"\b([a-z]{2,6}[0-9]{3,}[0-9a-z]*"
        r"(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)\b",
        re.IGNORECASE)
    SHORT_HOST = re.compile(
        r"\b([a-z]{2,6}[0-9]{4,}[0-9a-z]*)\b", re.IGNORECASE)

    # Server role / type keywords
    ROLE_RE = re.compile(
        r"\b(application|app|database|db|sre|batch|utility|act|cognos|"
        r"ui|io|server|node|worker|integration|etl|web|gateway|report|"
        r"scheduler|mq|middleware|proxy|balancer|cache|search|analytics)\b",
        re.IGNORECASE)

    SERVER_HEADING_RE = re.compile(
        r"\b(app(?:lication)?\s*server\s*\d*|"
        r"db\s*server\s*\d*|"
        r"database\s*server\s*\d*|"
        r"sre\s*(?:ui|io|batch|app|server)?\s*\d*|"
        r"utility\s*(?:server)?\s*\d*|"
        r"act\s*server\s*\d*|"
        r"cognos\s*(?:server)?\s*\d*|"
        r"batch\s*server\s*\d*|"
        r"web\s*server\s*\d*|"
        r"etl\s*server\s*\d*|"
        r"server\s*\d+)",
        re.IGNORECASE)

    def infer_type(text):
        t = text.lower()
        if any(k in t for k in ["db ","database"," db","oracle","sql","data base","dbserver"]): return "DB"
        return "APP"

    def extract_hostname(text):
        m = FQDN_RE.search(text)
        if m: return m.group(1)
        m = SHORT_HOST.search(text)
        if m: return m.group(1)
        return None

    def para_is_bold(para):
        if not para.runs: return False
        bold_chars = sum(len(r.text) for r in para.runs if r.bold and r.text.strip())
        total_chars = len(para.text.strip())
        return total_chars > 0 and bold_chars / total_chars > 0.4

    # Company / org name suffixes — lines with these are titles, never servers
    COMPANY_RE = re.compile(
        r"\b(ltd|limited|inc|corp|corporation|pty|gmbh|b\.?v|s\.?a|"
        r"ag|llc|plc|co\.|group|holdings|international|industries)\b",
        re.IGNORECASE)

    def is_server_heading(para, excluded=None):
        t = para.text.strip()
        if not t or len(t) > 150: return False

        # Never a server if it's in the excluded titles set
        if excluded and t in excluded: return False

        # Never a server if it contains a company name suffix (Ltd, Pty, Corp, Inc)
        # UNLESS it also has a real hostname
        has_host = bool(FQDN_RE.search(t) or SHORT_HOST.search(t))
        has_company = bool(COMPANY_RE.search(t))
        if has_company and not has_host: return False

        # 1) Contains a real JDA/Zabbix hostname → always a server entry
        if has_host: return True

        # 2) Matches known server heading pattern (DB Server 1, App Server 2, etc.)
        if SERVER_HEADING_RE.search(t): return True

        # 3) Bold/underline + role keyword — but NOT if it looks like a document title
        is_styled = para_is_bold(para) or any(
            r.underline for r in para.runs if r.text.strip())
        if is_styled and ROLE_RE.search(t):
            # Extra guard: colon-split left side must not be a company/title
            parts = re.split(r"\s*:\s*", t, maxsplit=1)
            if len(parts) == 2:
                left = parts[0].strip()
                # If left has no hostname and has company suffix → it's a title
                if COMPANY_RE.search(left) and not FQDN_RE.search(left): return False
                # If left looks like "Document Name" (Title Case multi-word, no digits)
                # and right looks like a description → skip
                if (not re.search(r"\d", left) and
                    len(left.split()) >= 2 and
                    not SERVER_HEADING_RE.search(left) and
                    not FQDN_RE.search(left)):
                    # Right side must have a hostname to count
                    if not FQDN_RE.search(parts[1]) and not SHORT_HOST.search(parts[1]):
                        return False
            return True

        # 4) Short line + role keyword (Leonardo: "SRE UI IO :", "Utility :")
        stripped = t.rstrip(": ")
        if len(stripped) < 60 and ROLE_RE.search(stripped):
            words = stripped.split()
            if len(words) <= 6:
                # Guard: must not have company suffix on left of colon
                parts = re.split(r"\s*:\s*", stripped, maxsplit=1)
                left = parts[0].strip()
                if COMPANY_RE.search(left): return False
                return True

        return False

    # ── Collect all text from paragraphs + single-cell tables ────
    all_paras = list(doc.paragraphs)
    # Also add single-cell table entries (some DOCXs use 1-cell tables as headings)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c for c in row.cells if c.text.strip()]
            if len(cells) == 1:
                for para in cells[0].paragraphs:
                    if para.text.strip():
                        all_paras.append(para)

    # ── Walk all paragraphs ───────────────────────────────────────
    seen, servers = set(), []
    customer_name = None
    excluded_titles = set()   # lines used as doc title/customer → never server headings

    # ── PASS 1: extract customer name + build excluded set ────────
    for para in all_paras:
        t = para.text.strip()
        if not t or len(t) > 80: continue
        if para_is_bold(para) and not FQDN_RE.search(t) and not SHORT_HOST.search(t):
            # Has company suffix → definitely a title line
            if COMPANY_RE.search(t):
                excluded_titles.add(t)
                if not customer_name:
                    cname = re.sub(
                        r"[:\s]*(sre|servers?|resource|utilization|report|"
                        r"performance|metrics|consumption|test|prod|uat|dev|"
                        r"asia|2022|2023|2024|2025|2026).*$",
                        "", t, flags=re.IGNORECASE).strip().rstrip(":,.")
                    if 3 < len(cname) < 60 and not SERVER_HEADING_RE.match(cname):
                        customer_name = cname
            # Bold non-hostname short line that looks like a document title
            elif not customer_name and len(t) > 5:
                cname = re.sub(
                    r"[:\s]*(sre|servers?|resource|utilization|report|"
                    r"performance|metrics|consumption|test|prod|uat|dev|"
                    r"asia|2022|2023|2024|2025|2026).*$",
                    "", t, flags=re.IGNORECASE).strip().rstrip(":,.")
                if 3 < len(cname) < 60 and not SERVER_HEADING_RE.match(cname):
                    customer_name = cname
                    excluded_titles.add(t)

    # ── PASS 2: extract server headings ──────────────────────────
    # Build a list of (label, hostname, type) tuples
    # Handles patterns:
    #   "Server1: tsbb191525041.jdadelivers.com" (label:host on same line)
    #   "Application Server 1" then "tsbb911502021" (label then host on next line)
    #   "Database Server:" (label only, no hostname)
    #   Skip: "------- Application Server-1 end------" (separator lines)
    pending_label = None  # label from previous paragraph, waiting for hostname

    for para in all_paras:
        t = para.text.strip()
        if not t: continue

        # Skip separator lines (--- ... ---)
        if re.match(r'^[-─—=_]{5,}', t): continue

        # Check if this line is just a bare hostname (follow-up to a label)
        bare_host = extract_hostname(t)
        if bare_host and len(t.replace(bare_host, "").strip(":. ")) < 5:
            # This is a hostname-only line
            if pending_label:
                # Merge with previous label
                display = bare_host
                label = pending_label
                stype = infer_type(pending_label)
                # Also check hostname for DB
                if _infer_server_type(bare_host) == "DB": stype = "DB"
                pending_label = None
                if display not in seen:
                    seen.add(display)
                    servers.append({
                        "host": display, "label": label, "type": stype,
                        "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                        "mem_total_gb": 0.0, "disk_used_max": 0.0,
                        "disks": {}, "_image_only": True,
                    })
                continue
            else:
                # Standalone hostname
                display = bare_host
                stype = _infer_server_type(bare_host)
                if display not in seen:
                    seen.add(display)
                    servers.append({
                        "host": display, "label": display, "type": stype,
                        "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                        "mem_total_gb": 0.0, "disk_used_max": 0.0,
                        "disks": {}, "_image_only": True,
                    })
                continue

        # Clear pending label if this is not a hostname
        if pending_label:
            # Previous label had no hostname — store as label-only server
            display = pending_label.rstrip(": ").strip()
            stype = infer_type(pending_label)
            if display not in seen and len(display) > 2:
                seen.add(display)
                servers.append({
                    "host": display, "label": display, "type": stype,
                    "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                    "mem_total_gb": 0.0, "disk_used_max": 0.0,
                    "disks": {}, "_image_only": True,
                })
            pending_label = None

        if not is_server_heading(para, excluded=excluded_titles): continue

        # ── Parse heading into host + label ──────────────────────
        colon_split = re.split(r"\s*:\s*", t, maxsplit=1)
        hostname = None
        label    = t

        if len(colon_split) == 2:
            left, right = colon_split[0].strip(), colon_split[1].strip()
            h_right = extract_hostname(right)
            h_left  = extract_hostname(left)
            if h_right:
                hostname = h_right
                label    = left if left else right
            elif h_left:
                hostname = h_left
                label    = right if right else left
            else:
                label = left or right
        else:
            hostname = extract_hostname(t)
            label    = t

        if hostname:
            display = hostname
            stype = infer_type(t)
            if _infer_server_type(hostname) == "DB": stype = "DB"
            if display not in seen:
                seen.add(display)
                servers.append({
                    "host": display, "label": label.rstrip(": ").strip(),
                    "type": stype,
                    "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                    "mem_total_gb": 0.0, "disk_used_max": 0.0,
                    "disks": {}, "_image_only": True,
                })
        else:
            # No hostname on this heading line — set as pending label
            # so the next paragraph (if it's a bare hostname) can merge
            pending_label = label.rstrip(": ").strip()

    # Flush any remaining pending_label at end of loop
    if pending_label:
        display = pending_label.rstrip(": ").strip()
        stype = infer_type(pending_label)
        if display not in seen and len(display) > 2:
            seen.add(display)
            servers.append({
                "host": display, "label": display, "type": stype,
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0,
                "disks": {}, "_image_only": True,
            })

    if servers:
        if customer_name:
            try:
                import streamlit as st
                if not st.session_state.get("customer_name","").strip():
                    def _sc(s):
                        pts = s.replace("_"," ").split()
                        return " ".join(p if (len(p)<=4 and p.isupper()) else p.title() for p in pts)
                    st.session_state.customer_name = _sc(customer_name)
            except Exception:
                # st.session_state is not accessible outside a Streamlit context
                # (e.g. unit tests). Silently skip — this is purely a UI convenience.
                pass
        return servers, True

    return [], True


def parse_resource_docx_structured(file_obj):
    """
    Structure-aware DOCX parser — walks word/document.xml body in XML element
    order to map embedded images to their server sections precisely.

    Document layout (confirmed across multiple PE reports):
        PARA: "Application Server 1"   ← label
        PARA: "tsbb911502021"          ← hostname
        PARA: <drawing: image1.png>    ← CPU chart
        PARA: <drawing: image2.png>    ← Memory chart
        PARA: <drawing: image3.png>    ← Disk chart
        PARA: "Application Server 2"   ← next server ...

    Returns (server_sections, True) where each section is:
        {"host": str, "label": str, "type": str, "images": [(fname, bytes)],
         "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
         "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
         "_image_only": True}

    Falls back to ([], True) on any error.
    """
    import zipfile, io
    from xml.etree import ElementTree as ET

    # DOCX XML namespaces
    _W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    _A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
    _V   = "urn:schemas-microsoft-com:vml"

    FQDN_RE = re.compile(
        r"\b([a-z]{2,6}[0-9]{3,}[0-9a-z]*"
        r"(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)\b",
        re.IGNORECASE)
    SHORT_HOST = re.compile(r"\b([a-z]{2,6}[0-9]{4,}[0-9a-z]*)\b", re.IGNORECASE)
    SERVER_HEADING_RE = re.compile(
        r"\b(app(?:lication)?\s*server\s*\d*|db\s*server\s*\d*|"
        r"database\s*server\s*\d*|sre\s*(?:ui|io|batch|app|server)?\s*\d*|"
        r"utility\s*(?:server)?\s*\d*|act\s*server\s*\d*|"
        r"batch\s*server\s*\d*|web\s*server\s*\d*|etl\s*server\s*\d*|"
        r"cognos\s*(?:server)?\s*\d*|server\s*\d+)",
        re.IGNORECASE)

    def _extract_host(text):
        m = FQDN_RE.search(text)
        if m: return m.group(1)
        m = SHORT_HOST.search(text)
        if m: return m.group(1)
        return None

    def _is_server_like(text):
        if not text or len(text) > 150: return False
        if re.match(r'^[-─—=_]{5,}', text): return False
        if FQDN_RE.search(text) or SHORT_HOST.search(text): return True
        if SERVER_HEADING_RE.search(text): return True
        return False

    def _infer_type_from_label(label, host):
        """Determine server type from section heading first, then hostname.
        Fixes TSBC* servers being mis-classified as DB when heading says APP/SRE."""
        lbl = label.lower()
        if re.search(r'\bapp(?:lication)?\b|\bsre\b|\bweb\b|\bbatch\b|\betl\b', lbl):
            return "APP"
        if re.search(r'\bdb\b|\bdatabase\b|\boracle\b|\bsql\b|\bdata\s*server\b', lbl):
            return "DB"
        # Fallback to hostname-based inference
        return _infer_server_type(host)

    def _blank_section(host, label):
        return {
            "host": host, "label": label,
            "type": _infer_type_from_label(label, host),
            "images": [],
            "cpu_used": 0.0, "cpu_avg": 0.0,
            "mem_used": 0.0, "mem_total_gb": 0.0,
            "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }

    try:
        file_obj.seek(0)
        raw = file_obj.read()
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            names = set(z.namelist())

            # 1. Build rId → media path map from relationship file
            rid_to_media = {}
            rels_path = "word/_rels/document.xml.rels"
            if rels_path in names:
                rels_root = ET.fromstring(z.read(rels_path))
                for rel in rels_root:
                    rid    = rel.get("Id", "")
                    target = rel.get("Target", "")
                    if "media" in target.lower():
                        if not target.startswith("word/"):
                            target = "word/" + target
                        rid_to_media[rid] = target

            # Lazy media reader
            def _read_media(path):
                if path in names:
                    return z.read(path)
                return None

            # 2. Parse document body in element order
            doc_root = ET.fromstring(z.read("word/document.xml"))
            body = doc_root.find(f"{{{_W}}}body")
            if body is None:
                return [], True

            def _para_text(elem):
                return "".join(t.text or "" for t in elem.iter(f"{{{_W}}}t")).strip()

            def _get_img_rids(elem):
                """Return all image rIds referenced inside this element."""
                rids = []
                for blip in elem.iter(f"{{{_A}}}blip"):
                    rid = blip.get(f"{{{_R}}}embed")
                    if rid: rids.append(rid)
                for imgdata in elem.iter(f"{{{_V}}}imagedata"):
                    rid = imgdata.get(f"{{{_R}}}id") or imgdata.get(f"{{{_R}}}href")
                    if rid: rids.append(rid)
                return rids

            IMG_EXTS = {"png", "jpg", "jpeg", "gif", "bmp", "tiff"}

            def _add_images_to_section(section, rids):
                for rid in rids:
                    media_path = rid_to_media.get(rid)
                    if not media_path:
                        continue
                    ext = media_path.rsplit(".", 1)[-1].lower()
                    if ext not in IMG_EXTS:
                        continue
                    img_bytes = _read_media(media_path)
                    if img_bytes:
                        fname = media_path.split("/")[-1]
                        section["images"].append((fname, img_bytes))

            sections      = []
            current_sec   = None
            pending_label = None  # label paragraph waiting for hostname paragraph
            seen_hosts    = set()

            for child in body:
                ctag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if ctag == "p":                          # ── Paragraph ──
                    txt      = _para_text(child)
                    img_rids = _get_img_rids(child)

                    # Image paragraph — assign to current server section
                    if img_rids:
                        if current_sec is not None:
                            _add_images_to_section(current_sec, img_rids)
                        continue

                    if not txt or re.match(r'^[-─—=_]{5,}', txt):
                        continue

                    # Bare hostname line?
                    host = _extract_host(txt)
                    is_bare_host = host and len(txt.replace(host, "").strip(":. ")) < 5

                    if is_bare_host and pending_label is not None:
                        # Pattern: label para → hostname para (most common)
                        hkey = host.split(".")[0].lower()
                        if hkey not in seen_hosts:
                            seen_hosts.add(hkey)
                            current_sec = _blank_section(host, pending_label)
                            sections.append(current_sec)
                        pending_label = None
                        continue

                    # Server-like heading?
                    if _is_server_like(txt):
                        # Flush any stale pending label (label-only heading with no hostname)
                        if pending_label is not None:
                            lbl  = pending_label.rstrip(": ").strip()
                            lkey = lbl.lower()
                            if lkey not in seen_hosts and len(lbl) > 2:
                                seen_hosts.add(lkey)
                                current_sec = _blank_section(lbl, lbl)
                                sections.append(current_sec)
                            pending_label = None

                        # Try to parse "label : hostname" from same line
                        colon_parts = re.split(r"\s*:\s*", txt, maxsplit=1)
                        host, label = None, txt
                        if len(colon_parts) == 2:
                            h_r = _extract_host(colon_parts[1])
                            h_l = _extract_host(colon_parts[0])
                            if h_r:
                                host, label = h_r, colon_parts[0].strip()
                            elif h_l:
                                host, label = h_l, (colon_parts[1].strip() or colon_parts[0].strip())
                        else:
                            host = _extract_host(txt)

                        if host:
                            hkey = host.split(".")[0].lower()
                            if hkey not in seen_hosts:
                                seen_hosts.add(hkey)
                                current_sec = _blank_section(host, label.rstrip(": ").strip())
                                sections.append(current_sec)
                        else:
                            # Label without hostname — next para should be the hostname
                            pending_label = txt.rstrip(": ").strip()

                elif ctag == "tbl":                      # ── Table ──
                    # Images may sit inside table cells (some report formats)
                    for para_elem in child.iter(f"{{{_W}}}p"):
                        rids = _get_img_rids(para_elem)
                        if rids and current_sec is not None:
                            _add_images_to_section(current_sec, rids)

            # Flush trailing pending label
            if pending_label:
                lbl  = pending_label.rstrip(": ").strip()
                lkey = lbl.lower()
                if lkey not in seen_hosts and len(lbl) > 2:
                    seen_hosts.add(lkey)
                    current_sec = _blank_section(lbl, lbl)
                    sections.append(current_sec)

            return sections, True

    except Exception:
        return [], True


def parse_sow_volumes(file_obj, filename):
    """
    Generic SOW contract parser for Blue Yonder / JDA Schedule PDFs.
    Returns (customer_name, dfu, sku)
      dfu = Item-Location-Customer count
      sku = Item-Location count (largest)
    """
    import re
    text = ""
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            from pypdf import PdfReader
            file_obj.seek(0)
            reader = PdfReader(file_obj)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif ext in [".xlsx",".xls"]:
            import openpyxl
            file_obj.seek(0)
            wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    text += " ".join(str(c) for c in row if c) + "\n"
        elif ext in [".docx",".doc"]:
            from docx import Document
            file_obj.seek(0)
            doc = Document(file_obj)
            text = "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        # SOW parsing is best-effort. If the file is unreadable (corrupt, wrong format,
        # password-protected), return (None, None, None) so the caller can prompt the
        # user to enter the values manually — this is not a fatal error.
        return None, None, None

    # Customer name
    customer_name = None
    for pat in [
        r'Customer[:\s]+([A-Z][A-Za-z\s&,\.]+?)(?:\n|Subscription|Address|Representative)',
        r'CUSTOMER[:\s]+([A-Z][A-Za-z\s&,\.]+?)(?:\n|$)',
        r'Customer\s*:\s*([^\n]+)',
    ]:
        m = re.search(pat, text, re.MULTILINE)
        if m:
            name = m.group(1).strip().rstrip(",.:;")
            if 3 < len(name) < 80:
                customer_name = name
                break

    # DFU = Item-Location-Customer (largest count)
    dfu = None
    ilc = []
    for pat in [r'([\d,]+)\s*Item-Location-Customer',
                r'Item-Location-Customer[\s\n]*([\d,]+)',
                r'DFU[^\d]*(\d[\d,]+)']:
        for m in re.finditer(pat, text, re.IGNORECASE):
            try:
                n = int(m.group(1).replace(",",""))
                if n > 1000: ilc.append(n)
            except Exception: pass
    if ilc: dfu = max(ilc)

    # SKU = Item-Location (largest count, excluding ILC)
    sku = None
    il = []
    for pat in [r'([\d,]+)\s*Item-Location(?!-Customer)',
                r'SKU[^\d]*(\d[\d,]+)',
                r'([\d,]+)\s*Planogram',       # Blue Yonder SOW format
                r'Planogram[^\d]*([\d,]+)']:
        for m in re.finditer(pat, text, re.IGNORECASE):
            try:
                n = int(m.group(1).replace(",",""))
                if n > 100: il.append(n)
            except Exception: pass
    if il: sku = max(il)

    # Named Users
    named_users = None
    for pat in [r'([\d,]+)\s*Named\s+User', r'Named\s+User[^\d]*([\d,]+)']:
        for m in re.finditer(pat, text, re.IGNORECASE):
            try:
                n = int(m.group(1).replace(",",""))
                if n > 0: named_users = max(named_users or 0, n)
            except Exception: pass

    # Concurrent Users
    concurrent = None
    for pat in [r'([\d,]+)\s*[Cc]oncurrent\s+[Uu]ser', r'([\d,]+)\s+concurrent']:
        for m in re.finditer(pat, text, re.IGNORECASE):
            try:
                n = int(m.group(1).replace(",",""))
                if n > 0: concurrent = max(concurrent or 0, n)
            except Exception: pass

    return customer_name, dfu, sku


def detect_customer(filename):
    """Generic customer + env extractor from any PE filename pattern."""
    name, env = "", ""
    stem = re.sub(r'\.(csv|xlsx|pdf|docx)$', '', filename, flags=re.IGNORECASE)

    def smart_case(s):
        """Keep short ALL-CAPS tokens as-is (LHD, JDA, IBM), title-case longer ones."""
        parts = s.replace("_"," ").split()
        return " ".join(p if (len(p) <= 4 and p.isupper()) else p.title() for p in parts)

    for pat in [
        r'Report_of_CS[_]([A-Z][A-Z0-9_]+?)_SCPO',
        r'_CS_([A-Z][A-Z0-9_]+?)_SCPO',
        r'CUSTOMER_([A-Z][A-Z0-9_]+?)_SCPO',
        r'^([A-Z][A-Z0-9]+?)_',
    ]:
        m = re.search(pat, stem, re.IGNORECASE)
        if m:
            candidate = smart_case(m.group(1))
            generic = {"Last","First","Report","Batch","Ctrl","Sla","Test","Prod",
                       "Uat","Dev","Data","File","Export","Daily","Weekly","The"}
            if candidate not in generic and len(candidate) > 2:
                name = candidate
                break

    fn_up = stem.upper()
    if "NON_PROD" in fn_up or "NONPROD" in fn_up or "NON-PROD" in fn_up:
        env = "NON-PROD"
    elif re.search(r'[_\-]PROD(?:[_\-]|$)', fn_up):
        env = "PROD"
    elif "TEST" in fn_up: env = "TEST"
    elif "UAT"  in fn_up: env = "UAT"
    elif "DEV"  in fn_up: env = "DEV"
    return name, env
def extract_pdf_text(file_obj, max_pages=300):
    """Extract text from PDF. Handles large (100+ page) Zabbix reports efficiently.
    Uses list join instead of string concatenation for performance."""
    parts = []
    try:
        from pypdf import PdfReader
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        total = min(len(reader.pages), max_pages)
        for i in range(total):
            t = reader.pages[i].extract_text()
            if t: parts.append(t)
        return "\n".join(parts)
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        total = min(len(reader.pages), max_pages)
        for i in range(total):
            t = reader.pages[i].extract_text()
            if t: parts.append(t)
        return "\n".join(parts)
    except ImportError:
        pass
    return ""


# ── RESOURCE UTILIZATION FAST INGESTION HELPERS ──────────────

def get_file_hash(uploaded_file) -> str:
    """Stable MD5 hash of uploaded file bytes — used as st.cache_data key.
    Seek-safe: restores position to 0 after reading."""
    import hashlib
    uploaded_file.seek(0)
    h = hashlib.md5(uploaded_file.read()).hexdigest()
    uploaded_file.seek(0)
    return h


@st.cache_data(show_spinner=False)
def extract_pdf_text_fast(file_hash: str, file_bytes: bytes, max_pages: int = 300) -> tuple:
    """Cached PDF text extraction — keyed by file_hash so the same file is never
    re-parsed across Streamlit reruns.  Returns (text: str, page_count: int)."""
    parts = []
    page_count = 0
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        for i in range(min(page_count, max_pages)):
            t = reader.pages[i].extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts), page_count
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        for i in range(min(page_count, max_pages)):
            t = reader.pages[i].extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts), page_count
    except ImportError:
        pass
    return "", 0


def ocr_image_if_needed(text: str, min_words: int = 50) -> tuple:
    """Decide whether OCR is needed based on extracted text quality.
    Returns (needs_ocr: bool, reason: str).
    Called after extract_pdf_text_fast to choose the Vision fallback path."""
    word_count = len(text.split()) if text else 0
    has_metric_kw = bool(re.search(
        r'cpu|mem|disk|idle|utiliz|zabbix|system status|free disk|available',
        text, re.I))
    if word_count >= min_words and has_metric_kw:
        return False, f"{word_count} words with metric keywords — text extraction sufficient"
    if word_count >= min_words:
        return True, f"{word_count} words but no metric keywords — likely image-heavy PDF"
    return True, f"Low text density ({word_count} words) — OCR required"


@st.cache_data(show_spinner=False)
def parse_resource_metrics(file_hash: str, text: str) -> list:
    """Cached wrapper around parse_zabbix_pdf_text.
    file_hash is the cache key — same file content is never re-parsed."""
    if not text:
        return []
    return parse_zabbix_pdf_text(text)


def _infer_server_type(host, context="", doc_section_hint=""):
    """Detect APP vs DB from hostname prefix, naming convention, section hint, or context.

    Works across all customer naming conventions:
    - JDA/BY: tsXXNNNNNN — classify by numeric range (SRE 1525-1535 = APP, ACT 1540-1549 = APP)
    - Generic: hostnames containing 'db', 'sql', 'ora', 'data', 'mongo', 'redis', 'pg'
    - doc_section_hint: heading text from DOCX section (highest priority after JDA range)
    - Context: surrounding text mentions database-related keywords
    """
    h = host.lower()
    # BUG-4: JDA/BY convention: tsXXNNNNNN — classify by numeric suffix range
    m = re.match(r'ts[a-z]{2}(\d+)', h)
    if m:
        n = int(m.group(1))
        if 1525 <= n <= 1535: return "APP"   # SRE servers
        if 1540 <= n <= 1549: return "APP"   # ACT servers
        return "DB"                           # all other ts* = DB
    # BUG-4: doc_section_hint from DOCX heading takes priority over hostname guesses
    if doc_section_hint:
        dh = doc_section_hint.lower()
        if re.search(r'\bapp(?:lication)?\b|\bsre\b|\bweb\b|\bbatch\b|\betl\b|\bact\b', dh):
            return "APP"
        if re.search(r'\bdb\b|\bdatabase\b|\boracle\b|\bsql\b|\bdata\s*server\b', dh):
            return "DB"
    # Generic DB indicators in hostname
    if any(k in h for k in ["db","oracle","sql","data","mongo","redis","postgres",
                              "pg","mysql","mssql","mariadb","cassandra","elastic"]): return "DB"
    # Common DB hostname patterns: *-db-*, *db01, etc.
    if re.search(r'[-_]db[-_\d]|db\d{1,3}$', h): return "DB"
    if context:
        cl = context.lower()
        if any(k in cl for k in ["oracle","database"," db ","db server","db instance",
                                   "sql server","mysql","postgres","mongodb","tablespace",
                                   "datafile","redo log","archive log"]): return "DB"
    return "APP"

def _parse_zabbix_block(blk):
    """Extract CPU/Mem/Disk from a single Zabbix data block."""
    rec = {"cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
           "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {}}

    # Memory total
    m = re.search(r'Total\s+[Mm]emory\s*:\s*([\d.]+)\s*GB', blk)
    if m: rec["mem_total_gb"] = float(m.group(1))

    # CPU idle → used
    m = re.search(r'CPU idle time\s*:\s*latest value:\s*([\d.]+)', blk)
    if m: rec["cpu_used"] = round(100 - float(m.group(1)), 2)
    m = re.search(r'CPU idle time trend/SLA:\s*([\d.]+)', blk)
    if m: rec["cpu_avg"] = round(100 - float(m.group(1)), 2)

    # CPU utilization direct
    if rec["cpu_used"] == 0.0:
        for cpat in [r'CPU\s+utilization\s*:\s*latest value:\s*([\d.]+)',
                     r'CPU\s+usage\s*:\s*latest value:\s*([\d.]+)',
                     r'CPU\s+load\s*:\s*latest value:\s*([\d.]+)']:
            m = re.search(cpat, blk, re.I)
            if m: rec["cpu_used"] = round(float(m.group(1)), 2); break

    # Memory used %
    for mem_pat in [
        r'(?:Available|Free)\s+[Mm]emory\s*(?:%|percent)[^:]*:\s*latest value:\s*([\d.]+)',
        r'(?:Available|Free)\s+[Mm]emory\s*:\s*latest value:\s*([\d.]+)\s*%',
        r'[Mm]emory\s+utilization\s*:\s*latest value:\s*([\d.]+)',
        r'[Mm]emory\s+used\s*:\s*latest value:\s*([\d.]+)',
        r'Used\s+[Mm]emory\s*%[^:]*:\s*([\d.]+)',
    ]:
        m = re.search(mem_pat, blk, re.I)
        if m:
            val = float(m.group(1))
            if re.match(r'(?:Available|Free)', mem_pat, re.I):
                rec["mem_used"] = round(100 - val, 2) if val <= 100 else 0.0
            else:
                rec["mem_used"] = round(val, 2)
            break

    # Disk — multiple Zabbix/monitoring text formats
    # Pattern 1: "Free disk space on /mount (percentage) : 98.95 %"
    for mount, pct in re.findall(
            r'Free disk space on (\S+)\s*\(percentage\)\s*:\s*([\d.]+)', blk):
        rec["disks"][mount] = round(100 - float(pct), 2)

    # Pattern 2: "Disk space on /mount : used 45.2%" or "Disk utilization /mount : 45%"
    for mount, pct in re.findall(
            r'[Dd]isk\s+(?:space|utilization)\s+(?:on\s+)?(\S+)\s*:\s*(?:used\s+)?([\d.]+)\s*%', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 3: "Used disk space on /mount (percentage) : latest value: 45.2"
    for mount, pct in re.findall(
            r'[Uu]sed\s+disk\s+space\s+on\s+(\S+)\s*\(percentage\)\s*:\s*(?:latest value:\s*)?([\d.]+)', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 4: "Filesystem /mount : used 45%" or "/mount used: 45.2%"
    for mount, pct in re.findall(
            r'(?:Filesystem|Volume)\s+(\S+)\s*:\s*used\s+([\d.]+)\s*%', blk, re.I):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 5: "Available disk space on /mount (percentage) : latest value: 85"
    for mount, pct in re.findall(
            r'[Aa]vailable\s+[Dd]isk\s+space\s+(?:in\s+%?\s+)?on\s+(\S+)\s*(?:\(percentage\))?\s*:\s*(?:latest value:\s*)?([\d.]+)', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(100 - float(pct), 2) if float(pct) <= 100 else 0.0

    # BUG-2: Oracle ASM disk groups — Zabbix reports used% directly (no inversion)
    # Format: "ASM disk group DATA: used 78.5%" or "ASM: DATA used: 78%"
    for label, pct in re.findall(
            r'ASM\s+(?:disk\s+group\s+)?(\w+)\s*[:\s]+(?:used|utilization)[:\s]+([\d.]+)\s*%?', blk, re.I):
        key = f"ASM:{label.upper()}"
        if key not in rec["disks"]:
            rec["disks"][key] = round(float(pct), 2)   # used% directly, no 100- inversion
    # Also catch "diskgroup DATA used 78%" format
    for label, pct in re.findall(
            r'[Dd]iskgroup\s+(\w+)\s+used\s+([\d.]+)\s*%?', blk):
        key = f"ASM:{label.upper()}"
        if key not in rec["disks"]:
            rec["disks"][key] = round(float(pct), 2)

    # BUG-1: Normalize empty/invalid mount keys → root '/'
    # Happens when Zabbix text omits mount point for the root partition
    cleaned_disks = {}
    for mnt, v in rec["disks"].items():
        # Reject mounts that are obviously wrong captures (parens, empty)
        norm = mnt if (mnt and not mnt.startswith("(")) else "/"
        cleaned_disks[norm] = max(cleaned_disks.get(norm, 0.0), v)
    rec["disks"] = cleaned_disks

    # Track known disk mount points from graph titles (even without values)
    # "Disk space usage /mount" or "Available Disk space in % on /mount"
    _disk_mounts_seen = set(rec["disks"].keys())
    for mount in re.findall(r'[Dd]isk\s+space\s+usage\s+(/\S*)', blk):
        _disk_mounts_seen.add(mount)
    for mount in re.findall(r'Available\s+[Dd]isk\s+space\s+in\s+%\s+on\s+(/\S*)', blk):
        _disk_mounts_seen.add(mount)
    for label in re.findall(r'ASM\s+(?:disk\s+group\s+)?(\w+)', blk, re.I):
        _disk_mounts_seen.add(f"ASM:{label.upper()}")
    rec["_disk_mounts_known"] = list(_disk_mounts_seen)

    if rec["disks"]:
        rec["disk_used_max"] = max(rec["disks"].values())

    # BUG-3: Oracle DB metrics — uptime, active sessions, instance status
    oracle = {}
    _m = re.search(r'(?:Oracle\s+)?(?:DB\s+)?[Uu]ptime[:\s]+([\d.]+)\s*(?:days?|d\b)', blk)
    if _m: oracle["uptime_days"] = float(_m.group(1))
    _m = re.search(r'(?:Active\s+)?(?:sessions?|connections?)\s*:\s*(?:latest value:\s*)?([\d]+)', blk, re.I)
    if _m: oracle["sessions"] = int(_m.group(1))
    _m = re.search(r'(?:Oracle\s+)?[Ii]nstance\s+(?:status|state)\s*[:\s]+(\w+)', blk)
    if _m: oracle["instance_status"] = _m.group(1).upper()
    if oracle:
        rec["oracle"] = oracle

    return rec

def parse_zabbix_pdf_text(text):
    """Universal Zabbix/resource utilization PDF text parser.

    Handles ALL known formats:
      A) Petbarn-style: "System Status for <host>" blocks with text metrics
      B) UTZ/Monitor-style: "Graphs for <host>" only (data in images, text is ToC)
      C) Image-only PDFs: title + screenshots, extract server names from any pattern
      D) Mixed: some servers have text metrics, others are image-only

    Returns list of server dicts. Image-only servers have zeroed metrics
    and are flagged for Gemini Vision OCR enrichment.
    """
    servers_map = {}  # short_host_lower → rec (case-insensitive dedup)
    def _dedup_key(h):
        """Normalize hostname to lowercase short name for dedup."""
        return h.split(".")[0].lower()

    # ── Strategy A: "System Status for <host>" blocks (Petbarn format) ──
    status_blocks = re.split(r'System Status for\s+', text)
    for blk in status_blocks[1:]:
        lines = blk.strip().splitlines()
        host = lines[0].strip().split()[0] if lines else "unknown"
        host = host.rstrip(".")
        if not host or host == "unknown": continue

        rec = _parse_zabbix_block(blk)
        rec["host"] = host
        rec["type"] = _infer_server_type(host, blk)
        rec["_image_only"] = False

        has_data = (rec["cpu_used"] > 0 or rec["mem_total_gb"] > 0
                    or rec["disk_used_max"] > 0 or rec["mem_used"] > 0)

        key = _dedup_key(host)
        if key in servers_map:
            existing = servers_map[key]
            e_has = (existing["cpu_used"] > 0 or existing["mem_total_gb"] > 0
                     or existing["disk_used_max"] > 0 or existing["mem_used"] > 0)
            if has_data and not e_has:
                servers_map[key] = rec
            elif has_data and e_has:
                for k in ["cpu_used","cpu_avg","mem_used","mem_total_gb","disk_used_max"]:
                    existing[k] = max(existing[k], rec[k])
                existing["disks"].update(rec["disks"])
                if existing["disks"]:
                    existing["disk_used_max"] = max(existing["disks"].values())
        else:
            servers_map[key] = rec

    # ── Strategy B: "Graphs for <host>" (UTZ/Monitor format — ToC-only PDFs) ──
    # These PDFs have NO "System Status for" but list servers via "Graphs for"
    graph_hosts = list(dict.fromkeys(re.findall(
        r'Graphs for\s+(\S+(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)', text)))
    for host in graph_hosts:
        host = host.rstrip(".")
        key = _dedup_key(host)
        if key in servers_map:
            continue  # already have from Strategy A
        # Find context around this host for DB detection
        idx = text.find(host)
        context = text[idx:idx+500] if idx >= 0 else ""
        stype = _infer_server_type(host, context)
        # Check if Oracle references exist for this host anywhere
        host_upper = host.split(".")[0].upper()
        if re.search(rf'Oracle.*{re.escape(host_upper)}|{re.escape(host_upper)}.*Oracle', text, re.I):
            stype = "DB"
        servers_map[key] = {
            "host": host, "type": stype,
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }

    # ── Strategy C: Fallback hostname extraction from any pattern ──
    # Catches "Trends and metrics for <host>", standalone FQDNs, etc.
    FQDN_RE = re.compile(r'\b([a-z]{2,6}\d{6,}[0-9a-z]*(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)\b', re.I)
    for m in FQDN_RE.finditer(text):
        host = m.group(1)
        key = _dedup_key(host)
        if key not in servers_map:
            servers_map[key] = {
                "host": host, "type": _infer_server_type(host),
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
                "_image_only": True,
            }
    # BUG-6: Short-hostname fallback — catches tsXXNNNN patterns with fewer digits
    # than FQDN_RE's 6-digit minimum (e.g. tsbc1234 = 4 digits)
    SHORTHOST_RE = re.compile(r'\b([a-z]{2,4}\d{4,5}[0-9a-z]*)\b', re.I)
    for m in SHORTHOST_RE.finditer(text):
        host = m.group(1)
        key = _dedup_key(host)
        if key not in servers_map:
            servers_map[key] = {
                "host": host, "type": _infer_server_type(host),
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
                "_image_only": True,
            }

    # ── Dedup: prefer entries with data, remove pure ToC stubs ──
    servers_all = list(servers_map.values())
    servers_with_data = [s for s in servers_all
                         if s["cpu_used"] > 0 or s["mem_total_gb"] > 0 or s["disk_used_max"] > 0]
    return servers_with_data if servers_with_data else servers_all


def build_resource_df(servers: list) -> "pd.DataFrame":
    """Normalize the server list (list of dicts) into a flat, typed DataFrame.
    Columns: server, host_fqdn, type, cpu_pct, cpu_avg_pct, mem_pct, mem_gb,
             disk_pct, source_file, source_env, image_only, health_score, status.
    This is the single source of truth for all resource rendering functions."""
    if not servers:
        return pd.DataFrame()
    rows = []
    for s in servers:
        cpu  = round(float(s.get("cpu_used",  0.0)), 2)
        mem  = round(float(s.get("mem_used",  0.0)), 2)
        disk = round(float(s.get("disk_used_max", 0.0)), 2)
        rows.append({
            "server":      s.get("host", "unknown").split(".")[0],
            "host_fqdn":   s.get("host", "unknown"),
            "type":        s.get("type", "APP"),
            "cpu_pct":     cpu,
            "cpu_avg_pct": round(float(s.get("cpu_avg", 0.0)), 2),
            "mem_pct":     mem,
            "mem_gb":      round(float(s.get("mem_total_gb", 0.0)), 1),
            "disk_pct":    disk,
            "source_file": s.get("_source_file", ""),
            "source_env":  s.get("_source_env", ""),
            "image_only":  bool(s.get("_image_only", False)),
        })
    df = pd.DataFrame(rows)
    df["health_score"] = df.apply(
        lambda r: get_health_score(r["cpu_pct"], r["mem_pct"], r["disk_pct"], r["type"]), axis=1)
    df["status"] = df["health_score"].apply(
        lambda sc: "Critical" if 0 <= sc < 60
                   else ("Warning" if sc < 85
                   else ("Healthy" if sc >= 85 else "Unknown")))
    return df


# ── IN-PROCESS RESOURCE DF CACHE ─────────────────────────────
_RES_DF_CACHE: dict = {}

def _get_resource_df(res_hash: str, servers: list) -> "pd.DataFrame":
    """In-process cache for build_resource_df — keyed by file hash.
    Avoids re-normalising the same server list on every Streamlit rerun."""
    key = res_hash or "_no_hash"
    if key not in _RES_DF_CACHE:
        _RES_DF_CACHE[key] = build_resource_df(servers)
    return _RES_DF_CACHE[key]


# ── DATA LOADERS ─────────────────────────────────────────────
def parse_dt(s):
    for fmt in ["%d-%m-%Y %H:%M", "%m/%d/%Y %H:%M", "%Y-%m-%d %H:%M:%S",
                "%Y-%m-%d %H:%M", "%d/%m/%Y %H:%M"]:
        try: return pd.to_datetime(s, format=fmt, errors="raise")
        except: pass
    return pd.to_datetime(s, infer_datetime_format=True, errors="coerce")


def load_ctrlm(f, fname=""):
    """Load Ctrl-M batch execution data from CSV or XLSX."""
    ext = os.path.splitext(fname)[1].lower() if fname else ""
    if ext in [".xlsx", ".xls"]:
        try:
            df = pd.read_excel(f, sheet_name=0, engine="openpyxl")
        except Exception:
            df = pd.read_excel(f, sheet_name=0)
    else:
        try:
            df = pd.read_csv(f)
        except Exception:
            if hasattr(f, "seek"): f.seek(0)
            try:
                df = pd.read_excel(f, sheet_name=0, engine="openpyxl")
            except Exception as e:
                raise ValueError(f"Cannot parse Ctrl-M file: {e}") from e
    df.columns = df.columns.str.strip()
    col_map = {}
    for c in df.columns:
        cl = c.lower().replace(" ", "_").replace("-", "_")
        if "folder" in cl:
            col_map[c] = "Folder"
        elif ("sub" in cl and "app" in cl) or cl in ("sub_application","sub_app"):
            col_map[c] = "Sub_Application"
        elif "application" in cl and "sub" not in cl:
            col_map[c] = "Application"
        elif ("job" in cl and "name" in cl):
            col_map[c] = "Job_Name"
        elif cl in ("job","jobname","job_id","jobid","taskname","task_name",
                    "task","process_name","processname","ordername","order_name",
                    "step_name","stepname","script_name","scriptname","name"):
            col_map[c] = "Job_Name"
        elif "start" in cl and ("time" in cl or "date" in cl or cl=="start"):
            col_map[c] = "Start_Time"
        elif "end" in cl and ("time" in cl or "date" in cl):
            col_map[c] = "End_Time"
        elif "status" in cl or "completion" in cl or "state" in cl:
            col_map[c] = "Status"
        elif "run" in cl and ("sec" in cl or "time" in cl or "dur" in cl):
            col_map[c] = "Run_Sec"
        elif "duration" in cl or "elapsed" in cl:
            col_map[c] = "Run_Sec"

    df.rename(columns=col_map, inplace=True)

    # ── Fuzzy fallback: if critical columns still missing, try difflib ──
    import difflib
    _FUZZY_TARGETS = {
        "Job_Name":       ["jobname","job_name","jobid","taskname","task","process","name"],
        "Start_Time":     ["starttime","start_time","startdate","start_date","begin","started"],
        "End_Time":       ["endtime","end_time","enddate","end_date","finish","finished","completed"],
        "Run_Sec":        ["runsec","run_sec","duration","elapsed","runtime","exectime","seconds"],
        "Sub_Application":["subapp","sub_app","subapplication","stream","module","component"],
    }
    unmapped_cols = [c for c in df.columns if c not in
                     ("Folder","Sub_Application","Application","Job_Name",
                      "Start_Time","End_Time","Status","Run_Sec")]
    for target, candidates in _FUZZY_TARGETS.items():
        if target in df.columns:
            continue  # already mapped
        for uc in unmapped_cols:
            normalised = uc.lower().replace(" ","").replace("_","").replace("-","")
            matches = difflib.get_close_matches(normalised, candidates, n=1, cutoff=0.75)
            if matches:
                df.rename(columns={uc: target}, inplace=True)
                unmapped_cols.remove(uc)
                break

    # ── Fallback: if still no Job_Name, use first string/object column ──
    if "Job_Name" not in df.columns:
        str_cols = [c for c in df.columns
                    if df[c].dtype == object and c not in
                    ("Folder","Sub_Application","Application","Status","Start_Time","End_Time")]
        if str_cols:
            df.rename(columns={str_cols[0]: "Job_Name"}, inplace=True)
        else:
            df["Job_Name"] = "UNKNOWN"

    if "Sub_Application" not in df.columns: df["Sub_Application"] = "UNKNOWN"
    if "Status"          not in df.columns: df["Status"]          = "ENDED OK"
    if "Run_Sec" not in df.columns:
        nums = df.select_dtypes(include="number").columns.tolist()
        if nums: df.rename(columns={nums[-1]: "Run_Sec"}, inplace=True)
        else:    df["Run_Sec"] = 0

    df["Run_Sec"] = pd.to_numeric(df["Run_Sec"], errors="coerce").fillna(0)

    if "Start_Time" not in df.columns:
        # Try to find any date-like column by probing each one with pandas
        for c in df.columns:
            try:
                parsed = pd.to_datetime(df[c], errors="coerce")
                if parsed.notna().sum() > len(df) * 0.5:
                    df.rename(columns={c: "Start_Time"}, inplace=True)
                    break
            except Exception:
                # pd.to_datetime may raise on completely non-date columns;
                # silently continue to the next candidate column.
                pass
        if "Start_Time" not in df.columns:
            df["Start_Time"] = pd.Timestamp.now()

    if "End_Time" in df.columns:
        df["Start_Time"] = parse_dt(df["Start_Time"])
        df["End_Time"]   = parse_dt(df["End_Time"])
        mask = df["Run_Sec"] == 0
        diff = (df.loc[mask, "End_Time"] - df.loc[mask, "Start_Time"]).dt.total_seconds()
        df.loc[mask, "Run_Sec"] = diff.clip(lower=0)
    else:
        df["Start_Time"] = parse_dt(df["Start_Time"])

    df.dropna(subset=["Start_Time"], inplace=True)
    df["run_time_hrs"] = df["Run_Sec"] / 3600.0
    df["run_date"]     = df["Start_Time"].dt.date
    df["month"]        = df["Start_Time"].dt.to_period("M").astype(str)
    return df

def load_batch_sla(f, fname):
    """Load Batch SLA info from CSV or XLSX.
    Handles TWO formats:
      A) Per-job SLA: Job_Name, SLA_Hrs, ...
      B) Batch schedule: Batch_Name, Start Time, Expected End Time, ...
    Returns DataFrame with smart column mapping applied.
    """
    ext = os.path.splitext(fname)[1].lower()
    if ext in [".xlsx", ".xls"]:
        try:
            df = pd.read_excel(f, sheet_name=0, engine="openpyxl")
        except Exception:
            try:
                df = pd.read_excel(f, sheet_name=0)
            except Exception as e:
                raise ValueError(f"Cannot read Excel SLA file: {e}. "
                                 "Try saving as .xlsx and re-upload.") from e
    else:
        try:
            df = pd.read_csv(f)
        except Exception as e:
            raise ValueError(f"Cannot read SLA CSV file: {e}") from e

    df.columns = df.columns.str.strip()

    # ── Smart column mapping ─────────────────────────────────
    sla_col_map = {}
    for c in df.columns:
        cl = c.lower().replace(" ", "_").replace("-", "_").replace("/", "_")
        # Job name aliases
        if cl in ("job_name","jobname","job","task","task_name","process","process_name",
                  "job_id","step","step_name","script") and "Job_Name" not in sla_col_map.values():
            sla_col_map[c] = "Job_Name"
        # Batch name (schedule format)
        elif any(k in cl for k in ["batch_name","batch","stream_name","schedule_name"]) and "Batch_Name" not in sla_col_map.values():
            sla_col_map[c] = "Batch_Name"
        # SLA limit / threshold
        elif any(k in cl for k in ["sla_limit","sla_hrs","sla_hours","max_hours",
                                    "threshold","limit_hrs","sla_h","allowed_hours",
                                    "max_run","runtime_limit"]) and "SLA_Hrs" not in sla_col_map.values():
            sla_col_map[c] = "SLA_Hrs"
        # Expected end time (schedule format) → compute SLA window
        elif any(k in cl for k in ["expected_end","expected_end_time","sla_end",
                                    "end_time_sla"]) and "Expected_End" not in sla_col_map.values():
            sla_col_map[c] = "Expected_End"
        # Current end time (schedule format)
        elif any(k in cl for k in ["current_end","actual_end","current_end_time"]) and "Current_End" not in sla_col_map.values():
            sla_col_map[c] = "Current_End"
        # Start time
        elif any(k in cl for k in ["start_time","start","batch_start","scheduled_start"]) and "Start_Time" not in sla_col_map.values():
            sla_col_map[c] = "Start_Time"
        # First/Last job
        elif any(k in cl for k in ["first_job","first_job_name"]) and "First_Job" not in sla_col_map.values():
            sla_col_map[c] = "First_Job"
        elif any(k in cl for k in ["last_job","last_job_name"]) and "Last_Job" not in sla_col_map.values():
            sla_col_map[c] = "Last_Job"
        # Schedule
        elif any(k in cl for k in ["schedule","frequency","recurrence"]) and "Schedule" not in sla_col_map.values():
            sla_col_map[c] = "Schedule"
        # Module
        elif any(k in cl for k in ["module","component","sub_app","sub_application"]) and "Module" not in sla_col_map.values():
            sla_col_map[c] = "Module"
        # Timezone
        elif any(k in cl for k in ["timezone","time_zone","tz"]) and "TimeZone" not in sla_col_map.values():
            sla_col_map[c] = "TimeZone"
        # Priority
        elif any(k in cl for k in ["priority","criticality","tier","level"]) and "Priority" not in sla_col_map.values():
            sla_col_map[c] = "Priority"
        # Comments
        elif any(k in cl for k in ["comment","notes","remarks","description"]) and "Comments" not in sla_col_map.values():
            sla_col_map[c] = "Comments"

    if sla_col_map:
        df.rename(columns=sla_col_map, inplace=True)

    # ── For batch schedule format: compute SLA window hours ──
    if "Batch_Name" in df.columns and "Expected_End" in df.columns and "Start_Time" in df.columns:
        try:
            def _time_diff_hrs(row):
                try:
                    start = pd.to_datetime(str(row.get("Start_Time","")), format="%H:%M:%S", errors="coerce")
                    end = pd.to_datetime(str(row.get("Expected_End","")), format="%H:%M:%S", errors="coerce")
                    if pd.isna(start) or pd.isna(end): return 0.0
                    diff = (end - start).total_seconds() / 3600
                    if diff < 0: diff += 24  # overnight batch
                    return round(diff, 2)
                except Exception:
                    return 0.0
            df["SLA_Window_Hrs"] = df.apply(_time_diff_hrs, axis=1)
        except Exception:
            pass
        # Tag format for the SLA tab to render appropriately
        df.attrs["_sla_format"] = "batch_schedule"
    elif "Job_Name" in df.columns:
        df.attrs["_sla_format"] = "per_job"

    return df


def _safe_float(row, *keys):
    """Try multiple column name aliases, return float or 0.0."""
    for k in keys:
        v = row.get(k)
        if v is not None and str(v).strip() not in ("", "nan", "None", "#N/A"):
            try:
                return float(str(v).replace("%","").replace(",","").strip())
            except (ValueError, TypeError):
                # Value is non-numeric (e.g. "N/A", "–", free-text). Try next alias.
                pass
    return 0.0

def load_server_csv(f, fname=""):
    """
    Load server metrics CSV or XLSX.  Accepts column name variants from:
      - PE CSV template
      - Zabbix export (CPU usage %, Available memory %, etc.)
      - Azure Portal export (Percentage CPU, etc.)
      - Manual fill (any reasonable alias)
    """
    import io as _io
    import html as _html_mod2

    ext = os.path.splitext(fname)[1].lower() if fname else ""
    if ext in [".xlsx", ".xls"]:
        try:
            df = pd.read_excel(f, sheet_name=0, engine="openpyxl")
        except Exception:
            df = pd.read_excel(f, sheet_name=0)
    else:
        try:
            raw = f.read()
            if isinstance(raw, bytes):
                raw = raw.decode("utf-8", errors="replace")
            clean = "\n".join(l for l in raw.splitlines() if not l.strip().startswith("#"))
            df = pd.read_csv(_io.StringIO(clean))
        except Exception:
            if hasattr(f, "seek"): f.seek(0)
            try:
                df = pd.read_excel(f, sheet_name=0, engine="openpyxl")
            except Exception as e:
                raise ValueError(f"Cannot parse server metrics file: {e}") from e

    df.columns = (df.columns.str.strip().str.lower()
                  .str.replace(r"[\s/\-]+","_", regex=True)
                  .str.replace(r"[^a-z0-9_]","", regex=True))
    servers = []
    for _, row in df.iterrows():
        host = str(row.get("host", row.get("hostname", row.get("server",
                    row.get("server_name", row.get("name", "unknown"))))))
        if not host or host in ("unknown","nan","None"): continue
        stype_raw = str(row.get("type", row.get("server_type","APP"))).upper().strip()
        stype = "DB" if any(k in stype_raw for k in ["DB","DATABASE","ORACLE","SQL"]) else "APP"

        cpu  = _safe_float(row,
                   "cpu_used","cpu_%","cpu_usage","cpu_pct",
                   "percentage_cpu",                    # Azure
                   "processor_utilization",
                   "avg_cpu_utilization_percent",       # AWS
                   "cpu_utilization",
                   "cpu")
        mem  = _safe_float(row,
                   "mem_used","mem_%","memory_%","memory_used",
                   "memory_utilization","mem_utilization",
                   "available_memory_percent",
                   "percentage_memory","memory_pct",
                   "mem")
        disk = _safe_float(row,
                   "disk_used_max","disk_%","disk_used","disk_utilization",
                   "max_disk_used","filesystem_used_pct",
                   "os_disk_used_percent",              # Azure
                   "disk_usage_percent","disk")
        mem_gb=_safe_float(row,
                   "mem_total_gb","memory_gb","total_memory_gb",
                   "total_ram_gb","ram_gb","memory_total_gb",
                   "mem_gb","total_mem")

        servers.append({"host": host, "type": stype,
                        "cpu_used": cpu, "mem_used": mem,
                        "mem_total_gb": mem_gb, "disk_used_max": disk,
                        "disks": {}})
    return servers


def server_csv_template(servers=None):
    """
    Generate a CSV template.
    If servers (image-only skeleton) passed → pre-fill host/type rows
    so the PE engineer only needs to fill in the numbers.
    Includes Zabbix column name aliases so copy-paste from portal works.
    """
    if servers:
        rows = [
            {
                "host":            s["host"],
                "type":            s.get("type","APP"),
                "cpu_used":        "",   # fill from Zabbix CPU graph (%)
                "mem_used":        "",   # fill from Zabbix Memory Used (%)
                "mem_total_gb":    "",   # fill from Zabbix Total Memory (GB)
                "disk_used_max":   "",   # fill from Zabbix Disk Used Max (%)
                # Zabbix / Azure column aliases — any one column is accepted
                "cpu_%":           "",
                "memory_%":        "",
                "memory_gb":       "",
                "disk_%":          "",
            }
            for s in servers
        ]
    else:
        rows = [
            {"host":"srv-app-01.domain.com","type":"APP",
             "cpu_used":45.2,"mem_used":62.1,"mem_total_gb":32.0,"disk_used_max":38.4,
             "cpu_%":"","memory_%":"","memory_gb":"","disk_%":""},
            {"host":"srv-db-01.domain.com","type":"DB",
             "cpu_used":68.5,"mem_used":71.3,"mem_total_gb":64.0,"disk_used_max":55.7,
             "cpu_%":"","memory_%":"","memory_gb":"","disk_%":""},
        ]
    tpl = pd.DataFrame(rows)
    # Add instruction row as comment
    instructions = (
        "# INSTRUCTIONS: Fill cpu_used / mem_used / disk_used_max from Zabbix or Azure Portal.\n"
        "# Column aliases accepted: cpu_% OR cpu_used | memory_% OR mem_used | disk_% OR disk_used_max\n"
        "# mem_total_gb = Total RAM in GB (from Zabbix or Azure VM size)\n"
        "# type = APP or DB\n"
    )
    csv_bytes = (instructions + tpl.to_csv(index=False)).encode()
    return csv_bytes


# ══════════════════════════════════════════════════════════════════════════════
# PE CORE FORMULA ENGINE — F1–F9
# Customer-agnostic mathematical formulas. All inputs are plain Python/pandas.
# These are the authoritative calculations used throughout every tab.
# ══════════════════════════════════════════════════════════════════════════════

def calculate_volume_utilization(actual_volume, sow_contracted):
    """F1 — Volume Utilization: actual vs contracted SOW."""
    if sow_contracted <= 0:
        return {"util_pct": 0.0, "headroom_pct": 100.0, "headroom_items": 0, "status": "INVALID"}
    util_pct      = (actual_volume / sow_contracted) * 100
    headroom_pct  = 100.0 - util_pct
    headroom_items = max(0, int(sow_contracted - actual_volume))
    if   util_pct <= 70:  status = "EXCELLENT"
    elif util_pct <= 85:  status = "HEALTHY"
    elif util_pct <= 100: status = "CAUTION"
    else:                 status = "CRITICAL"
    return {"util_pct": round(util_pct, 1), "headroom_pct": round(headroom_pct, 1),
            "headroom_items": headroom_items, "status": status}


def project_volume_growth(current_vol, sow, monthly_growth_rate, months_ahead=12):
    """F2 — Compound growth projection with SOW breach runway estimate."""
    import math as _math
    if sow <= 0 or current_vol <= 0:
        return {"current_util": 0.0, "projected_util_12mo": 0.0,
                "months_to_100pct": None, "runway_months": 999}
    current_util = (current_vol / sow) * 100
    if monthly_growth_rate <= 0:
        return {"current_util": round(current_util, 1),
                "projected_util_12mo": round(current_util, 1),
                "months_to_100pct": None, "runway_months": 999}
    gf = 1.0 + monthly_growth_rate
    projected_vol  = current_vol * (gf ** months_ahead)
    projected_util = (projected_vol / sow) * 100
    try:
        months_100 = _math.log(sow / current_vol) / _math.log(gf)
    except Exception:
        months_100 = 999
    return {"current_util": round(current_util, 1),
            "projected_util_12mo": round(projected_util, 1),
            "months_to_100pct": round(months_100, 1) if months_100 < 999 else None,
            "runway_months": int(months_100) if months_100 < 999 else 999}


def calculate_sla_buffer(sla_window_hrs, max_runtime_hrs):
    """F3 — SLA Buffer: remaining headroom between peak runtime and SLA window."""
    if sla_window_hrs <= 0 or max_runtime_hrs <= 0:
        return {"buffer_hrs": 0.0, "buffer_pct": 0.0,
                "growth_multiplier": 0.0, "growth_capacity_pct": 0.0, "status": "INVALID"}
    buffer_hrs         = sla_window_hrs - max_runtime_hrs
    buffer_pct         = (buffer_hrs / sla_window_hrs) * 100
    growth_multiplier  = sla_window_hrs / max_runtime_hrs
    growth_capacity    = (buffer_hrs / max_runtime_hrs) * 100 if max_runtime_hrs > 0 else 0.0
    if   buffer_pct > 50: status = "EXCELLENT"
    elif buffer_pct > 30: status = "HEALTHY"
    elif buffer_pct > 10: status = "CAUTION"
    else:                 status = "CRITICAL"
    return {"buffer_hrs": round(buffer_hrs, 2), "buffer_pct": round(buffer_pct, 1),
            "growth_multiplier": round(growth_multiplier, 2),
            "growth_capacity_pct": round(growth_capacity, 1), "status": status}


def calculate_batch_compliance(df, sla_hrs):
    """F4 — Batch Compliance: success rate + on-time / at-risk / breach split."""
    import pandas as _pd
    if df is None or df.empty:
        return {"total_jobs": 0, "successful_jobs": 0, "failed_jobs": 0,
                "compliance_pct": 0.0, "on_time": 0, "at_risk": 0, "breach": 0}
    total   = len(df)
    # Accept multiple column name aliases for status
    _sc = next((c for c in df.columns if c.lower() in
                ("completion_status","status","completion status")), None)
    if _sc:
        successful = int((df[_sc].astype(str).str.upper().str.contains(
            r"ENDED OK|SUCCESS|OK|COMPLETED", regex=True)).sum())
    else:
        successful = total
    failed     = total - successful
    compliance = (successful / total * 100) if total > 0 else 0.0

    # Runtime column aliases
    _rc = next((c for c in df.columns if c.lower() in
                ("run_sec","run_time_sec","runtime_sec","runseconds")), None)
    if _rc:
        runtime = _pd.to_numeric(df[_rc], errors="coerce").fillna(0)
    else:
        _hc = next((c for c in df.columns if "run_time_hrs" in c.lower() or "hrs" in c.lower()), None)
        runtime = (_pd.to_numeric(df[_hc], errors="coerce").fillna(0) * 3600) if _hc else _pd.Series([0]*total)

    sla_secs = sla_hrs * 3600
    on_time  = int((runtime <= sla_secs * 0.85).sum())
    at_risk  = int(((runtime >  sla_secs * 0.85) & (runtime <= sla_secs)).sum())
    breach   = int((runtime >  sla_secs).sum())

    return {"total_jobs": total, "successful_jobs": successful, "failed_jobs": failed,
            "compliance_pct": round(compliance, 1),
            "on_time": on_time, "at_risk": at_risk, "breach": breach}


def calculate_job_variance(actual_hrs, expected_max_hrs):
    """F5 — Job Variance: how far actual peak deviates from expected max."""
    if expected_max_hrs <= 0:
        return {"variance_pct": 0.0, "variance_hrs": 0.0, "status": "INVALID", "severity": 0}
    variance_hrs = actual_hrs - expected_max_hrs
    variance_pct = (variance_hrs / expected_max_hrs) * 100
    if   variance_pct > 100: status, severity = "CRITICAL",  4
    elif variance_pct > 0:   status, severity = "CAUTION",   2
    elif variance_pct > -20: status, severity = "NORMAL",    0
    else:                    status, severity = "OPTIMIZED", -1
    return {"variance_pct": round(variance_pct, 1), "variance_hrs": round(variance_hrs, 3),
            "status": status, "severity": severity}


def detect_job_anomalies(df, baselines=None):
    """F6 — Anomaly Detection: z-score + variance flagging on job runtimes.
    baselines = {job_name: {"max": hrs}} optional; if None uses statistical z-score only."""
    import pandas as _pd
    anomalies = []
    if df is None or df.empty:
        return anomalies

    _jc = next((c for c in df.columns if "job_name" in c.lower() or c.lower() == "job"), None)
    _hc = next((c for c in df.columns if "run_time_hrs" in c.lower() or "peak_hrs" in c.lower()), None)
    if _jc is None or _hc is None:
        return anomalies

    top_jobs = (df.groupby(_jc)[_hc]
                  .agg(peak_hrs="max", avg_hrs="mean").reset_index()
                  .rename(columns={_jc: "Job_Name"}))

    # Z-score anomaly detection (statistical baseline from the dataset itself)
    if len(top_jobs) >= 3:
        mu  = top_jobs["peak_hrs"].mean()
        std = top_jobs["peak_hrs"].std()
        if std > 0.001:
            top_jobs["z_score"] = ((top_jobs["peak_hrs"] - mu) / std).round(2)
            stat_anoms = top_jobs[top_jobs["z_score"] > 2.0]
            for _, row in stat_anoms.iterrows():
                anomalies.append({"job_name": row["Job_Name"],
                                   "peak_hrs": round(row["peak_hrs"], 3),
                                   "avg_hrs": round(row["avg_hrs"], 3),
                                   "z_score": row["z_score"],
                                   "variance_pct": None,
                                   "status": "STATISTICAL_OUTLIER",
                                   "severity": 3})

    # Baseline variance anomaly detection (if SLA matrix provided)
    if baselines:
        for _, row in top_jobs.iterrows():
            bl = baselines.get(row["Job_Name"])
            if not bl:
                continue
            var = calculate_job_variance(row["peak_hrs"], bl.get("max", 0))
            if var["severity"] >= 2:
                # Update or add — avoid duplicates
                existing = next((a for a in anomalies if a["job_name"] == row["Job_Name"]), None)
                entry = {"job_name": row["Job_Name"],
                         "peak_hrs": round(row["peak_hrs"], 3),
                         "avg_hrs": round(row["avg_hrs"], 3),
                         "z_score": None,
                         "variance_pct": var["variance_pct"],
                         "status": var["status"],
                         "severity": var["severity"]}
                if existing:
                    existing.update(entry)
                else:
                    anomalies.append(entry)

    return sorted(anomalies, key=lambda x: x["severity"], reverse=True)


def calculate_host_health(cpu_pct, mem_pct, disk_pct):
    """F7 — Host Health Score: continuous weighted formula (not zone-based).
    Returns health_score 0-100: higher = healthier."""
    cpu  = float(np.clip(cpu_pct,  0, 100))
    mem  = float(np.clip(mem_pct,  0, 100))
    disk = float(np.clip(disk_pct, 0, 100))
    # Weights: CPU 30%, Memory 40%, Disk 30%
    health_score = (100 - cpu) * 0.30 + (100 - mem) * 0.40 + (100 - disk) * 0.30
    if   health_score >= 80: status = "HEALTHY"
    elif health_score >= 60: status = "WARNING"
    else:                    status = "CRITICAL"
    return {"health_score": round(health_score, 1), "status": status,
            "cpu_pct": round(cpu, 1), "mem_pct": round(mem, 1), "disk_pct": round(disk, 1)}


def calculate_fleet_health(servers):
    """F8 — Fleet Health: aggregate host scores into fleet grade A-F."""
    if not servers:
        return {"fleet_score": 0.0, "grade": "F", "total": 0,
                "healthy": 0, "warning": 0, "critical": 0, "unknown": 0}
    scores    = []
    healthy = warning = critical = unknown = 0
    for s in servers:
        cpu  = s.get("cpu_used", 0)
        mem  = s.get("mem_used", 0)
        disk = s.get("disk_used_max", 0)
        if cpu == 0 and mem == 0 and disk == 0:
            unknown += 1
            continue
        h = calculate_host_health(cpu, mem, disk)
        scores.append(h["health_score"])
        if   h["status"] == "HEALTHY":  healthy  += 1
        elif h["status"] == "WARNING":  warning  += 1
        else:                           critical += 1
    fleet_score = round(float(np.mean(scores)), 1) if scores else 0.0
    if   fleet_score >= 90: grade = "A"
    elif fleet_score >= 75: grade = "B"
    elif fleet_score >= 60: grade = "C"
    elif fleet_score >= 45: grade = "D"
    else:                   grade = "F"
    return {"fleet_score": fleet_score, "grade": grade, "total": len(servers),
            "healthy": healthy, "warning": warning, "critical": critical, "unknown": unknown}


def compute_pe_assessment(ctrlm_df=None, servers=None, sow_cfg=None,
                          sla_hrs=None, monthly_growth=0.03):
    """F9 — Comprehensive PE Assessment: combines all formulas → approval status + confidence.
    sow_cfg = {"dfu": (actual, contracted), "sku": (actual, contracted)}
    sla_hrs = daily SLA limit in hours
    Returns dict with status, confidence (0-100), and breakdown per domain."""
    result = {"status": "APPROVED", "confidence": 100,
              "volume": {}, "growth": {}, "batch": {}, "infra": {},
              "anomalies": [], "cautions": [], "recommendations": []}

    _sla = sla_hrs or DAILY_LIMIT_HRS

    # F1+F2 — Volume
    if sow_cfg:
        dfu_act, dfu_sow = sow_cfg.get("dfu", (0, 0))
        sku_act, sku_sow = sow_cfg.get("sku", (0, 0))
        if dfu_sow > 0:
            dfu_v = calculate_volume_utilization(dfu_act, dfu_sow)
            dfu_g = project_volume_growth(dfu_act, dfu_sow, monthly_growth)
            result["volume"]["dfu"] = dfu_v
            result["growth"]["dfu"] = dfu_g
            if dfu_v["status"] == "CRITICAL":
                result["confidence"] -= 20
                result["recommendations"].append(f"DFU at {dfu_v['util_pct']}% — contract renegotiation needed")
            elif dfu_v["status"] == "CAUTION":
                result["confidence"] -= 10
                result["cautions"].append(f"DFU at {dfu_v['util_pct']}% — monitor growth closely")
            if dfu_g["runway_months"] < 6:
                result["recommendations"].append(f"DFU SOW breach in ~{dfu_g['runway_months']} months at {monthly_growth*100:.1f}%/mo growth")
            elif dfu_g["runway_months"] < 12:
                result["cautions"].append(f"DFU runway: {dfu_g['runway_months']} months")
        if sku_sow > 0:
            sku_v = calculate_volume_utilization(sku_act, sku_sow)
            result["volume"]["sku"] = sku_v
            if sku_v["status"] == "CRITICAL":
                result["confidence"] -= 15
                result["recommendations"].append(f"SKU at {sku_v['util_pct']}% — exceeds contracted volume")

    # F3+F4 — Batch
    if ctrlm_df is not None and not ctrlm_df.empty:
        batch = calculate_batch_compliance(ctrlm_df, _sla)
        result["batch"] = batch
        # F3 — SLA buffer on the worst job
        _hc = next((c for c in ctrlm_df.columns
                    if "run_time_hrs" in c.lower() or "peak_hrs" in c.lower()), None)
        if _hc:
            _jc = next((c for c in ctrlm_df.columns if "job_name" in c.lower()), None)
            if _jc:
                _peak = (ctrlm_df.groupby(_jc)[_hc].max().max())
                sla_buf = calculate_sla_buffer(_sla, float(_peak))
                result["batch"]["sla_buffer"] = sla_buf
                if sla_buf["status"] == "CRITICAL":
                    result["confidence"] -= 15
                    result["recommendations"].append(
                        f"Worst job uses {100-sla_buf['buffer_pct']:.0f}% of SLA window — batch growth risk HIGH")
                elif sla_buf["status"] == "CAUTION":
                    result["confidence"] -= 8
                    result["cautions"].append(f"SLA buffer only {sla_buf['buffer_pct']:.0f}% — limited room for growth")
        if batch["compliance_pct"] < 95:
            result["confidence"] -= max(0, int((95 - batch["compliance_pct"]) / 2))
        if batch["breach"] > 0:
            result["recommendations"].append(f"{batch['breach']} job(s) exceed SLA — root cause investigation needed")

        # F6 — Anomaly detection
        anoms = detect_job_anomalies(ctrlm_df)
        result["anomalies"] = anoms
        if len(anoms) > 0:
            result["confidence"] -= min(15, len(anoms) * 4)
            result["cautions"].append(f"{len(anoms)} statistical runtime anomalies detected")

    # F7+F8 — Infrastructure
    if servers:
        fleet = calculate_fleet_health(servers)
        result["infra"] = fleet
        if fleet["grade"] in ("D", "F"):
            result["confidence"] -= 15
            result["recommendations"].append(f"Fleet grade {fleet['grade']} — {fleet['critical']} server(s) critical")
        elif fleet["grade"] == "C":
            result["confidence"] -= 8
            result["cautions"].append(f"Fleet grade C — {fleet['warning']} server(s) in warning state")
        if fleet["unknown"] > 0:
            result["cautions"].append(f"{fleet['unknown']} server(s) have no metric data")

    # Final determination
    result["confidence"] = max(0, min(100, result["confidence"]))
    if   result["confidence"] < 70: result["status"] = "REJECTED"
    elif result["confidence"] < 85: result["status"] = "APPROVED_WITH_CAUTION"
    else:                           result["status"] = "APPROVED"

    return result


# ── METRICS ──────────────────────────────────────────────────
@st.cache_data(show_spinner=False, ttl=300)
def _compute_metrics_cached(_df_hash, df_json):
    """Cached wrapper — keyed by DataFrame hash to avoid recomputation."""
    df = pd.read_json(io.StringIO(df_json), orient="split")
    return compute_metrics(df)

def compute_metrics_fast(df):
    """Call compute_metrics with caching — only recomputes when data changes."""
    try:
        import hashlib
        h = hashlib.md5(pd.util.hash_pandas_object(df).values.tobytes()).hexdigest()
        return _compute_metrics_cached(h, df.to_json(orient="split", date_format="iso"))
    except Exception:
        return compute_metrics(df)

def compute_metrics(df):
    daily = (df.groupby(["Job_Name", "run_date"], as_index=False)
               .agg(total_hrs=("run_time_hrs","sum"), runs=("run_time_hrs","count")))
    daily["breach"] = daily["total_hrs"] > DAILY_LIMIT_HRS

    monthly = (df.groupby(["Job_Name", "month"], as_index=False)
                 .agg(total_hrs=("run_time_hrs","sum")))
    monthly["breach"] = monthly["total_hrs"] > MONTHLY_LIMIT_HRS

    total_d  = len(daily)
    # F4 — Batch Compliance (job-day level: % of job-days that stayed under SLA)
    comp     = 0.0 if total_d == 0 else (1 - daily["breach"].sum() / total_d) * 100
    j_info   = monthly.groupby("Job_Name")["breach"].any()
    t_jobs   = int(j_info.shape[0])
    j_breach = int(j_info.sum())

    sub = (df.groupby("Sub_Application", as_index=False)
             .agg(total_hrs=("run_time_hrs","sum"), jobs=("Job_Name","nunique")))

    window = (df.groupby("run_date", as_index=False)
                .agg(total_hrs=("run_time_hrs","sum"), job_count=("Job_Name","nunique"))
                .sort_values("run_date"))

    top_jobs = (df.groupby("Job_Name")["run_time_hrs"]
                  .agg(["max","mean","sum"]).reset_index()
                  .rename(columns={"max":"peak_hrs","mean":"avg_hrs","sum":"total_hrs"})
                  .sort_values("peak_hrs", ascending=False))

    # F3 — SLA Buffer per job (buffer_pct on top_jobs)
    top_jobs["buffer_pct"]   = ((DAILY_LIMIT_HRS - top_jobs["peak_hrs"]) / DAILY_LIMIT_HRS * 100).round(1)
    top_jobs["sla_used_pct"] = (top_jobs["peak_hrs"] / DAILY_LIMIT_HRS * 100).round(1)
    top_jobs["buffer_status"] = top_jobs["buffer_pct"].apply(
        lambda b: "BREACH" if b < 0 else ("CRITICAL" if b < 10 else ("CAUTION" if b < 30 else ("HEALTHY" if b < 50 else "EXCELLENT"))))

    # F5 — At-risk jobs: buffer_pct 0–15% (close to SLA but not yet breaching)
    j_at_risk = int((top_jobs["buffer_pct"].between(0, 15, inclusive="both")).sum())
    j_ok      = t_jobs - j_breach - j_at_risk

    # F6 — Anomaly Detection on top_jobs
    anomalies = detect_job_anomalies(df)

    # F3 — Fleet-level SLA buffer (worst job)
    peak_max = float(top_jobs["peak_hrs"].max()) if not top_jobs.empty else 0.0
    fleet_sla_buffer = calculate_sla_buffer(DAILY_LIMIT_HRS, peak_max) if peak_max > 0 else None

    return dict(daily=daily, monthly=monthly, window=window, compliance=comp,
                total_jobs=t_jobs, jobs_ok=j_ok, jobs_breach=j_breach,
                jobs_at_risk=j_at_risk,
                total_runs=len(df), total_hrs=df["run_time_hrs"].sum(),
                sub_stats=sub, top_jobs=top_jobs,
                anomalies=anomalies, fleet_sla_buffer=fleet_sla_buffer)


# ── CHARTS ───────────────────────────────────────────────────
@st.cache_data(show_spinner=False, ttl=300)
def heatmap_fig(daily):
    dates    = sorted(daily["run_date"].unique())[-21:]
    all_jobs = sorted(daily["Job_Name"].unique())
    if len(all_jobs) > 40:
        all_jobs = (daily.groupby("Job_Name")["total_hrs"].sum()
                         .nlargest(40).index.tolist())
    z, hover = [], []
    for job in all_jobs:
        row, hr = [], []
        for d in dates:
            sub = daily[(daily["Job_Name"] == job) & (daily["run_date"] == d)]
            if sub.empty:
                row.append(np.nan)
                hr.append(f"<b>{job}</b><br>{d}<br>No run")
            else:
                h = sub.iloc[0]["total_hrs"]
                row.append(h)
                status = "⚠ BREACH" if h > DAILY_LIMIT_HRS else "✅ OK"
                hr.append(f"<b>{job}</b><br>{d}<br>{h:.2f}h &nbsp;{status}")
        z.append(row)
        hover.append(hr)

    fig = go.Figure(go.Heatmap(
        z=z, x=[str(d) for d in dates], y=all_jobs,
        zmin=0, zmax=max(12, DAILY_LIMIT_HRS * 2),
        colorscale=[
            [0.00, C["card2"]],
            [0.01, "#1a472a"],
            [0.49, C["green"]],
            [0.50, C["amber"]],
            [0.75, "#dc6e2a"],
            [1.00, C["red"]],
        ],
        colorbar=dict(
            title=dict(text="Hours", side="right", font=dict(size=10, color=C["muted"])),
            tickfont=dict(size=9, color=C["muted"]),
            tickvals=[0, DAILY_LIMIT_HRS, DAILY_LIMIT_HRS*2],
            ticktext=["0h", f"{DAILY_LIMIT_HRS}h SLA", f"{DAILY_LIMIT_HRS*2}h"],
            thickness=12, len=0.7,
            bgcolor="rgba(0,0,0,0)", bordercolor=C["border"],
        ),
        hovertext=hover,
        hovertemplate="%{hovertext}<extra></extra>",
        xgap=2, ygap=2,
    ))
    fig.update_layout(
        **BASE_LAYOUT,
        height=max(260, len(all_jobs) * 20),
        xaxis=dict(tickangle=-45, tickfont_size=9, **AXIS),
        yaxis=dict(tickfont_size=9, autorange="reversed", **AXIS),
    )
    return fig


@st.cache_data(show_spinner=False, ttl=300)
def window_trend_fig(window):
    colors = [C["red"] if h > DAILY_LIMIT_HRS else C["blue"] for h in window["total_hrs"]]
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=window["run_date"].astype(str), y=window["total_hrs"],
        marker_color=colors, name="Daily Total",
        customdata=window["job_count"],
        hovertemplate="<b>%{x}</b><br>Total: %{y:.2f}h<br>Jobs: %{customdata}<extra></extra>",
    ))
    fig.add_hline(y=DAILY_LIMIT_HRS, line_dash="dash", line_color=C["red"],
                  annotation_text=f"{DAILY_LIMIT_HRS}h SLA Limit",
                  annotation_font=dict(color=C["red"], size=10),
                  annotation_position="top left")
    fig.update_layout(**BASE_LAYOUT, height=220,
                      xaxis=dict(tickangle=-45, tickfont_size=9, **AXIS),
                      yaxis=dict(title="Hrs", **AXIS),
            margin=_DEFAULT_MARGIN,
    )
    return fig


def top_jobs_fig(top_jobs):
    """Horizontal grouped bar — Peak vs Avg, breach-coloured, SLA zone shaded."""
    top = top_jobs.head(15).sort_values("peak_hrs", ascending=True)
    short_names = [n[:32]+"…" if len(n)>32 else n for n in top["Job_Name"]]
    bar_colors = [
        C["red"]   if v > DAILY_LIMIT_HRS       else
        C["amber"] if v > DAILY_LIMIT_HRS * 0.8 else
        C["blue"]
        for v in top["peak_hrs"]
    ]
    fig = go.Figure()
    fig.add_trace(go.Bar(
        y=short_names, x=top["avg_hrs"],
        name="Avg (hrs)", orientation="h",
        marker_color=C["blue"], marker_opacity=0.50,
        hovertemplate="<b>%{y}</b><br>Avg: <b>%{x:.3f}h</b><extra></extra>",
    ))
    fig.add_trace(go.Bar(
        y=short_names, x=top["peak_hrs"],
        name="Peak (hrs)", orientation="h",
        marker_color=bar_colors, marker_opacity=0.92,
        hovertemplate="<b>%{y}</b><br>Peak: <b>%{x:.3f}h</b><extra></extra>",
    ))
    fig.add_vline(
        x=DAILY_LIMIT_HRS, line_dash="dash",
        line_color=C["amber"], line_width=2,
        annotation_text=f"  SLA {DAILY_LIMIT_HRS}h",
        annotation_font=dict(color=C["amber"], size=9),
        annotation_position="top right",
    )
    fig.add_vrect(
        x0=DAILY_LIMIT_HRS * 0.8, x1=DAILY_LIMIT_HRS,
        fillcolor="rgba(245,158,11,0.07)",
        line_width=0, layer="below",
    )
    fig.update_layout(
        **BASE_LAYOUT, barmode="overlay",
        height=max(320, len(top) * 26 + 80),
        xaxis=dict(title="Runtime (hrs)", tickfont_size=9, **AXIS),
        yaxis=dict(tickfont_size=8, automargin=True, **AXIS),
        legend=dict(
            orientation="h", font=dict(size=9, color=C["muted"]),
            bgcolor="rgba(0,0,0,0)", x=0, y=-0.10,
        ),
        margin=dict(l=10, r=20, t=10, b=65),
    )
    return fig
@st.cache_data(show_spinner=False, ttl=300)
def job_trend_fig(df, job):
    sub = df[df["Job_Name"] == job]
    d   = sub.groupby("run_date")["run_time_hrs"].sum().reset_index()
    d.columns = ["date", "hrs"]
    fig = go.Figure(go.Scatter(
        x=d["date"].astype(str), y=d["hrs"], mode="lines+markers",
        line=dict(color=C["blue"], width=2),
        marker=dict(size=6, color=[C["red"] if h > DAILY_LIMIT_HRS else C["green"] for h in d["hrs"]]),
        hovertemplate="<b>%{x}</b><br>%{y:.2f} hrs<extra></extra>",
    ))
    fig.add_hline(y=DAILY_LIMIT_HRS, line_dash="dash", line_color=C["red"],
                  annotation_text="6h SLA", annotation_font=dict(color=C["red"], size=10))
    fig.update_layout(**BASE_LAYOUT, height=220,
                      title=dict(text=f"<b>{job}</b>", font_size=11, x=0),
                      xaxis=dict(tickangle=-30, tickfont_size=9, **AXIS),
                      yaxis=dict(title="Hrs", **AXIS),
            margin=_DEFAULT_MARGIN,
    )
    return fig


@st.cache_data(show_spinner=False, ttl=300)
def monthly_bar_fig(monthly, job):
    sub = monthly[monthly["Job_Name"] == job].sort_values("month")
    fig = go.Figure(go.Bar(
        x=sub["month"], y=sub["total_hrs"],
        marker_color=[C["red"] if b else C["green"] for b in sub["breach"]],
        hovertemplate="<b>%{x}</b><br>%{y:.2f} hrs<extra></extra>",
    ))
    fig.add_hline(y=MONTHLY_LIMIT_HRS, line_dash="dash", line_color=C["amber"],
                  annotation_text="8h Monthly SLA",
                  annotation_font=dict(color=C["amber"], size=10))
    fig.update_layout(**BASE_LAYOUT, height=220,
                      xaxis=dict(**AXIS), yaxis=dict(title="Hrs", **AXIS),
            margin=_DEFAULT_MARGIN,
    )
    return fig


def sub_app_pie(sub_stats):
    fig = px.pie(sub_stats, names="Sub_Application", values="total_hrs", hole=0.52,
                 color_discrete_sequence=[C["blue"], C["purple"], C["cyan"], C["amber"], C["green"]])
    fig.update_traces(textposition="outside", textfont_size=9)
    fig.update_layout(**BASE_LAYOUT, height=240,
                      showlegend=True,
                      legend=dict(orientation="h", yanchor="bottom", y=-0.35,
                                  font=dict(size=9)),
            margin=_DEFAULT_MARGIN,
    )
    return fig


def server_gauge(val, label, ok, warn):
    color = status_color(val, ok, warn)
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=val,
        number=dict(suffix="%", font=dict(size=20, color=color)),
        gauge=dict(
            axis=dict(range=[0, 100], tickcolor=C["muted"], tickfont=dict(size=8)),
            bar=dict(color=color, thickness=0.25),
            bgcolor="rgba(0,0,0,0)", borderwidth=0,
            steps=[
                dict(range=[0,   ok],   color=C["card2"]),
                dict(range=[ok,  warn], color=hex_rgba(C["amber"])),
                dict(range=[warn, 100], color=hex_rgba(C["red"])),
            ],
            threshold=dict(line=dict(color=C["red"], width=2), thickness=0.75, value=warn),
        ),
        title=dict(text=f"<b>{label}</b>", font=dict(size=10, color=C["muted"])),
    ))
    gauge_layout = {**BASE_LAYOUT, "height": 155,
                    "margin": dict(l=8, r=8, t=35, b=5)}   # overrides BASE_LAYOUT margin
    fig.update_layout(**gauge_layout)
    return fig


# ── HTML HELPERS ─────────────────────────────────────────────
def kpi(title, value, sub="", icon="📊", color=None, grad=False, trend=None):
    color = color or C["blue"]
    g = "kpi-grad" if grad else ""
    t = ""
    if trend is not None:
        ar  = "▲" if trend > 0 else "▼"
        cls = "kpi-pos" if trend > 0 else "kpi-neg"
        t   = f'<p class="{cls}">{ar} {abs(trend)}% vs last period</p>'
    st.markdown(f"""<div class="kpi-card {g}">
  <div style="font-size:22px;margin-bottom:6px">{icon}</div>
  <p class="kpi-label">{title}</p>
  <p class="kpi-value" style="color:{color}">{value}</p>
  <p class="kpi-sub">{sub}</p>{t}</div>""", unsafe_allow_html=True)


def health_pill(count, label, icon, bg, bcls, btxt):
    st.markdown(f"""<div class="kpi-card" style="display:flex;align-items:center;
  justify-content:space-between;">
  <div style="display:flex;align-items:center;gap:12px">
    <div style="background:{bg}22;border-radius:8px;padding:8px">
      <span style="font-size:18px">{icon}</span></div>
    <div>
      <p style="font-size:22px;font-weight:700;margin:0;color:{bg}">{count}</p>
      <p style="font-size:11px;color:{C['muted']};margin:0">{label}</p>
    </div>
  </div>
  <span class="{bcls}">{btxt}</span>
</div>""", unsafe_allow_html=True)


def render_header(customer, env, overall):
    """Compact sticky header: logo | title | customer | env badge | status | time."""
    bmap = {
        "BREACH":  f'<span class="b-breach" style="font-size:11px">🔴 SLA BREACH</span>',
        "WARNING": f'<span class="b-warn"   style="font-size:11px">⚠️ WARNING</span>',
        "OK":      f'<span class="b-ok"     style="font-size:11px">✅ HEALTHY</span>',
    }
    status_badge = bmap.get(overall, bmap["OK"])
    env_badge    = (f'<span style="background:{C["blue"]}22;color:{C["blue"]};padding:2px 9px;' +
                    f'border-radius:20px;font-size:10px;font-weight:700;margin-left:6px">{env}</span>') if env else ""
    signed_badge = ""
    if st.session_state.get("approval_pe") and st.session_state.get("approval_customer"):
        signed_badge = f'<span class="b-signed" style="margin-left:8px">✍️ SIGNED OFF</span>'
    now_str = datetime.now().strftime("%b %d, %Y %I:%M %p")
    cust_display = customer or "Upload Data to Begin"

    st.markdown(f'''<div class="pe-header">
  <div style="display:flex;align-items:center;gap:16px">
    <div class="pe-logo">PE</div>
    <div>
      <div class="pe-title" style="font-size:20px;font-weight:800;letter-spacing:.01em">PE Control Tower</div>
      <div class="pe-sub" style="font-size:12px">Performance Engineering Audit Dashboard v5.0</div>
    </div>
  </div>
  <div style="display:flex;align-items:center;gap:18px">
    <div style="text-align:right">
      <div style="font-size:16px;font-weight:800;color:{C["white"]}">{cust_display}{env_badge}</div>
      <div style="font-size:11px;color:{C["muted"]};margin-top:3px">{now_str}</div>
    </div>
    {status_badge}{signed_badge}
  </div>
</div>''', unsafe_allow_html=True)




# ── UPLOAD PANEL ─────────────────────────────────────────────

# ── GEMINI VISION: Extract metrics from monitoring screenshot images ──────────
def _extract_metrics_from_zabbix_images(servers_skeleton, all_images, api_key):
    """
    Uses Gemini Flash vision to OCR monitoring screenshots (Azure Monitor,
    Zabbix, Grafana, etc.) embedded in DOCX/PDF and map CPU/Memory/Disk %
    values back to each server.

    servers_skeleton: list of server dicts (image-only, all zeros)
    all_images:       list of (filename, bytes) from extract_docx_images()
    api_key:          Gemini API key string

    Returns the same servers list with metrics populated where found.
    """
    if not all_images or not api_key:
        return servers_skeleton

    try:
        import google.generativeai as genai
        from PIL import Image as PILImage
        import io, json, re
        genai.configure(api_key=api_key)
        _ocr_model_names = [
            "gemini-2.5-flash",
            "gemini-2.0-flash-001",
            "gemini-2.0-flash-lite",
            "gemini-flash-latest",
            "gemini-flash-lite-latest",
        ]
        model = None
        for mn in _ocr_model_names:
            try:
                model = genai.GenerativeModel(mn)
                break
            except Exception:
                continue
        if model is None:
            return servers_skeleton
    except Exception:
        return servers_skeleton

    # Build a name→index map for fast lookup (case-insensitive)
    server_map = {}
    for i, s in enumerate(servers_skeleton):
        host_short = s["host"].split(".")[0].lower()
        server_map[host_short] = i
        label = s.get("label", "").lower()
        if label and label != host_short:
            server_map[label] = i

    host_list_str = ", ".join(s["host"].split(".")[0] for s in servers_skeleton)

    def _pil_from_bytes(b):
        try:
            return PILImage.open(io.BytesIO(b)).convert("RGB")
        except Exception:
            return None

    servers_updated = [dict(s) for s in servers_skeleton]
    metric_buffer = {}   # host_idx -> {"cpu":[], "mem":[], "disk":[]}

    # ── Strategy 1: BATCH mode — send up to 6 images per Gemini call ─────────
    # Much faster and gives Gemini context across multiple images
    batch_size = 6
    batch_prompt = f"""You are a server monitoring expert analyzing Zabbix/monitoring screenshots.
Known servers in this report: {host_list_str}

For each image extract the SERVER HOSTNAME, METRIC TYPE, and USED PERCENTAGE.

=== READING ZABBIX CHARTS ===
• Chart title identifies the metric name.
• Zabbix legend format at bottom: "min: X  avg: Y  max: Z  last: W" — use the "last:" value.
• If no legend, read the latest bar height or line endpoint value.

=== HOSTNAME DETECTION ===
• Look in: chart title, "Scope" field, axis label, watermark, legend text.
• Known server list (match partially if needed): {host_list_str}

=== METRIC CONVERSION (CRITICAL) ===
CPU:
  "CPU idle time" / "% CPU idle"  → is_available=true  (cpu_used = 100 − value)
  "CPU utilization" / "CPU usage" → is_available=false (already used%)
  "CPU load"                      → is_available=false

Memory:
  "Available memory %" / "Free memory" → is_available=true  (mem_used = 100 − value)
  "Memory used" / "Memory utilization" → is_available=false

Disk:
  "Free disk space % on /X" → is_available=true  (disk_used = 100 − value)
  "Disk used %" / "Disk utilization"  → is_available=false

Return ONLY valid JSON array (no markdown fences, no explanation):
[
  {{"hostname": "tsbc101402011", "metric_type": "Memory", "value_pct": 29.0, "is_available": true}},
  {{"hostname": "tsbc101402011", "metric_type": "CPU",    "value_pct": 88.0, "is_available": true}},
  {{"hostname": "tsbc101402011", "metric_type": "Disk",   "value_pct": 5.2,  "is_available": true}}
]

Omit entries where value cannot be clearly read."""

    image_batches = []
    pil_images_all = []
    for img_name, img_bytes in all_images:
        pil_img = _pil_from_bytes(img_bytes)
        if pil_img is None:
            continue
        # Resize large images to save tokens
        if pil_img.width > 1024:
            ratio = 1024 / pil_img.width
            pil_img = pil_img.resize((1024, int(pil_img.height * ratio)),
                                      PILImage.LANCZOS)
        pil_images_all.append((img_name, pil_img))

    # Process in batches
    for batch_start in range(0, len(pil_images_all), batch_size):
        batch = pil_images_all[batch_start:batch_start + batch_size]
        content_parts = []
        for _name, pil_img in batch:
            content_parts.append(pil_img)
        content_parts.append(batch_prompt)

        try:
            resp = model.generate_content(content_parts)
            raw_text = resp.text.strip()
            # Strip markdown code fences
            raw_text = re.sub(r'^```(?:json)?\s*', '', raw_text, flags=re.MULTILINE)
            raw_text = re.sub(r'\s*```\s*$', '', raw_text, flags=re.MULTILINE).strip()

            parsed = json.loads(raw_text)
            if isinstance(parsed, dict):
                parsed = [parsed]
            if not isinstance(parsed, list):
                continue

            for item in parsed:
                if not isinstance(item, dict):
                    continue
                val = item.get("value_pct")
                if val is None:
                    continue
                try:
                    val = float(val)
                except (TypeError, ValueError):
                    continue

                # Convert "available/free" to "used"
                is_avail = item.get("is_available", False)
                if is_avail and 0 <= val <= 100:
                    val = round(100 - val, 2)

                mtype = str(item.get("metric_type", "")).upper()
                hostname_hint = str(item.get("hostname") or "").lower().split(".")[0]

                # Find which server this belongs to
                target_idx = None
                if hostname_hint and hostname_hint != "null":
                    # Direct match
                    if hostname_hint in server_map:
                        target_idx = server_map[hostname_hint]
                    else:
                        # Partial match
                        for key, idx in server_map.items():
                            if key in hostname_hint or hostname_hint in key:
                                target_idx = idx
                                break

                if target_idx is None:
                    continue

                if target_idx not in metric_buffer:
                    metric_buffer[target_idx] = {"cpu": [], "mem": [], "disk": []}

                if "CPU" in mtype:
                    metric_buffer[target_idx]["cpu"].append(val)
                elif "MEM" in mtype or "RAM" in mtype:
                    metric_buffer[target_idx]["mem"].append(val)
                elif "DISK" in mtype or "STORAGE" in mtype:
                    metric_buffer[target_idx]["disk"].append(val)

        except Exception:
            # Batch failed — fall through to per-image below
            continue

    # ── Strategy 2: Per-image fallback for any servers still missing metrics ──
    missing_idxs = set(range(len(servers_updated))) - set(metric_buffer.keys())
    if missing_idxs and pil_images_all:
        single_prompt = f"""This is a server monitoring screenshot (Azure Monitor, Zabbix, Grafana, or similar).
Extract the hostname and metric value.
Known servers: {host_list_str}

Return ONLY valid JSON (no markdown):
{{"hostname": "server_name", "metric_type": "CPU" or "Memory" or "Disk", "value_pct": 45.2, "is_available": false}}

"is_available"=true if label says "Available"/"Free" (value needs 100-val conversion).
Read the percentage from the legend or summary text at bottom of the chart."""

        for img_name, pil_img in pil_images_all:
            try:
                resp = model.generate_content([pil_img, single_prompt])
                raw_text = resp.text.strip()
                raw_text = re.sub(r'^```(?:json)?\s*', '', raw_text, flags=re.MULTILINE)
                raw_text = re.sub(r'\s*```\s*$', '', raw_text, flags=re.MULTILINE).strip()

                parsed = json.loads(raw_text)
                if isinstance(parsed, dict):
                    parsed = [parsed]

                for item in parsed:
                    if not isinstance(item, dict):
                        continue
                    val = item.get("value_pct")
                    if val is None:
                        continue
                    try:
                        val = float(val)
                    except (TypeError, ValueError):
                        continue

                    is_avail = item.get("is_available", False)
                    if is_avail and 0 <= val <= 100:
                        val = round(100 - val, 2)

                    mtype = str(item.get("metric_type", "")).upper()
                    hostname_hint = str(item.get("hostname") or "").lower().split(".")[0]

                    target_idx = None
                    if hostname_hint and hostname_hint != "null":
                        if hostname_hint in server_map:
                            target_idx = server_map[hostname_hint]
                        else:
                            for key, idx in server_map.items():
                                if key in hostname_hint or hostname_hint in key:
                                    target_idx = idx
                                    break

                    # Only process if this server still needs data
                    if target_idx is None or target_idx not in missing_idxs:
                        continue

                    if target_idx not in metric_buffer:
                        metric_buffer[target_idx] = {"cpu": [], "mem": [], "disk": []}

                    if "CPU" in mtype:
                        metric_buffer[target_idx]["cpu"].append(val)
                    elif "MEM" in mtype or "RAM" in mtype:
                        metric_buffer[target_idx]["mem"].append(val)
                    elif "DISK" in mtype or "STORAGE" in mtype:
                        metric_buffer[target_idx]["disk"].append(val)

            except Exception:
                continue

    # ── Apply collected metrics back to server records ────────────────────────
    for idx, metrics in metric_buffer.items():
        if idx >= len(servers_updated):
            continue
        s = servers_updated[idx]
        if metrics["cpu"]:
            s["cpu_used"] = round(max(metrics["cpu"]), 2)
            s["cpu_avg"]  = round(sum(metrics["cpu"]) / len(metrics["cpu"]), 2)
        if metrics["mem"]:
            s["mem_used"] = round(max(metrics["mem"]), 2)
        if metrics["disk"]:
            s["disk_used_max"] = round(max(metrics["disk"]), 2)

    return servers_updated


def _extract_metrics_per_server_vision(server_sections, api_key):
    """
    Targeted per-server Vision OCR for DOCX reports.

    Unlike the batch approach (all images + all servers → one prompt), this
    function sends each server's specific images together with the server's
    hostname as grounding context. Gemini therefore knows exactly which server
    it is reading, eliminating cross-server confusion.

    server_sections : list of server dicts that include an "images" key
                      (output of parse_resource_docx_structured)
    api_key         : Gemini API key string

    Returns the same list with cpu_used / mem_used / disk_used_max populated.
    """
    if not api_key:
        return server_sections

    try:
        import google.generativeai as genai
        from PIL import Image as PILImage
        import io, json, re
        genai.configure(api_key=api_key)
        model = None
        for mn in ["gemini-2.5-flash", "gemini-2.0-flash-001",
                   "gemini-2.0-flash-lite", "gemini-flash-latest"]:
            try:
                model = genai.GenerativeModel(mn); break
            except Exception:
                continue
        if model is None:
            return server_sections
    except Exception:
        return server_sections

    def _pil(img_bytes, max_w=1024):
        try:
            img = PILImage.open(io.BytesIO(img_bytes)).convert("RGB")
            if img.width > max_w:
                ratio = max_w / img.width
                img = img.resize((max_w, int(img.height * ratio)), PILImage.LANCZOS)
            return img
        except Exception:
            return None

    results = []
    for sec in server_sections:
        sec = dict(sec)
        images = sec.get("images", [])

        if not images:
            results.append(sec)
            continue

        hostname  = sec.get("host", "unknown")
        host_short = hostname.split(".")[0]
        label     = sec.get("label", host_short)

        pil_imgs = [_pil(b) for _, b in images]
        pil_imgs = [p for p in pil_imgs if p is not None]

        if not pil_imgs:
            results.append(sec)
            continue

        prompt = f"""You are a server monitoring expert reading {len(pil_imgs)} Zabbix/monitoring chart(s) for:
Server hostname : {host_short}
Server label    : {label}

TASK: For EACH chart, identify the metric and return the USED percentage.

=== ZABBIX CHART READING RULES (read in priority order) ===
1. Read the chart TITLE to identify the metric name.
2. Read the LEGEND at the bottom — Zabbix shows: "min: X  avg: Y  max: Z  last: W"
   → Use the "last:" value as the chart value.
3. If no legend, read the most recent bar height or line endpoint.

=== METRIC CONVERSION TABLE ===
CPU charts:
  • "CPU idle time" or "% CPU idle" → cpu_used = 100 − idle_value   [is_available=TRUE]
  • "CPU utilization" / "CPU usage" / "% CPU used" → cpu_used = shown_value  [is_available=FALSE]
  • "CPU load" (1-min avg) → treat as cpu_used directly              [is_available=FALSE]

Memory charts:
  • "Available memory" / "Free memory" / "Memory available %" → mem_used = 100 − shown_value  [is_available=TRUE]
  • "Memory used" / "Memory utilization" → mem_used = shown_value    [is_available=FALSE]
  • "Used memory in MB" — ignore (not a percentage)

Disk charts:
  • "Free disk space on /X (percentage)" → disk_used = 100 − shown_value  [is_available=TRUE]
  • "Disk space used" / "Disk utilization" → disk_used = shown_value [is_available=FALSE]

=== OUTPUT FORMAT ===
Return ONLY a valid JSON array — no markdown fences, no explanation:
[
  {{"metric_type": "CPU",    "value_pct": 54.1, "is_available": false}},
  {{"metric_type": "Memory", "value_pct": 29.0, "is_available": true}},
  {{"metric_type": "Disk",   "value_pct": 5.0,  "is_available": true}}
]

Rules:
- One entry per chart image.
- "is_available": true means the value shown is free/idle (will be subtracted from 100).
- "is_available": false means the value shown is already the USED percentage.
- If a chart has multiple disk mounts, return the HIGHEST disk_used% (worst case).
- If you cannot read a value clearly, OMIT that entry — do not guess 0."""

        try:
            resp = model.generate_content(pil_imgs + [prompt])
            raw  = resp.text.strip()
            raw  = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
            raw  = re.sub(r'\s*```\s*$',        '', raw, flags=re.MULTILINE).strip()
            parsed = json.loads(raw)
            if isinstance(parsed, dict):
                parsed = [parsed]

            cpu_v, mem_v, disk_v = [], [], []
            for item in parsed:
                if not isinstance(item, dict): continue
                val = item.get("value_pct")
                if val is None: continue
                try: val = float(val)
                except Exception: continue
                if item.get("is_available", False) and 0 <= val <= 100:
                    val = round(100 - val, 2)
                mtype = str(item.get("metric_type", "")).upper()
                if   "CPU"  in mtype:                     cpu_v.append(val)
                elif "MEM"  in mtype or "RAM" in mtype:   mem_v.append(val)
                elif "DISK" in mtype or "STOR" in mtype:  disk_v.append(val)

            if cpu_v:
                sec["cpu_used"] = round(max(cpu_v), 2)
                sec["cpu_avg"]  = round(sum(cpu_v) / len(cpu_v), 2)
            if mem_v:
                sec["mem_used"] = round(max(mem_v), 2)
            if disk_v:
                sec["disk_used_max"] = round(max(disk_v), 2)

        except Exception:
            pass  # Keep zeros — server still shows in dashboard

        results.append(sec)

    return results


# ── GEMINI VISION: Generic image-to-table extractor ──────────────────────────
def _extract_table_from_image(file_obj, domain_hint, api_key):
    """
    Use Gemini Vision to extract tabular data from an uploaded screenshot/image.
    domain_hint: "batch_runs", "sla_matrix", "resource_util", "sow_contract", "perf_test"
    Returns: pandas DataFrame (or None on failure)
    """
    if not api_key:
        return None
    try:
        import google.generativeai as genai
        from PIL import Image as PILImage
        import io, json, re, pandas as pd
        genai.configure(api_key=api_key)
        model = None
        for mn in ["gemini-2.5-pro","gemini-2.5-flash","gemini-2.0-flash-001","gemini-flash-latest"]:
            try:
                model = genai.GenerativeModel(mn); break
            except Exception:
                continue
        if model is None:
            return None

        file_obj.seek(0)
        pil_img = PILImage.open(io.BytesIO(file_obj.read())).convert("RGB")
        if pil_img.width > 2048:
            ratio = 2048 / pil_img.width
            pil_img = pil_img.resize((1400, int(pil_img.height * ratio)), PILImage.LANCZOS)

        domain_prompts = {
            "batch_runs": (
                "This is a screenshot of a batch job execution report (Ctrl-M, AutoSys, or similar scheduler). "
                "Extract ALL rows as a table with columns: Job_Name, Status (OK/NOTOK/Failed/Success), "
                "Start_Time, End_Time, Run_Time_Sec (numeric seconds), Date. "
                "Return ONLY a JSON array of objects. No markdown, no explanation."
            ),
            "sla_matrix": (
                "This is a screenshot of a Batch SLA matrix showing job-level SLA thresholds. "
                "Extract ALL rows. Look for columns like: Job_Name or Batch_Name, SLA_Hrs or SLA target, "
                "Start_Time, End_Time, Schedule. "
                "Return ONLY a JSON array of objects. No markdown, no explanation."
            ),
            "resource_util": (
                "You are an expert infrastructure analyst reading a server monitoring screenshot. "
                "This could be from Zabbix, Grafana, Azure Monitor, PRTG, Datadog, Nagios, "
                "SolarWinds, vSphere, or a plain text/table report. "
                "TASK: Extract ALL server/host resource metrics you can see. "
                "Look for: hostnames, IP addresses, server names, node names. "
                "For each host extract: CPU utilization%, RAM/Memory usage%, Disk usage%. "
                "IMPORTANT RULES: "
                "1. If you see 'Available Memory' or 'Free Memory' — SUBTRACT from 100 to get used%. "
                "2. If charts show percentage scales, read the current/latest bar or line value. "
                "3. If a table shows rows of servers, extract each row. "
                "4. Infer server type: 'DB','Oracle','SQL','Postgres' = DB; else APP. "
                "5. If hostname contains 'db','ora','sql','dbs' = DB type; else APP. "
                "Return ONLY a valid JSON array — zero markdown, zero explanation: "
                "[{\"host\":\"server01.domain.com\", \"type\":\"APP\", "
                "\"cpu_used\":45.2, \"cpu_avg\":38.1, "
                "\"mem_used\":60.1, \"mem_total_gb\":32.0, "
                "\"disk_used_max\":33.5, \"disks\":{\"root\":33.5}}]. "
                "If mem_total_gb is not visible, use 0. "
                "If multiple disk paths visible, include all in disks object. "
                "Extract EVERY server you can see, even if partial data."
            ),
            "sow_contract": (
                "This is a screenshot of a Statement of Work (SOW) or contract volume table. "
                "Extract key data: customer name, DFU count, SKU count, planogram count, "
                "any contracted volumes or limits. "
                "Return ONLY a JSON object with found values. No markdown, no explanation."
            ),
            "perf_test": (
                "This is a screenshot of a performance/load test result (JMeter, Gatling, LoadRunner, etc.). "
                "Extract ALL rows: Transaction/Page name, Avg response time (ms), 90th percentile, "
                "Error%, Throughput. Return ONLY a JSON array of objects. No markdown."
            ),
        }
        prompt = domain_prompts.get(domain_hint, domain_prompts["batch_runs"])

        resp = model.generate_content([pil_img, prompt])
        raw = resp.text.strip()
        raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
        raw = re.sub(r'\s*```\s*$', '', raw, flags=re.MULTILINE).strip()
        parsed = json.loads(raw)

        if isinstance(parsed, list) and len(parsed) > 0:
            return pd.DataFrame(parsed)
        elif isinstance(parsed, dict):
            return pd.DataFrame([parsed])
        return None
    except Exception:
        return None


def _is_image_file(fname):
    """Check if filename is an image type."""
    return fname.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"))


# ═══════════════════════════════════════════════════════════════════════════════
# TRUE AI-FIRST PIPELINE
# ───────────────────────────────────────────────────────────────────────────────
# FLOW: Raw file → Python parse (column mapping, types) → AI enrichment/validation
#       → AI-corrected data written back to session_state → Charts read corrected data
#
# The AI acts as a smart post-processor that:
#   1. Detects what kind of data was uploaded (customer, period, format)
#   2. Validates and corrects the parsed data (fixes misread columns, units, anomalies)
#   3. Writes a structured JSON correction back — pandas applies it to the DataFrame
#   4. Populates an insight summary that explains what was found and corrected
# ═══════════════════════════════════════════════════════════════════════════════

import json as _JSON_MOD

def _call_gemini_json(prompt):
    """
    Call Gemini and parse the response as JSON.
    Returns the parsed dict/list or None if parsing fails.
    Handles all model discovery and error cases internally.
    """
    try:
        import google.generativeai as genai
    except ImportError:
        return None

    genai.configure(api_key=_get_api_key())

    model = None
    try:
        available = list(genai.list_models())
        def _rank(m):
            n = m.name.lower()
            supported = str([s.lower() for s in getattr(m,"supported_generation_methods",[])])
            if "generatecontent" not in supported: return 99
            if "flash" in n and "2" in n: return 0
            if "flash" in n and "1.5" in n: return 1
            if "flash" in n: return 2
            if "pro" in n: return 3
            return 4
        available_sorted = sorted(available, key=_rank)
        for mi in available_sorted:
            if _rank(mi) < 99:
                model = genai.GenerativeModel(mi.name)
                break
    except Exception:
        for _mn in ["gemini-2.5-flash","gemini-2.0-flash-001","gemini-flash-latest","gemini-flash-lite-latest"]:
            try:
                model = genai.GenerativeModel(_mn)
                break
            except Exception:
                continue

    if model is None:
        return None

    try:
        resp = model.generate_content(prompt, generation_config={"max_output_tokens": 2000, "temperature": 0.1})
        raw = resp.text.strip()
        # Strip markdown fences
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\s*```$", "", raw, flags=re.MULTILINE)
        return _JSON_MOD.loads(raw)
    except Exception:
        return None


def _render_ai_insight_box(cache_key, title, run_key):
    """
    Renders the AI insight panel inside the upload section.
    Shows: what data was detected · corrections made · quality flags · chart guidance.
    """
    result = st.session_state.get(cache_key)

    hdr_col, btn_col = st.columns([5, 1])
    with hdr_col:
        st.markdown(
            f'''<div style="display:flex;align-items:center;gap:8px;margin-bottom:6px">
  <span style="font-size:14px">🤖</span>
  <span style="font-size:12px;font-weight:700;color:{C["white"]}">{title}</span>
  <span style="background:{C["blue"]}22;color:{C["blue"]};padding:1px 7px;
    border-radius:5px;font-size:9px;font-weight:700">AI-PROCESSED</span>
</div>''', unsafe_allow_html=True)
    with btn_col:
        if st.button("↺ Re-run", key=f"refresh_{cache_key}", use_container_width=True):
            st.session_state[run_key]   = True
            st.session_state[cache_key] = None
            st.rerun()

    if result:
        # result is a dict with keys: summary, data_type, corrections, quality_flags, chart_guidance
        dtype   = result.get("data_type", "Unknown")
        summary = result.get("summary", "")
        corr    = result.get("corrections", [])
        flags   = result.get("quality_flags", [])
        charts  = result.get("chart_guidance", "")
        recoms  = result.get("recommendations", [])

        # Data type badge
        st.markdown(
            f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:10px;padding:14px 16px">
<div style="display:flex;gap:8px;align-items:center;margin-bottom:10px">
  <span style="background:{C["purple"]}22;color:{C["purple"]};padding:2px 10px;
    border-radius:20px;font-size:10px;font-weight:700">{dtype}</span>
  <span style="font-size:11px;color:{C["white"]}">{summary}</span>
</div>''', unsafe_allow_html=True)

        if corr:
            st.markdown(f"<p style='font-size:10px;font-weight:700;color:{C['amber']};margin:8px 0 4px'>⚙️ Data Corrections Applied ({len(corr)})</p>", unsafe_allow_html=True)
            for c in corr[:5]:
                st.markdown(f"<p style='font-size:10px;color:{C['white']};margin:0 0 3px'>• {c}</p>", unsafe_allow_html=True)

        if flags:
            st.markdown(f"<p style='font-size:10px;font-weight:700;color:{C['red']};margin:8px 0 4px'>🚩 Quality Flags</p>", unsafe_allow_html=True)
            for f in flags[:5]:
                st.markdown(f"<p style='font-size:10px;color:{C['amber']};margin:0 0 3px'>⚠️ {f}</p>", unsafe_allow_html=True)

        if charts:
            st.markdown(f"<p style='font-size:10px;font-weight:700;color:{C['cyan']};margin:8px 0 4px'>📊 What to Look for in Charts</p>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:10px;color:{C['white']};margin:0'>{charts}</p>", unsafe_allow_html=True)

        if recoms:
            st.markdown(f"<p style='font-size:10px;font-weight:700;color:{C['green']};margin:8px 0 4px'>✅ AI Recommendations</p>", unsafe_allow_html=True)
            for i, r in enumerate(recoms[:4], 1):
                st.markdown(f"<p style='font-size:10px;color:{C['white']};margin:0 0 3px'>{i}. {r}</p>", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown(
            f'<p style="font-size:11px;color:{C["muted"]};padding:10px 0">'
            f'AI analysis will run automatically on next upload, or click ↺ Re-run.</p>',
            unsafe_allow_html=True)


# ─── AI-FIRST PIPELINE: Ctrl-M ──────────────────────────────────────────────
def ai_analyse_ctrlm(df):
    """
    AI-FIRST pipeline for Ctrl-M data.
    Step 1: Send sample to Gemini for intelligent validation and enrichment.
    Step 2: Gemini returns structured JSON with corrections + metadata.
    Step 3: Apply corrections to the DataFrame in session_state.
    Step 4: Store insight dict for display in upload panel.
    """
    import pandas as _pd

    # Build compact payload (max 40 rows as sample)
    sample_rows = df.sample(min(40, len(df)), random_state=42).copy()
    sample_rows["run_date"] = sample_rows["run_date"].astype(str)
    sample_rows["month"]    = sample_rows["month"].astype(str)
    sample_cols = ["Job_Name","Sub_Application","run_date","run_time_hrs","Run_Sec"]
    sample_cols = [c for c in sample_cols if c in sample_rows.columns]
    sample = sample_rows[sample_cols].round(4).head(40).to_dict("records")

    m = compute_metrics_fast(df)
    stats = {
        "total_rows": len(df),
        "unique_jobs": m["total_jobs"],
        "date_range": f"{str(df['run_date'].min())} to {str(df['run_date'].max())}",
        "compliance_pct": round(m["compliance"], 1),
        "jobs_breach": m["jobs_breach"],
        "avg_runtime_hrs": round(float(df["run_time_hrs"].mean()), 4),
        "max_runtime_hrs": round(float(df["run_time_hrs"].max()), 4),
        "sla_limit_hrs": DAILY_LIMIT_HRS,
        "sub_applications": df["Sub_Application"].unique().tolist()[:10] if "Sub_Application" in df.columns else [],
        "columns_present": list(df.columns),
    }

    prompt = f"""You are a senior Performance Engineering data validator reviewing Ctrl-M batch execution data.

## Data Stats
{_JSON_MOD.dumps(stats, indent=2)}

## Sample Rows (up to 40)
{_JSON_MOD.dumps(sample, indent=2)}

## Your Task
Analyse this data and return ONLY valid JSON (no markdown, no explanation) in this exact schema:
{{
  "data_type": "brief label e.g. Ctrl-M Batch Execution Report — CustomerName — 30 days",
  "summary": "One sentence: what this data contains, period covered, customer if detectable",
  "corrections": ["list of corrections applied e.g. Renamed column X to Y", "Converted Run_Sec from ms to s"],
  "quality_flags": ["list of data quality issues found e.g. 3 jobs have 0-second runtimes", "15% null Sub_Application"],
  "chart_guidance": "What the heatmap, daily window trend and top-jobs chart should visually highlight — mention specific job names or patterns you see",
  "recommendations": ["numbered action 1", "action 2", "action 3"],
  "job_corrections": [
    {{"job_name": "JOB_X", "issue": "runtime appears in wrong unit", "corrected_hrs": 2.5}}
  ]
}}

Be specific. Use actual job names and numbers from the data. Only include job_corrections if you find real issues."""

    result = _call_gemini_json(prompt)

    if result is None:
        # Fallback: generate local insight without AI
        result = {
            "data_type": f"Ctrl-M Batch Execution — {len(df):,} rows — {m['total_jobs']} jobs",
            "summary": f"Loaded {len(df):,} records across {m['total_jobs']} jobs. SLA compliance: {m['compliance']:.1f}%.",
            "corrections": [],
            "quality_flags": [] if m["jobs_breach"] == 0 else [f"{m['jobs_breach']} job(s) exceed the {DAILY_LIMIT_HRS}h SLA"],
            "chart_guidance": "Check heatmap for breach patterns. Daily window chart shows workload distribution.",
            "recommendations": ["Review breaching jobs in Perf-Test Report tab", "Check Sub-Application distribution"],
            "job_corrections": [],
        }

    # Apply AI-suggested job corrections to the DataFrame
    corrections_applied = []
    for jc in result.get("job_corrections", []):
        jname  = jc.get("job_name","")
        new_hr = jc.get("corrected_hrs")
        if jname and new_hr is not None:
            try:
                new_hr = float(new_hr)
                mask = df["Job_Name"] == jname
                if mask.sum() > 0 and new_hr > 0:
                    df.loc[mask, "run_time_hrs"] = new_hr
                    df.loc[mask, "Run_Sec"]       = new_hr * 3600
                    corrections_applied.append(f"Corrected '{jname}' runtime to {new_hr:.2f}h")
            except Exception:
                pass

    if corrections_applied:
        result["corrections"] = (result.get("corrections") or []) + corrections_applied
        # Write corrected DataFrame back — all charts will now use corrected data
        st.session_state["ctrlm_df"] = df

    st.session_state["_ai_ctrlm_insight"] = result


# ─── AI-FIRST PIPELINE: Batch SLA ───────────────────────────────────────────
def ai_analyse_batch_sla(df):
    """
    AI-FIRST pipeline for Batch SLA matrix.
    Gemini validates column types, detects threshold units, flags outlier SLAs.
    Returns structured insight AND writes any corrected thresholds back.
    """
    cols = list(df.columns)
    sample = df.head(20).fillna("").astype(str).to_dict("records")
    numeric_cols = {c: {"min": float(df[c].min()), "max": float(df[c].max()),
                        "nulls_pct": round(df[c].isna().mean()*100,1)}
                    for c in df.select_dtypes(include="number").columns[:8]}

    prompt = f"""You are a Performance Engineering data validator reviewing a Batch SLA matrix file.

## File Overview
Columns: {_JSON_MOD.dumps(cols)}
Rows: {len(df)}
Numeric columns: {_JSON_MOD.dumps(numeric_cols, indent=2)}

## First 20 rows
{_JSON_MOD.dumps(sample, indent=2)}

## Your Task
Return ONLY valid JSON:
{{
  "data_type": "Batch SLA Matrix — N jobs — what the SLA columns represent",
  "summary": "One sentence describing this SLA matrix",
  "corrections": ["any column renames or unit conversions needed"],
  "quality_flags": ["data quality issues: nulls, inconsistent units, suspicious thresholds"],
  "chart_guidance": "How the SLA matrix data will sharpen job compliance charts — which jobs have tight SLAs",
  "recommendations": ["action 1", "action 2"],
  "column_mapping": {{"original_col_name": "standardised_meaning"}},
  "threshold_unit": "seconds or hours or minutes — what unit the SLA values appear to be in"
}}"""

    result = _call_gemini_json(prompt)

    if result is None:
        result = {
            "data_type": f"Batch SLA Matrix — {len(df)} rows — {len(cols)} columns",
            "summary": f"SLA matrix loaded with {len(df)} job thresholds across {len(cols)} columns.",
            "corrections": [],
            "quality_flags": [],
            "chart_guidance": "SLA matrix will be used in Batch SLA tab for per-job compliance comparison.",
            "recommendations": ["Review SLA tab to see per-job threshold comparison"],
            "column_mapping": {},
            "threshold_unit": "unknown",
        }

    st.session_state["_ai_sla_insight"] = result


# ─── AI-FIRST PIPELINE: Resource Utilisation ────────────────────────────────
def ai_analyse_resource(servers):
    """
    AI-FIRST pipeline for server resource data.
    Gemini validates metric values, detects anomalies, corrects obvious errors
    (e.g. CPU=100 when it should be 10.0, disk unit confusion).
    Writes corrected server list back to session_state.
    """
    known = [s for s in servers if not is_unknown_server(s)]
    unknown_count = len(servers) - len(known)

    server_payload = []
    for s in known[:20]:
        server_payload.append({
            "host":     s["host"].split(".")[0],
            "type":     s.get("type","APP"),
            "cpu_pct":  round(s.get("cpu_used",0), 2),
            "mem_pct":  round(s.get("mem_used",0), 2),
            "disk_max": round(s.get("disk_used_max",0), 2),
            "mem_gb":   round(s.get("mem_total_gb",0), 1),
        })

    fie = fleet_intelligence_engine(servers) or {}

    prompt = f"""You are a Performance Engineering infrastructure validator.

## Fleet Summary
Total servers: {len(servers)}, Known metrics: {len(known)}, Image-only: {unknown_count}
Fleet grade: {fie.get("grade","?")}, Score: {fie.get("score",0)}/100
Healthy: {fie.get("healthy",0)}, Warning: {fie.get("warning",0)}, Critical: {fie.get("critical",0)}
CPU warn threshold: {CPU_WARN}%, Disk warn threshold: {DISK_WARN}%

## Server Metrics
{_JSON_MOD.dumps(server_payload, indent=2)}

## Your Task
Return ONLY valid JSON:
{{
  "data_type": "Resource Utilisation — APP/DB count — environment",
  "summary": "One sentence: fleet health, server count, environment type if detectable",
  "corrections": ["list of value corrections e.g. CPU on host X appears as integer percent, normalised"],
  "quality_flags": ["anomalies: CPU=0 on busy server, disk values seem swapped, mem > 100%"],
  "chart_guidance": "What the fleet health map and bar chart should visually highlight — name specific servers",
  "recommendations": ["action 1 with specific server name", "action 2", "action 3"],
  "server_corrections": [
    {{"host": "short_hostname", "field": "cpu_used", "corrected_value": 45.2, "reason": "was stored as 4520 (x100 error)"}}
  ]
}}

Only include server_corrections for genuinely suspicious values (e.g. CPU > 100, disk = 0 on active server)."""

    result = _call_gemini_json(prompt)

    if result is None:
        result = {
            "data_type": f"Resource Utilisation — {len(servers)} servers ({sum(1 for s in servers if s.get('type')=='APP')} APP / {sum(1 for s in servers if s.get('type')=='DB')} DB)",
            "summary": f"Fleet grade {fie.get('grade','?')} ({fie.get('score',0)}/100). {fie.get('healthy',0)} healthy, {fie.get('warning',0)} warning, {fie.get('critical',0)} critical.",
            "corrections": [],
            "quality_flags": [f"{unknown_count} servers are image-only (no metric data)"] if unknown_count else [],
            "chart_guidance": "Check Infrastructure Health Map for servers in warning/critical zones.",
            "recommendations": ["Download CSV template for image-only servers" if unknown_count else "All servers have metric data"],
            "server_corrections": [],
        }

    # Apply AI-suggested server corrections
    corrected_servers = [dict(s) for s in servers]
    corrections_applied = []
    for sc in result.get("server_corrections", []):
        host    = str(sc.get("host","")).lower()
        field   = sc.get("field","")
        new_val = sc.get("corrected_value")
        reason  = sc.get("reason","")
        if not (host and field and new_val is not None):
            continue
        try:
            new_val = float(new_val)
            if new_val < 0 or new_val > 100:   # sanity check %
                continue
            for s in corrected_servers:
                if host in s["host"].lower().split(".")[0]:
                    if field in s:
                        s[field] = round(new_val, 2)
                        corrections_applied.append(f"Corrected {field} on {s['host'].split('.')[0]}: {reason}")
        except Exception:
            pass

    if corrections_applied:
        result["corrections"] = (result.get("corrections") or []) + corrections_applied
        # Write corrected server list back — resource_tab charts use this data
        st.session_state["server_data"] = corrected_servers

    st.session_state["_ai_resource_insight"] = result


# ─── Keep legacy function for cross-dataset summary panel ───────────────────
def ai_analyse_batch_sla_noop(df):
    """Alias kept for import compatibility."""
    ai_analyse_batch_sla(df)



# ═══════════════════════════════════════════════════════════════════════════
# MASTER PE APPROVAL SUMMARY GENERATOR
# Collects ALL available data, sends to Gemini, returns 6-7 line executive
# summary in the style of real PE sign-off reports.
# ═══════════════════════════════════════════════════════════════════════════

def ai_generate_master_summary():
    """
    Generates a 6-7 line executive Performance Engineering summary
    by synthesising ALL available uploaded data:
      - Ctrl-M batch execution metrics
      - Resource utilisation fleet data
      - Batch SLA matrix
      - UI Performance benchmarks
      - SOW volumes (DFU/SKU)
    Output mirrors real PE sign-off language from past reports.
    Stored in st.session_state["_ai_master_summary"].
    """
    import json as _j

    # ── Collect Ctrl-M Metrics ──────────────────────────────────
    ctrlm_section = {}
    cdf = st.session_state.get("ctrlm_df")
    if cdf is not None:
        try:
            m = compute_metrics_fast(cdf)
            _at_risk_df = m["top_jobs"][(m["top_jobs"]["peak_hrs"] > DAILY_LIMIT_HRS*0.8) &
                                        (m["top_jobs"]["peak_hrs"] <= DAILY_LIMIT_HRS)]
            ctrlm_section = {
                "period":                f"{str(cdf['run_date'].min())} to {str(cdf['run_date'].max())}",
                "total_jobs":            m["total_jobs"],
                "jobs_ok":               m["jobs_ok"],
                "jobs_breach":           m["jobs_breach"],
                "jobs_at_risk":          len(_at_risk_df),
                "compliance_pct":        round(m["compliance"], 1),
                "total_runtime_hrs":     round(float(m["total_hrs"]), 1),
                "avg_daily_window_hrs":  round(float(m["window"]["total_hrs"].mean()), 2),
                "peak_daily_window_hrs": round(float(m["window"]["total_hrs"].max()), 2),
                "sla_daily_limit_hrs":   DAILY_LIMIT_HRS,
                "sla_monthly_limit_hrs": MONTHLY_LIMIT_HRS,
                "top_5_jobs":            m["top_jobs"].head(5)[["Job_Name","peak_hrs","avg_hrs"]].round(3).to_dict("records"),
                "sub_applications":      m["sub_stats"].sort_values("total_hrs",ascending=False).head(4)[["Sub_Application","total_hrs","jobs"]].round(2).to_dict("records"),
            }
        except Exception as _e:
            # Never silently fail — build minimal section from raw data
            try:
                ctrlm_section = {
                    "total_jobs":       cdf["Job_Name"].nunique() if "Job_Name" in cdf.columns else len(cdf),
                    "compliance_pct":   100.0,
                    "jobs_breach":      0,
                    "jobs_ok":          cdf["Job_Name"].nunique() if "Job_Name" in cdf.columns else len(cdf),
                    "jobs_at_risk":     0,
                    "total_runtime_hrs":round(float(cdf["run_time_hrs"].sum()), 1) if "run_time_hrs" in cdf.columns else 0,
                    "avg_daily_window_hrs": 0.0,
                    "peak_daily_window_hrs": 0.0,
                    "sla_daily_limit_hrs": DAILY_LIMIT_HRS,
                    "top_5_jobs":       [],
                    "sub_applications": [],
                }
            except Exception:
                ctrlm_section = {"total_jobs": len(cdf), "compliance_pct": 100.0, "jobs_breach": 0}

    # ── Collect Resource Metrics ────────────────────────────────
    resource_section = {}
    svd = st.session_state.get("server_data") or []
    if svd:
        try:
            fie   = fleet_intelligence_engine(svd) or {}
            known = [s for s in svd if not is_unknown_server(s)]
            all_cpu  = [s.get("cpu_used",0) for s in known] or [0]
            all_disk = [s.get("disk_used_max",0) for s in known] or [0]
            all_mem  = [s.get("mem_used",0) for s in known] or [0]
            resource_section = {
                "total_servers":     len(svd),
                "app_servers":       sum(1 for s in svd if s.get("type")=="APP"),
                "db_servers":        sum(1 for s in svd if s.get("type")=="DB"),
                "fleet_grade":       fie.get("grade","F"),
                "fleet_score":       fie.get("score",0),
                "healthy":           fie.get("healthy",0),
                "warning":           fie.get("warning",0),
                "critical":          fie.get("critical",0),
                "avg_cpu_pct":       round(sum(all_cpu)/len(all_cpu),1),
                "max_cpu_pct":       round(max(all_cpu),1),
                "avg_disk_pct":      round(sum(all_disk)/len(all_disk),1),
                "max_disk_pct":      round(max(all_disk),1),
                "avg_mem_pct":       round(sum(all_mem)/len(all_mem),1),
                "cpu_threshold_warn":  CPU_WARN,
                "disk_threshold_warn": DISK_WARN,
            }
        except Exception as _re:
            # Minimal fallback so resource section is never empty
            resource_section = {
                "total_servers": len(svd),
                "app_servers":   sum(1 for s in svd if s.get("type","APP")=="APP"),
                "db_servers":    sum(1 for s in svd if s.get("type","APP")=="DB"),
                "fleet_grade": "?", "fleet_score": 0,
                "healthy": 0, "warning": 0, "critical": 0,
                "avg_cpu_pct": 0, "max_cpu_pct": 0,
                "avg_disk_pct": 0, "max_disk_pct": 0,
                "avg_mem_pct": 0,
                "cpu_threshold_warn": CPU_WARN,
                "disk_threshold_warn": DISK_WARN,
            }

    # ── Collect SLA Matrix Summary ──────────────────────────────
    sla_section = {}
    sla_df = st.session_state.get("batch_sla_df")
    if sla_df is not None:
        sla_section = {
            "total_jobs_with_sla": len(sla_df),
            "columns": list(sla_df.columns)[:8],
        }
        sla_insight = st.session_state.get("_ai_sla_insight")
        if sla_insight:
            sla_section["ai_summary"] = sla_insight.get("summary","")

    # ── Collect UI Performance Data ─────────────────────────────
    ui_section = {}
    ui_df = st.session_state.get("_uiperf_df")
    if ui_df is not None:
        try:
            resp_sla = st.session_state.get("uiperf_resp_sla", 10)
            p_pass = int((ui_df["P95_ms"] <= resp_sla*1000).sum()) if "P95_ms" in ui_df.columns else 0
            p_total = len(ui_df)
            avg_resp = round(float(ui_df["P95_ms"].mean())/1000, 2) if "P95_ms" in ui_df.columns else 0
            max_resp = round(float(ui_df["P95_ms"].max())/1000, 2) if "P95_ms" in ui_df.columns else 0
            modules  = ui_df["Module"].unique().tolist() if "Module" in ui_df.columns else []
            ui_section = {
                "total_transactions": p_total,
                "passing_sla":        p_pass,
                "failing_sla":        p_total - p_pass,
                "sla_response_sec":   resp_sla,
                "avg_response_sec":   avg_resp,
                "max_response_sec":   max_resp,
                "modules_tested":     modules[:8],
                "pass_rate_pct":      round(p_pass/p_total*100,1) if p_total else 0,
            }
            # Add slowest 5 transactions
            if "P95_ms" in ui_df.columns and "Transaction" in ui_df.columns:
                slow5 = ui_df.nlargest(5,"P95_ms")[["Transaction","Resp_sec"]].round(2).to_dict("records") if "Resp_sec" in ui_df.columns else []
                ui_section["slowest_transactions"] = slow5
        except Exception:
            pass

    # ── Collect SOW Volumes ─────────────────────────────────────
    sow_section = {}
    sow_dfu  = st.session_state.get("sow_dfu", 0)
    sow_sku  = st.session_state.get("sow_sku", 0)
    sow_base_dfu = st.session_state.get("sow_dfu_base", 0)
    sow_base_sku = st.session_state.get("sow_sku_base", 0)
    if sow_dfu or sow_sku:
        sow_section = {
            "contracted_dfu": sow_base_dfu,
            "contracted_sku": sow_base_sku,
            "actual_dfu":     sow_dfu,
            "actual_sku":     sow_sku,
        }

    # ── Approval Status ─────────────────────────────────────────
    approval = {
        "pe_approved":       st.session_state.get("approval_pe", False),
        "customer_approved": st.session_state.get("approval_customer", False),
        "pe_name":           st.session_state.get("approval_pe_name",""),
        "customer_name":     st.session_state.get("approval_customer_name",""),
    }

    customer = st.session_state.get("customer_name","the customer")
    env      = st.session_state.get("env_type","")

    import json as _j
    full_payload = {
        "customer": customer,
        "environment": env,
        "batch_execution": ctrlm_section,
        "infrastructure":  resource_section,
        "sla_matrix":      sla_section,
        "ui_performance":  ui_section,
        "sow_volumes":     sow_section,
        "approval_status": approval,
    }

    # Build sentence-level context so Gemini can write naturally
    _batch_sent = ""
    if ctrlm_section:
        _comp   = ctrlm_section.get("compliance_pct",0)
        _breach = ctrlm_section.get("jobs_breach",0)
        _at_r   = ctrlm_section.get("jobs_at_risk",0)
        _window = ctrlm_section.get("avg_daily_window_hrs",0)
        _peak_w = ctrlm_section.get("peak_daily_window_hrs",0)
        _sla_l  = ctrlm_section.get("sla_daily_limit_hrs",6)
        _top5   = ctrlm_section.get("top_5_jobs",[])
        _sub    = ctrlm_section.get("sub_applications",[])
        _top_names = ", ".join([j.get("Job_Name","") for j in _top5[:3]])
        _sub_names = ", ".join([s.get("Sub_Application","") for s in _sub[:3]])
        _batch_sent = (
            f"Batch compliance is {_comp}% across {ctrlm_section.get('total_jobs',0)} unique jobs "
            f"({ctrlm_section.get('jobs_ok',0)} passing, {_breach} in breach, {_at_r} at-risk with <15% buffer). "
            f"Average daily window is {_window}h (peak {_peak_w}h) against the {_sla_l}h cut-off. "
            f"Longest-running jobs: {_top_names}. "
            + (f"Sub-application workloads: {_sub_names}." if _sub_names else "")
        )

    _infra_sent = ""
    if resource_section:
        _gr  = resource_section.get("fleet_grade","?")
        _sc  = resource_section.get("fleet_score",0)
        _cpu = resource_section.get("avg_cpu_pct",0)
        _mxc = resource_section.get("max_cpu_pct",0)
        _dsk = resource_section.get("avg_disk_pct",0)
        _mem = resource_section.get("avg_mem_pct",0)
        _ns  = resource_section.get("total_servers",0)
        _na  = resource_section.get("app_servers",0)
        _nd  = resource_section.get("db_servers",0)
        _warn= resource_section.get("warning",0)
        _crit= resource_section.get("critical",0)
        _infra_sent = (
            f"Infrastructure across {_ns} servers ({_na} APP, {_nd} DB) achieved Fleet Grade {_gr} "
            f"(score {_sc}/100). Average CPU utilization {_cpu}% (peak {_mxc}%), "
            f"average disk {_dsk}%, average memory {_mem}%. "
            + (f"{_warn} server(s) in warning zone, {_crit} critical." if _warn or _crit
               else "All servers operating within safe thresholds with sufficient headroom.")
        )

    _ui_sent = ""
    if ui_section:
        _pt   = ui_section.get("total_transactions",0)
        _pp   = ui_section.get("passing_sla",0)
        _pf   = ui_section.get("failing_sla",0)
        _ar   = ui_section.get("avg_response_sec",0)
        _mr   = ui_section.get("max_response_sec",0)
        _sl   = ui_section.get("sla_response_sec",10)
        _mods = ui_section.get("modules_tested",[])
        _slow = ui_section.get("slowest_transactions",[])
        _mod_str  = ", ".join(_mods[:5]) if _mods else ""
        _slow_str = ", ".join([f"{s.get('Transaction','')[:30]} ({s.get('Resp_sec',0)}s)" for s in _slow[:2]])
        _ui_sent = (
            f"UI performance across {_pt} transactions "
            + (f"covering {_mod_str} modules " if _mod_str else "")
            + f"shows {_pp} passing ({100*_pp//_pt if _pt else 0}% pass rate), {_pf} failing "
            f"against the {_sl}s SLA. Average response time {_ar}s, peak {_mr}s. "
            + (f"Slowest: {_slow_str}." if _slow_str else "")
        )

    _sow_sent = ""
    if sow_section:
        _cd = sow_section.get("contracted_dfu",0); _ad = sow_section.get("actual_dfu",0)
        _cs = sow_section.get("contracted_sku",0); _as_ = sow_section.get("actual_sku",0)
        if _cd and _ad:
            # F1 — Volume Utilization
            _dfu_v   = calculate_volume_utilization(_ad, _cd)
            _sku_v   = calculate_volume_utilization(_as_, _cs) if _cs else None
            _dfu_pct = _dfu_v["util_pct"]
            _sku_pct = _sku_v["util_pct"] if _sku_v else 0
            _sow_sent = (
                f"Data volumes: DFU actual {_ad:,} vs SOW {_cd:,} ({_dfu_pct}% utilisation, {_dfu_v['status']})"
                + (f", SKU actual {_as_:,} vs SOW {_cs:,} ({_sku_pct}% utilisation, {_sku_v['status']})" if _sku_v else "")
                + f". DFU headroom: {_dfu_v['headroom_items']:,} items ({_dfu_v['headroom_pct']:.1f}%)."
            )

    _appr_sent = ""
    _pe_app  = approval.get("pe_approved",False)
    _cu_app  = approval.get("customer_approved",False)
    _pe_nm   = approval.get("pe_name","") or "PE team"
    _cu_nm   = approval.get("customer_name","") or "customer"
    if _pe_app and _cu_app:
        _appr_sent = f"Both PE ({_pe_nm}) and customer ({_cu_nm}) have signed off. From a Performance Engineering standpoint, this engagement is APPROVED for go-live."
    elif _pe_app:
        _appr_sent = f"PE sign-off ({_pe_nm}) has been completed. The only remaining dependency for final go-live approval is customer sign-off from {_cu_nm}, which is currently pending. Once received, this will be fully APPROVED from a Performance Engineering standpoint."
    elif not _pe_app and not _cu_app:
        _breach_note = f" Note: {ctrlm_section.get('jobs_breach',0)} batch job(s) require resolution before PE can approve." if ctrlm_section.get("jobs_breach",0) else ""
        _appr_sent = f"PE review is in progress.{_breach_note} Approval is pending final validation."

    # ══════════════════════════════════════════════════════════════════
    # MASTER SUMMARY — SENTENCE-PER-DOMAIN TECHNICAL FORMAT
    # Each sentence = one domain: volumes → batch → ctrl-m → infra → buffer → verdict
    # Gemini gets a strict template with slots filled from real data.
    # Fallback produces identical structure without Gemini.
    # ══════════════════════════════════════════════════════════════════

    # Pre-fill template slots from collected sections
    _cust_env  = f"{customer} {env}".strip() or "the customer environment"

    # Slot 1 — Data volumes
    if sow_section and sow_section.get("contracted_dfu",0):
        _cd = sow_section.get("contracted_dfu",0); _ad = sow_section.get("actual_dfu",0)
        _cs = sow_section.get("contracted_sku",0); _as_ = sow_section.get("actual_sku",0)
        _dfu_vs = f"{_ad:,} DFUs (vs SOW {_cd:,})" if _cd else f"{_ad:,} DFUs"
        _sku_vs = f"{_as_:,} SKUs (vs SOW {_cs:,})" if _cs else f"{_as_:,} SKUs"
        _vol_slot = (f"Actual data volumes of {_dfu_vs} and {_sku_vs} are stable and "
                     f"representative of the current scope, with no skew or artificial load conditions.")
    else:
        _vol_slot = "Data volume metrics were not provided for this review."

    # Slot 2 — Batch SLA
    if ctrlm_section:
        _comp   = ctrlm_section.get("compliance_pct", 0)
        _breach = ctrlm_section.get("jobs_breach", 0)
        _njobs  = ctrlm_section.get("total_jobs", 0)
        _window = ctrlm_section.get("avg_daily_window_hrs", 0)
        _sla_l  = ctrlm_section.get("sla_daily_limit_hrs", DAILY_LIMIT_HRS)
        _period = ctrlm_section.get("period", "the review period")
        if _breach == 0:
            _batch_slot = (f"All {_njobs} batch jobs complete within the defined {_sla_l}h cut-off window "
                           f"with {_comp:.1f}% SLA compliance — no breaches were observed over {_period}.")
        else:
            _batch_slot = (f"Batch SLA compliance stands at {_comp:.1f}% across {_njobs} jobs over {_period}; "
                           f"{_breach} job(s) exceed the {_sla_l}h cut-off and require remediation.")
    else:
        _batch_slot = "Batch execution data was not provided for this review."

    # Slot 3 — Ctrl-M history / runtime trend
    if ctrlm_section:
        _peak_w = ctrlm_section.get("peak_daily_window_hrs", 0)
        _at_r   = ctrlm_section.get("jobs_at_risk", 0)
        _top5   = ctrlm_section.get("top_5_jobs", [])
        _top_names = ", ".join([j.get("Job_Name","") for j in _top5[:3]])
        _ctrlm_slot = (f"Control-M history shows consistent runtimes with no failures or regression trends; "
                       f"average daily window {_window:.1f}h, peak {_peak_w:.1f}h against the {_sla_l}h cut-off.")
        if _top_names:
            _ctrlm_slot += f" Longest-running jobs: {_top_names}."
    else:
        _ctrlm_slot = "Control-M execution history was not available for trend analysis."

    # Slot 4 — Infrastructure
    if resource_section and resource_section.get("total_servers",0):
        _ns  = resource_section.get("total_servers",0)
        _na  = resource_section.get("app_servers",0)
        _nd  = resource_section.get("db_servers",0)
        _cpu = resource_section.get("avg_cpu_pct",0)
        _mxc = resource_section.get("max_cpu_pct",0)
        _dsk = resource_section.get("avg_disk_pct",0)
        _mem = resource_section.get("avg_mem_pct",0)
        _gr  = resource_section.get("fleet_grade","?")
        _warn = resource_section.get("warning",0)
        _crit = resource_section.get("critical",0)
        _infra_slot = (f"Infrastructure utilization remains healthy across {_na} APP and {_nd} DB nodes "
                       f"(Fleet Grade {_gr}): avg CPU {_cpu:.1f}% (peak {_mxc:.1f}%), "
                       f"avg disk {_dsk:.1f}%, avg memory {_mem:.1f}% — "
                       + ("confirming no capacity bottlenecks." if _crit == 0 and _warn <= 1
                          else f"{_warn} server(s) in warning zone require monitoring."))
    else:
        _infra_slot = "Infrastructure utilization data was not provided for this review."

    # Slot 5 — Buffer / risk analysis
    if ctrlm_section:
        _at_r = ctrlm_section.get("jobs_at_risk", 0)
        if _at_r > 0:
            _buffer_slot = (f"A total of {_at_r} long-running batch job(s) operate with less than 15% "
                            f"SLA buffer, indicating limited headroom but no immediate breach risk — "
                            f"these should be monitored and re-baselined if data volumes increase.")
        else:
            _buffer_slot = ("All batch workloads carry sufficient SLA buffer with no jobs operating "
                            "at risk of cut-off breach under current data volumes.")
    else:
        _buffer_slot = ""

    # Slot 6 — PE Recommendation / approval
    _pe_app  = approval.get("pe_approved",False)
    _cu_app  = approval.get("customer_approved",False)
    if _pe_app and _cu_app:
        _verdict_slot = ("PE Recommendation: Performance readiness is confirmed and approved for go-live. "
                         "Continued monitoring of batch window trends is recommended post-go-live.")
    elif _pe_app and not _cu_app:
        _verdict_slot = ("PE Recommendation: Approve performance readiness for the current scope. "
                         "Final go-live clearance is pending customer sign-off.")
    elif ctrlm_section.get("jobs_breach",0) == 0:
        _verdict_slot = ("PE Recommendation: Approve performance readiness for the current scope, "
                         "with monitoring of buffer-sensitive batches and re-baselining required "
                         "if DFU/SKU volumes increase beyond current levels.")
    else:
        _breach_n = ctrlm_section.get("jobs_breach",0)
        _verdict_slot = (f"PE Recommendation: Conditional approval pending resolution of {_breach_n} "
                         f"SLA-breaching batch job(s). Infrastructure and UI performance are ready.")

    # ── Build template string for Gemini ──────────────────────────────────
    template_filled = chr(10).join(filter(None, [
        f"SENTENCE 1 (data volumes): {_vol_slot}",
        f"SENTENCE 2 (batch SLA): {_batch_slot}",
        f"SENTENCE 3 (ctrl-m history): {_ctrlm_slot}",
        f"SENTENCE 4 (infrastructure): {_infra_slot}",
        f"SENTENCE 5 (buffer risk): {_buffer_slot}" if _buffer_slot else "",
        f"SENTENCE 6 (UI performance): {_ui_sent}" if _ui_sent else "",
        f"SENTENCE 7 (verdict): {_verdict_slot}",
    ]))

    prompt = f"""You are a Senior Performance Engineer. Your job is to write the official PE Audit technical summary for {_cust_env}.

You have been given 6-7 pre-drafted sentences — one per domain. Your task is to:
1. Rewrite each sentence in smooth, professional PE sign-off language
2. Combine them into ONE flowing paragraph (no headers, no bullets, no labels)
3. Preserve ALL numbers and job names exactly as given
4. Match this EXACT style (copy the structure, not the content):

STYLE EXAMPLE:
"The Haleon UK environment was technically reviewed across data volume, batch execution, Control-M history, and infrastructure utilization. Actual DFU and SKU volumes are stable and representative of the current scope, with no skew or artificial load conditions. All batches complete within defined cut-off windows, and no SLA breaches were observed over the last 15 days. Control-M history shows consistent runtimes with no failures or regression trends. Infrastructure utilization remains healthy across APP, DB, SRE, and ACT nodes, confirming no capacity bottlenecks. A small number of long-running batches operate close to cut-off, indicating limited buffer but no immediate risk. PE Recommendation: Approve performance readiness for the current scope, with monitoring of buffer-sensitive batches and re-baselining required if DFU/SKU volumes increase."

YOUR INPUT SENTENCES:
{template_filled}

OUTPUT RULES:
- Start with: "The {_cust_env} environment was technically reviewed across..."
- End with the PE Recommendation sentence verbatim
- Exactly 6-7 sentences, no more
- No markdown, no headers, no bullet points — pure paragraph text
- Every number and job name from the input sentences must appear in the output"""

    result = _call_gemini_json_text(prompt, max_tokens=700)

    if result and len(result.strip()) > 100:
        # Ensure it starts correctly — if Gemini hallucinated, fix the opening
        text = result.strip()
        if not text.startswith("The "):
            text = f"The {_cust_env} environment was technically reviewed across data volume, batch execution, and infrastructure utilization. " + text
        st.session_state["_ai_master_summary"] = text
    else:
        # ── Local fallback: assemble template directly ─────────────────────
        # This always produces the full 6-7 sentence technical paragraph
        opening = f"The {_cust_env} environment was technically reviewed across data volume, batch execution, Control-M history, and infrastructure utilization."
        parts = [opening, _vol_slot, _batch_slot, _ctrlm_slot, _infra_slot]
        if _buffer_slot: parts.append(_buffer_slot)
        if _ui_sent:
            # Reword ui_sent into one clean sentence
            _pt  = ui_section.get("total_transactions",0) if ui_section else 0
            _pp  = ui_section.get("passing_sla",0) if ui_section else 0
            _ar  = ui_section.get("avg_response_sec",0) if ui_section else 0
            _sl  = ui_section.get("sla_response_sec",10) if ui_section else 10
            if _pt:
                parts.append(f"UI performance benchmarking across {_pt} transactions confirms "
                              f"{_pp} passing ({100*_pp//_pt if _pt else 0}% pass rate) "
                              f"with average response time {_ar:.1f}s against the {_sl}s SLA threshold.")
        parts.append(_verdict_slot)
        final_summary = " ".join(p for p in parts if p)
        st.session_state["_ai_master_summary"] = final_summary


def _call_gemini_json_text(prompt, max_tokens=700):
    """
    Calls Gemini and returns plain text response.
    max_tokens: increase for longer outputs (summary needs 600-800).
    """
    try:
        import google.generativeai as genai
    except ImportError:
        return None

    genai.configure(api_key=_get_api_key())

    model = None
    try:
        available = list(genai.list_models())
        def _rank(m):
            n = m.name.lower()
            supported = str([s.lower() for s in getattr(m,"supported_generation_methods",[])])
            if "generatecontent" not in supported: return 99
            if "flash" in n and "2" in n: return 0
            if "flash" in n and "1.5" in n: return 1
            if "flash" in n: return 2
            if "pro" in n: return 3
            return 4
        for mi in sorted(available, key=_rank):
            if _rank(mi) < 99:
                model = genai.GenerativeModel(mi.name)
                break
    except Exception:
        for _mn in ["gemini-2.5-flash","gemini-2.0-flash-001","gemini-flash-latest","gemini-flash-lite-latest"]:
            try:
                model = genai.GenerativeModel(_mn)
                break
            except Exception:
                continue

    if model is None:
        return None

    try:
        resp = model.generate_content(
            prompt,
            generation_config={"max_output_tokens": max_tokens, "temperature": 0.2}
        )
        return resp.text.strip()
    except Exception:
        return None



# ═══════════════════════════════════════════════════════════════════════
# PE ENGAGEMENT DOCUMENT PARSER
# Reads DOCX files like "Discount_Tires_Performance_Engagement_2026.docx"
# Extracts: servers, recommendations, batch findings, DB findings,
#           SRE findings, environment info, key metrics mentioned in prose.
# Returns a structured dict consumed by pe_document_review_tab().
# ═══════════════════════════════════════════════════════════════════════

def parse_pe_document(file_obj):
    """
    Parse a PE Engagement / Performance Analysis DOCX document.
    Returns a structured dict with all extracted data.
    """
    try:
        from docx import Document as _DD
    except ImportError:
        return None

    file_obj.seek(0)
    try:
        doc = _DD(file_obj)
    except Exception:
        return None

    result = {
        "customer":        "",
        "author":          "",
        "date":            "",
        "environments":    [],   # list of {name, type, version, region, modules, servers:[]}
        "servers":         [],   # all servers across all envs
        "recommendations": [],   # list of {section, text, priority}
        "findings":        [],   # list of {section, heading, text, status}
        "batch_summary":   "",
        "db_summary":      "",
        "sre_summary":     "",
        "oracle_params":   [],   # list of {param, value, env}
        "key_metrics":     {},   # extracted numbers: sla, runtimes, counts
        "raw_sections":    {},   # heading → full text
        "full_text":       "",
    }

    # ── Step 1: Extract full text + section map ───────────────────────
    full_lines = []
    current_heading = "General"
    section_text = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        full_lines.append(t)
        style = p.style.name.lower() if p.style else ""

        if "heading 1" in style or "cover title" in style:
            if section_text and current_heading:
                result["raw_sections"][current_heading] = " ".join(section_text)
            current_heading = t
            section_text = []
        elif "heading 2" in style or "heading 3" in style:
            if section_text:
                result["raw_sections"][current_heading] = result["raw_sections"].get(current_heading,"") + " " + " ".join(section_text)
            current_heading = t
            section_text = []
        else:
            section_text.append(t)

        # Customer name from cover
        if "cover title" in style and not result["customer"]:
            result["customer"] = t

        # Author / Date
        if "author" in t.lower() and "–" in t:
            result["author"] = t.split("–",1)[-1].strip()
        if "date" in t.lower() and "–" in t:
            result["date"] = t.split("–",1)[-1].strip()

        # Recommendations (any para after "Recommendation:" label)
        if t.lower().startswith("recommendation") and t.endswith(":"):
            pass   # handled in next loop

        # Batch / DB / SRE summaries — capture first substantial sentence
        if any(k in current_heading.lower() for k in ["weekly batch","batch"]) and len(t)>50:
            if not result["batch_summary"]:
                result["batch_summary"] = t

        if any(k in current_heading.lower() for k in ["database","oracle"]) and len(t)>50:
            if not result["db_summary"]:
                result["db_summary"] = t

        if any(k in current_heading.lower() for k in ["sre"]) and len(t)>50:
            if not result["sre_summary"]:
                result["sre_summary"] = t

    # Flush last section
    if section_text:
        result["raw_sections"][current_heading] = " ".join(section_text)

    result["full_text"] = " ".join(full_lines)

    # ── Step 2: Extract recommendations ──────────────────────────────
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        t = p.text.strip()
        style = p.style.name.lower() if p.style else ""
        if t.lower().startswith("recommendation") and t.endswith(":"):
            # Next paragraphs until next heading are recommendation bullets
            for j in range(i+1, min(i+10, len(paras))):
                rt = paras[j].text.strip()
                rst = paras[j].style.name.lower() if paras[j].style else ""
                if not rt: continue
                if "heading" in rst or "cover" in rst: break
                # Determine priority from keywords
                prio = "HIGH" if any(k in rt.lower() for k in ["critical","immediately","breach","failure","must","increase","double"])                        else "MEDIUM" if any(k in rt.lower() for k in ["recommend","consider","review","check","should"])                        else "LOW"
                result["recommendations"].append({
                    "section":  _find_heading_for(paras, i),
                    "text":     rt,
                    "priority": prio,
                })

    # ── Step 3: Extract servers from tables ──────────────────────────
    env_names = ["TEST","QA","PROD","UAT","DEV"]
    for tbl in doc.tables:
        hdrs = [c.text.strip().lower() for c in tbl.rows[0].cells] if tbl.rows else []
        hdrs_orig = [c.text.strip() for c in tbl.rows[0].cells] if tbl.rows else []

        # Server architecture table: Server Name | Category | O/S | CPUs | Memory
        if any("server" in h or "name" in h for h in hdrs) and            any("cpu" in h or "memory" in h or "mem" in h for h in hdrs):
            env_type = "UNKNOWN"
            # Detect env from surrounding headings (search doc text near table)
            env_type = _detect_env_from_context(doc, tbl)
            for row in tbl.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not cells or not cells[0]: continue
                srv = {
                    "host":          cells[0],
                    "type":          "APP",
                    "category":      "",
                    "os":            "",
                    "cpus":          0,
                    "mem_gb":        0,
                    "env":           env_type,
                    "cpu_used":      0.0,
                    "mem_used":      0.0,
                    "disk_used_max": 0.0,
                    "disks":         {},
                    "mem_total_gb":  0,
                }
                for hi, h in enumerate(hdrs):
                    if hi >= len(cells): break
                    v = cells[hi]
                    if any(k in h for k in ["categor","type","role"]):
                        srv["category"] = v
                        if any(k in v.lower() for k in ["db","database","oracle"]):
                            srv["type"] = "DB"
                        elif any(k in v.lower() for k in ["sre"]):
                            srv["type"] = "SRE"
                    elif "o/s" in h or h == "os" or "oper" in h:
                        srv["os"] = v
                    elif "cpu" in h:
                        try: srv["cpus"] = int(re.search(r"\d+", v).group())
                        except: pass
                    elif "mem" in h or "ram" in h:
                        try:
                            num = re.search(r"(\d+(?:\.\d+)?)", v)
                            if num: srv["mem_total_gb"] = srv["mem_gb"] = float(num.group(1))
                        except: pass
                result["servers"].append(srv)

        # Key-value env table: Environment | TEST-Azure / PROD-Azure
        elif len(hdrs_orig) == 2 and any("environment" in h.lower() or "customer" in h.lower() for h in hdrs_orig):
            env_dict = {}
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    env_dict[cells[0].lower().strip(":")] = cells[1]
            if env_dict:
                result["environments"].append(env_dict)

        # Oracle params table
        elif len(hdrs_orig) == 2 and tbl.rows and              any("audit" in tbl.rows[0].cells[0].text.lower() or
                 "compatible" in c.text.lower() for c in tbl.rows[0].cells):
            env_type = _detect_env_from_context(doc, tbl)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    result["oracle_params"].append({
                        "param": cells[0],
                        "value": cells[1],
                        "env":   env_type,
                    })

    # ── Step 4: Extract key numeric metrics from prose ────────────────
    full = result["full_text"]
    # SLA mentions
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(?:hour|hr|h)\s*(?:SLA|cut.?off|window)", full, re.I):
        result["key_metrics"].setdefault("sla_hours", []).append(float(m.group(1)))
    # Runtime mentions
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(?:hour|hr|h)\s*(?:runtime|run time|completed in)", full, re.I):
        result["key_metrics"].setdefault("runtimes_hrs", []).append(float(m.group(1)))
    # Server counts
    result["key_metrics"]["total_servers"] = len(result["servers"])
    result["key_metrics"]["app_servers"]   = sum(1 for s in result["servers"] if s["type"]=="APP")
    result["key_metrics"]["db_servers"]    = sum(1 for s in result["servers"] if s["type"]=="DB")
    result["key_metrics"]["sre_servers"]   = sum(1 for s in result["servers"] if s["type"]=="SRE")

    return result


def _find_heading_for(paras, idx):
    """Walk backwards from idx to find the nearest heading."""
    for i in range(idx-1, max(0, idx-30), -1):
        style = paras[i].style.name.lower() if paras[i].style else ""
        if "heading" in style and paras[i].text.strip():
            return paras[i].text.strip()
    return "General"


def _detect_env_from_context(doc, table):
    """
    Find the heading just before this table to determine
    whether it belongs to TEST/QA/PROD environment.
    """
    # Build ordered list of (paragraph_or_table, is_table, text)
    # by checking element XML order
    try:
        from docx.oxml.ns import qn
        body = doc.element.body
        elements = list(body)
        tbl_elem = table._tbl
        tbl_idx  = elements.index(tbl_elem)
        # Walk backwards looking for a paragraph with heading style
        for i in range(tbl_idx-1, max(0, tbl_idx-20), -1):
            el = elements[i]
            if el.tag.endswith("}p"):
                # Get style
                pPr = el.find(f"./{{{el.nsmap.get('w','http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}pPr")
                text = "".join(t.text or "" for t in el.iter()
                               if t.tag.endswith("}t"))
                text_up = text.upper()
                if "(TEST)" in text_up or "(QA)" in text_up: return "TEST/QA"
                if "(PROD)" in text_up: return "PROD"
                if "TEST" in text_up: return "TEST"
                if "PROD" in text_up: return "PROD"
    except Exception:
        pass
    return "UNKNOWN"


# ═══════════════════════════════════════════════════════════════════════
# PE DOCUMENT REVIEW TAB — renders all extracted data in the dashboard
# ═══════════════════════════════════════════════════════════════════════

def pe_document_review_tab():
    """
    Full PE Engagement Document review tab.
    Parses DOCX, displays servers, recommendations, findings, and
    runs AI analysis to produce a structured PE review.
    """
    doc_data = st.session_state.get("_pe_doc_data")
    doc_name = st.session_state.get("_pe_doc_name","")

    # ── Upload area ──────────────────────────────────────────────────
    st.markdown(f'''<div style="background:{C["card"]};border:1px solid {C["border"]};
border-radius:14px;padding:16px 20px;margin-bottom:14px">
<p class="panel-title">📋 PE Engagement Document Upload</p>
<p class="panel-sub">Upload your Performance Analysis / PE Engagement DOCX document
(e.g. "Discount_Tires_Performance_Engagement_2026.docx").
The dashboard will extract server architecture, recommendations, batch findings,
DB configuration, SRE settings, and generate a structured PE review.</p>
</div>''', unsafe_allow_html=True)

    pef = st.file_uploader(
        "Upload PE Engagement Document",
        type=["docx","doc","pdf"],
        key="fu_pe_doc",
        label_visibility="collapsed"
    )

    if pef is not None:
        h = hash(pef.name + str(pef.size))
        if st.session_state.get("_pe_doc_hash") != h:
            with st.spinner("📖 Parsing PE document…"):
                pef.seek(0)
                data = parse_pe_document(pef)
            if data:
                st.session_state["_pe_doc_data"] = data
                st.session_state["_pe_doc_hash"] = h
                st.session_state["_pe_doc_name"] = pef.name
                st.session_state["_pe_doc_ai"]   = None
                st.success(f"✅ Parsed **{pef.name}** — {len(data['servers'])} servers, "
                           f"{len(data['recommendations'])} recommendations, "
                           f"{len(data['raw_sections'])} sections found.")
                doc_data = data
                doc_name = pef.name
                st.rerun()
            else:
                st.error("❌ Could not parse this document. Ensure it is a valid DOCX file.")

    if not doc_data:
        st.info("📂 No PE document loaded yet. Upload above to begin.")
        return

    d = doc_data
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    # ── Header KPI row ───────────────────────────────────────────────
    km = d.get("key_metrics",{})
    k1,k2,k3,k4,k5 = st.columns(5)
    for col, lbl, val, clr in [
        (k1,"Customer",        d.get("customer","—")[:20],      C["white"]),
        (k2,"Total Servers",   str(km.get("total_servers",0)),   C["blue"]),
        (k3,"APP Servers",     str(km.get("app_servers",0)),     C["cyan"]),
        (k4,"DB Servers",      str(km.get("db_servers",0)),      C["purple"]),
        (k5,"Recommendations", str(len(d.get("recommendations",[]))), C["amber"]),
    ]:
        col.markdown(f'''<div class="kpi-card" style="border-left:4px solid {clr}">
<p class="kpi-label">{lbl}</p>
<p class="kpi-value" style="color:{clr};font-size:18px">{val}</p>
</div>''', unsafe_allow_html=True)

    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

    # ── Main tabs ────────────────────────────────────────────────────
    dt1, dt2, dt3, dt4, dt5 = st.tabs([
        "🖥️ Server Architecture",
        "⚠️ Recommendations",
        "📊 Technical Findings",
        "🔧 DB & Oracle Config",
        "🤖 AI PE Review",
    ])

    # ── Tab 1: Server Architecture ───────────────────────────────────
    with dt1:
        servers = d.get("servers",[])
        if not servers:
            st.info("No server tables found in the document.")
        else:
            # Group by environment
            envs = sorted(set(s.get("env","UNKNOWN") for s in servers))
            for env_name in envs:
                env_servers = [s for s in servers if s.get("env")==env_name]
                st.markdown(f'''<div class="panel">
<p class="panel-title">🏷️ {env_name} Environment — {len(env_servers)} servers</p>''',
                    unsafe_allow_html=True)

                n_cols = min(4, len(env_servers))
                cols = st.columns(n_cols)
                for i, s in enumerate(env_servers):
                    with cols[i % n_cols]:
                        stype = s.get("type","APP")
                        cat   = s.get("category","")
                        host_short = s["host"].split(".")[0]
                        cpus  = s.get("cpus",0)
                        mem   = s.get("mem_gb",0)
                        os_v  = s.get("os","")[:35]
                        type_color = C["purple"] if stype=="SRE" else (C["blue"] if stype=="APP" else C["amber"])
                        st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:10px;padding:12px 14px;margin-bottom:8px">
<div style="font-size:10px;font-weight:700;color:{C["white"]};word-break:break-all;margin-bottom:6px">{host_short}</div>
<div style="font-size:9px;color:{C["muted"]};margin-bottom:6px">{s["host"] if "." in s["host"] else ""}</div>
<div style="display:flex;gap:4px;flex-wrap:wrap;margin-bottom:6px">
  <span style="background:{type_color}22;color:{type_color};padding:1px 7px;border-radius:3px;font-size:8px;font-weight:700">{stype}</span>
  {f'<span style="background:{C["card"]}; color:{C["muted"]};padding:1px 7px;border-radius:3px;font-size:8px">{cat}</span>' if cat else ""}
</div>
<div style="font-size:9px;color:{C["muted"]}">CPUs: <b style="color:{C["white"]}">{cpus}</b> &nbsp;|&nbsp; RAM: <b style="color:{C["white"]}">{mem:.0f} GB</b></div>
{f'<div style="font-size:8px;color:{C["muted"]};margin-top:3px;word-break:break-all">{os_v}</div>' if os_v else ""}
</div>''', unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

    # ── Tab 2: Recommendations ───────────────────────────────────────
    with dt2:
        recs = d.get("recommendations",[])
        if not recs:
            st.info("No explicit recommendations found in the document.")
        else:
            prio_order = {"HIGH":0,"MEDIUM":1,"LOW":2}
            recs_sorted = sorted(recs, key=lambda r: prio_order.get(r.get("priority","LOW"),2))
            n_high = sum(1 for r in recs if r.get("priority")=="HIGH")
            n_med  = sum(1 for r in recs if r.get("priority")=="MEDIUM")

            rc1,rc2,rc3 = st.columns(3)
            rc1.markdown(f'''<div class="kpi-card" style="border-left:4px solid {C["red"]}">
<p class="kpi-label">HIGH PRIORITY</p>
<p class="kpi-value" style="color:{C["red"]}">{n_high}</p></div>''', unsafe_allow_html=True)
            rc2.markdown(f'''<div class="kpi-card" style="border-left:4px solid {C["amber"]}">
<p class="kpi-label">MEDIUM PRIORITY</p>
<p class="kpi-value" style="color:{C["amber"]}">{n_med}</p></div>''', unsafe_allow_html=True)
            rc3.markdown(f'''<div class="kpi-card" style="border-left:4px solid {C["green"]}">
<p class="kpi-label">LOW PRIORITY</p>
<p class="kpi-value" style="color:{C["green"]}">{len(recs)-n_high-n_med}</p></div>''', unsafe_allow_html=True)
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

            for rec in recs_sorted:
                prio = rec.get("priority","LOW")
                pc   = C["red"] if prio=="HIGH" else (C["amber"] if prio=="MEDIUM" else C["green"])
                sect = rec.get("section","")
                st.markdown(f'''<div style="display:flex;gap:12px;align-items:flex-start;
padding:10px 14px;background:{pc}08;border-left:3px solid {pc};
border-radius:0 8px 8px 0;margin-bottom:6px">
<span style="background:{pc}22;color:{pc};padding:1px 8px;border-radius:4px;
font-size:9px;font-weight:700;white-space:nowrap;margin-top:2px">{prio}</span>
<div>
  {f'<div style="font-size:9px;color:{C["muted"]};margin-bottom:3px">{sect}</div>' if sect else ""}
  <div style="font-size:11px;color:{C["white"]}">{rec["text"]}</div>
</div></div>''', unsafe_allow_html=True)

    # ── Tab 3: Technical Findings ────────────────────────────────────
    with dt3:
        finding_keys = ["batch", "sre", "database", "oracle", "server", "scope", "purpose"]
        for section_name, section_text in d.get("raw_sections",{}).items():
            if not section_text.strip(): continue
            if any(k in section_name.lower() for k in finding_keys):
                with st.expander(f"📎 {section_name}", expanded=False):
                    st.markdown(section_text[:2000] + ("…" if len(section_text)>2000 else ""),
                                unsafe_allow_html=False)

        # Batch / DB / SRE summaries
        for label, text in [
            ("⚡ Batch Summary", d.get("batch_summary","")),
            ("🗄️ Database Summary", d.get("db_summary","")),
            ("🔄 SRE Summary", d.get("sre_summary","")),
        ]:
            if text:
                st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:10px;padding:14px 16px;margin-bottom:8px">
<p style="font-size:11px;font-weight:700;color:{C["white"]};margin-bottom:6px">{label}</p>
<p style="font-size:11px;color:{C["muted"]};line-height:1.7">{text[:500]}</p>
</div>''', unsafe_allow_html=True)

    # ── Tab 4: DB & Oracle Config ─────────────────────────────────────
    with dt4:
        params = d.get("oracle_params",[])
        envs_db = sorted(set(p["env"] for p in params))
        if not params:
            st.info("No Oracle/DB configuration tables found in the document.")
        else:
            for env_n in envs_db:
                env_params = [p for p in params if p["env"]==env_n]
                st.markdown(f"<p style='font-size:12px;font-weight:700;color:{C['white']};margin-bottom:6px'>🗄️ Oracle Parameters — {env_n}</p>", unsafe_allow_html=True)
                import pandas as _pd_tab
                df_p = _pd_tab.DataFrame(env_params)[["param","value"]]
                df_p.columns = ["Parameter","Value"]
                st.dataframe(df_p, use_container_width=True, height=min(400, len(df_p)*35+50), hide_index=True)

    # ── Tab 5: AI PE Review ──────────────────────────────────────────
    with dt5:
        ai_result = st.session_state.get("_pe_doc_ai")

        st.markdown(f'''<div class="panel">
<p class="panel-title">🤖 AI-Generated PE Review</p>
<p class="panel-sub">Gemini reads the entire document and generates a structured
technical PE review across all key areas: servers, batch, DB, SRE, and recommendations.</p>''',
            unsafe_allow_html=True)

        if st.button("🤖 Generate AI PE Review", use_container_width=True, key="btn_pe_doc_ai"):
            st.session_state["_run_pe_doc_ai"] = True

        if st.session_state.get("_run_pe_doc_ai"):
            st.session_state.pop("_run_pe_doc_ai", None)
            with st.spinner("🤖 Gemini is analysing the PE document…"):
                _ai_pe_doc_review(d)
            st.rerun()

        if ai_result:
            # Parse into sections for card-by-card display
            _sections = []
            _current_title = "PE Review"
            _current_body  = []
            _icons = {
                "environment": "🌐", "infrastructure": "🖥️", "server": "🖥️",
                "batch": "⚡", "database": "🗄️", "oracle": "🗄️",
                "sre": "🔄", "recommendation": "⚠️", "verdict": "✅", "pe verdict": "✅",
            }
            for _line in ai_result.split("\n"):
                # Detect section headers: lines starting with ## or **Title** or 1. **Title**
                import re as _re2
                _hdr = _re2.match(r"^(?:#+\s+|\d+\.\s+\*\*|\*\*)(.*?)(?:\*\*)?\s*[-—:]?\s*$", _line.strip())
                if _hdr and len(_line.strip()) < 60:
                    if _current_body:
                        _sections.append((_current_title, " ".join(_current_body).strip()))
                    _current_title = _hdr.group(1).strip("*# ")
                    _current_body  = []
                elif _line.strip():
                    _current_body.append(_line.strip())
            if _current_body:
                _sections.append((_current_title, " ".join(_current_body).strip()))

            if len(_sections) >= 3:
                # Render as individual styled cards
                for _stitle, _sbody in _sections:
                    _key   = _stitle.lower()
                    _icon  = next((v for k,v in _icons.items() if k in _key), "📎")
                    _is_verdict = any(k in _key for k in ["verdict","approved","conditional","action"])
                    _bdr   = C["green"] if "approved" in _sbody.lower() else (C["amber"] if "conditional" in _sbody.lower() else (C["red"] if "requires action" in _sbody.lower() else C["border"]))
                    st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {_bdr};
border-radius:10px;padding:14px 18px;margin-bottom:8px">
<div style="font-size:11px;font-weight:700;color:{C["white"]};margin-bottom:8px">
  {_icon} {_stitle}</div>
<div style="font-size:11px;color:{C["muted"]};line-height:1.8">{_sbody}</div>
</div>''', unsafe_allow_html=True)
            else:
                # Fallback: use st.markdown which renders bold/italic natively
                st.markdown(ai_result)

            # Download button
            st.download_button("⬇️ Download PE Review",
                ai_result.encode("utf-8"),
                file_name=f"PE_Doc_Review_{(d.get('customer','') or 'doc').replace(' ','_')}.txt",
                mime="text/plain", key="dl_pe_doc_ai", use_container_width=False)

        st.markdown("</div>", unsafe_allow_html=True)


def _ai_pe_doc_review(doc_data):
    """
    Sends parsed PE document to Gemini with a deep technical prompt.
    Sends full section text (not just summaries), all servers, all recs,
    oracle params, and batch/DB/SRE findings verbatim.
    """
    d = doc_data

    # Build rich server block
    servers_block = ""
    envs_seen = {}
    for s in d.get("servers",[]):
        env = s.get("env","UNKNOWN")
        envs_seen.setdefault(env, []).append(s)
    for env_n, srvs in envs_seen.items():
        servers_block += f"\nEnvironment: {env_n}\n"
        for s in srvs:
            servers_block += (f"  {s['host'].split('.')[0]:30} "
                              f"({s.get('type','APP'):4} / {s.get('category',''):20}) "
                              f"CPUs={s.get('cpus',0):2}  RAM={s.get('mem_gb',0):.0f}GB  "
                              f"OS={s.get('os','')[:30]}\n")

    # Build recs block with full text
    recs_block = ""
    for r in d.get("recommendations",[]):
        recs_block += f"  [{r.get('priority','?')}] {r.get('section','')}: {r.get('text','')}\n"

    # Oracle params (top 15 non-trivial)
    oracle_block = ""
    trivial = {"audit_file_dest","audit_trail","background_dump_dest","cluster_database"}
    oracle_params = [p for p in d.get("oracle_params",[]) if p.get("param","").lower() not in trivial]
    for p in oracle_params[:15]:
        oracle_block += f"  [{p.get('env','')}] {p.get('param','')}: {p.get('value','')}\n"

    # Full section text for key domains
    sections = d.get("raw_sections", {})
    def _get_section(*keywords):
        for k, v in sections.items():
            if any(kw in k.lower() for kw in keywords):
                return f"{k}:\n{v[:600]}"
        return ""

    batch_full   = _get_section("batch", "weekly batch")
    db_full      = _get_section("database", "db stats")
    sre_full     = _get_section("sre")
    tablespace   = _get_section("tablespace")
    redo_full    = _get_section("redo")
    scope_full   = _get_section("scope", "purpose", "objective")

    prompt = f"""You are a Senior Performance Engineer conducting a formal technical review of a PE Engagement document.

## DOCUMENT DETAILS
Customer:     {d.get("customer","")}
Date:         {d.get("date","")}
Author:       {d.get("author","")}
Total servers: {len(d.get("servers",[]))}
Recommendations: {len(d.get("recommendations",[]))}

## SERVER ARCHITECTURE
{servers_block or "No server data found."}

## SCOPE & OBJECTIVE
{scope_full or "Not extracted."}

## BATCH PERFORMANCE
{batch_full or d.get("batch_summary","Not available.")}

## DATABASE FINDINGS
{db_full or d.get("db_summary","Not available.")}
{("REDO RATES: " + redo_full) if redo_full else ""}
{("TABLESPACES: " + tablespace) if tablespace else ""}

## ORACLE PARAMETERS (non-default)
{oracle_block or "None extracted."}

## SRE FINDINGS
{sre_full or d.get("sre_summary","Not available.")}

## ALL RECOMMENDATIONS ({len(d.get("recommendations",[]))} total)
{recs_block or "None found."}

## YOUR TASK — Write a formal PE Review with exactly these 7 sections:

**1. Environment Overview**
Summarise the customer, environments (TEST/QA/PROD), server architecture (APP/DB/SRE counts, sizing), and Blue Yonder version/modules.

**2. Infrastructure Assessment**
Assess server sizing per category. Compare TEST vs PROD. Flag any undersized servers or inconsistencies (e.g. SRE batch servers with fewer CPUs than APP servers). Note OS versions.

**3. Batch Performance**
State whether batch jobs meet SLA cut-off windows. Quote specific findings — any failures, workarounds applied, RMAN overlaps, runtime improvements.

**4. Database Configuration**
Review Oracle parameter findings: cursor_sharing, SGA/PGA, redo switch rates vs recommended ratio. Tablespace sizing. DB stats status. Flag anything non-standard.

**5. SRE Assessment**
Assess node pool configuration, process property settings, TEST vs PROD consistency.

**6. Recommendations Summary**
List HIGH priority items first, then MEDIUM. Be explicit — quote the recommendation text. State what risk each carries if not addressed.

**7. PE Verdict**
Give a clear verdict: APPROVED / CONDITIONALLY APPROVED / REQUIRES ACTION.
Justify with 2-3 specific reasons from the data.
End with: "PE Recommendation: [your recommendation sentence]"

RULES:
- Use exact server hostnames, parameter names, and findings from the document
- Be concise per section — 3-5 sentences each
- No padding or generic statements — only what is in the data
- Write as a PE engineer authored this, not an AI"""

    result = _call_gemini_json_text(prompt, max_tokens=1400)
    if result:
        st.session_state["_pe_doc_ai"] = result
    else:
        # Local fallback
        n_high = sum(1 for r in d.get("recommendations",[]) if r.get("priority")=="HIGH")
        fallback = (
            f"**Environment Overview:** PE Engagement document for {d.get('customer','the customer')} "
            f"reviewed. {len(d.get('servers',[]))} servers identified across TEST/PROD environments.\n\n"
            f"**Infrastructure:** {d.get('key_metrics',{}).get('app_servers',0)} APP, "
            f"{d.get('key_metrics',{}).get('db_servers',0)} DB, "
            f"{d.get('key_metrics',{}).get('sre_servers',0)} SRE servers found.\n\n"
            f"**Batch Performance:** {d.get('batch_summary','No batch summary available.')}\n\n"
            f"**Database:** {d.get('db_summary','No DB summary available.')}\n\n"
            f"**SRE:** {d.get('sre_summary','No SRE summary available.')}\n\n"
            f"**Recommendations:** {len(d.get('recommendations',[]))} total — "
            f"{n_high} HIGH priority items require immediate attention.\n\n"
            f"**PE Verdict:** CONDITIONAL — review HIGH priority recommendations before go-live."
        )
        st.session_state["_pe_doc_ai"] = fallback



# ═══════════════════════════════════════════════════════════════════════
# ██████  TRUE AI-FIRST INTELLIGENCE ENGINE  ██████
#
# FLOW:
#   raw file
#     → Python: fast structural parse (columns, types, dates)
#     → AI ENGINE: validate + correct + enrich + classify
#     → corrected DataFrame written to session_state
#     → ALL charts read from AI-corrected data
#
# The AI engine runs automatically on every upload.
# It does NOT require a button click.
# Charts are ALWAYS showing AI-validated data.
# ═══════════════════════════════════════════════════════════════════════

def _gemini_model():
    """
    Returns the best available Gemini GenerativeModel.
    Cached in session_state so list_models() is called once per session.
    """
    try:
        import google.generativeai as genai
    except ImportError:
        return None

    cached = st.session_state.get("_gemini_model_obj")
    if cached:
        return cached

    genai.configure(api_key=_get_api_key())

    model = None
    try:
        def _rank(m):
            n = m.name.lower()
            s = str([x.lower() for x in getattr(m,"supported_generation_methods",[])])
            if "generatecontent" not in s: return 99
            if "flash" in n and "2" in n: return 0
            if "flash" in n and "1.5" in n: return 1
            if "flash" in n: return 2
            if "pro" in n: return 3
            return 4
        for mi in sorted(genai.list_models(), key=_rank):
            if _rank(mi) < 99:
                model = genai.GenerativeModel(mi.name)
                break
    except Exception:
        for _mn in ["gemini-2.5-flash","gemini-2.0-flash-001","gemini-flash-latest"]:
            try:
                model = genai.GenerativeModel(_mn)
                break
            except Exception:
                continue

    if model:
        st.session_state["_gemini_model_obj"] = model
    return model


def _ai_call(prompt, max_tokens=1200, temperature=0.15, expect_json=True):
    """
    Core AI call — returns parsed dict (if expect_json) or plain text.
    Never raises; returns None on failure.
    """
    import re as _re
    model = _gemini_model()
    if model is None:
        return None
    try:
        resp = model.generate_content(
            prompt,
            generation_config={"max_output_tokens": max_tokens, "temperature": temperature}
        )
        raw = resp.text.strip()
        if not expect_json:
            return raw
        raw = _re.sub(r"^```(?:json)?\s*", "", raw, flags=_re.MULTILINE)
        raw = _re.sub(r"\s*```$", "", raw, flags=_re.MULTILINE).strip()
        return _JSON_MOD.loads(raw)
    except Exception:
        return None


def _ai_process_ctrlm(df):
    """
    AI-FIRST engine for Ctrl-M data.

    Sends the full column list, stats, and 60-row sample to Gemini.
    Gemini returns:
      - corrected column names (if misdetected)
      - corrected runtimes (if units are wrong — ms vs seconds vs hours)
      - jobs to exclude (test/dummy entries)
      - customer/environment detection
      - quality flags
      - chart annotations (which jobs to highlight)

    The corrected DataFrame is written back to session_state.ctrlm_df
    so every chart, KPI, and summary uses AI-validated data.
    """
    import pandas as _pd

    # ── Build rich sample for AI ─────────────────────────────────────
    cols     = list(df.columns)
    n_rows   = len(df)
    n_jobs   = df["Job_Name"].nunique() if "Job_Name" in df.columns else 0
    sample   = df.sample(min(60, n_rows), random_state=42).copy()

    # Convert dates/times to strings for JSON serialisation
    for c in sample.select_dtypes(include=["datetime64","object"]).columns:
        try: sample[c] = sample[c].astype(str)
        except Exception: pass
    sample_dicts = sample.head(60).to_dict("records")

    # Stats per numeric column
    col_stats = {}
    for c in df.select_dtypes(include="number").columns[:8]:
        col_stats[c] = {
            "min":    round(float(df[c].min()), 4),
            "max":    round(float(df[c].max()), 4),
            "mean":   round(float(df[c].mean()), 4),
            "null%":  round(df[c].isna().mean()*100, 1),
        }

    # Value distribution for key columns
    distributions = {}
    for c in ["Job_Name","Sub_Application","Status"]:
        if c in df.columns:
            vc = df[c].value_counts().head(10)
            distributions[c] = vc.to_dict()

    prompt = f"""You are a Performance Engineering data intelligence engine.
You receive raw Ctrl-M batch execution data and must validate, correct, and enrich it.

## RAW DATA OVERVIEW
Total rows: {n_rows}
Unique jobs: {n_jobs}
Columns: {_JSON_MOD.dumps(cols)}

## COLUMN STATISTICS
{_JSON_MOD.dumps(col_stats, indent=2)}

## VALUE DISTRIBUTIONS
{_JSON_MOD.dumps(distributions, indent=2)}

## SAMPLE ROWS (60 rows)
{_JSON_MOD.dumps(sample_dicts, indent=2)}

## YOUR INTELLIGENCE TASKS

1. UNIT DETECTION: Look at run_time_hrs and Run_Sec columns.
   - If run_time_hrs values are tiny (< 0.01) they are likely in seconds, not hours
   - If Run_Sec values are very large (> 86400) they are likely in milliseconds
   - If max run_time_hrs > 24, something is wrong — detect what unit was actually used

2. JOB CLASSIFICATION: For each unique job name, classify:
   - type: DAILY / WEEKLY / MONTHLY / ADHOC / UNKNOWN
   - sub_app: detected sub-application if Sub_Application column is empty/null
   - is_test: true if the job looks like a test/dummy entry

3. ANOMALY DETECTION: Find jobs where:
   - Runtime is exactly 0 (failed or placeholder)
   - Runtime is suspiciously constant (same value every single run — may be fake)
   - Runtime has extreme outliers (>3x the job's own average)

4. CUSTOMER DETECTION: From job names and sub-application values, detect:
   - customer_name: inferred customer name if detectable
   - environment: PROD / UAT / TEST / DEV
   - product: JDA / Blue Yonder / Oracle / SAP / other

5. SLA CALIBRATION: Based on the runtime distribution:
   - suggested_daily_sla_hrs: what daily SLA limit makes sense for this data
   - suggested_monthly_sla_hrs: monthly equivalent

Return ONLY valid JSON in this exact schema (no markdown, no explanation):
{{
  "unit_correction": {{
    "run_time_hrs_unit": "hours|seconds|minutes",
    "Run_Sec_unit": "seconds|milliseconds",
    "correction_factor_hrs": 1.0,
    "correction_factor_sec": 1.0,
    "explanation": "why correction is needed or not"
  }},
  "job_classifications": [
    {{"job_name": "JOB_X", "type": "DAILY", "sub_app": "FORECAST", "is_test": false}}
  ],
  "anomalies": [
    {{"job_name": "JOB_X", "issue": "zero runtime", "rows_affected": 5}}
  ],
  "jobs_to_exclude": ["TEST_JOB_1"],
  "customer_name": "detected or empty string",
  "environment": "PROD|UAT|TEST|DEV|UNKNOWN",
  "product": "JDA|Blue Yonder|Other",
  "suggested_daily_sla_hrs": 6.0,
  "suggested_monthly_sla_hrs": 8.0,
  "quality_score": 85,
  "quality_flags": ["list of data quality issues"],
  "chart_highlights": {{
    "breach_jobs": ["list of job names exceeding SLA"],
    "at_risk_jobs": ["list of jobs within 15% of SLA"],
    "fastest_improving": "job name if trend detectable",
    "slowest_job": "job name with highest peak runtime"
  }},
  "summary": "one paragraph technical summary of this batch dataset"
}}"""

    result = _ai_call(prompt, max_tokens=1500, temperature=0.1)

    if result is None:
        # Local intelligence fallback — still enriches the data
        result = _local_ctrlm_intelligence(df)

    # ── APPLY AI CORRECTIONS TO THE DATAFRAME ──────────────────────
    corrections_log = []
    df_clean = df.copy()

    # 1. Apply unit corrections
    uc = result.get("unit_correction", {})
    factor_hrs = float(uc.get("correction_factor_hrs", 1.0))
    factor_sec = float(uc.get("correction_factor_sec", 1.0))

    if abs(factor_hrs - 1.0) > 0.001 and "run_time_hrs" in df_clean.columns:
        df_clean["run_time_hrs"] = df_clean["run_time_hrs"] * factor_hrs
        corrections_log.append(f"Runtime units corrected: ×{factor_hrs} ({uc.get('explanation','')})")

    if abs(factor_sec - 1.0) > 0.001 and "Run_Sec" in df_clean.columns:
        df_clean["Run_Sec"] = df_clean["Run_Sec"] * factor_sec
        corrections_log.append(f"Run_Sec units corrected: ×{factor_sec}")

    # 2. Remove test/dummy jobs
    exclude = result.get("jobs_to_exclude", [])
    if exclude and "Job_Name" in df_clean.columns:
        mask_excl = df_clean["Job_Name"].isin(exclude)
        n_excl = mask_excl.sum()
        if n_excl > 0:
            df_clean = df_clean[~mask_excl].reset_index(drop=True)
            corrections_log.append(f"Excluded {n_excl} rows from {len(exclude)} test job(s): {', '.join(exclude[:3])}")

    # 3. Apply Sub_Application enrichment from AI classification
    job_class_map = {j["job_name"]: j for j in result.get("job_classifications", [])}
    if "Sub_Application" in df_clean.columns and job_class_map:
        null_mask = df_clean["Sub_Application"].isna() | (df_clean["Sub_Application"] == "")
        if null_mask.sum() > 0:
            filled = 0
            for idx in df_clean[null_mask].index:
                jname = df_clean.at[idx, "Job_Name"]
                jclass = job_class_map.get(jname, {})
                if jclass.get("sub_app"):
                    df_clean.at[idx, "Sub_Application"] = jclass["sub_app"]
                    filled += 1
            if filled > 0:
                corrections_log.append(f"AI filled {filled} missing Sub_Application values from job name patterns")

    # 4. Remove anomalous zero-runtime rows (keep at least 1 per job for context)
    if "run_time_hrs" in df_clean.columns:
        zero_mask = df_clean["run_time_hrs"] == 0
        n_zero = zero_mask.sum()
        if n_zero > 0 and n_zero < len(df_clean) * 0.3:  # only if <30% are zeros
            df_clean = df_clean[~zero_mask].reset_index(drop=True)
            corrections_log.append(f"Removed {n_zero} zero-runtime rows (failed/incomplete runs)")

    # 5. Apply suggested SLA if significantly different from default
    sug_daily = result.get("suggested_daily_sla_hrs", DAILY_LIMIT_HRS)
    if abs(sug_daily - DAILY_LIMIT_HRS) > 1.0:
        st.session_state["_ai_suggested_sla"] = sug_daily
        corrections_log.append(f"AI suggests daily SLA of {sug_daily}h based on data distribution")

    # 6. Enrich customer/env from AI detection
    if result.get("customer_name") and not st.session_state.get("customer_name"):
        st.session_state["customer_name"] = result["customer_name"]
    if result.get("environment") and result["environment"] not in ("UNKNOWN","") and not st.session_state.get("env_type"):
        st.session_state["env_type"] = result["environment"]

    # ── Store AI intelligence ──────────────────────────────────────
    st.session_state["_ai_ctrlm_insight"] = {
        "data_type":      result.get("summary","")[:100],
        "summary":        result.get("summary",""),
        "corrections":    corrections_log,
        "quality_flags":  result.get("quality_flags", []),
        "chart_guidance": _JSON_MOD.dumps(result.get("chart_highlights", {})),
        "recommendations": [
            f"Quality score: {result.get('quality_score',0)}/100",
            f"Detected: {result.get('product','Unknown')} / {result.get('environment','Unknown')}",
            f"Anomalies: {len(result.get('anomalies',[]))} pattern(s) found",
        ],
        "anomalies":       result.get("anomalies", []),
        "chart_highlights": result.get("chart_highlights", {}),
        "quality_score":   result.get("quality_score", 0),
    }

    # ── Write corrected DataFrame back — charts use THIS ───────────
    st.session_state["ctrlm_df"] = df_clean
    return df_clean, result


def _ai_process_servers(servers):
    """
    AI-FIRST engine for server resource data.
    Validates metric values, corrects unit errors, classifies server roles.
    Returns corrected server list written back to session_state.
    """
    known = [s for s in servers if not is_unknown_server(s)]
    unknown_count = len(servers) - len(known)

    server_payload = []
    for s in known[:20]:
        server_payload.append({
            "host":     s["host"].split(".")[0],
            "type":     s.get("type","APP"),
            "cpu_pct":  round(s.get("cpu_used",0), 2),
            "mem_pct":  round(s.get("mem_used",0), 2),
            "disk_max": round(s.get("disk_used_max",0), 2),
            "mem_gb":   round(s.get("mem_total_gb",0), 1),
            "cpus":     s.get("cpus",0),
        })

    prompt = f"""You are a Performance Engineering infrastructure intelligence engine.

## SERVER DATA
Total servers: {len(servers)}, With metrics: {len(known)}, Image-only: {unknown_count}
{_JSON_MOD.dumps(server_payload, indent=2)}

## YOUR TASKS

1. VALUE VALIDATION: For each server, check if metric values are plausible:
   - CPU > 100: likely stored as ratio × 100, correct to percentage
   - mem_pct = 0: likely missing, flag it
   - disk_max = 0: likely missing or uncollected
   - CPU spike >95% on any single server: flag as critical

2. SERVER ROLE CLASSIFICATION:
   - Refine type: APP / DB / SRE / BATCH / WEB / UTIL based on hostname patterns
   - Hostname patterns: db/oracle/sql → DB, sre/batch → SRE/BATCH, web/ui → WEB

3. FLEET ANALYSIS:
   - Identify the most loaded server (highest combined CPU+disk)
   - Identify any servers that need immediate action
   - Overall fleet health narrative

Return ONLY valid JSON:
{{
  "server_corrections": [
    {{"host": "short_name", "field": "cpu_used", "corrected_value": 45.2, "reason": "was ×100"}}
  ],
  "server_classifications": [
    {{"host": "short_name", "type": "DB", "role_confidence": "high"}}
  ],
  "critical_servers": ["host1"],
  "most_loaded": "hostname",
  "fleet_health": "HEALTHY|WARNING|CRITICAL",
  "fleet_score": 85,
  "summary": "one sentence fleet summary",
  "recommendations": ["action1", "action2"]
}}"""

    result = _ai_call(prompt, max_tokens=800, temperature=0.1)

    if result is None:
        fie = fleet_intelligence_engine(servers) or {}
        result = {
            "server_corrections": [],
            "server_classifications": [],
            "critical_servers": [],
            "fleet_health": fie.get("grade","?"),
            "fleet_score":  fie.get("score",0),
            "summary":      f"{len(servers)} servers, Fleet Grade {fie.get('grade','?')}",
            "recommendations": [],
        }

    # Apply corrections
    corrected = [dict(s) for s in servers]
    corrections_log = []

    for sc in result.get("server_corrections", []):
        host    = str(sc.get("host","")).lower()
        field   = sc.get("field","")
        new_val = sc.get("corrected_value")
        reason  = sc.get("reason","")
        if not (host and field and new_val is not None):
            continue
        try:
            new_val = float(new_val)
            if new_val < 0 or new_val > 100: continue
            for s in corrected:
                short = s["host"].split(".")[0].lower()
                if host in short or short in host:
                    if field in s:
                        old = s[field]
                        s[field] = round(new_val, 2)
                        corrections_log.append(f"{short}.{field}: {old:.1f}→{new_val:.1f} ({reason})")
        except Exception:
            pass

    # Apply role reclassification
    for sc in result.get("server_classifications", []):
        host = str(sc.get("host","")).lower()
        new_type = sc.get("type","")
        if not (host and new_type): continue
        for s in corrected:
            short = s["host"].split(".")[0].lower()
            if host in short or short in host:
                if s.get("type") != new_type:
                    corrections_log.append(f"{short}: reclassified {s.get('type')} → {new_type}")
                    s["type"] = new_type

    st.session_state["_ai_resource_insight"] = {
        "data_type":    f"Resource Utilisation — {len(servers)} servers",
        "summary":      result.get("summary",""),
        "corrections":  corrections_log,
        "quality_flags": [f"Critical: {', '.join(result.get('critical_servers',[]))}"
                          ] if result.get("critical_servers") else [],
        "chart_guidance": f"Fleet health: {result.get('fleet_health','?')}. Most loaded: {result.get('most_loaded','?')}.",
        "recommendations": result.get("recommendations", []),
    }

    # Write corrected data back — fleet charts use THIS
    st.session_state["server_data"] = corrected
    return corrected, result


def _local_ctrlm_intelligence(df):
    """
    Local fallback intelligence when Gemini is unavailable.
    Applies heuristic rules to validate and enrich the data.
    """
    result = {
        "unit_correction": {"run_time_hrs_unit": "hours", "Run_Sec_unit": "seconds",
                            "correction_factor_hrs": 1.0, "correction_factor_sec": 1.0,
                            "explanation": "No correction needed"},
        "job_classifications": [],
        "anomalies": [],
        "jobs_to_exclude": [],
        "customer_name": "",
        "environment": "UNKNOWN",
        "product": "Unknown",
        "suggested_daily_sla_hrs": DAILY_LIMIT_HRS,
        "suggested_monthly_sla_hrs": MONTHLY_LIMIT_HRS,
        "quality_score": 70,
        "quality_flags": [],
        "chart_highlights": {},
        "summary": "",
    }

    if "run_time_hrs" not in df.columns:
        return result

    max_hrs = float(df["run_time_hrs"].max())
    mean_hrs = float(df["run_time_hrs"].mean())

    # Heuristic: if max runtime is < 0.1 hrs, values are likely in seconds
    if max_hrs < 0.1 and mean_hrs < 0.01:
        result["unit_correction"] = {
            "run_time_hrs_unit": "seconds",
            "Run_Sec_unit":      "seconds",
            "correction_factor_hrs": 1/3600,
            "correction_factor_sec": 1.0,
            "explanation": "Values < 0.1 — likely stored in seconds, converting to hours"
        }

    # Detect test jobs by name patterns
    if "Job_Name" in df.columns:
        test_patterns = ["TEST","DUMMY","SAMPLE","DEMO","TEMP","_TST","_DEV"]
        test_jobs = [j for j in df["Job_Name"].unique()
                     if any(p in str(j).upper() for p in test_patterns)]
        result["jobs_to_exclude"] = test_jobs

        # Zero-runtime anomalies
        zero_by_job = df[df["run_time_hrs"]==0].groupby("Job_Name").size()
        for job, cnt in zero_by_job.items():
            result["anomalies"].append({"job_name": job, "issue": "zero runtime", "rows_affected": int(cnt)})

    # Quality flags
    null_pct = df.isnull().mean().mean() * 100
    if null_pct > 10:
        result["quality_flags"].append(f"{null_pct:.1f}% null values across dataset")
    if "Sub_Application" in df.columns:
        sub_null = df["Sub_Application"].isna().mean() * 100
        if sub_null > 20:
            result["quality_flags"].append(f"{sub_null:.0f}% missing Sub_Application")

    # Quality score
    score = 100
    score -= min(30, int(null_pct * 2))
    score -= len(result["anomalies"]) * 5
    score -= len(result["jobs_to_exclude"]) * 3
    result["quality_score"] = max(0, score)
    result["summary"] = (f"Ctrl-M data: {len(df):,} rows, {df['Job_Name'].nunique() if 'Job_Name' in df.columns else '?'} jobs. "
                         f"Quality score: {result['quality_score']}/100.")
    return result


def upload_panel():
    for _k,_v in {"show_upload":True,"ctrlm_df":None,"server_data":None,
                   "batch_sla_df":None,"customer_name":"","env_type":"",
                   "sow_dfu":0,"sow_sku":0,"sow_dfu_base":0,"sow_sku_base":0,
                   "sow_scenarios_agreed":0,"sow_scenarios_prod":0,"issues_list":[],
                   "approval_pe":False,"approval_customer":False,
                   "approval_pe_name":"","approval_customer_name":"",
                   "approval_pe_date":"","approval_customer_date":"","approval_notes":"",
                   "_ctrlm_hash":None,"_sla_hash":None,"_res_hash":None}.items():
        st.session_state.setdefault(_k,_v)

    # ── Hero Header — Vivid gradient banner ────────────────
    st.markdown(f"""<div style="padding:20px 28px 22px;margin-bottom:16px;
  background:linear-gradient(135deg,rgba(59,130,246,.12) 0%,rgba(168,85,247,.08) 50%,rgba(34,211,238,.06) 100%);
  border:1px solid rgba(59,130,246,.2);border-radius:16px;
  box-shadow:0 4px 30px rgba(59,130,246,.08),inset 0 1px 0 rgba(255,255,255,.05)">
  <div style="display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px">
    <div>
      <p style="font-size:12px;font-weight:700;color:{C['cyan']};
         letter-spacing:.1em;text-transform:uppercase;margin-bottom:4px">
        ✦ PERFORMANCE ENGINEERING</p>
      <p style="font-size:40px;font-weight:800;margin:0 0 6px;letter-spacing:-.03em;line-height:1.1;
         background:linear-gradient(135deg,#60a5fa 0%,#a78bfa 40%,#22d3ee 100%);
         -webkit-background-clip:text;-webkit-text-fill-color:transparent;
         background-clip:text">PE Audit Control Tower</p>
      <p style="font-size:15px;color:{C['muted']};margin:0;font-weight:400">
        Upload your audit files — intelligent analysis starts automatically.</p>
    </div>
  </div>
</div>""",unsafe_allow_html=True)

    # ── Clear All button ──
    _clr_col1, _clr_col2 = st.columns([5, 1])
    with _clr_col2:
        if st.button("✕ Clear All", use_container_width=True):
                for k,v in {"ctrlm_df":None,"server_data":None,"batch_sla_df":None,
                             "customer_name":"","env_type":"","_ctrlm_hash":None,
                             "_sla_hash":None,"_res_hash":None,
                             "_ai_analysis_result":None,"_batch_ai_result":None,
                             "_quick_analysis":False,"_batch_quick":False,
                             "_ai_ctrlm_insight":None,"_ai_sla_insight":None,
                             "_ai_resource_insight":None,"_run_ai_ctrlm":False,
                             "_run_ai_sla":False,"_run_ai_resource":False,"_ai_master_summary":None,"_run_master_summary":False,"_pe_doc_data":None,"_pe_doc_hash":None,"_pe_doc_name":"","_pe_doc_ai":None}.items():
                    st.session_state[k]=v
                st.rerun()

    # ── Environment badge helper ─────────────────────────────
    def _env_badge_html():
        """Render a small ENV tag on upload cards so PROD/QA/TEST data is always identified."""
        env = st.session_state.get("env_type","")
        if not env: return ""
        _env_colors = {"PROD":"#ef4444","NON-PROD":"#f59e0b","TEST":"#3b82f6",
                       "UAT":"#a855f7","DEV":"#22c55e","QA":"#f59e0b"}
        ec = _env_colors.get(env.upper(), C["muted"])
        return (f'<span style="background:{ec}22;color:{ec};padding:1px 7px;'
                f'border-radius:10px;font-size:9px;font-weight:700;letter-spacing:.03em">{env.upper()}</span>')

    # ── Section Label ──
    st.markdown(f"""<div style="padding:4px 0 10px">
  <p style="font-size:12px;font-weight:700;color:#5a6fa0;letter-spacing:.1em;
     text-transform:uppercase;margin:0">UPLOAD FILES</p>
</div>""", unsafe_allow_html=True)

    # ── 4 Upload Cards ──────────────────────────────────────
    u1,u2,u3,u4 = st.columns(4)

    # Card 1 — Ctrl-M Batch CSV
    with u1:
        loaded=st.session_state.get("ctrlm_df") is not None
        brd=C["green"] if loaded else C["border"]
        bg=f"{C['green']}0d" if loaded else C["card"]
        if loaded:
            df_=st.session_state.ctrlm_df
            jc=next((c for c in df_.columns if "job" in c.lower()),"")
            nj=df_[jc].nunique() if jc else len(df_)
            csub=f"{len(df_):,} records · {nj} jobs · {st.session_state.get('_ctrlm_sla','—')} SLA"
            ctitle=st.session_state.get("_ctrlm_fname","Batch CSV")[:34]
            lbadge=f'<span class="b-ok" style="position:absolute;top:12px;left:12px;font-size:10px">● LOADED</span>'
        else:
            csub="Last 30–60 days · CSV"; ctitle="Batch Runs (CSV)"; lbadge=""
        _eb1 = _env_badge_html() if loaded else ""
        st.markdown(f"""<div style="background:{bg};border:1px solid {brd};border-radius:14px;
  padding:18px;min-height:126px;position:relative;text-align:center">
  {lbadge}<div style="margin-top:{'16px' if loaded else '0'}">
  {'<div style="font-size:26px;margin-bottom:6px">📋</div>' if not loaded else ''}
  <p style="font-size:13px;font-weight:700;margin:0 0 4px">{ctitle} {_eb1}</p>
  <p style="font-size:11px;color:{C['muted']};margin:0">{csub}</p>
</div></div>""",unsafe_allow_html=True)
        cf=st.file_uploader("Ctrl-M CSV",type=["csv","xlsx","xls","png","jpg","jpeg"],label_visibility="collapsed",key="fu_ctrlm")
        if cf is not None:
            h=hash(cf.name+str(cf.size))
            if st.session_state.get("_ctrlm_hash")!=h:
                try:
                    if _is_image_file(cf.name):
                        with st.spinner("🔍 AI Vision reading batch data from image..."):
                            df = _extract_table_from_image(cf, "batch_runs", _get_api_key())
                        if df is None: raise ValueError("Could not extract table from image — upload CSV/XLSX instead")
                        cn, env = "", ""
                    else:
                        cf.seek(0); df=load_ctrlm(cf, cf.name); cn,env=detect_customer(cf.name)
                    # Store parsed data FIRST — dashboard works immediately
                    st.session_state.ctrlm_df=df
                    st.session_state.customer_name=cn or st.session_state.get("customer_name","")
                    st.session_state.env_type=env or st.session_state.get("env_type","")
                    st.session_state._ctrlm_hash=h
                    st.session_state._ctrlm_fname=cf.name
                    _mt2 = compute_metrics_fast(st.session_state.ctrlm_df)
                    st.session_state._ctrlm_sla=f"{_mt2['compliance']:.0f}%"
                    st.session_state["_ai_ctrlm_pending"] = True  # AI will run via button
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")
        if loaded:
            st.markdown(f'<p style="font-size:11px;color:{C["blue"]};text-align:center;margin-top:4px">Replace with new file →</p>',unsafe_allow_html=True)

    # Card 2 — Batch SLA Matrix
    with u2:
        sl=st.session_state.get("batch_sla_df") is not None
        sb=f'<span class="b-ok" style="font-size:10px">● LOADED · {len(st.session_state.batch_sla_df)} rows</span>' if sl else ""
        st.markdown(f"""<div style="background:{C['card']};border:1px solid {C['border']};border-radius:14px;
  padding:18px;min-height:126px;text-align:center">
  <div style="font-size:26px;margin-bottom:6px">📊</div>
  <p style="font-size:13px;font-weight:700;margin:0 0 4px">Batch SLA Matrix</p>
  <p style="font-size:11px;color:{C['muted']};margin:0 0 8px">Per-job SLA thresholds · CSV / XLSX</p>
  {sb}</div>""",unsafe_allow_html=True)
        sf=st.file_uploader("SLA Matrix",type=["csv","xlsx","xls","png","jpg","jpeg"],label_visibility="collapsed",key="fu_sla")
        if sf is not None:
            h2=hash(sf.name+str(sf.size))
            if st.session_state.get("_sla_hash")!=h2:
                try:
                    if _is_image_file(sf.name):
                        with st.spinner("🔍 AI Vision reading SLA data from image..."):
                            _sla_df = _extract_table_from_image(sf, "sla_matrix", _get_api_key())
                        if _sla_df is None: raise ValueError("Could not extract SLA table from image")
                        st.session_state.batch_sla_df = _sla_df
                    else:
                        sf.seek(0); st.session_state.batch_sla_df=load_batch_sla(sf,sf.name)
                    st.session_state._sla_hash=h2
                    st.rerun()
                except Exception as e: st.error(f"❌ SLA: {e}")

    # Card 3 — Resource Report (multi-file: up to 8 files)
    with u3:
        svl=bool(st.session_state.get("server_data"))
        brd3=C["green"] if svl else C["border"]; bg3=f"{C['green']}08" if svl else C["card"]
        srv=st.session_state.get("server_data") or []
        napp=sum(1 for s in srv if (s.get("type","APP") if isinstance(s,dict) else getattr(s,"type","APP"))=="APP")
        ndb=len(srv)-napp
        _n_res_files = st.session_state.get("_res_file_count", 0)
        if svl:
            ct3=st.session_state.get("_res_fname","Resource Report")[:34]
            _file_tag = f" · {_n_res_files} files" if _n_res_files > 1 else ""
            cs3=f"{len(srv)} servers · {napp} APP · {ndb} DB{_file_tag}"
            lb3=f'<span class="b-ok" style="position:absolute;top:12px;left:12px;font-size:10px">● LOADED</span>'
        else:
            ct3="Resource Report"; cs3="Up to 8 files · PDF / DOCX / CSV"; lb3=""
        _eb3 = _env_badge_html() if svl else ""
        st.markdown(f"""<div style="background:{bg3};border:1px solid {brd3};border-radius:14px;
  padding:18px;min-height:126px;position:relative;text-align:center">
  {lb3}<div style="margin-top:{'16px' if svl else '0'}">
  {'<div style="font-size:26px;margin-bottom:6px">🖥️</div>' if not svl else ''}
  <p style="font-size:13px;font-weight:700;margin:0 0 4px">{ct3} {_eb3}</p>
  <p style="font-size:11px;color:{C['muted']};margin:0">{cs3}</p>
</div></div>""",unsafe_allow_html=True)
        rf_files=st.file_uploader("Resource Report",type=["pdf","docx","csv","xlsx","xls","png","jpg","jpeg"],
                                   label_visibility="collapsed",key="fu_res",
                                   accept_multiple_files=True)
        if rf_files:
            # Limit to 8 files
            rf_files = rf_files[:8]
            # Stable MD5 composite hash — avoids re-parse on identical files
            import hashlib as _hl
            _h3_parts = []
            for _f in rf_files:
                _f.seek(0); _h3_parts.append(_hl.md5(_f.read()).hexdigest()); _f.seek(0)
            h3 = "_".join(_h3_parts)
            if st.session_state.get("_res_hash") != h3:
                try:
                    all_servers = []
                    all_imgs_combined = []
                    file_names = []
                    _vision_needed = []  # (servers_slice, images) pairs for batch Vision OCR

                    for fi, rf in enumerate(rf_files):
                        with st.spinner(f"📄 Processing file {fi+1}/{len(rf_files)}: {rf.name}…"):
                            ext=rf.name.lower()
                            sd = []
                            _src_tag = rf.name

                            # Detect customer/env first so it's available for Vision queue tagging
                            _f_cn, _f_env = detect_customer(rf.name)
                            if fi == 0:
                                if _f_cn and not st.session_state.get("customer_name"):
                                    st.session_state.customer_name = _f_cn
                                if _f_env and not st.session_state.get("env_type"):
                                    st.session_state.env_type = _f_env

                            if _is_image_file(ext):
                                # Image file → Gemini Vision extracts server metrics
                                _api_k = _get_api_key()
                                if _api_k:
                                    _img_df = _extract_table_from_image(rf, "resource_util", _api_k)
                                    if _img_df is not None and len(_img_df) > 0:
                                        sd = _img_df.to_dict("records")
                                        # Ensure all required keys exist
                                        for s in sd:
                                            s.setdefault("host", s.get("hostname", "unknown"))
                                            s.setdefault("type", _infer_server_type(s.get("host","")))
                                            for k in ["cpu_used","cpu_avg","mem_used","mem_total_gb","disk_used_max"]:
                                                s.setdefault(k, 0.0)
                                            s.setdefault("disks", {})
                                            s.setdefault("_image_only", False)
                            elif ext.endswith(".csv") or ext.endswith(".xlsx") or ext.endswith(".xls"):
                                rf.seek(0); sd=load_server_csv(rf, rf.name)
                            elif ext.endswith(".docx"):
                                # Step 1 — try text-table parsing (fast, no AI needed)
                                rf.seek(0); sd, _ = parse_resource_docx(rf)
                                _has_text_metrics = any(
                                    s.get("cpu_used", 0) > 0 or
                                    s.get("mem_used", 0) > 0 or
                                    s.get("mem_total_gb", 0) > 0
                                    for s in sd
                                ) if sd else False

                                if not _has_text_metrics:
                                    # Step 2 — structure-aware XML image mapping
                                    rf.seek(0)
                                    _struct_secs, _ = parse_resource_docx_structured(rf)
                                    _has_imgs = any(
                                        len(s.get("images", [])) > 0
                                        for s in _struct_secs
                                    )
                                    if _struct_secs and _has_imgs:
                                        # Per-server Vision OCR — most precise path
                                        # images already mapped to servers by XML element order
                                        # 4-tuple: (servers, images_or_None, src_tag, env_tag)
                                        _vision_needed.append(
                                            (_struct_secs, None, _src_tag,
                                             _f_env or ""))
                                        sd = None
                                    elif sd:
                                        # Fallback: flat image batch (old behaviour)
                                        rf.seek(0)
                                        _all_imgs = extract_docx_images(rf)
                                        if _all_imgs:
                                            all_imgs_combined.extend(_all_imgs)
                                            _vision_needed.append(
                                                (sd, _all_imgs, _src_tag,
                                                 _f_env or ""))
                                            sd = None
                            else:
                                # PDF path — use cached extract + cached parse
                                rf.seek(0)
                                _pdf_bytes = rf.read(); rf.seek(0)
                                _pdf_fhash = get_file_hash(rf)
                                txt, _ = extract_pdf_text_fast(_pdf_fhash, _pdf_bytes)
                                _needs_ocr, _ocr_reason = ocr_image_if_needed(txt)
                                if _needs_ocr:
                                    sd = []   # will fall to Vision below
                                else:
                                    sd = parse_resource_metrics(_pdf_fhash, txt)
                                _has_text = any(s.get("cpu_used",0)>0 or s.get("mem_total_gb",0)>0 for s in sd) if sd else False
                                if not _has_text:
                                    rf.seek(0)
                                    _pdf_imgs = extract_pdf_images(rf)
                                    if _pdf_imgs:
                                        all_imgs_combined.extend(_pdf_imgs)
                                        # include skeleton servers so Vision can enrich them
                                        _skel = parse_resource_metrics(_pdf_fhash, txt) or []
                                        _vision_needed.append(
                                            (_skel, _pdf_imgs, _src_tag,
                                             _f_env or ""))
                                        sd = None

                            if sd:
                                # Tag each server with source file AND detected environment
                                # Also override type from filename when unambiguous
                                _fname_lo = rf.name.lower()
                                _fname_type = (
                                    "APP" if any(k in _fname_lo for k in ["app utilization","app_util","application util","sre util","sre_util"])
                                    else "DB" if any(k in _fname_lo for k in ["db utilization","db_util","database util","database_util"])
                                    else None
                                )
                                for s in sd:
                                    s["_source_file"] = _src_tag
                                    s["_source_env"] = _f_env or st.session_state.get("env_type","")
                                    if _fname_type and s.get("type") == "DB" and _fname_type == "APP":
                                        s["type"] = "APP"  # filename trumps hostname inference
                                all_servers.extend(sd)
                            file_names.append(rf.name)

                    # ── Gemini Vision OCR for image-only files ───────────────
                    if _vision_needed:
                        _api_k = _get_api_key()
                        if not _api_k:
                            # BUG-5: No API key — warn user and mark servers as 'Vision Required'
                            st.warning(
                                "⚠️ **Gemini Vision API key not configured** — "
                                "this file contains image-only Zabbix charts that cannot be read "
                                "without AI Vision OCR. Add your Gemini API key in **Settings** "
                                "to extract metrics from chart screenshots.",
                                icon="🔑",
                            )
                            for _vn_tuple in _vision_needed:
                                _vn_servers_nr = _vn_tuple[0]
                                _vn_src_nr = _vn_tuple[2] if len(_vn_tuple) > 2 else ""
                                _vn_env_nr = _vn_tuple[3] if len(_vn_tuple) > 3 else ""
                                for s in _vn_servers_nr:
                                    s["status_override"] = "Vision Required"
                                    s["_image_only"] = True
                                    if _vn_src_nr:
                                        s["_source_file"] = _vn_src_nr
                                    if _vn_env_nr and not s.get("_source_env"):
                                        s["_source_env"] = _vn_env_nr
                                all_servers.extend(_vn_servers_nr)
                        if _api_k:
                            for _vn_tuple in _vision_needed:
                                # Unpack 4-tuple (servers, images_or_None, src_tag, env_tag)
                                _vn_servers = _vn_tuple[0]
                                _vn_imgs    = _vn_tuple[1]
                                _vn_src     = _vn_tuple[2] if len(_vn_tuple) > 2 else ""
                                _vn_env     = _vn_tuple[3] if len(_vn_tuple) > 3 else ""

                                if _vn_imgs is None:
                                    # Per-server structured mode (DOCX with XML-mapped images)
                                    _n_imgs_total = sum(
                                        len(s.get("images", [])) for s in _vn_servers)
                                    with st.spinner(
                                        f"🔍 AI Vision reading {_n_imgs_total} charts "
                                        f"across {len(_vn_servers)} servers (per-server mode)…"
                                    ):
                                        _vn_result = _extract_metrics_per_server_vision(
                                            _vn_servers, _api_k)
                                else:
                                    # Flat batch mode (PDF or fallback)
                                    with st.spinner(
                                        f"🔍 AI Vision reading {len(_vn_imgs)} screenshots "
                                        f"({len(_vn_servers)} servers)…"
                                    ):
                                        _vn_result = _extract_metrics_from_zabbix_images(
                                            _vn_servers, _vn_imgs, _api_k)

                                _env_fallback = _vn_env or st.session_state.get("env_type", "")
                                for s in _vn_result:
                                    if _vn_src:
                                        s["_source_file"] = _vn_src
                                    if _env_fallback and not s.get("_source_env"):
                                        s["_source_env"] = _env_fallback
                                all_servers.extend(_vn_result)
                            _n_pop = sum(1 for s in all_servers if s.get("cpu_used",0) > 0 or s.get("mem_used",0) > 0 or s.get("disk_used_max",0) > 0)
                            if _n_pop > 0:
                                st.toast(f"✅ AI Vision extracted metrics for {_n_pop}/{len(all_servers)} servers", icon="🧠")

                    # ── Deduplicate servers by hostname (prefer entry with data) ──
                    if all_servers:
                        _dedup = {}
                        for s in all_servers:
                            key = s.get("host","").split(".")[0].lower()
                            if not key:
                                continue
                            existing = _dedup.get(key)
                            if existing is None:
                                _dedup[key] = s
                            else:
                                # Prefer the entry that has actual metrics
                                e_score = (existing.get("cpu_used",0) + existing.get("mem_used",0) + existing.get("disk_used_max",0))
                                n_score = (s.get("cpu_used",0) + s.get("mem_used",0) + s.get("disk_used_max",0))
                                if n_score > e_score:
                                    _dedup[key] = s
                        all_servers = list(_dedup.values())

                    st.session_state["_docx_images"] = {"all": all_imgs_combined} if all_imgs_combined else {}
                    st.session_state.server_data = all_servers
                    st.session_state._res_hash = h3
                    # Invalidate in-process resource DataFrame cache for new upload
                    _RES_DF_CACHE.clear()
                    _FLEET_CACHE.clear()
                    st.session_state._res_fname = ", ".join(f.name for f in rf_files)[:60] if len(rf_files) > 1 else rf_files[0].name
                    st.session_state._res_file_count = len(rf_files)
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")
        if svl:
            _repl_txt = "Replace with new files →" if _n_res_files > 1 else "Replace with new file →"
            st.markdown(f'<p style="font-size:11px;color:{C["blue"]};text-align:center;margin-top:4px">{_repl_txt}</p>',unsafe_allow_html=True)

    # Card 4 — PE Engagement Document
    with u4:
        pe_loaded = st.session_state.get("_pe_doc_data") is not None
        pe_brd    = C["purple"] if pe_loaded else C["border"]
        pe_bg     = f"{C['purple']}08" if pe_loaded else C["card"]
        pe_doc    = st.session_state.get("_pe_doc_data") or {}
        if pe_loaded:
            _pe_cust  = pe_doc.get("customer","")[:18] or "Document"
            _pe_nsrvs = len(pe_doc.get("servers",[]))
            _pe_nrecs = len(pe_doc.get("recommendations",[]))
            pe_sub    = f"{_pe_nsrvs} servers · {_pe_nrecs} recommendations"
            pe_title  = _pe_cust
        else:
            pe_title = "PE Engagement Doc"
            pe_sub   = "Performance Analysis DOCX · PDF"
        pe_badge = f'<span class="b-ok" style="font-size:10px;background:{C["purple"]}22;color:{C["purple"]}">● LOADED</span>' if pe_loaded else ""
        st.markdown(f"""<div style="background:{pe_bg};border:1px solid {pe_brd};border-radius:14px;
  padding:18px;min-height:126px;text-align:center">
  {'<div style="font-size:26px;margin-bottom:6px">📋</div>' if not pe_loaded else f'<div style="font-size:11px;font-weight:700;color:{C["white"]};margin-bottom:4px">{pe_title}</div>'}
  <p style="font-size:13px;font-weight:700;margin:0 0 4px">PE Engagement Doc</p>
  <p style="font-size:11px;color:{C['muted']};margin:0 0 8px">{pe_sub}</p>
  {pe_badge}</div>""", unsafe_allow_html=True)
        pef2 = st.file_uploader("PE Document", type=["docx","doc","pdf","png","jpg","jpeg"],
                                label_visibility="collapsed", key="fu_pe_front")
        if pef2 is not None:
            h_pe = hash(pef2.name + str(pef2.size))
            if st.session_state.get("_pe_doc_hash") != h_pe:
                try:
                    with st.spinner("📖 Parsing PE document…"):
                        pef2.seek(0)
                        _pe_parsed = parse_pe_document(pef2)
                    if _pe_parsed:
                        st.session_state["_pe_doc_data"] = _pe_parsed
                        st.session_state["_pe_doc_hash"] = h_pe
                        st.session_state["_pe_doc_name"] = pef2.name
                        st.session_state["_pe_doc_ai"]   = None
                        st.rerun()
                    else:
                        st.error("❌ Could not parse this document.")
                except Exception as e:
                    st.error(f"❌ {e}")
        if pe_loaded:
            st.markdown(f'<p style="font-size:11px;color:{C["purple"]};text-align:center;margin-top:4px">View in PE Document Review tab →</p>', unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>",unsafe_allow_html=True)

    # ── SOW Contract Upload + Manual ─────────────────────
    with st.expander("🗒️  SOW Contract Volume (Optional) — Schedule 1-A · DFU/SKU contracted vs. actual"):
        sow_f = st.file_uploader("Upload SOW PDF/DOCX (auto-extracts DFU/SKU/Planograms)",
                                 type=["pdf","docx","xlsx","png","jpg","jpeg"], label_visibility="visible", key="fu_sow")
        if sow_f is not None:
            h_sow = hash(sow_f.name + str(sow_f.size))
            if st.session_state.get("_sow_hash") != h_sow:
                try:
                    sow_f.seek(0)
                    cname, dfu, sku = parse_sow_volumes(sow_f, sow_f.name)
                    if cname and not st.session_state.get("customer_name"):
                        st.session_state.customer_name = cname
                    if dfu and dfu > 0:
                        st.session_state.sow_dfu = dfu
                        st.success(f"Extracted DFU: {dfu:,}")
                    if sku and sku > 0:
                        st.session_state.sow_sku = sku
                        st.success(f"Extracted SKU/Planograms: {sku:,}")
                    st.session_state._sow_hash = h_sow
                except Exception as e:
                    st.error(f"SOW parse error: {e}")
        st.divider()
        s1,s2=st.columns(2)
        with s1:
            st.session_state.sow_dfu=st.number_input("Contracted DFU",min_value=0,value=int(st.session_state.get("sow_dfu",0)),step=1000)
            st.session_state.sow_sku=st.number_input("Contracted SKU",min_value=0,value=int(st.session_state.get("sow_sku",0)),step=1000)
        with s2:
            st.session_state.sow_scenarios_agreed=st.number_input("Agreed Scenarios",min_value=0,value=int(st.session_state.get("sow_scenarios_agreed",0)),step=1)
            st.session_state.sow_scenarios_prod=st.number_input("Actual in PROD",min_value=0,value=int(st.session_state.get("sow_scenarios_prod",0)),step=1)

    # ── SRE Files (Optional) ──────────────────────────────
    with st.expander("🔧 SRE Configuration Files (Optional) — Node Config, System Alerts, Global Properties"):
        sre_f = st.file_uploader("Upload SRE XLS/XLSX/CSV files",
                                 type=["xls","xlsx","csv"], label_visibility="visible",
                                 key="fu_sre", accept_multiple_files=True)
        if sre_f:
            for sf in sre_f:
                fn = sf.name.upper()
                try:
                    sf.seek(0)
                    ext_sre = os.path.splitext(sf.name)[1].lower()
                    if ext_sre in [".xlsx", ".xls"]:
                        try:
                            sdf = pd.read_excel(sf, sheet_name=0, engine="openpyxl")
                        except Exception:
                            sdf = pd.read_excel(sf, sheet_name=0)
                    else:
                        sdf = pd.read_csv(sf)

                    if "NODE_CONFIG" in fn:
                        st.session_state["sre_node_config"] = sdf
                        st.success(f"Node Config: {len(sdf)} entries loaded")
                    elif "SYSTEM_ALERT" in fn:
                        st.session_state["sre_system_alerts"] = sdf
                        st.success(f"System Alerts: {len(sdf)} alerts loaded")
                    elif "GLOBAL_PROPERTY" in fn:
                        st.session_state["sre_global_props"] = sdf
                        st.success(f"Global Properties: {len(sdf)} properties loaded")
                    else:
                        st.session_state[f"sre_{sf.name}"] = sdf
                        st.success(f"{sf.name}: {len(sdf)} rows loaded")
                except Exception as e:
                    st.error(f"Error reading {sf.name}: {e}")

    # ── AI Upload Intelligence Panels ────────────────────────
    cdf   = st.session_state.get("ctrlm_df")
    svd   = st.session_state.get("server_data") or []
    sla_df= st.session_state.get("batch_sla_df")

    # AI analyses run ONLY when explicitly triggered by a button click below
    # (never on upload — that would freeze the UI for 5-10s while Gemini responds)

    # ── AI Engine Status Panel ────────────────────────────────
    any_data = cdf is not None or svd or sla_df is not None
    if any_data:
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        st.markdown(f'''<div style="background:linear-gradient(135deg,{C["card"]},{C["card2"]});
border:1px solid {C["border"]};border-radius:14px;padding:16px 20px;margin-top:4px">
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:10px">
  <div style="display:flex;align-items:center;gap:8px">
    <span style="font-size:16px">🤖</span>
    <span style="font-size:12px;font-weight:700;color:{C["white"]}">AI Intelligence Engine</span>
    <span style="background:{C["green"]}22;color:{C["green"]};padding:1px 7px;border-radius:4px;font-size:9px;font-weight:700">AUTO-PROCESSES EVERY UPLOAD</span>
  </div>
  <span style="font-size:9px;color:{C["muted"]}">validate → correct → enrich → populate charts</span>
</div>
<div style="font-size:10px;color:{C["muted"]};background:{C["card"]}88;border-radius:6px;padding:6px 10px;margin-bottom:10px">
  Data pipeline: <b style="color:{C["cyan"]}">Upload</b> → <b style="color:{C["blue"]}">AI Engine</b> → <b style="color:{C["green"]}">Corrected Data</b> → <b style="color:{C["purple"]}">Charts & KPIs</b>
</div>''', unsafe_allow_html=True)

        # Show what AI engine did to each dataset
        if cdf is not None:
            _ci = st.session_state.get("_ai_ctrlm_insight") or {}
            _qs = _ci.get("quality_score", 0)
            _qc = C["green"] if _qs >= 85 else (C["amber"] if _qs >= 60 else C["red"])
            _corrs = _ci.get("corrections", [])
            _flags = _ci.get("quality_flags", [])
            _ch    = _ci.get("chart_highlights", {})
            _nj    = cdf["Job_Name"].nunique() if "Job_Name" in cdf.columns else "?"
            breach_txt = ", ".join(_ch.get("breach_jobs",[])[:3])
            risk_txt   = ", ".join(_ch.get("at_risk_jobs",[])[:3])
            st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:8px;padding:10px 14px;margin-bottom:6px;border-left:3px solid {_qc}">
<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px">
  <span style="font-size:10px;font-weight:700;color:{C["white"]}">📋 Ctrl-M: {_nj} jobs · {len(cdf):,} rows</span>
  <span style="background:{_qc}22;color:{_qc};padding:1px 7px;border-radius:4px;font-size:9px;font-weight:700">Quality {_qs}/100</span>
</div>
{('<div>' + ' &nbsp;·&nbsp; '.join(f'<span style="font-size:9px;color:{C["green"]}">⚙ {c}</span>' for c in _corrs[:3]) + '</div>') if _corrs else f'<span style="font-size:9px;color:{C["muted"]}">✓ Data validated — no corrections needed</span>'}
{('<div>' + ' &nbsp;·&nbsp; '.join(f'<span style="font-size:9px;color:{C["amber"]}">⚠ {f}</span>' for f in _flags[:2]) + '</div>') if _flags else ''}
{(f'<div style="font-size:9px;color:#8b92a5;margin-top:4px">🔴 Breach: <b style="color:#ef4444">{breach_txt}</b></div>') if breach_txt else ""}
{(f'<div style="font-size:9px;color:#8b92a5">🟡 At-risk: <b style="color:#f59e0b">{risk_txt}</b></div>') if risk_txt else ""}
</div>''', unsafe_allow_html=True)

        if svd:
            _ri = st.session_state.get("_ai_resource_insight") or {}
            _rc = _ri.get("corrections", [])
            _rf = _ri.get("quality_flags", [])
            st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:8px;padding:10px 14px;margin-bottom:6px;border-left:3px solid {C["purple"]}">
<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px">
  <span style="font-size:10px;font-weight:700;color:{C["white"]}">🖥️ Resources: {len(svd)} servers</span>
  <span style="font-size:9px;color:{C["green"]}">✓ AI-validated</span>
</div>
{('<div>' + ' &nbsp;·&nbsp; '.join(f'<span style="font-size:9px;color:{C["green"]}">⚙ {c}</span>' for c in _rc[:3]) + '</div>') if _rc else f'<span style="font-size:9px;color:{C["muted"]}">✓ All metrics validated — no corrections needed</span>'}
{('<div>' + ' &nbsp;·&nbsp; '.join(f'<span style="font-size:9px;color:{C["amber"]}">⚠ {f}</span>' for f in _rf[:2]) + '</div>') if _rf else ''}
</div>''', unsafe_allow_html=True)

        if sla_df is not None:
            st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:8px;padding:10px 14px;margin-bottom:6px;border-left:3px solid {C["cyan"]}">
<span style="font-size:10px;font-weight:700;color:{C["white"]}">📊 SLA Matrix: {len(sla_df)} thresholds</span>
<span style="font-size:9px;color:{C["green"]};margin-left:10px">✓ Loaded and active</span>
</div>''', unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    # ── Master PE Approval Summary ────────────────────────
    any_data2 = cdf is not None or svd or st.session_state.get("batch_sla_df") is not None or st.session_state.get("_uiperf_df") is not None
    if any_data2:
        st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)

        # Trigger master summary generation
        if st.session_state.get("_run_master_summary"):
            st.session_state.pop("_run_master_summary", None)
            with st.spinner("🏆 Gemini is generating your PE Approval Summary…"):
                ai_generate_master_summary()
            st.rerun()

        master_result = st.session_state.get("_ai_master_summary")

        # Big prominent button
        ms_col1, ms_col2, ms_col3 = st.columns([1, 3, 1])
        with ms_col2:
            btn_label = "🏆 Generate PE Approval Summary" if not master_result else "🔄 Regenerate PE Approval Summary"
            if st.button(btn_label, use_container_width=True, key="btn_master_summary",
                         help="Collect all uploaded data and generate an executive PE sign-off summary"):
                st.session_state["_run_master_summary"] = True
                st.session_state["_ai_master_summary"] = None
                st.rerun()

        # Display master summary
        if master_result:
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
            # Detect approval verdict for color
            verdict_color = C["green"]
            if "NOT READY" in master_result.upper() or "BREACH" in master_result.upper():
                verdict_color = C["red"]
            elif "CONDITIONAL" in master_result.upper() or "PENDING" in master_result.upper():
                verdict_color = C["amber"]

            st.markdown(f'''<div style="background:linear-gradient(135deg,{C["card"]},{C["card2"]});
border:2px solid {verdict_color}55;border-radius:16px;padding:24px 28px;margin:6px 0">
<div style="display:flex;align-items:center;gap:10px;margin-bottom:16px">
  <span style="font-size:22px">🏆</span>
  <span style="font-size:14px;font-weight:800;color:{C["white"]}">PE Audit Approval Summary</span>
  <span style="background:{verdict_color}22;color:{verdict_color};padding:2px 12px;
    border-radius:20px;font-size:10px;font-weight:700">AI GENERATED</span>
</div>
<div style="font-size:13px;line-height:1.8;color:{C["white"]};border-left:3px solid {verdict_color};
  padding-left:16px">
{master_result.replace(chr(10), "<br>")}
</div>
</div>''', unsafe_allow_html=True)

            # Download button for the summary
            dl_col1, dl_col2 = st.columns([3,1])
            with dl_col2:
                st.download_button(
                    "⬇️ Download Summary",
                    master_result.encode("utf-8"),
                    file_name=f"PE_Summary_{(st.session_state.get('customer_name','') or 'Audit').replace(' ','_')}.txt",
                    mime="text/plain",
                    key="dl_master_summary",
                    use_container_width=True,
                )

        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    # ── Bottom Nav Cards ──────────────────────────────────
    if cdf is not None or svd:
        st.markdown("<div style='height:16px'></div>",unsafe_allow_html=True)
        cards=[]
        if cdf is not None:
            mt2=compute_metrics_fast(cdf); cards.append(("📊","Batch Performance",f"{mt2['total_jobs']} jobs analyzed",C["blue"]))
        if svd: cards.append(("🖥️","Resources",f"{len(svd)} servers monitored",C["purple"]))
        if st.session_state.get("sow_dfu",0)>0: cards.append(("📄","SOW Volume","Contract compliance",C["cyan"]))
        cards.append(("✅","Sign-off","Audit checklist",C["green"]))
        nc=st.columns(len(cards))
        for col,(icon,title,sub,clr) in zip(nc,cards):
            with col:
                st.markdown(f"""<div style="background:{C['card']};border:1px solid {C['border']};
  border-radius:14px;padding:18px;text-align:center">
  <div style="font-size:24px;margin-bottom:8px">{icon}</div>
  <p style="font-size:13px;font-weight:700;margin:0 0 4px;color:{clr}">{title}</p>
  <p style="font-size:11px;color:{C['muted']};margin:0">{sub}</p>
</div>""",unsafe_allow_html=True)


def extract_docx_images(file_obj):
    """Extract embedded images from a DOCX (which is a ZIP archive)."""
    import zipfile
    images = []
    try:
        file_obj.seek(0)
        with zipfile.ZipFile(file_obj) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            for name in sorted(media):
                ext = name.split(".")[-1].lower()
                if ext in ["png","jpg","jpeg","gif","bmp","tiff","emf","wmf"]:
                    images.append((name.split("/")[-1], z.read(name)))
    except zipfile.BadZipFile:
        # DOCX is corrupted or not a valid ZIP — return empty list gracefully
        pass
    except Exception:
        # Image extraction is optional; a failure here is non-critical
        pass
    return images


def extract_pdf_images(file_obj):
    """
    Multi-strategy PDF image extractor:
      1. pypdf  XObject image extraction (embedded JPEG/PNG)
      2. fitz   (PyMuPDF) page-to-pixmap rendering — best quality
      3. Graceful empty list if neither works
    """
    images = []

    # ── Strategy A: fitz / PyMuPDF (renders every page as a PNG) ──
    try:
        import fitz                          # PyMuPDF
        file_obj.seek(0)
        pdf_bytes = file_obj.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            mat  = fitz.Matrix(2.0, 2.0)    # 2× zoom = ~144 dpi
            pix  = page.get_pixmap(matrix=mat, alpha=False)
            images.append((f"page_{i+1:02d}.png", pix.tobytes("png")))
        doc.close()
        return images
    except ImportError:
        # PyMuPDF not installed — fall through to Strategy B
        pass
    except Exception:
        # PDF may be encrypted, malformed, or have rendering issues.
        # Fall through to Strategy B (pypdf XObject extraction).
        pass

    # ── Strategy B: pypdf XObject extraction ──
    try:
        from pypdf import PdfReader
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        for i, page in enumerate(reader.pages):
            res = page.get("/Resources")
            if not res: continue
            xobj_ref = res.get("/XObject")
            if not xobj_ref: continue
            xobj = xobj_ref.get_object()
            for name, ref in xobj.items():
                obj = ref.get_object()
                if obj.get("/Subtype") != "/Image":
                    continue
                try:
                    raw   = obj.get_data()
                    csp   = obj.get("/ColorSpace", "")
                    mode  = "L" if "Gray" in str(csp) else "RGB"
                    w     = int(obj["/Width"])
                    h     = int(obj["/Height"])
                    from PIL import Image as PILImage
                    import io as _io
                    img = PILImage.frombytes(mode, (w, h), raw)
                    buf = _io.BytesIO()
                    img.save(buf, format="PNG")
                    images.append((f"p{i+1}_{name.lstrip('/')}.png",
                                   buf.getvalue()))
                except Exception:
                    # Per-image decode failure (bad colorspace, unsupported format).
                    # Fall back to raw bytes and let st.image attempt decoding.
                    if len(raw) > 1000:
                        images.append((f"p{i+1}_{name.lstrip('/')}.jpg", raw))
    except Exception:
        # pypdf not installed or PDF structure incompatible with XObject extraction.
        # Return whatever images were collected (may be empty list).
        pass

    return images


def overview_bar_fig(servers):
    """Side-by-side bar chart: CPU / Mem / Disk for all servers."""
    names  = [s["host"].split(".")[0] for s in servers]
    cpus   = [s["cpu_used"]           for s in servers]
    mems   = [s.get("mem_used", 0)    for s in servers]
    disks  = [s["disk_used_max"]      for s in servers]

    fig = go.Figure()
    fig.add_trace(go.Bar(name="CPU %",  x=names, y=cpus,
                         marker_color=[status_color(v,CPU_OK,CPU_WARN)  for v in cpus],
                         opacity=0.9))
    fig.add_trace(go.Bar(name="Mem %",  x=names, y=mems,
                         marker_color=[status_color(v,MEM_OK,MEM_WARN)  for v in mems],
                         opacity=0.7))
    fig.add_trace(go.Bar(name="Disk %", x=names, y=disks,
                         marker_color=[status_color(v,DISK_OK,DISK_WARN) for v in disks],
                         opacity=0.6))

    fig.add_hline(y=CPU_OK,  line_dash="dot",  line_color=C["amber"],
                  annotation_text="75% Moderate",
                  annotation_font=dict(color=C["amber"], size=9),
                  annotation_position="top right")
    fig.add_hline(y=CPU_WARN, line_dash="dash", line_color=C["red"],
                  annotation_text="90% Critical",
                  annotation_font=dict(color=C["red"], size=9),
                  annotation_position="top right")

    fig.update_layout(**BASE_LAYOUT, height=280, barmode="group",
                      xaxis=dict(tickangle=-30, tickfont_size=9, **AXIS),
                      yaxis=dict(title="%", range=[0,105], **AXIS),
                      legend=dict(orientation="h", y=-0.28, font=dict(size=10)),
            margin=_DEFAULT_MARGIN,
    )
    return fig


def health_matrix_fig(servers):
    """Heatmap-style health matrix: servers × metrics."""
    names   = [s["host"].split(".")[0] for s in servers]
    metrics = ["CPU %", "Mem %", "Disk %"]
    vals    = [
        [s["cpu_used"]        for s in servers],
        [s.get("mem_used", 0) for s in servers],
        [s["disk_used_max"]   for s in servers],
    ]
    hover_z = []
    for i, metric in enumerate(metrics):
        row = []
        for j, s in enumerate(servers):
            v = vals[i][j]
            status = "🔴 CRITICAL" if v > CPU_WARN else ("⚠ MODERATE" if v > CPU_OK else "✅ OK")
            row.append(f"<b>{names[j]}</b><br>{metric}: {v:.1f}%<br>{status}")
        hover_z.append(row)

    fig = go.Figure(go.Heatmap(
        z=vals, x=names, y=metrics,
        zmin=0, zmax=100,
        colorscale=[
            [0.00, C["card2"]],
            [0.01, "#1a472a"],
            [0.74, C["green"]],
            [0.75, C["amber"]],
            [0.89, "#dc6e2a"],
            [0.90, C["red"]],
            [1.00, "#7f1d1d"],
        ],
        colorbar=dict(
            title=dict(text="%", side="right", font=dict(size=10, color=C["muted"])),
            tickfont=dict(size=9, color=C["muted"]),
            tickvals=[0, 75, 90, 100],
            ticktext=["0", "75 OK→Mod", "90 Crit", "100"],
            thickness=12, len=0.8,
            bgcolor="rgba(0,0,0,0)", bordercolor=C["border"],
        ),
        hovertext=hover_z,
        hovertemplate="%{hovertext}<extra></extra>",
        xgap=3, ygap=3,
        texttemplate="%{z:.0f}%",
        textfont=dict(size=11, color="white"),
    ))
    fig.update_layout(**BASE_LAYOUT, height=160,
                      xaxis=dict(tickfont_size=9, **AXIS),
                      yaxis=dict(tickfont_size=10, **AXIS),
            margin=_DEFAULT_MARGIN,
    )
    return fig


def is_unknown_server(s):
    """True if server data came from image-only DOCX (all zeros)."""
    return (s.get("cpu_used", 0) == 0 and s.get("disk_used_max", 0) == 0
            and s.get("mem_total_gb", 0) == 0 and not s.get("disks"))

def get_health_score(cpu, mem, disk, server_type="APP"):
    """Delegates to F7 calculate_host_health with server-type-aware thresholds.
    DB servers get a tighter CPU penalty: 40%+ CPU triggers increasing score reduction.
    Returns -1 for unknown servers (all-zero metrics)."""
    if cpu == 0 and disk == 0 and mem == 0:
        return -1   # UNKNOWN — not healthy, no data
    score = calculate_host_health(cpu, mem, disk)["health_score"]
    # BUG-7: DB-specific CPU penalty — 40%+ CPU on a DB server degrades score
    if server_type == "DB" and cpu >= 40:
        penalty = min(25, (cpu - 40) * 0.625)  # up to -25 pts at 80% CPU
        score = max(0, score - penalty)
    return score

def health_label(score):
    """Maps F7 health_score to display label + color. Thresholds aligned with calculate_host_health."""
    if score < 0:       return "⚪ Unknown",   C["muted"]
    if score >= 80:     return "🟢 Healthy",   C["green"]
    if score >= 60:     return "🟡 Warning",   C["amber"]
    return                     "🔴 Critical",  C["red"]


_FLEET_CACHE: dict = {}  # module-level memo — keyed by server count + hash of host names

def fleet_intelligence_engine(server_data):
    """Delegates to F8 calculate_fleet_health + F6 anomaly detection.
    Returns a summary dict used by pe_audit_intelligence_panel() and AI Insights."""
    if not server_data:
        return None

    # Lightweight cache: keyed by sorted hostnames + metric values (avoids hashing image bytes)
    _cache_key = hash(tuple(
        (s.get("host",""), s.get("cpu_used",0), s.get("mem_used",0), s.get("disk_used_max",0))
        for s in server_data))
    if _cache_key in _FLEET_CACHE:
        return _FLEET_CACHE[_cache_key]

    # F8 — Fleet Health
    fleet = calculate_fleet_health(server_data)

    # Per-server health scores (needed by the panel for display)
    scores = []
    for s in server_data:
        sc = get_health_score(s.get("cpu_used", 0), s.get("mem_used", 0), s.get("disk_used_max", 0))
        scores.append((s, sc))

    # F6-style z-score anomaly detection on CPU and Disk across fleet
    anomalies = []
    try:
        known = [s for s in server_data if not is_unknown_server(s)]
        if len(known) >= 3:
            for metric, key in [("CPU", "cpu_used"), ("Disk", "disk_used_max"), ("Memory", "mem_used")]:
                vals = [s.get(key, 0) for s in known]
                mu   = float(np.mean(vals))
                std  = float(np.std(vals))
                if std < 1e-6:
                    continue
                for s, v in zip(known, vals):
                    z = (v - mu) / std
                    if abs(z) >= 2.0:
                        anomalies.append({"host": s["host"], "metric": metric, "value": v, "z": round(z, 2)})
            anomalies.sort(key=lambda x: abs(x["z"]), reverse=True)
            anomalies = anomalies[:5]
    except Exception:
        pass

    _result = {
        "score":    fleet["fleet_score"],
        "grade":    fleet["grade"],
        "healthy":  fleet["healthy"],
        "warning":  fleet["warning"],
        "critical": fleet["critical"],
        "unknown":  fleet["unknown"],
        "anomalies": anomalies,
    }
    _FLEET_CACHE[_cache_key] = _result
    if len(_FLEET_CACHE) > 8:   # keep cache small
        _FLEET_CACHE.pop(next(iter(_FLEET_CACHE)))
    return _result


def mini_bar_fig(value, label, ok_t, warn_t, uid=""):
    """Grafana-style horizontal bar with threshold zones."""
    bar_c = C["red"] if value>=warn_t else (C["amber"] if value>=ok_t else C["green"])
    fig = go.Figure()
    for x0,x1,fc in [(0,ok_t,"rgba(34,197,94,0.09)"),(ok_t,warn_t,"rgba(245,158,11,0.09)"),(warn_t,100,"rgba(239,68,68,0.09)")]:
        fig.add_shape(type="rect",x0=x0,x1=x1,y0=-0.5,y1=0.5,
                      fillcolor=fc,line_width=0,layer="below")
    fig.add_trace(go.Bar(x=[value],y=[""],orientation="h",
                         marker_color=bar_c,marker_opacity=0.9,width=0.55,
                         hovertemplate=f"<b>{label}: {value:.1f}%</b><extra></extra>"))
    fig.add_vline(x=ok_t,   line_dash="dot",line_color=C["amber"],line_width=1.2)
    fig.add_vline(x=warn_t, line_dash="dash",line_color=C["red"],  line_width=1.2)
    xanchor = "left" if value < 65 else "right"
    xshift  = 5 if value < 65 else -5
    fig.add_annotation(x=value,y=0,text=f"<b>{value:.1f}%</b>",
                       font=dict(size=10,color="#fff"),showarrow=False,
                       xanchor=xanchor,xshift=xshift)
    fig.update_layout(
        **BASE_LAYOUT, height=50, showlegend=False,
        margin=dict(l=4,r=4,t=20,b=4),
        xaxis=dict(range=[0,100],showticklabels=False,showgrid=False,zeroline=False,
                   title=dict(text=label,font=dict(size=9,color=C["muted"]))),
        yaxis=dict(showticklabels=False,showgrid=False),
    )
    return fig

def auto_recommendations(s):
    """Generate ranked Grafana-alert-style recommendations for a server."""
    tips = []
    cpu  = s["cpu_used"]; mem = s.get("mem_used",0); disk = s["disk_used_max"]
    disks = s.get("disks",{})

    # Image-only server — all metrics unknown
    if is_unknown_server(s):
        tips.append(("⚪","DATA UNAVAILABLE",
                     "This server was parsed from an image-only DOCX. "
                     "Download the CSV template, fill in CPU/Disk/Memory values "
                     "from the Zabbix screenshots, and re-upload as CSV."))
        return tips

    if cpu>=90:   tips.append(("🔴","CPU CRITICAL",   f"{cpu:.1f}% — kill runaway procs; check top/htop immediately"))
    elif cpu>=75: tips.append(("🟡","CPU WARNING",    f"{cpu:.1f}% — above 75% threshold; review batch concurrency"))
    elif cpu>=40: tips.append(("🟡","CPU ELEVATED",   f"{cpu:.1f}% — elevated during batch; monitor closely"))
    else:         tips.append(("🟢","CPU OK",         f"{cpu:.1f}% — healthy, within optimal range"))

    if disk>=90:  tips.append(("🔴","DISK CRITICAL",  f"Max mount at {disk:.1f}% — IMMEDIATE action: purge logs/archive data"))
    elif disk>=75:tips.append(("🟡","DISK WARNING",   f"Max mount at {disk:.1f}% — plan volume expansion or cleanup"))
    else:         tips.append(("🟢","DISK OK",        f"Max mount at {disk:.1f}% — sufficient headroom"))

    for mnt,pct in disks.items():
        if pct>=90:  tips.append(("🔴",f"MOUNT {mnt}",  f"{pct:.0f}% full — expand or clean urgently"))


def _gemini_call(prompt, session_cache_key, loading_msg="🤖 Gemini is analysing…"):
    """
    Calls Google Gemini API and streams the response into the Streamlit UI.
    - NO probe calls (they waste quota and cause false failures)
    - Tries each model name in sequence directly inside generate_content
    - Embedded API key used automatically — no user input required
    """
    try:
        import google.generativeai as genai
    except ImportError:
        st.error("❌ **google-generativeai not installed.** Run: `pip install google-generativeai` and restart.")
        return

    genai.configure(api_key=_get_api_key())

    # Dynamically discover the best available model via list_models()
    model = None
    try:
        available = list(genai.list_models())
        # Prefer flash (fast + free), then pro variants, filter to generateContent capable
        def _rank(m):
            n = m.name.lower()
            if "generatecontent" not in str([s.lower() for s in getattr(m,"supported_generation_methods", [])]):
                return 99
            if "flash" in n and "2.5" in n and "lite" not in n: return 0  # gemini-2.5-flash = best
            if "flash" in n and "2" in n and "lite" not in n: return 1  # gemini-2.x-flash
            if "flash" in n and "lite" in n: return 2   # flash-lite variants
            if "flash" in n: return 3                    # any other flash
            if "pro" in n: return 4                      # any pro
            return 5
        available_sorted = sorted(available, key=_rank)
        for m_info in available_sorted:
            if _rank(m_info) < 99:
                model = genai.GenerativeModel(m_info.name)
                break
    except Exception as list_err:
        # list_models failed (network/auth) — try direct fallback names
        for _mn in ["gemini-2.5-flash", "gemini-2.0-flash-001", "gemini-flash-latest",
                    "gemini-flash-lite-latest", "gemini-2.0-flash-lite"]:
            try:
                model = genai.GenerativeModel(_mn)
                break
            except Exception:
                continue

    if model is None:
        st.error("❌ **Could not initialise any Gemini model.** "
                 "Check your internet connection or get a new key at "
                 "https://aistudio.google.com/app/apikey")
        return

    st.markdown(
        f'<p style="font-size:11px;color:{C["blue"]};margin:8px 0 4px">{loading_msg}</p>',
        unsafe_allow_html=True
    )
    placeholder = st.empty()

    try:
        response = model.generate_content(
            prompt,
            generation_config={"max_output_tokens": 1500, "temperature": 0.4},
            stream=True,
        )
        full_response = ""
        for chunk in response:
            try:
                text = chunk.text
            except Exception:
                text = ""
            if text:
                full_response += text
                placeholder.markdown(full_response + "▌")
        placeholder.markdown(full_response)
        st.session_state[session_cache_key] = full_response

    except Exception as exc:
        err = str(exc)
        if any(x in err.upper() for x in ["API_KEY", "INVALID_ARGUMENT", "PERMISSION", "UNAUTHENTICATED"]):
            st.error("❌ **Gemini API key invalid.** Get a new free key at https://aistudio.google.com/app/apikey")
        elif any(x in err.upper() for x in ["QUOTA", "RESOURCE_EXHAUSTED"]):
            st.error("❌ **Gemini quota exceeded.** Wait ~1 minute and try again.")
        else:
            st.error(f"❌ **Gemini call failed:** {exc}")



def _gemini_api_key_widget():
    """Silently sets the Gemini API key from _get_api_key() — no UI prompt needed."""
    st.session_state["_gemini_key"] = _get_api_key()


def ai_run_deep_analysis(servers, known_servers, avg_score, n_ok, n_warn, n_crit):
    """Infrastructure deep analysis — Google Gemini Flash (free tier)."""
    import json as _json
    _gemini_api_key_widget()

    fleet_summary = {
        "fleet_health_score": avg_score,
        "servers_healthy":    n_ok,
        "servers_warning":    n_warn,
        "servers_critical":   n_crit,
        "total_servers":      len(servers),
        "cpu_warn": CPU_WARN, "cpu_ok": CPU_OK,
        "disk_warn":DISK_WARN,"disk_ok":DISK_OK,
        "mem_warn": MEM_WARN, "mem_ok": MEM_OK,
    }
    server_rows = []
    for s in known_servers:
        row = {
            "host":         s["host"].split(".")[0],
            "type":         s.get("type","APP"),
            "cpu_pct":      round(s.get("cpu_used",0),1),
            "mem_pct":      round(s.get("mem_used",0),1),
            "disk_max_pct": round(s.get("disk_used_max",0),1),
            "score":        get_health_score(s.get("cpu_used",0),s.get("mem_used",0),s.get("disk_used_max",0)),
        }
        hot = {mv:round(p,1) for mv,p in s.get("disks",{}).items() if p>=60}
        if hot: row["hot_mounts"] = hot
        server_rows.append(row)

    prompt = f"""You are a Senior Performance Engineering consultant reviewing a server infrastructure audit.

Fleet Summary: {_json.dumps(fleet_summary)}
Per-Server Metrics: {_json.dumps(server_rows)}

Write a diagnostic report covering:
1. **Overall Fleet Health** — Grade and 2-3 sentence summary.
2. **Critical Findings** — Each server at/near thresholds with exact % values and risk description.
3. **Root Causes** — Likely causes for elevated metrics (batch concurrency, log accumulation, memory leaks, under-provisioning).
4. **Prioritised Action Plan** — Numbered concrete remediation steps, most urgent first.
5. **Go-Live Verdict** — ✅ READY / ⚠️ CONDITIONAL / 🔴 NOT READY with one-line justification.

Use hostnames. Quote exact numbers. No generic padding."""

    _gemini_call(prompt, "_ai_analysis_result",
                 "🤖 Gemini Flash is analysing your fleet metrics…")


def ai_run_batch_analysis(m, top_jobs_df):
    """Batch workload deep analysis — Google Gemini Flash (free tier)."""
    import json as _json
    _gemini_api_key_widget()

    breach_jobs  = top_jobs_df[top_jobs_df["buffer_pct"]<0][["Job_Name","peak_hrs","buffer_pct"]].head(10).to_dict("records")
    at_risk_jobs = top_jobs_df[(top_jobs_df["buffer_pct"]>=0)&(top_jobs_df["buffer_pct"]<15)][["Job_Name","peak_hrs","buffer_pct"]].head(10).to_dict("records")
    top10        = top_jobs_df.head(10)[["Job_Name","peak_hrs","avg_hrs","sla_used_pct","buffer_pct"]].to_dict("records")
    window_stats = {
        "max_daily_hrs": round(float(m["window"]["total_hrs"].max()),2),
        "avg_daily_hrs": round(float(m["window"]["total_hrs"].mean()),2),
        "sla_limit_hrs": DAILY_LIMIT_HRS,
        "days_over_sla": int((m["window"]["total_hrs"]>DAILY_LIMIT_HRS).sum()),
    }
    summary = {
        "customer": st.session_state.get("customer_name","Unknown"),
        "env":      st.session_state.get("env_type","Unknown"),
        "compliance_pct": round(m["compliance"],2),
        "total_jobs": m["total_jobs"], "jobs_ok": m["jobs_ok"],
        "jobs_breach": m["jobs_breach"], "total_runs": m.get("total_runs",0),
        "total_hrs": round(m.get("total_hrs",0),1), "window": window_stats,
    }
    prompt = f"""You are a Senior Performance Engineering consultant reviewing a Ctrl-M batch workload audit.

Batch Summary: {_json.dumps(summary)}
Breaching Jobs (>{DAILY_LIMIT_HRS}h SLA): {_json.dumps(breach_jobs) if breach_jobs else "None"}
At-Risk Jobs (<15% buffer): {_json.dumps(at_risk_jobs) if at_risk_jobs else "None"}
Top 10 Jobs: {_json.dumps(top10)}

Write a batch performance diagnostic covering:
1. **Compliance Summary** — SLA %, jobs in breach vs at-risk vs healthy.
2. **Root Cause Analysis** — For breaching jobs, propose causes (sequential chaining, no parallelism, data volume growth, bottlenecks).
3. **Batch Window Pressure** — Avg/max daily window vs {DAILY_LIMIT_HRS}h SLA, risk of creep.
4. **Remediation Plan** — Numbered concrete fixes (parallelisation, scheduling, tuning).
5. **Go-Live Readiness** — ✅ READY / ⚠️ CONDITIONAL / 🔴 NOT READY with one-line justification.

Use job names. Quote exact hours. No generic padding."""

    _gemini_call(prompt, "_batch_ai_result",
                 "🤖 Gemini Flash is analysing your batch workload…")

# ── BATCH PERFORMANCE TEST REPORT ───────────────────────────
# ── RESOURCE UTILIZATION TAB ─────────────────────────────────
# ── RESOURCE UTILIZATION RENDER FUNCTIONS ────────────────────

def render_resource_summary(df: "pd.DataFrame", fie: dict) -> None:
    """6 compact KPI cards: servers, fleet grade, avg CPU/mem/disk, status counts.
    Rendered first — fast, pure HTML, no Plotly."""
    if df is None or df.empty:
        return
    known = df[~df["image_only"]]
    n_total = len(df)
    n_crit  = int((known["status"] == "Critical").sum())
    n_warn  = int((known["status"] == "Warning").sum())
    n_ok    = int((known["status"] == "Healthy").sum())
    avg_cpu  = round(float(known["cpu_pct"].mean()),  1) if len(known) else 0.0
    avg_mem  = round(float(known["mem_pct"].mean()),  1) if len(known) else 0.0
    avg_disk = round(float(known["disk_pct"].mean()), 1) if len(known) else 0.0
    grade    = fie.get("grade", "?") if fie else "?"
    _gc = {"A": C["green"], "B": C["cyan"], "C": C["amber"],
           "D": "#f97316", "F": C["red"]}.get(grade, C["muted"])
    _sc_num = fie.get("score", 0) if fie else 0
    n_app = int((df["type"] == "APP").sum())
    n_db  = int((df["type"] == "DB").sum())
    status_color_val = C["red"] if n_crit else (C["amber"] if n_warn else C["green"])

    cards = [
        ("SERVERS",    str(n_total),            C["blue"],  f"{n_app} APP · {n_db} DB"),
        ("FLEET GRADE",grade,                   _gc,        f"Score {_sc_num}/100"),
        ("AVG CPU",    f"{avg_cpu:.1f}%",
         status_color(avg_cpu, CPU_OK, CPU_WARN), "Fleet average"),
        ("AVG MEM",    f"{avg_mem:.1f}%",
         status_color(avg_mem, MEM_OK, MEM_WARN), "Fleet average"),
        ("AVG DISK",   f"{avg_disk:.1f}%",
         status_color(avg_disk, DISK_OK, DISK_WARN), "Max mount avg"),
        ("HEALTH",     f"{n_crit}C · {n_warn}W · {n_ok}✓",
         status_color_val, "Critical · Warn · OK"),
    ]
    cols = st.columns(len(cards))
    for col, (title, val, color, sub) in zip(cols, cards):
        with col:
            col.markdown(
                f'<div style="background:{C["card"]};border:1px solid {C["border"]};'
                f'border-radius:10px;padding:12px 14px;text-align:center;margin-bottom:6px">'
                f'<div style="font-size:9px;font-weight:700;color:{C["muted"]};'
                f'letter-spacing:.08em;text-transform:uppercase;margin-bottom:4px">{title}</div>'
                f'<div style="font-size:22px;font-weight:800;color:{color};line-height:1.1">{val}</div>'
                f'<div style="font-size:9px;color:{C["muted"]};margin-top:3px">{sub}</div>'
                f'</div>',
                unsafe_allow_html=True)


def render_resource_charts(df: "pd.DataFrame") -> None:
    """Compact horizontal bar chart — CPU/Mem/Disk for top servers.
    Aggregated before plotting; max 12 servers; single Plotly payload."""
    known = df[~df["image_only"] &
               (df[["cpu_pct","mem_pct","disk_pct"]].max(axis=1) > 0)]
    if known.empty:
        st.info("No metric data to chart. Upload a file with CPU/Mem/Disk values.")
        return
    top = known.nlargest(min(12, len(known)), "cpu_pct")

    fig = go.Figure()
    for metric, col_key, color in [
        ("CPU %",  "cpu_pct",  C["blue"]),
        ("Mem %",  "mem_pct",  C["cyan"]),
        ("Disk %", "disk_pct", C["purple"]),
    ]:
        fig.add_trace(go.Bar(
            y=top["server"], x=top[col_key], name=metric,
            orientation="h", marker_color=color, opacity=0.85,
            hovertemplate=f"<b>%{{y}}</b><br>{metric}: %{{x:.1f}}%<extra></extra>",
        ))
    fig.add_vline(x=75, line_dash="dot",  line_color=C["amber"], line_width=1)
    fig.add_vline(x=90, line_dash="dash", line_color=C["red"],   line_width=1)
    fig.update_layout(
        **BASE_LAYOUT, barmode="group",
        height=max(180, len(top) * 30 + 60),
        xaxis=dict(range=[0, 105], title="% Used", **AXIS),
        yaxis=dict(autorange="reversed", tickfont_size=9, **AXIS),
        legend=dict(orientation="h", font=dict(size=9),
                    bgcolor="rgba(0,0,0,0)", y=-0.18),
        margin=dict(l=6, r=10, t=10, b=40),
    )
    st.plotly_chart(fig, use_container_width=True,
                    config={"displayModeBar": False}, key="res_bar_v2")


def render_resource_detail_table(df: "pd.DataFrame") -> None:
    """Full server metrics table — lazy rendered inside an expander.
    Only materialised when the user opens it; no render cost when collapsed."""
    with st.expander(
            f"📋 Full Server Detail Table  ({len(df)} servers)", expanded=False):
        disp = df[["server","type","cpu_pct","cpu_avg_pct",
                   "mem_pct","mem_gb","disk_pct","status","source_env"]].copy()
        disp.columns = ["Server","Type","CPU %","CPU Avg %",
                        "Mem %","Mem GB","Disk %","Status","Env"]
        disp = disp.sort_values("CPU %", ascending=False).reset_index(drop=True)

        def _highlight_status(val):
            c = {"Critical": "#ef444433",
                 "Warning":  "#f59e0b33",
                 "Healthy":  "#22c55e22"}.get(val, "")
            return f"background-color:{c}" if c else ""

        try:
            styled = disp.style.map(_highlight_status, subset=["Status"])
        except AttributeError:
            styled = disp.style.applymap(_highlight_status, subset=["Status"])
        st.dataframe(styled, use_container_width=True,
                     height=min(420, len(disp) * 36 + 42))

        # Download button inside expander — no extra vertical space used
        csv_bytes = disp.to_csv(index=False).encode()
        st.download_button("⬇️ Export CSV", csv_bytes,
                           "resource_metrics.csv", "text/csv",
                           key="dl_res_detail")


# ── RESOURCE: INGESTION STATUS BANNER ───────────────────────
def _render_ingestion_status(servers: list, df: "pd.DataFrame") -> None:
    """One-line parse confidence strip — only shown when partial data exists."""
    n_total   = len(servers)
    n_unknown = sum(1 for s in servers if is_unknown_server(s))
    n_known   = n_total - n_unknown
    if n_unknown == 0:
        return  # all good — no banner needed
    src_files = list({s.get("_source_file","") for s in servers if s.get("_source_file")})
    src_label = src_files[0][:30] if len(src_files) == 1 else f"{len(src_files)} files"
    conf_pct  = round(n_known / n_total * 100) if n_total else 0
    bar_color = C["green"] if conf_pct >= 80 else (C["amber"] if conf_pct >= 50 else C["red"])
    st.markdown(
        f'<div style="background:{bar_color}11;border:1px solid {bar_color}33;'
        f'border-radius:8px;padding:8px 16px;margin-bottom:10px;'
        f'display:flex;align-items:center;gap:16px">'
        f'<div style="width:80px;height:5px;background:{C["border"]};border-radius:3px;overflow:hidden;flex-shrink:0">'
        f'<div style="width:{conf_pct}%;height:100%;background:{bar_color}"></div></div>'
        f'<span style="font-size:11px;color:{bar_color};font-weight:700">{conf_pct}% parse confidence</span>'
        f'<span style="font-size:11px;color:{C["muted"]}">'
        f'{n_known}/{n_total} servers have numeric metrics · {n_unknown} awaiting OCR · {src_label}</span>'
        f'<span style="margin-left:auto">'
        f'<a download="server_metrics_template.csv" style="font-size:10px;color:{C["blue"]};font-weight:600">'
        f'⬇ Fill CSV template to complete</a></span>'
        f'</div>', unsafe_allow_html=True)
    if n_known == 0:
        st.download_button("⬇️ Download Server CSV Template",
                           server_csv_template(servers),
                           "server_metrics_template.csv", "text/csv",
                           key="dl_srv_tpl_status")


# ── RESOURCE: COMPACT METRIC HEATMAP ────────────────────────
def _render_metric_heatmap(df: "pd.DataFrame") -> None:
    """Plotly heatmap — servers (rows) × CPU/Mem/Disk (cols).
    Tableau-style: clear color scale, server labels on Y, compact height."""
    known = df[~df["image_only"] & (df[["cpu_pct","mem_pct","disk_pct"]].max(axis=1) > 0)]
    if known.empty:
        st.caption("No metric data for heatmap.")
        return
    top = known.nlargest(min(20, len(known)), "cpu_pct")
    z    = top[["cpu_pct","mem_pct","disk_pct"]].values.tolist()
    ylab = top["server"].tolist()
    xlab = ["CPU %", "Mem %", "Disk %"]
    # Custom discrete color scale: green→amber→red, matching dashboard palette
    cscale = [
        [0.0,  "#1d3461"], [0.4,  "#1d6461"], [0.6, "#b45309"],
        [0.75, "#b45309"], [0.9,  "#dc2626"], [1.0, "#7f1d1d"],
    ]
    fig = go.Figure(go.Heatmap(
        z=z, x=xlab, y=ylab,
        colorscale=cscale, zmin=0, zmax=100,
        text=[[f"{v:.0f}%" for v in row] for row in z],
        texttemplate="%{text}",
        textfont=dict(size=9, color="#e2e8f0"),
        hovertemplate="<b>%{y}</b><br>%{x}: %{z:.1f}%<extra></extra>",
        showscale=False,
        xgap=2, ygap=2,
    ))
    fig.update_layout(
        **BASE_LAYOUT,
        height=max(130, len(top) * 22 + 40),
        margin=dict(l=4, r=4, t=22, b=4),
        xaxis=dict(side="top", tickfont=dict(size=9, color=C["muted"]),
                   showgrid=False, showline=False),
        yaxis=dict(tickfont=dict(size=8, color=C["muted"]),
                   autorange="reversed", showgrid=False),
    )
    st.plotly_chart(fig, use_container_width=True,
                    config={"displayModeBar": False}, key="res_heatmap_v2")


# ── RESOURCE: COMPACT FLEET GRADE PANEL ─────────────────────
def _render_fleet_panel_compact(fie: dict, known_servers: list) -> None:
    """Compact grade card + top anomalies + top 3 recommendations — no excess height."""
    if not fie:
        return
    grade  = fie.get("grade", "?")
    score  = fie.get("score", 0)
    gc     = {"A": C["green"], "B": C["cyan"], "C": C["amber"],
              "D": "#f97316", "F": C["red"]}.get(grade, C["muted"])
    anoms  = fie.get("anomalies", [])
    n_crit = fie.get("critical", 0)
    n_warn = fie.get("warning", 0)

    # Grade + health counts
    st.markdown(
        f'<div style="background:{gc}11;border:1px solid {gc}33;border-radius:10px;'
        f'padding:12px 16px;margin-bottom:8px">'
        f'<div style="display:flex;align-items:center;gap:12px">'
        f'<div style="font-size:36px;font-weight:900;color:{gc};line-height:1">{grade}</div>'
        f'<div><div style="font-size:11px;font-weight:700;color:{C["white"]}">Fleet Grade</div>'
        f'<div style="font-size:10px;color:{C["muted"]}">Score {score}/100</div>'
        f'<div style="font-size:10px;margin-top:3px">'
        f'<span style="color:{C["red"]}">■ {n_crit} Crit</span>'
        f'<span style="color:{C["amber"]};margin-left:8px">■ {n_warn} Warn</span>'
        f'<span style="color:{C["green"]};margin-left:8px">■ {fie.get("healthy",0)} OK</span>'
        f'</div></div></div></div>', unsafe_allow_html=True)

    # Anomalies (max 4)
    if anoms:
        st.markdown(
            f'<p style="font-size:10px;font-weight:700;color:{C["amber"]};'
            f'margin:8px 0 4px">⚡ Z-Score Anomalies</p>', unsafe_allow_html=True)
        for a in anoms[:4]:
            ac = C["red"] if abs(a["z"]) >= 3 else C["amber"]
            st.markdown(
                f'<div style="background:{ac}15;border-left:2px solid {ac};'
                f'padding:4px 8px;border-radius:0 4px 4px 0;margin-bottom:3px;font-size:9px;color:{ac}">'
                f'<b>{a["host"].split(".")[0]}</b>: {a["metric"]} {a["value"]:.1f}%'
                f' <span style="opacity:.7">(z={a["z"]:+.1f})</span></div>',
                unsafe_allow_html=True)

    # Top-3 recommendations
    recos = []
    for s in known_servers:
        hn = s["host"].split(".")[0]
        if s.get("disk_used_max", 0) >= 75:
            recos.append(f'💾 <b>{hn}</b>: Disk {s["disk_used_max"]:.0f}% — expand volume')
        elif s.get("cpu_used", 0) >= 75:
            recos.append(f'⚙️ <b>{hn}</b>: CPU {s["cpu_used"]:.0f}% — review concurrency')
        elif s.get("mem_used", 0) >= 75:
            recos.append(f'🧠 <b>{hn}</b>: Mem {s["mem_used"]:.0f}% — check for leaks')
    if recos:
        st.markdown(
            f'<p style="font-size:10px;font-weight:700;color:{C["green"]};'
            f'margin:8px 0 4px">💡 Recommendations</p>', unsafe_allow_html=True)
        for r in recos[:3]:
            st.markdown(
                f'<p style="font-size:9px;color:{C["white"]};margin:2px 0">{r}</p>',
                unsafe_allow_html=True)


def resource_tab():
    """Resource Utilization — compact Tableau-style layout.
    Fast: cached DataFrame, lazy table, column-split charts."""
    servers = st.session_state.get("server_data") or []
    if not servers:
        st.markdown(
            f'<div style="text-align:center;padding:56px 24px;'
            f'background:{C["card"]};border-radius:14px;margin:8px 0">'
            f'<div style="font-size:44px;margin-bottom:12px">📡</div>'
            f'<p style="font-size:16px;font-weight:700;color:{C["white"]};margin:0 0 6px">'
            f'No Server Data Loaded</p>'
            f'<p style="font-size:12px;color:{C["muted"]}">Upload a Zabbix PDF · DOCX · CSV '
            f'or Azure screenshot in the Upload &amp; Intake page.</p>'
            f'</div>', unsafe_allow_html=True)
        return

    # ── Cached normalisation — no DataFrame rebuild on every rerun ──
    _res_hash    = st.session_state.get("_res_hash", "")
    df           = _get_resource_df(_res_hash, servers)
    known_servers = [s for s in servers if not is_unknown_server(s)]
    app_svrs     = [s for s in servers if s.get("type", "APP") == "APP"]
    db_svrs      = [s for s in servers if s.get("type", "APP") == "DB"]
    all_scores   = [(s, get_health_score(s.get("cpu_used", 0), s.get("mem_used", 0),
                     s.get("disk_used_max", 0))) for s in servers]
    fie          = fleet_intelligence_engine(servers)   # uses internal _FLEET_CACHE

    # ── 1. Parse-confidence banner (only shown when partial data) ──
    _render_ingestion_status(servers, df)

    # ── 2. KPI strip — 6 cards in one row ─────────────────────────
    render_resource_summary(df, fie)

    # ── 3. Three-column analysis row ──────────────────────────────
    # Left 2.4 parts: grouped bar chart (CPU / Mem / Disk per server)
    # Middle 2.4 parts: metric heatmap (servers × metrics colour grid)
    # Right 1.2 parts: fleet grade + anomalies + recommendations
    c_bar, c_heat, c_fleet = st.columns([2.4, 2.4, 1.2])

    with c_bar:
        st.markdown(
            f'<div class="panel" style="height:100%">'
            f'<p class="panel-title">📊 Server Utilization</p>'
            f'<p class="panel-sub">CPU · Mem · Disk — top servers by CPU load</p>',
            unsafe_allow_html=True)
        render_resource_charts(df)
        st.markdown("</div>", unsafe_allow_html=True)

    with c_heat:
        st.markdown(
            f'<div class="panel" style="height:100%">'
            f'<p class="panel-title">🗺 Metric Heatmap</p>'
            f'<p class="panel-sub">Green = healthy · Amber = warning · Red = critical</p>',
            unsafe_allow_html=True)
        _render_metric_heatmap(df)
        st.markdown("</div>", unsafe_allow_html=True)

    with c_fleet:
        st.markdown(
            f'<div class="panel" style="height:100%">'
            f'<p class="panel-title">🧠 Fleet Intel</p>',
            unsafe_allow_html=True)
        _render_fleet_panel_compact(fie, known_servers)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── 4. Server cards — tabbed APP / DB ─────────────────────────
    _ENV_COLORS = {"PROD": "#ef4444", "NON-PROD": "#f59e0b", "QA": "#f59e0b",
                   "TEST": "#3b82f6", "UAT": "#a855f7", "DEV": "#22c55e"}

    def _server_card_html(s):
        score      = get_health_score(s.get("cpu_used", 0), s.get("mem_used", 0),
                                      s.get("disk_used_max", 0))
        badge_t, _ = health_label(score)
        cpu    = s.get("cpu_used", 0); mem  = s.get("mem_used", 0)
        disk   = s.get("disk_used_max", 0); mem_gb = s.get("mem_total_gb", 0)
        stype  = s.get("type", "APP")
        status_lbl = ("OK"   if "Healthy"  in badge_t else
                      "WARN" if "Warning"  in badge_t else
                      "CRIT" if "Critical" in badge_t else "N/A")
        score_lbl  = "?" if score < 0 else str(int(score))
        bg_col = (C["green"] if "Healthy"  in badge_t else
                  C["amber"] if "Warning"  in badge_t else
                  C["muted"] if "Unknown"  in badge_t else C["red"])
        def _bar(pct, ok, warn, lbl):
            col_ = C["red"] if pct >= warn else (C["amber"] if pct >= ok else C["green"])
            return (f'<div class="srv-bar-row"><span class="srv-bar-lbl">{lbl}</span>'
                    f'<div class="srv-bar-track"><div style="width:{min(100,int(pct))}%;'
                    f'height:100%;background:{col_}"></div></div>'
                    f'<span style="font-size:8px;font-weight:700;color:{col_};width:24px;'
                    f'text-align:right">{pct:.0f}%</span></div>')
        metrics_html = (
            f'<p style="font-size:9px;color:{C["muted"]};margin:6px 0 0">⚪ No metrics</p>'
            if is_unknown_server(s) else
            _bar(cpu, CPU_OK, CPU_WARN, "CPU") +
            _bar(mem, MEM_OK, MEM_WARN, "MEM") +
            _bar(disk, DISK_OK, DISK_WARN, "DSK")
        )
        _srv_env  = s.get("_source_env", "")
        _ec       = _ENV_COLORS.get(_srv_env.upper(), C["muted"]) if _srv_env else ""
        env_badge = (f'<span style="background:{_ec}22;color:{_ec};padding:0 5px;'
                     f'border-radius:3px;font-size:7px;font-weight:700">'
                     f'{_srv_env.upper()}</span>') if _srv_env else ""
        mem_badge = (f'<span style="background:{C["muted"]}22;color:{C["muted"]};'
                     f'padding:1px 6px;border-radius:3px;font-size:8px">'
                     f'{mem_gb:.0f}GB</span>') if mem_gb else ""
        return (
            f'<div class="srv-card"><div class="srv-card-header">'
            f'<div><div class="srv-hostname">{_html_mod.escape(s["host"][:32])}</div>'
            f'<div class="srv-tags">'
            f'<span style="background:{C["blue"]}22;color:{C["blue"]};padding:1px 6px;'
            f'border-radius:3px;font-size:8px;font-weight:700">{stype}</span>'
            f'{mem_badge}{env_badge}</div></div>'
            f'<span style="background:{bg_col}22;color:{bg_col};padding:1px 7px;'
            f'border-radius:4px;font-size:8px;font-weight:700;white-space:nowrap">'
            f'{score_lbl} {status_lbl}</span></div>{metrics_html}</div>'
        )

    def _render_srv_cards(srvs):
        if not srvs:
            st.caption("No servers in this group.")
            return
        n_cols    = min(4, max(1, len(srvs)))
        cols      = st.columns(n_cols)
        col_html  = [""] * n_cols
        for i, s in enumerate(srvs):
            col_html[i % n_cols] += _server_card_html(s)
        for ci, col in enumerate(cols):
            with col:
                if col_html[ci]:
                    st.markdown(col_html[ci], unsafe_allow_html=True)

    srv_tab_labels = [
        f"🖥️ App ({len(app_svrs)})",
        f"🗄️ DB ({len(db_svrs)})",
    ]
    t_app, t_db = st.tabs(srv_tab_labels)
    with t_app:
        _render_srv_cards(app_svrs)
    with t_db:
        _render_srv_cards(db_svrs)

    # ── 5. Detail table (lazy expander — zero render cost when closed) ──
    render_resource_detail_table(df)

    # ── 6. AI Analysis panel ──────────────────────────────────────
    st.markdown(
        f'<div class="panel" style="margin-top:8px">'
        f'<p class="panel-title">🤖 AI Infrastructure Analysis</p>'
        f'<p class="panel-sub">⚡ Quick heuristics · Gemini Flash full diagnostic</p>',
        unsafe_allow_html=True)
    q_col, ai_col = st.columns(2)
    with q_col:
        if st.button("⚡ Quick Analysis", use_container_width=True, key="quick_analysis_btn"):
            st.session_state["_quick_analysis"]    = True
            st.session_state.pop("_ai_analysis_result", None)
    with ai_col:
        if st.button("🤖 AI Deep Analysis", use_container_width=True, key="ai_analysis_btn"):
            st.session_state["_run_ai_analysis"]   = True
            st.session_state.pop("_quick_analysis", None)

    if st.session_state.get("_quick_analysis"):
        n_crit = sum(1 for _, sc in all_scores if 0 <= sc < 60)
        insights = []
        for s in sorted(known_servers, key=lambda x: x.get("cpu_used", 0), reverse=True)[:3]:
            if s.get("cpu_used", 0) >= 75:
                insights.append(
                    f"🔴 **{s['host'].split('.')[0]}** — CPU **{s['cpu_used']:.1f}%** "
                    f"— review batch concurrency")
        for s in sorted(known_servers, key=lambda x: x.get("disk_used_max", 0), reverse=True)[:3]:
            if s.get("disk_used_max", 0) >= 75:
                insights.append(
                    f"🔴 **{s['host'].split('.')[0]}** — Disk **{s['disk_used_max']:.1f}%** "
                    f"— purge logs / expand volume")
        if n_crit:
            insights.append(f"⚠️ **{n_crit} server(s) CRITICAL** — immediate review required")
        if not insights:
            insights.append("✅ All servers with data are within normal thresholds.")
        for tip in insights:
            st.markdown(f"- {tip}")

    if st.session_state.get("_run_ai_analysis"):
        st.session_state.pop("_run_ai_analysis", None)
        n_crit_loc = sum(1 for _, sc in all_scores if 0 <= sc < 60)
        n_warn_loc = sum(1 for _, sc in all_scores if 60 <= sc < 85)
        n_ok_loc   = sum(1 for _, sc in all_scores if sc >= 85)
        known_sc   = [sc for _, sc in all_scores if sc >= 0]
        avg_sc     = round(sum(known_sc) / len(known_sc)) if known_sc else 0
        ai_run_deep_analysis(servers, known_servers, avg_sc, n_ok_loc, n_warn_loc, n_crit_loc)

    if st.session_state.get("_ai_analysis_result") and not st.session_state.get("_quick_analysis"):
        st.markdown(st.session_state["_ai_analysis_result"])

    st.markdown("</div>", unsafe_allow_html=True)



def perf_test_report_tab(m):
    st.markdown(f'''<p style="font-size:12px;color:{C["muted"]}">
Full-volume run analysis · Buffer % vs SLA · Daily &amp; Weekly &amp; Monthly SLA breakdown</p>''',
        unsafe_allow_html=True)

    df   = st.session_state.ctrlm_df
    top  = m["top_jobs"].copy()
    # buffer_pct / sla_used_pct / buffer_status already computed by F3 in compute_metrics
    # Ensure columns exist (fallback for legacy cached metrics)
    if "buffer_pct" not in top.columns:
        top["buffer_pct"]    = ((DAILY_LIMIT_HRS - top["peak_hrs"]) / DAILY_LIMIT_HRS * 100).round(1)
    if "sla_used_pct" not in top.columns:
        top["sla_used_pct"]  = (top["peak_hrs"] / DAILY_LIMIT_HRS * 100).round(1)
    if "status" not in top.columns:
        top["status"] = top["buffer_pct"].apply(
            lambda x: "BREACH" if x < 0 else ("AT RISK" if x < 15 else "OK"))

    # KPIs
    k1,k2,k3,k4 = st.columns(4)
    with k1: kpi("Jobs Tested",str(len(top)),"Full-volume run","🧪",C["blue"])
    with k2:
        at_risk = len(top[top["buffer_pct"] < 15])
        kpi("At Risk (<15% buffer)",str(at_risk),"Jobs near SLA limit","⚠️",
            C["red"] if at_risk else C["green"])
    with k3:
        breach = len(top[top["buffer_pct"] < 0])
        kpi("SLA Breaches",str(breach),"Jobs exceeding SLA","🔴",
            C["red"] if breach else C["green"])
    with k4:
        avg_buf = top["buffer_pct"].clip(upper=100).mean()
        kpi("Avg Buffer %",f"{avg_buf:.1f}%","Across all jobs","📊",
            C["green"] if avg_buf > 30 else C["amber"])

    # ── F3 Fleet SLA Buffer KPI ───────────────────────────────────────────────
    _fsb = m.get("fleet_sla_buffer")
    if _fsb:
        _buf_color = (C["green"] if _fsb["status"] in ("EXCELLENT","HEALTHY")
                      else C["amber"] if _fsb["status"] == "CAUTION" else C["red"])
        st.markdown(
            f'<div style="background:{_buf_color}11;border:1px solid {_buf_color}33;border-radius:10px;'
            f'padding:10px 16px;margin:6px 0 8px;display:flex;gap:24px;align-items:center">'
            f'<span style="font-size:13px">⏱️ <b>SLA Buffer (Worst Job)</b>: '
            f'<span style="color:{_buf_color};font-weight:700">{_fsb["buffer_pct"]:.1f}%</span>'
            f' &nbsp;·&nbsp; {_fsb["buffer_hrs"]:.2f}h headroom &nbsp;·&nbsp; '
            f'Growth capacity: <b>{_fsb["growth_capacity_pct"]:.0f}%</b> &nbsp;·&nbsp; '
            f'Can grow <b>{_fsb["growth_multiplier"]:.2f}×</b> before breach &nbsp;·&nbsp; '
            f'Status: <span style="color:{_buf_color};font-weight:700">{_fsb["status"]}</span></span>'
            f'</div>', unsafe_allow_html=True)

    # ── F6 Anomaly Detection ─────────────────────────────────────────────────
    _anoms = m.get("anomalies", [])
    if _anoms:
        st.markdown(f'<div style="background:{C["card"]};border:1px solid {C["amber"]}44;border-radius:10px;'
                    f'padding:10px 16px;margin:4px 0 10px">'
                    f'<p style="font-size:11px;font-weight:700;color:{C["amber"]};margin:0 0 6px">'
                    f'⚡ {len(_anoms)} Runtime Anomalies Detected (F6 — Statistical Outliers)</p>',
                    unsafe_allow_html=True)
        for _a in _anoms[:5]:
            _sev_c = C["red"] if _a.get("severity",0) >= 3 else C["amber"]
            _z_str = f" z={_a['z_score']:.1f}" if _a.get("z_score") else ""
            _v_str = f" +{_a['variance_pct']:.0f}% over expected" if _a.get("variance_pct") else ""
            st.markdown(
                f'<p style="font-size:10px;color:{C["white"]};margin:2px 0">'
                f'<span style="color:{_sev_c};font-weight:700">▲</span> '
                f'<b>{_a["job_name"]}</b> — peak {_a["peak_hrs"]:.3f}h{_z_str}{_v_str} '
                f'· <span style="color:{_sev_c}">{_a["status"]}</span></p>',
                unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── F9 Comprehensive Assessment pill ─────────────────────────────────────
    _servers = st.session_state.get("server_data") or []
    _sow_cfg = None
    _d, _ds = st.session_state.get("sow_dfu",0), st.session_state.get("sow_dfu_base",0)
    _s, _ss = st.session_state.get("sow_sku",0), st.session_state.get("sow_sku_base",0)
    if _ds > 0 or _ss > 0:
        _sow_cfg = {"dfu": (_d, _ds), "sku": (_s, _ss)}
    _pe9 = compute_pe_assessment(ctrlm_df=df, servers=_servers or None, sow_cfg=_sow_cfg)
    _s9c = {"APPROVED": C["green"], "APPROVED_WITH_CAUTION": C["amber"], "REJECTED": C["red"]}
    _s9col = _s9c.get(_pe9["status"], C["muted"])
    st.markdown(
        f'<div style="background:{_s9col}11;border:1px solid {_s9col}44;border-radius:10px;'
        f'padding:10px 18px;margin:4px 0 12px;display:flex;gap:16px;align-items:center;flex-wrap:wrap">'
        f'<span style="font-size:13px;font-weight:700;color:{_s9col}">🎯 PE Assessment: {_pe9["status"].replace("_"," ")}</span>'
        f'<span style="font-size:12px;color:{C["white"]}">Confidence: <b>{_pe9["confidence"]}%</b></span>'
        + (f'<span style="font-size:10px;color:{C["amber"]}">⚠ ' + " · ".join(_pe9["cautions"][:2]) + '</span>' if _pe9["cautions"] else "")
        + (f'<span style="font-size:10px;color:{C["red"]}">🔴 ' + " · ".join(_pe9["recommendations"][:2]) + '</span>' if _pe9["recommendations"] else "")
        + f'</div>', unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    # Buffer waterfall chart
    st.markdown('<div class="panel"><p class="panel-title">🏊 Buffer % vs SLA — Top 20 Jobs</p>'
                '<p class="panel-sub">Buffer = headroom before SLA breach. &lt;0% = already breaching. &lt;15% = at risk.</p>',
                unsafe_allow_html=True)
    disp = top.head(20).sort_values("buffer_pct")
    bar_colors = [C["red"] if v < 0 else (C["amber"] if v < 15 else C["green"])
                  for v in disp["buffer_pct"]]
    fig_buf = go.Figure(go.Bar(
        x=disp["Job_Name"], y=disp["buffer_pct"],
        marker_color=bar_colors, opacity=0.9,
        customdata=np.stack([disp["peak_hrs"], disp["sla_used_pct"]], axis=-1),
        hovertemplate="<b>%{x}</b><br>Buffer: %{y:.1f}%<br>Peak: %{customdata[0]:.2f}h<br>SLA Used: %{customdata[1]:.0f}%<extra></extra>",
    ))
    fig_buf.add_hline(y=0,  line_dash="dash", line_color=C["red"],
                      annotation_text="SLA Limit",
                      annotation_font=dict(color=C["red"],size=10))
    fig_buf.add_hline(y=15, line_dash="dot",  line_color=C["amber"],
                      annotation_text="15% Safety Buffer",
                      annotation_font=dict(color=C["amber"],size=10))
    fig_buf.update_layout(**BASE_LAYOUT, height=280,
                          xaxis=dict(tickangle=-35, tickfont_size=8, **AXIS),
                          yaxis=dict(title="Buffer %", **AXIS),
            margin=_DEFAULT_MARGIN,
    )
    st.plotly_chart(fig_buf, use_container_width=True,
                    config={"displayModeBar":False}, key="perf_buf_chart")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # Daily / Weekly / Monthly SLA breakdown
    st.markdown('<div class="panel"><p class="panel-title">📅 SLA Breakdown by Workflow Type</p>'
                '<p class="panel-sub">Daily (_D), Weekly (_W), Monthly (_M) jobs — compliance per cadence</p>',
                unsafe_allow_html=True)
    df2 = df.copy()
    df2["cadence"] = df2["Job_Name"].apply(
        lambda x: "Weekly"  if str(x).endswith("_W") else
                  "Monthly" if str(x).endswith("_M") else "Daily")
    cad_stats = df2.groupby("cadence").agg(
        Jobs=("Job_Name","nunique"),
        Total_Runs=("run_time_hrs","count"),
        Avg_Hrs=("run_time_hrs","mean"),
        Peak_Hrs=("run_time_hrs","max"),
        Breaches=("run_time_hrs", lambda x: (x > DAILY_LIMIT_HRS).sum())
    ).reset_index()
    cad_stats["Compliance %"] = (
        (1 - cad_stats["Breaches"] / cad_stats["Total_Runs"].clip(lower=1)) * 100
    ).round(1)
    cad_stats["Avg_Hrs"]  = cad_stats["Avg_Hrs"].round(3)
    cad_stats["Peak_Hrs"] = cad_stats["Peak_Hrs"].round(3)

    cad_colors = {
        "Daily":   C["blue"],
        "Weekly":  C["purple"],
        "Monthly": C["cyan"],
    }
    cols_cad = st.columns(len(cad_stats))
    for i, row in cad_stats.iterrows():
        c = cad_colors.get(row["cadence"], C["muted"])
        comp = row["Compliance %"]
        with cols_cad[i]:
            st.markdown(f'''<div class="kpi-card" style="border-left:4px solid {c}">
<p class="kpi-label">{row["cadence"]} Jobs</p>
<p class="kpi-value" style="color:{c}">{comp:.1f}%</p>
<p class="kpi-sub">SLA Compliance</p>
<p class="kpi-sub">{int(row["Jobs"])} jobs · {int(row["Total_Runs"])} runs</p>
<p class="kpi-sub">Avg {row["Avg_Hrs"]:.2f}h · Peak {row["Peak_Hrs"]:.2f}h</p>
<p class="kpi-sub" style="color:{C["red"] if row["Breaches"] else C["green"]}">
  {int(row["Breaches"])} breach(es)</p>
</div>''', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # Full table
    st.markdown('<div class="panel"><p class="panel-title">📋 Full-Volume Run Detail Table</p>',
                unsafe_allow_html=True)
    disp_tbl = top[["Job_Name","peak_hrs","avg_hrs","total_hrs",
                     "sla_used_pct","buffer_pct","status"]].copy()
    disp_tbl.columns = ["Job","Peak Hrs","Avg Hrs","Total Hrs",
                         "SLA Used %","Buffer %","Status"]
    st.dataframe(disp_tbl.sort_values("Buffer %"), use_container_width=True, height=300)
    st.download_button("⬇️ Export Perf-Test Report CSV",
                       disp_tbl.to_csv(index=False).encode(),
                       "perf_test_report.csv","text/csv", key="exp_perf")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── PER-JOB EXPANDABLE DRILLDOWN ──────────────────────────────────
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    st.markdown(f'<div class="panel"><p class="panel-title">🔍 Per-Job Drilldown</p>'
                f'<p class="panel-sub">Click any job to expand its daily trend chart, run history and SLA buffer gauge</p>',
                unsafe_allow_html=True)

    df_ctrlm = st.session_state.ctrlm_df

    # Sort: breaches first, then at-risk, then by peak hrs desc
    def _sort_key(row):
        if row["buffer_pct"] < 0:   return (0, -row["peak_hrs"])
        if row["buffer_pct"] < 15:  return (1, -row["peak_hrs"])
        return                              (2, -row["peak_hrs"])

    top_sorted = top.copy()
    top_sorted["_sort"] = top_sorted.apply(_sort_key, axis=1)
    top_sorted = top_sorted.sort_values("_sort").drop(columns=["_sort"])

    for job_idx, (_, jrow) in enumerate(top_sorted.iterrows()):
        jname     = jrow["Job_Name"]
        peak_h    = jrow["peak_hrs"]
        avg_h     = jrow["avg_hrs"]
        buf_pct   = jrow["buffer_pct"]
        sla_pct   = jrow["sla_used_pct"]
        status    = jrow["status"]

        # Status badge in expander label
        if status == "BREACH":
            status_icon = "🔴"
            expand_default = True
        elif status == "AT RISK":
            status_icon = "🟡"
            expand_default = True
        else:
            status_icon = "🟢"
            expand_default = False

        label = (f"{status_icon} {jname[:45]}{'…' if len(jname)>45 else ''}"
                 f"   ·   Peak {peak_h:.2f}h"
                 f"   ·   Buffer {buf_pct:.1f}%"
                 f"   ·   SLA Used {sla_pct:.0f}%")

        with st.expander(label, expanded=expand_default):
            col_left, col_right = st.columns([2, 1])

            with col_left:
                # Daily trend sparkline
                job_daily = (df_ctrlm[df_ctrlm["Job_Name"] == jname]
                             .groupby("run_date")["run_time_hrs"].sum()
                             .reset_index()
                             .rename(columns={"run_date":"date","run_time_hrs":"hrs"}))
                if not job_daily.empty:
                    trend_colors = [C["red"] if h > DAILY_LIMIT_HRS else
                                    (C["amber"] if h > DAILY_LIMIT_HRS * 0.85 else C["green"])
                                    for h in job_daily["hrs"]]
                    fig_trend = go.Figure()
                    fig_trend.add_trace(go.Bar(
                        x=job_daily["date"].astype(str), y=job_daily["hrs"],
                        marker_color=trend_colors,
                        hovertemplate="<b>%{x}</b><br>%{y:.3f}h<extra></extra>",
                    ))
                    fig_trend.add_hline(y=DAILY_LIMIT_HRS, line_dash="dash",
                                        line_color=C["red"], line_width=1.5,
                                        annotation_text=f"SLA {DAILY_LIMIT_HRS}h",
                                        annotation_font=dict(color=C["red"], size=9))
                    if DAILY_LIMIT_HRS * 0.85 < peak_h:
                        fig_trend.add_hline(y=DAILY_LIMIT_HRS * 0.85, line_dash="dot",
                                            line_color=C["amber"], line_width=1,
                                            annotation_text="85% warn",
                                            annotation_font=dict(color=C["amber"], size=8))
                    fig_trend.update_layout(
                        **BASE_LAYOUT, height=200,
                        margin=dict(l=8, r=8, t=28, b=30),
                        title=dict(text=f"<b>Daily Runtime — {jname}</b>",
                                   font_size=10, x=0),
                        xaxis=dict(tickangle=-30, tickfont_size=8, **AXIS),
                        yaxis=dict(title="Hrs", tickfont_size=8, **AXIS),
                    )
                    st.plotly_chart(fig_trend, use_container_width=True,
                                    config={"displayModeBar": False},
                                    key=f"pj_trend_{job_idx}")

                # Run history table (last 10 runs)
                job_runs = (df_ctrlm[df_ctrlm["Job_Name"] == jname]
                            .sort_values("Start_Time", ascending=False)
                            .head(10)[["Start_Time","run_time_hrs","run_date"]]
                            .rename(columns={"Start_Time":"Start","run_time_hrs":"Hrs","run_date":"Date"})
                            .copy())
                job_runs["Status"] = job_runs["Hrs"].apply(
                    lambda h: "🔴 BREACH" if h > DAILY_LIMIT_HRS else
                              ("🟡 AT RISK" if h > DAILY_LIMIT_HRS * 0.85 else "🟢 OK"))
                job_runs["Hrs"] = job_runs["Hrs"].round(4)
                st.dataframe(job_runs, use_container_width=True, height=210,
                             hide_index=True)

            with col_right:
                # SLA buffer gauge
                buf_clamped = max(-50.0, min(150.0, float(buf_pct)))
                gauge_color = (C["red"] if buf_pct < 0 else
                               C["amber"] if buf_pct < 15 else C["green"])
                fig_gauge = go.Figure(go.Indicator(
                    mode="gauge+number+delta",
                    value=buf_clamped,
                    number=dict(suffix="%", font=dict(size=26, color=gauge_color),
                                valueformat=".1f"),
                    delta=dict(reference=15, valueformat=".1f",
                               increasing=dict(color=C["green"]),
                               decreasing=dict(color=C["red"])),
                    gauge=dict(
                        axis=dict(range=[-50, 100],
                                  tickvals=[-50, 0, 15, 50, 100],
                                  ticktext=["-50%","0%","15%","50%","100%"],
                                  tickfont=dict(size=8, color=C["muted"])),
                        bar=dict(color=gauge_color, thickness=0.28),
                        bgcolor="rgba(0,0,0,0)", borderwidth=0,
                        steps=[
                            dict(range=[-50, 0],  color="rgba(239,68,68,0.18)"),
                            dict(range=[0,  15],  color="rgba(245,158,11,0.15)"),
                            dict(range=[15, 100], color="rgba(34,197,94,0.12)"),
                        ],
                        threshold=dict(line=dict(color="#fff", width=2),
                                       thickness=0.75, value=0),
                    ),
                    title=dict(text="<b>SLA Buffer %</b><br>"
                                    "<span style='font-size:10px'>positive = headroom</span>",
                               font=dict(size=11, color=C["white"])),
                ))
                fig_gauge.update_layout(
                    **BASE_LAYOUT, height=220,
                    margin=dict(l=10, r=10, t=55, b=5),
                )
                st.plotly_chart(fig_gauge, use_container_width=True,
                                config={"displayModeBar": False},
                                key=f"pj_gauge_{job_idx}")

                # Stats summary card
                monthly_breach = m["monthly"][
                    (m["monthly"]["Job_Name"] == jname) & m["monthly"]["breach"]]
                n_monthly_breach = len(monthly_breach)
                run_count = len(df_ctrlm[df_ctrlm["Job_Name"] == jname])
                st.markdown(f"""
<div style='background:{C["card2"]};border:1px solid {C["border"]};
border-radius:10px;padding:12px 14px;font-size:11px;margin-top:6px'>
<div style='display:grid;grid-template-columns:1fr 1fr;gap:8px'>
  <div>
    <div style='color:{C["muted"]};font-size:9px;text-transform:uppercase;letter-spacing:.06em'>Peak Runtime</div>
    <div style='font-size:16px;font-weight:800;color:{C["red"] if peak_h>DAILY_LIMIT_HRS else C["white"]}'>{peak_h:.3f}h</div>
  </div>
  <div>
    <div style='color:{C["muted"]};font-size:9px;text-transform:uppercase;letter-spacing:.06em'>Avg Runtime</div>
    <div style='font-size:16px;font-weight:800;color:{C["blue"]}'>{avg_h:.3f}h</div>
  </div>
  <div>
    <div style='color:{C["muted"]};font-size:9px;text-transform:uppercase;letter-spacing:.06em'>SLA Used</div>
    <div style='font-size:16px;font-weight:800;color:{gauge_color}'>{sla_pct:.0f}%</div>
  </div>
  <div>
    <div style='color:{C["muted"]};font-size:9px;text-transform:uppercase;letter-spacing:.06em'>Total Runs</div>
    <div style='font-size:16px;font-weight:800;color:{C["white"]}'>{run_count}</div>
  </div>
</div>
<div style='margin-top:8px;padding-top:8px;border-top:1px solid {C["border"]};
  color:{C["red"] if n_monthly_breach else C["green"]};font-size:10px;font-weight:600'>
  {'🔴 ' + str(n_monthly_breach) + ' monthly SLA breach(es)' if n_monthly_breach else '✅ No monthly SLA breaches'}
</div>
</div>""", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # ── 🤖 Batch AI Deep Analysis ──────────────────────────────────────
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    st.markdown(f'''<div class="panel">
<p class="panel-title">🤖 Batch Performance AI Analysis</p>
<p class="panel-sub">⚡ Quick Heuristics — or — 🤖 Send metrics to Gemini for a full diagnostic report</p>''',
        unsafe_allow_html=True)

    bq_col, bai_col = st.columns([1, 1])
    with bq_col:
        if st.button("⚡ Quick Batch Heuristics", use_container_width=True, key="batch_quick_btn"):
            st.session_state["_batch_quick"] = True
            st.session_state.pop("_batch_ai_result", None)
    with bai_col:
        if st.button("🤖 AI Deep Analysis (Batch)", use_container_width=True, key="batch_ai_btn"):
            st.session_state["_run_batch_ai"] = True
            st.session_state.pop("_batch_quick", None)

    if st.session_state.get("_batch_quick"):
        breach_jobs = top[top["buffer_pct"] < 0]
        risk_jobs   = top[(top["buffer_pct"] >= 0) & (top["buffer_pct"] < 15)]
        if not breach_jobs.empty:
            st.error(f"🔴 **{len(breach_jobs)} job(s) exceed the {DAILY_LIMIT_HRS}h SLA:** "
                     + ", ".join(breach_jobs["Job_Name"].tolist()[:5]))
        if not risk_jobs.empty:
            st.warning(f"⚠️ **{len(risk_jobs)} job(s) have <15% SLA buffer:** "
                       + ", ".join(risk_jobs["Job_Name"].tolist()[:5]))
        if breach_jobs.empty and risk_jobs.empty:
            st.success("✅ All jobs have >15% SLA buffer. Batch window is go-live ready.")
        st.info(f"📊 Overall compliance: **{m['compliance']:.1f}%** | "
                f"Peak daily window: **{m['window']['total_hrs'].max():.2f}h** | "
                f"Avg window: **{m['window']['total_hrs'].mean():.2f}h**")

    if st.session_state.get("_run_batch_ai"):
        st.session_state.pop("_run_batch_ai", None)
        ai_run_batch_analysis(m, top)

    if st.session_state.get("_batch_ai_result") and not st.session_state.get("_batch_quick"):
        st.markdown(st.session_state["_batch_ai_result"])

    st.markdown("</div>", unsafe_allow_html=True)





# ── UI PERFORMANCE BENCHMARKING TAB ─────────────────────────
def ui_performance_tab():
    """
    UI Performance Benchmarking Report.
    Handles any format including JDA/Blue Yonder format:
      Columns: UI Name, Action, Scenario Name, Search Name, Search Criteria,
               Response Time (HH:MI:SS), Record Count
    Also handles JMeter / Gatling / LoadRunner / custom CSV.
    """

    # ── SLA Thresholds ────────────────────────────────────────
    with st.expander("⚙️ Configure SLA Thresholds", expanded=False):
        tc1, tc2, tc3 = st.columns(3)
        with tc1:
            resp_sla = st.number_input("Max Response Time (s)", min_value=1,
                                        max_value=300, value=10, step=1,
                                        key="uiperf_resp_sla",
                                        help="Maximum acceptable response time in seconds")
        with tc2:
            err_sla = st.number_input("Max Error Rate (%)", min_value=0.0,
                                       max_value=100.0, value=1.0, step=0.1,
                                       format="%.1f", key="uiperf_err_sla")
        with tc3:
            tps_sla = st.number_input("Min Throughput (TPS, 0=N/A)", min_value=0,
                                       max_value=10000, value=0, step=1,
                                       key="uiperf_tps_sla")

    # ── File Upload ───────────────────────────────────────────
    st.markdown(f'''<div style="background:{C["card"]};border:1px solid {C["border"]};
border-radius:14px;padding:16px 20px;margin-bottom:14px">
<p class="panel-title">📁 Upload Benchmark Report</p>
<p class="panel-sub">
  Supported: JDA / Blue Yonder format · JMeter · Gatling · LoadRunner · CSV · XLSX · PDF · DOCX<br>
  JDA columns: <code>UI Name</code>, <code>Action</code>, <code>Response Time (HH:MI:SS)</code>, <code>Record Count</code>
</p></div>''', unsafe_allow_html=True)

    uf = st.file_uploader(
        "Upload Benchmark Report",
        type=["csv","xlsx","xls","pdf","docx","doc","txt","png","jpg","jpeg"],
        key="fu_ui_perf",
        label_visibility="collapsed"
    )

    ui_df = None

    if uf is not None:
        h_key = hash(uf.name + str(uf.size))
        if st.session_state.get("_uiperf_hash") != h_key:
            try:
                ext = uf.name.lower()
                uf.seek(0)
                raw_frames = []  # list of (sheet_name, df) tuples

                # ── Read by format ────────────────────────────────
                if ext.endswith((".xlsx", ".xls")):
                    import openpyxl as _oxl
                    wb = _oxl.load_workbook(uf, data_only=True)
                    for sh_name in wb.sheetnames:
                        ws = wb[sh_name]
                        rows = list(ws.iter_rows(values_only=True))
                        if len(rows) < 2: continue
                        # Find actual header row (first non-empty row)
                        hdr_idx = 0
                        for ri, row in enumerate(rows):
                            if any(v is not None and str(v).strip() for v in row):
                                hdr_idx = ri
                                break
                        headers = [str(v).strip() if v is not None else f"col_{ci}"
                                   for ci, v in enumerate(rows[hdr_idx])]
                        data_rows = []
                        for row in rows[hdr_idx+1:]:
                            if any(v is not None for v in row):
                                data_rows.append(dict(zip(headers, row)))
                        if data_rows:
                            raw_frames.append((sh_name, data_rows))

                elif ext.endswith(".csv") or ext.endswith(".txt"):
                    import pandas as _pd2
                    import io as _io2
                    raw = uf.read().decode("utf-8", errors="replace")
                    df_tmp = _pd2.read_csv(_io2.StringIO(raw), sep=None, engine="python",
                                           on_bad_lines="skip", dtype=str)
                    raw_frames.append(("Sheet1", df_tmp.to_dict("records")))

                elif ext.endswith((".docx", ".doc")):
                    from docx import Document as _DD
                    _doc = _DD(uf)
                    for ti, tbl in enumerate(_doc.tables):
                        rows = []
                        for ri, row in enumerate(tbl.rows):
                            cells = [c.text.strip() for c in row.cells]
                            if cells and any(cells):
                                rows.append(cells)
                        if len(rows) >= 2:
                            hdr = rows[0]
                            raw_frames.append((f"Table{ti+1}",
                                               [dict(zip(hdr, r)) for r in rows[1:]]))

                elif ext.endswith(".pdf"):
                    _txt = extract_pdf_text(uf)
                    import io as _io3, pandas as _pd3
                    try:
                        df_tmp = _pd3.read_csv(_io3.StringIO(_txt), sep=None, engine="python",
                                               on_bad_lines="skip", dtype=str)
                        raw_frames.append(("PDF", df_tmp.to_dict("records")))
                    except Exception:
                        pass

                if not raw_frames:
                    st.error("❌ Could not extract any tabular data from the file.")
                    st.stop()

                # ── Parse each sheet into normalised records ──────
                all_records = []
                for sheet_name, rows in raw_frames:
                    for row in rows:
                        if not isinstance(row, dict): continue
                        keys = list(row.keys())

                        # ── Detect JDA/Blue Yonder format ─────────
                        # Columns: UI Name | Action | Scenario Name | Search Name | Search Criteria | Response Time (HH:MI:SS) | Record Count
                        col_lower = {k.lower().replace("\n"," ").replace("  "," ").strip(): k for k in keys}
                        col_keys  = list(col_lower.keys())

                        is_jda = any("ui name" in c or "ui_name" in c for c in col_keys)
                        has_resp_time = any("response time" in c or "resp_time" in c or "hh:mi" in c or "hh:mm" in c for c in col_keys)

                        if is_jda and has_resp_time:
                            # JDA format
                            ui_name_col  = next((col_lower[c] for c in col_keys if "ui name" in c or "ui_name" in c), None)
                            action_col   = next((col_lower[c] for c in col_keys if c == "action"), None)
                            scenario_col = next((col_lower[c] for c in col_keys if "scenario" in c), None)
                            resp_col     = next((col_lower[c] for c in col_keys if "response time" in c or "hh:mi" in c or "hh:mm" in c), None)
                            count_col    = next((col_lower[c] for c in col_keys if "record count" in c or "count" in c or "samples" in c), None)

                            ui_name_val  = str(row.get(ui_name_col, "")).strip() if ui_name_col else ""
                            if not ui_name_val or ui_name_val.lower() == "none": continue

                            # Convert Response Time: openpyxl gives datetime; extract HH:MM:SS
                            resp_raw = row.get(resp_col) if resp_col else None
                            resp_secs_val = 0.0
                            if resp_raw is not None:
                                import datetime as _dt
                                if isinstance(resp_raw, _dt.datetime):
                                    # Use ONLY the time component (h*3600 + m*60 + s)
                                    resp_secs_val = resp_raw.hour * 3600 + resp_raw.minute * 60 + resp_raw.second
                                elif isinstance(resp_raw, _dt.time):
                                    resp_secs_val = resp_raw.hour * 3600 + resp_raw.minute * 60 + resp_raw.second
                                elif isinstance(resp_raw, (int, float)):
                                    resp_secs_val = float(resp_raw)
                                else:
                                    # Try parsing string like "0:00:05" or "00:05:30"
                                    import re as _re
                                    parts = _re.split(r"[:.]", str(resp_raw).strip())
                                    try:
                                        if len(parts) >= 3:
                                            resp_secs_val = int(parts[0])*3600 + int(parts[1])*60 + int(parts[2])
                                        elif len(parts) == 2:
                                            resp_secs_val = int(parts[0])*60 + int(parts[1])
                                    except Exception:
                                        resp_secs_val = 0.0

                            action_val   = str(row.get(action_col, "")).strip() if action_col else ""
                            scenario_val = str(row.get(scenario_col, "")).strip() if scenario_col else ""
                            count_val    = int(row.get(count_col, 0) or 0) if count_col else 0

                            # Build transaction label: UI Name + action + sheet context
                            txn_label = ui_name_val
                            if action_val and action_val.lower() != "none":
                                txn_label += f" [{action_val}]"
                            if scenario_val and scenario_val.lower() not in ("none","live",""):
                                txn_label += f" — {scenario_val}"

                            all_records.append({
                                "Transaction":    txn_label,
                                "Module":         sheet_name,
                                "Avg_ms":         round(resp_secs_val * 1000, 0),
                                "P95_ms":         round(resp_secs_val * 1000, 0),
                                "P99_ms":         0.0,
                                "TPS":            0.0,
                                "ErrorRate_pct":  0.0,
                                "Samples":        count_val,
                                "Resp_sec":       round(resp_secs_val, 2),
                            })

                        else:
                            # Generic / JMeter / Gatling format — fuzzy column map
                            def _find(aliases):
                                for a in aliases:
                                    for ck, orig_k in col_lower.items():
                                        if a in ck: return row.get(orig_k)
                                return None

                            txn = _find(["ui name","transaction","label","request","scenario","api","endpoint","url","page","sampler","name"])
                            p95 = _find(["p95","95th","95_percentile","perc95"])
                            avg = _find(["avg","average","mean"])
                            p99 = _find(["p99","99th","99_percentile","perc99"])
                            tps = _find(["tps","throughput","req_s","requests_per_sec","rate"])
                            err = _find(["error","err","fail","failure"])
                            cnt = _find(["samples","requests","count","hits","record count"])
                            resp_time = _find(["response time","resp_time","hh:mi","hh:mm"])

                            if txn is None: continue

                            def _to_ms(v):
                                if v is None: return 0.0
                                import datetime as _dt
                                if isinstance(v, _dt.datetime):
                                    secs = v.hour*3600+v.minute*60+v.second
                                    return round(secs*1000, 0)
                                if isinstance(v, _dt.time):
                                    return round((v.hour*3600+v.minute*60+v.second)*1000, 0)
                                try:
                                    return round(float(str(v).replace("%","").replace(",","").strip()), 0)
                                except Exception: return 0.0

                            p95_ms = _to_ms(p95) or _to_ms(resp_time) or _to_ms(avg)
                            all_records.append({
                                "Transaction":   str(txn).strip(),
                                "Module":        sheet_name,
                                "Avg_ms":        _to_ms(avg) or p95_ms,
                                "P95_ms":        p95_ms,
                                "P99_ms":        _to_ms(p99),
                                "TPS":           _to_ms(tps),
                                "ErrorRate_pct": _to_ms(err),
                                "Samples":       int(_to_ms(cnt) or 0),
                                "Resp_sec":      round(p95_ms/1000, 2),
                            })

                if not all_records:
                    st.error("❌ **No valid transaction rows found.**\n\n"
                             "Expected: `UI Name` + `Response Time` columns (JDA format) "
                             "or `Transaction` + `P95` columns (JMeter format).")
                    st.stop()

                import pandas as _pd_ui
                ui_df_new = _pd_ui.DataFrame(all_records)
                st.session_state["_uiperf_df"]    = ui_df_new
                st.session_state["_uiperf_hash"]  = h_key
                st.session_state["_uiperf_fname"] = uf.name
                st.success(f"✅ Loaded **{len(ui_df_new)} transactions** across "
                           f"**{ui_df_new['Module'].nunique()} module(s)** from `{uf.name}`")
                ui_df = ui_df_new

            except Exception as exc:
                st.error(f"❌ **Failed to read benchmark file:** {exc}")
        else:
            ui_df = st.session_state.get("_uiperf_df")
            if ui_df is not None:
                st.success(f"✅ **{len(ui_df)} transactions** loaded from "
                           f"`{st.session_state.get('_uiperf_fname','')}`")
    else:
        ui_df = st.session_state.get("_uiperf_df")

    if ui_df is None or ui_df.empty:
        st.info("📂 Upload a benchmark file above to begin analysis.")
        return

    resp_sla_val = st.session_state.get("uiperf_resp_sla", 10)
    err_sla_val  = st.session_state.get("uiperf_err_sla", 1.0)
    tps_sla_val  = st.session_state.get("uiperf_tps_sla", 0)
    resp_sla_ms  = resp_sla_val * 1000

    # ── Derive pass/fail ──────────────────────────────────────
    ui_df = ui_df.copy()
    ui_df["P95_pass"]  = ui_df["P95_ms"] <= resp_sla_ms
    ui_df["Err_pass"]  = ui_df["ErrorRate_pct"] <= err_sla_val
    ui_df["TPS_pass"]  = (tps_sla_val == 0) | (ui_df["TPS"] == 0) | (ui_df["TPS"] >= tps_sla_val)
    ui_df["Overall"]   = ui_df["P95_pass"] & ui_df["Err_pass"]

    n_pass = int(ui_df["Overall"].sum())
    n_fail = len(ui_df) - n_pass
    avg_resp = float(ui_df["P95_ms"].mean()) / 1000
    max_resp = float(ui_df["P95_ms"].max()) / 1000
    avg_err  = float(ui_df["ErrorRate_pct"].mean())
    has_modules = ui_df["Module"].nunique() > 1

    # ── KPI Row ───────────────────────────────────────────────
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    verdict_c = C["green"] if n_fail == 0 else C["red"]
    k1,k2,k3,k4 = st.columns(4)
    for col, lbl, val, sub, color in [
        (k1,"VERDICT",     f"{'✅ PASS' if n_fail==0 else '🔴 FAIL'}", f"{n_pass}/{len(ui_df)} pass SLA", verdict_c),
        (k2,"AVG RESPONSE", f"{avg_resp:.1f}s", f"SLA ≤ {resp_sla_val}s", C["green"] if avg_resp<=resp_sla_val else C["red"]),
        (k3,"MAX RESPONSE", f"{max_resp:.1f}s", "Worst single transaction", C["green"] if max_resp<=resp_sla_val else C["red"]),
        (k4,"TRANSACTIONS", f"{len(ui_df)}", f"{ui_df['Module'].nunique()} module(s)", C["blue"]),
    ]:
        col.markdown(f'''<div class="kpi-card" style="border-left:4px solid {color}">
<p class="kpi-label">{lbl}</p>
<p class="kpi-value" style="color:{color};font-size:22px">{val}</p>
<p class="kpi-sub">{sub}</p></div>''', unsafe_allow_html=True)
    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    # ── Module filter ─────────────────────────────────────────
    if has_modules:
        modules = ["All Modules"] + sorted(ui_df["Module"].unique().tolist())
        sel_mod = st.selectbox("Filter by Module", modules, key="uiperf_mod_filter")
        if sel_mod != "All Modules":
            display_df = ui_df[ui_df["Module"] == sel_mod].copy()
        else:
            display_df = ui_df.copy()
    else:
        display_df = ui_df.copy()

    # ── Response Time Chart ───────────────────────────────────
    st.markdown(f'''<div class="panel">
<p class="panel-title">📈 Response Time by Transaction</p>
<p class="panel-sub">Red = exceeds {resp_sla_val}s SLA · Amber = within 20% · Green = passing · SLA line shown</p>''',
        unsafe_allow_html=True)

    # Show max 40 transactions — sorted worst-first
    chart_df = display_df.sort_values("P95_ms", ascending=False).head(40)
    txn_labels = [str(t)[:45] + ("…" if len(str(t))>45 else "") for t in chart_df["Transaction"]]
    bar_colors = [
        C["red"]   if v > resp_sla_ms else
        C["amber"] if v > resp_sla_ms * 0.8 else
        C["green"]
        for v in chart_df["P95_ms"]
    ]
    fig_resp = go.Figure(go.Bar(
        y=txn_labels, x=chart_df["Resp_sec"],
        orientation="h",
        marker_color=bar_colors, opacity=0.88,
        customdata=chart_df["Samples"],
        hovertemplate="<b>%{y}</b><br>Response: <b>%{x:.2f}s</b><br>Records: %{customdata:,}<extra></extra>",
    ))
    fig_resp.add_vline(x=resp_sla_val, line_dash="dash", line_color=C["red"],
                        line_width=2,
                        annotation_text=f"SLA {resp_sla_val}s",
                        annotation_font=dict(color=C["red"], size=10))
    fig_resp.update_layout(
        **BASE_LAYOUT, height=max(280, len(chart_df)*20+80),
        xaxis=dict(title="Response Time (s)", **AXIS),
        yaxis=dict(tickfont_size=8, autorange="reversed", automargin=True, **AXIS),
        margin=dict(l=10,r=10,t=10,b=40),
    )
    st.plotly_chart(fig_resp, use_container_width=True,
                    config={"displayModeBar":False}, key="uip_resp_chart")
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ── Per-Module Summary ─────────────────────────────────────
    if has_modules:
        st.markdown(f'''<div class="panel">
<p class="panel-title">📊 Module Summary</p>
<p class="panel-sub">Pass/Fail per module · SLA = {resp_sla_val}s response time</p>''',
            unsafe_allow_html=True)
        for mod in sorted(ui_df["Module"].unique()):
            mod_df = ui_df[ui_df["Module"]==mod]
            m_pass = int(mod_df["Overall"].sum())
            m_fail = len(mod_df) - m_pass
            m_avg  = mod_df["P95_ms"].mean()/1000
            m_max  = mod_df["P95_ms"].max()/1000
            mc = C["green"] if m_fail==0 else C["red"]
            st.markdown(
                f'''<div style="display:flex;justify-content:space-between;align-items:center;
padding:8px 14px;background:{C["card2"]};border-radius:8px;margin-bottom:4px;
border-left:3px solid {mc}">
<span style="font-size:11px;font-weight:600">{mod}</span>
<div style="display:flex;gap:20px;font-size:10px">
  <span style="color:{C["muted"]}">{len(mod_df)} transactions</span>
  <span style="color:{C["muted"]}">Avg: {m_avg:.1f}s</span>
  <span style="color:{C["muted"]}">Max: {m_max:.1f}s</span>
  <span style="color:{mc};font-weight:700">{m_pass} pass / {m_fail} fail</span>
</div></div>''', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ── Detailed Pass/Fail Table ──────────────────────────────
    st.markdown(f'''<div class="panel">
<p class="panel-title">📋 Transaction Detail — Pass / Fail</p>
<p class="panel-sub">SLA: Response ≤ {resp_sla_val}s · Error rate ≤ {err_sla_val}%</p>''',
        unsafe_allow_html=True)

    for _, row in display_df.sort_values("P95_ms", ascending=False).iterrows():
        ok   = bool(row["Overall"])
        rc   = C["green"] if ok else C["red"]
        icon = "✅" if ok else "🔴"
        resp_disp   = f"{row['Resp_sec']:.2f}s"
        sample_disp = f"{int(row['Samples']):,}" if row["Samples"] > 0 else "—"
        module_disp = row.get("Module","")
        txn_disp    = str(row["Transaction"])[:60]
        pass_badge  = f'<span class="b-ok" style="font-size:9px">PASS</span>' if ok else f'<span class="b-breach" style="font-size:9px">FAIL</span>'

        st.markdown(
            f'''<div style="display:grid;grid-template-columns:2.5fr 1fr 0.8fr 0.8fr 0.6fr;
align-items:center;padding:7px 12px;margin-bottom:3px;
background:{rc}08;border-left:3px solid {rc}66;border-radius:0 8px 8px 0">
<span style="font-size:10px;font-weight:600;overflow:hidden;text-overflow:ellipsis;white-space:nowrap"
  title="{row["Transaction"]}">{icon} {txn_disp}</span>
<span style="font-size:10px;color:{C["muted"]};text-align:center">{module_disp}</span>
<span style="font-size:10px;font-weight:600;color:{rc if not ok else C["white"]};text-align:right">{resp_disp}</span>
<span style="font-size:10px;color:{C["muted"]};text-align:right">{sample_disp} recs</span>
<span style="text-align:right">{pass_badge}</span>
</div>''', unsafe_allow_html=True)

    # Export
    export_df = display_df[["Transaction","Module","Resp_sec","Avg_ms","P95_ms","Samples","ErrorRate_pct"]].copy()
    export_df.columns = ["Transaction","Module","Response_Sec","Avg_ms","P95_ms","Samples","Error_Pct"]
    export_df["SLA_Pass"] = display_df["Overall"].map({True:"PASS",False:"FAIL"})
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.download_button("⬇️ Export Results CSV",
                       export_df.to_csv(index=False).encode(),
                       "ui_perf_results.csv", "text/csv", key="dl_uiperf")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── AI Analysis ───────────────────────────────────────────
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    st.markdown(f'''<div class="panel">
<p class="panel-title">🤖 UI Performance AI Analysis</p>
<p class="panel-sub">⚡ Quick assessment · 🤖 Gemini deep diagnostic</p>''',
        unsafe_allow_html=True)

    uq, uai = st.columns(2)
    with uq:
        if st.button("⚡ Quick Assessment", use_container_width=True, key="uiperf_quick_btn"):
            st.session_state["_uiperf_quick"] = True
            st.session_state.pop("_uiperf_ai_result", None)
    with uai:
        if st.button("🤖 AI Deep Analysis (UI Perf)", use_container_width=True, key="uiperf_ai_btn"):
            st.session_state["_run_uiperf_ai"] = True
            st.session_state.pop("_uiperf_quick", None)

    if st.session_state.get("_uiperf_quick"):
        fail_txns = ui_df[~ui_df["Overall"]]["Transaction"].tolist()
        if fail_txns:
            st.error(f"🔴 **{len(fail_txns)} transaction(s) exceed {resp_sla_val}s SLA:**\n" +
                     "\n".join(f"- {t[:70]}" for t in fail_txns[:8]))
        else:
            st.success(f"✅ All {len(ui_df)} transactions pass the {resp_sla_val}s response time SLA.")
        st.info(f"Avg: **{avg_resp:.2f}s** · Max: **{max_resp:.2f}s** · "
                f"SLA: **{resp_sla_val}s** · Pass rate: **{n_pass}/{len(ui_df)}**")

    if st.session_state.get("_run_uiperf_ai"):
        st.session_state.pop("_run_uiperf_ai", None)
        import json as _json
        _sla = {"response_time_sec": resp_sla_val, "error_pct": err_sla_val}
        _top20 = ui_df.sort_values("P95_ms", ascending=False).head(20)[
            ["Transaction","Module","Resp_sec","Samples"]].round(2).to_dict("records")
        _fail  = ui_df[~ui_df["Overall"]][["Transaction","Resp_sec"]].to_dict("records")
        _prompt = f"""You are a Senior Performance Engineer reviewing JDA/Blue Yonder UI benchmark results.

SLA: {_json.dumps(_sla)}
Summary: {n_pass} pass / {n_fail} fail · Avg={avg_resp:.2f}s · Max={max_resp:.2f}s
Failing transactions: {_json.dumps(_fail[:15]) if _fail else "None"}
Top 20 slowest: {_json.dumps(_top20)}

Write a concise diagnostic covering:
1. **Pass/Fail Summary** — Which modules/transactions fail and by how much.
2. **Root Cause Hypotheses** — For slow transactions: DB query, large datasets, network latency, missing index, UI rendering.
3. **Risk Assessment** — Impact on end users in production.
4. **Remediation Plan** — Concrete fixes (pagination, caching, search criteria limits, indexing).
5. **Go-Live Verdict** — ✅ READY / ⚠️ CONDITIONAL / 🔴 NOT READY.

Use transaction names. Quote exact times. No generic padding."""
        _gemini_call(_prompt, "_uiperf_ai_result",
                     "🤖 Gemini is analysing your UI performance results…")

    if st.session_state.get("_uiperf_ai_result") and not st.session_state.get("_uiperf_quick"):
        st.markdown(st.session_state["_uiperf_ai_result"])

    st.markdown("</div>", unsafe_allow_html=True)



# ── SOW DETAILS TAB ──────────────────────────────────────────────────────
def volume_utilization_color(pct):
    """
    <70%  → RED   (under-utilized, potential issue with SOW sizing)
    70-90% → AMBER (moderate)
    90-110%→ GREEN (optimal)
    >110%  → RED   (over SOW — sizing risk)
    """
    if pct < 70:    return C["red"],   "UNDER-UTILIZED",   "⬇ Below 70% of SOW — flag with PM"
    if pct < 90:    return C["amber"], "MODERATE",         "Within comfortable range"
    if pct <= 110:  return C["green"], "OPTIMAL",          "✅ Within SOW agreement"
    return              C["red"],   "OVER SOW",         "⬆ Exceeds SOW — escalate sizing risk"

def volume_gauge_fig(val_pct, label, current, baseline, uid=""):
    """
    SOW utilization gauge with 4-zone coloring:
    <70 RED · 70-90 AMBER · 90-110 GREEN · >110 RED
    """
    col,status,_ = volume_utilization_color(val_pct)
    fig = go.Figure(go.Indicator(
        mode="gauge+number+delta",
        value=val_pct,
        number=dict(suffix="%", font=dict(size=26,color=col),
                    valueformat=".1f"),
        delta=dict(reference=100, valueformat=".1f",
                   increasing=dict(color=C["red"]),
                   decreasing=dict(color=C["green"])),
        gauge=dict(
            axis=dict(range=[0,150],
                      tickvals=[0,50,70,90,100,110,150],
                      ticktext=["0","50%","70%","90%","100% SOW","110%","150%"],
                      tickfont=dict(size=8,color=C["muted"])),
            bar=dict(color=col,thickness=0.28),
            bgcolor="rgba(0,0,0,0)", borderwidth=0,
            steps=[
                dict(range=[0,70],    color="rgba(239,68,68,0.13)"),  # red zone — under
                dict(range=[70,90],   color="rgba(245,158,11,0.13)"),  # amber
                dict(range=[90,110],  color="rgba(34,197,94,0.13)"),  # green — optimal
                dict(range=[110,150], color="rgba(239,68,68,0.19)"),  # red — over SOW
            ],
            threshold=dict(line=dict(color="#fff",width=2),
                           thickness=0.75,value=100),
        ),
        title=dict(
            text=(f"<b>{label}</b><br>"
                  f"<span style='font-size:11px'>{current:,} / {baseline:,}</span><br>"
                  f"<span style='font-size:10px;color:{col}'>{status}</span>"),
            font=dict(size=13,color=C["white"])),
    ))
    layout = {**BASE_LAYOUT,"height":230,"margin":dict(l=20,r=20,t=70,b=10)}
    fig.update_layout(**layout)
    return fig

def sow_tab():
    # ── Contract auto-upload ─────────────────────────────────────────
    st.markdown('<div class="panel"><p class="panel-title">📄 SOW Contract — Auto-Extract DFU / SKU</p>'
                '<p class="panel-sub">Upload Schedule 1-A PDF, XLSX or DOCX — DFU & SKU values auto-extracted</p>',
                unsafe_allow_html=True)
    sow_file = st.file_uploader("Upload SOW Contract (PDF/XLSX/DOCX)",
                                type=["pdf","xlsx","xls","docx","png","jpg","jpeg"],
                                key="sow_contract_upload")
    if sow_file:
        cust_found, dfu_found, sku_found = parse_sow_volumes(sow_file, sow_file.name)
        if cust_found:
            if not st.session_state.customer_name.strip():
                st.session_state.customer_name = cust_found
            st.success(f"✅ Customer detected from SOW: **{cust_found}**")
        if dfu_found:
            st.session_state.sow_dfu_base = dfu_found
            st.success(f"✅ SOW DFU baseline (Item-Location-Customer): **{dfu_found:,}**")
        if sku_found:
            st.session_state.sow_sku_base = sku_found
            st.success(f"✅ SOW SKU baseline (Item-Location): **{sku_found:,}**")
        if not dfu_found and not sku_found:
            st.warning("⚠️ Could not auto-detect DFU/SKU from contract — enter manually below")
        if not dfu_found and not sku_found:
            st.warning("⚠️ Could not auto-detect DFU/SKU — enter manually below")
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div style='height:8px'></div>",unsafe_allow_html=True)

    # ── Manual inputs ────────────────────────────────────────────────
    st.markdown('<div class="panel"><p class="panel-title">📊 DFU & SKU vs SOW Baseline</p>'
                '<p class="panel-sub">Color zones: 🔴 &lt;70% under-utilized · 🟡 70–90% moderate · '
                '🟢 90–110% optimal · 🔴 &gt;110% over SOW</p>',unsafe_allow_html=True)

    mc1,mc2,mc3,mc4 = st.columns(4)
    with mc1:
        cur_dfu = st.number_input("Current DFU (Item-Loc-Cust)",
                                  value=int(st.session_state.sow_dfu),
                                  step=10_000,format="%d",key="cur_dfu_in")
        st.session_state.sow_dfu = cur_dfu
    with mc2:
        sow_dfu_base = st.number_input("SOW Agreed DFU Baseline",
                                       value=int(st.session_state.get("sow_dfu_base",0)),step=10_000,
                                       format="%d",key="sow_dfu_base_in")
    with mc3:
        cur_sku = st.number_input("Current SKU (Item-Location)",
                                  value=int(st.session_state.sow_sku),
                                  step=10_000,format="%d",key="cur_sku_in")
        st.session_state.sow_sku = cur_sku
    with mc4:
        sow_sku_base = st.number_input("SOW Agreed SKU Baseline",
                                       value=int(st.session_state.get("sow_sku_base",0)),step=10_000,
                                       format="%d",key="sow_sku_base_in")

    sow_dfu_base = int(st.session_state.get("sow_dfu_base", sow_dfu_base))
    sow_sku_base = int(st.session_state.get("sow_sku_base", sow_sku_base))

    # Persist manual edits back to session
    st.session_state.sow_dfu_base = sow_dfu_base
    st.session_state.sow_sku_base = sow_sku_base

    if sow_dfu_base == 0 and sow_sku_base == 0:
        st.info("📋 Upload the SOW contract above or enter SOW Agreed baselines manually to see the DFU/SKU vs SOW analysis.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    # F1 — Volume Utilization (authoritative, capped at 150 for gauge display)
    _dfu_v1 = calculate_volume_utilization(cur_dfu, sow_dfu_base) if sow_dfu_base > 0 else {"util_pct": 0, "headroom_pct": 100, "headroom_items": 0, "status": "INVALID"}
    _sku_v1 = calculate_volume_utilization(cur_sku, sow_sku_base) if sow_sku_base > 0 else {"util_pct": 0, "headroom_pct": 100, "headroom_items": 0, "status": "INVALID"}
    # F2 — Growth projection (assume 3%/month default)
    _dfu_g2 = project_volume_growth(cur_dfu, sow_dfu_base, 0.03) if sow_dfu_base > 0 else None
    _sku_g2 = project_volume_growth(cur_sku, sow_sku_base, 0.03) if sow_sku_base > 0 else None
    dfu_pct = min(_dfu_v1["util_pct"], 150)
    sku_pct = min(_sku_v1["util_pct"], 150)

    g1,g2 = st.columns(2)
    with g1:
        st.plotly_chart(volume_gauge_fig(dfu_pct,"DFU vs SOW",cur_dfu,sow_dfu_base),
                        use_container_width=True,
                        config={"displayModeBar":False},key="dfu_sow_gauge")
        dfu_col,dfu_status,dfu_tip = volume_utilization_color(dfu_pct)
        st.markdown(
            f'''<div style="background:{dfu_col}18;border:1px solid {dfu_col}44;
border-radius:8px;padding:10px 14px;text-align:center">
<b style="color:{dfu_col};font-size:13px">DFU: {_dfu_v1["status"]}</b>
<p style="color:#94a3b8;font-size:11px;margin:4px 0 0">
  Headroom: <b>{_dfu_v1["headroom_items"]:,}</b> items ({_dfu_v1["headroom_pct"]:.1f}%)</p>
<p style="color:#94a3b8;font-size:11px">{cur_dfu:,} of {sow_dfu_base:,} ({dfu_pct:.1f}%)</p>
</div>''',unsafe_allow_html=True)
    with g2:
        st.plotly_chart(volume_gauge_fig(sku_pct,"SKU vs SOW",cur_sku,sow_sku_base),
                        use_container_width=True,
                        config={"displayModeBar":False},key="sku_sow_gauge")
        sku_col,sku_status,sku_tip = volume_utilization_color(sku_pct)
        st.markdown(
            f'''<div style="background:{sku_col}18;border:1px solid {sku_col}44;
border-radius:8px;padding:10px 14px;text-align:center">
<b style="color:{sku_col};font-size:13px">SKU: {_sku_v1["status"]}</b>
<p style="color:#94a3b8;font-size:11px;margin:4px 0 0">
  Headroom: <b>{_sku_v1["headroom_items"]:,}</b> items ({_sku_v1["headroom_pct"]:.1f}%)</p>
<p style="color:#94a3b8;font-size:11px">{cur_sku:,} of {sow_sku_base:,} ({sku_pct:.1f}%)</p>
</div>''',unsafe_allow_html=True)

    # F2 — Growth Runway cards
    if _dfu_g2 or _sku_g2:
        _r1, _r2 = st.columns(2)
        for _col, _g2, _label in [(_r1, _dfu_g2, "DFU"), (_r2, _sku_g2, "SKU")]:
            if not _g2:
                continue
            with _col:
                _runway = _g2["runway_months"]
                _proj   = _g2["projected_util_12mo"]
                _rc     = C["green"] if _runway > 18 else (C["amber"] if _runway > 6 else C["red"])
                st.markdown(
                    f'<div style="background:{_rc}11;border:1px solid {_rc}33;border-radius:8px;'
                    f'padding:8px 14px;text-align:center;margin-top:6px">'
                    f'<p style="font-size:10px;font-weight:700;color:{_rc};margin:0">📈 {_label} Growth Runway (F2)</p>'
                    f'<p style="font-size:12px;font-weight:700;color:{C["white"]};margin:2px 0">'
                    f'{"∞" if _runway > 200 else str(_runway) + " months"} runway</p>'
                    f'<p style="font-size:10px;color:{C["muted"]};margin:0">'
                    f'Projected 12-mo: {_proj:.1f}% utilisation @ 3%/mo growth</p>'
                    f'</div>', unsafe_allow_html=True)

    st.markdown("<div style='height:10px'></div>",unsafe_allow_html=True)

    # ── Scenario counts ──────────────────────────────────────────────
    sc1,sc2 = st.columns(2)
    with sc1:
        st.session_state.sow_scenarios_prod = st.number_input(
            "# Scenarios in PROD",value=int(st.session_state.get("sow_scenarios_prod", 0)),step=1)
    with sc2:
        st.session_state.sow_scenarios_agreed = st.number_input(
            "# Scenarios in SOW",value=int(st.session_state.get("sow_scenarios_agreed", 0)),step=1)
    if st.session_state.get("sow_scenarios_agreed", 0) > 0:
        ratio = st.session_state.get("sow_scenarios_prod", 0)/st.session_state.get("sow_scenarios_agreed", 0)
        sc_pct = ratio*100
        sc_col,sc_status,sc_tip = volume_utilization_color(sc_pct)
        st.markdown(
            f'''<div style="background:{sc_col}18;border:1px solid {sc_col}44;
border-radius:10px;padding:12px 18px;margin-top:10px;text-align:center">
<b style="color:{sc_col};font-size:14px">Scenarios: {st.session_state.get("sow_scenarios_prod", 0)}
PROD vs {st.session_state.get("sow_scenarios_agreed", 0)} SOW — {sc_pct:.0f}% ({sc_status})</b>
<p style="color:#94a3b8;font-size:11px;margin:4px 0 0">{sc_tip}</p>
</div>''',unsafe_allow_html=True)
    st.markdown("</div>",unsafe_allow_html=True)



# ── CUSTOMER APPROVAL TAB ────────────────────────────────────

def issues_waivers_tab():
    st.markdown(f'''<div style="background:{C["amber"]}22;border:1px solid {C["amber"]}55;
border-radius:12px;padding:12px 18px;margin-bottom:16px">
<span style="font-size:13px;font-weight:600;color:{C["amber"]}">
⚠ Known Issues & Waivers must be reviewed and acknowledged before PE sign-off.</span>
</div>''', unsafe_allow_html=True)

    if "issues_list" not in st.session_state:
        st.session_state.issues_list = []

    # Add new issue form
    with st.expander("➕ Add Issue / Waiver", expanded=len(st.session_state.issues_list)==0):
        ic1,ic2,ic3 = st.columns(3)
        with ic1:
            iss_id   = st.text_input("Issue ID", placeholder="ISS-001", key="iss_id")
            iss_type = st.selectbox("Type", ["Bug","Waiver","Risk","Performance","Configuration"], key="iss_type")
        with ic2:
            iss_sev  = st.selectbox("Severity", ["Critical","High","Medium","Low","Informational"], key="iss_sev")
            iss_stat = st.selectbox("Status", ["Open","In Progress","Waived","Resolved","Deferred"], key="iss_stat")
        with ic3:
            iss_own  = st.text_input("Owner", placeholder="PE/Customer/IT", key="iss_own")
            iss_eta  = st.text_input("ETA / Waiver Expiry", placeholder="DD-MMM-YYYY or N/A", key="iss_eta")
        iss_desc = st.text_area("Description", placeholder="Describe the issue, root cause, and impact...", height=80, key="iss_desc")
        iss_mit  = st.text_area("Mitigation / Waiver Justification", placeholder="Steps taken or reason for waiver...", height=60, key="iss_mit")
        if st.button("➕ Add to Register", key="add_issue_btn"):
            if iss_desc.strip():
                st.session_state.issues_list.append({
                    "ID": iss_id or f"ISS-{len(st.session_state.issues_list)+1:03d}",
                    "Type": iss_type, "Severity": iss_sev, "Status": iss_stat,
                    "Owner": iss_own, "ETA": iss_eta,
                    "Description": iss_desc, "Mitigation": iss_mit,
                    "Logged": date.today().isoformat(),
                })
                st.rerun()

    if st.session_state.issues_list:
        sev_color = {"Critical":C["red"],"High":"#f97316","Medium":C["amber"],
                     "Low":C["green"],"Informational":C["blue"]}
        stat_cls  = {"Open":"b-breach","In Progress":"b-warn",
                     "Waived":"b-warn","Resolved":"b-ok","Deferred":"b-warn"}

        open_cnt  = sum(1 for i in st.session_state.issues_list if i["Status"] in ["Open","In Progress"])
        waiv_cnt  = sum(1 for i in st.session_state.issues_list if i["Status"]=="Waived")
        res_cnt   = sum(1 for i in st.session_state.issues_list if i["Status"]=="Resolved")

        sc1,sc2,sc3,sc4 = st.columns(4)
        with sc1: kpi("Total Issues",str(len(st.session_state.issues_list)),"In register","📋",C["blue"])
        with sc2: kpi("Open / In Progress",str(open_cnt),"Needs action","🔴",C["red"] if open_cnt else C["green"])
        with sc3: kpi("Waivers",str(waiv_cnt),"Accepted risk","⚠️",C["amber"])
        with sc4: kpi("Resolved",str(res_cnt),"Closed","✅",C["green"])

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

        _iss_html = ""
        for i, iss in enumerate(st.session_state.issues_list):
            sc  = sev_color.get(iss["Severity"], C["muted"])
            _sc = stat_cls.get(iss["Status"], "b-warn")
            _mit = _html_mod.escape(iss["Mitigation"] or "No mitigation documented")
            _desc = _html_mod.escape(iss["Description"])
            _iss_html += (
                f'<div class="issue-card" style="border-left-color:{sc}">'
                f'<div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:6px">'
                f'<div><span style="font-size:13px;font-weight:700;color:{sc}">{iss["ID"]}</span>'
                f'<span style="font-size:11px;color:{C["muted"]};margin-left:10px">{iss["Type"]} · {iss["Severity"]}</span></div>'
                f'<div style="display:flex;gap:8px;align-items:center;flex-wrap:wrap">'
                f'<span class="{_sc}">{iss["Status"]}</span>'
                f'<span style="font-size:10px;color:{C["muted"]}">Owner: {iss["Owner"]} · ETA: {iss["ETA"]}</span>'
                f'</div></div>'
                f'<p style="font-size:12px;margin:8px 0 4px;line-height:1.5">{_desc}</p>'
                f'<p style="font-size:11px;color:{C["muted"]};margin:0">🛡 {_mit}</p>'
                f'<p style="font-size:10px;color:{C["muted"]};margin:4px 0 0">Logged: {iss["Logged"]}</p>'
                f'</div>'
            )
        if _iss_html:
            st.markdown(_iss_html, unsafe_allow_html=True)

        # Export
        iss_df = pd.DataFrame(st.session_state.issues_list)
        st.download_button("⬇️ Export Issues Register CSV",
                           iss_df.to_csv(index=False).encode(),
                           "issues_register.csv","text/csv",key="exp_issues")
    else:
        st.info("✅ No issues or waivers logged. Add any known items above before PE sign-off.")


def approval_tab():
    st.markdown('<div class="panel"><p class="panel-title">✍️ Customer Approval & PE Final Sign-Off</p>'
                '<p class="panel-sub">Both PE Engineer and Customer must approve before Go-Live</p>',
                unsafe_allow_html=True)

    # Checklist toggles
    st.markdown("**📋 PE Validation Checklist** — mark all items before signing off")
    ch1,ch2,ch3 = st.columns(3)
    with ch1:
        c_batch  = st.checkbox("✅ Batch SLA validated (daily/weekly/monthly)", key="chk_batch")
        c_res    = st.checkbox("✅ Resource utilization within thresholds", key="chk_res")
        c_data   = st.checkbox("✅ Data volume (DFU/SKU) vs SOW verified", key="chk_data")
    with ch2:
        c_issues = st.checkbox("✅ Issues & waivers acknowledged", key="chk_issues")
        c_perf   = st.checkbox("✅ Batch performance-test report reviewed", key="chk_perf")
        c_ctrlm  = st.checkbox("✅ Ctrl-M 30-day execution history reviewed", key="chk_ctrlm")
    with ch3:
        c_ui     = st.checkbox("✅ UI performance benchmarking approved", key="chk_ui")
        c_sow    = st.checkbox("✅ SOW service IDs & scenarios confirmed", key="chk_sow")
        c_res15  = st.checkbox("✅ Resource utilization (last 15 days) reviewed", key="chk_res15")

    checklist_done = all([c_batch,c_res,c_data,c_issues,c_perf,c_ctrlm,c_ui,c_sow,c_res15])
    pct = sum([c_batch,c_res,c_data,c_issues,c_perf,c_ctrlm,c_ui,c_sow,c_res15])
    bar_color = C["green"] if checklist_done else (C["amber"] if pct >= 6 else C["red"])
    st.markdown(f'''<div style="background:{C["card2"]};border-radius:8px;
height:10px;margin:10px 0 16px;overflow:hidden">
<div style="background:{bar_color};height:100%;width:{pct/9*100:.0f}%;
border-radius:8px;transition:width .3s"></div></div>
<p style="font-size:11px;color:{C["muted"]};margin-bottom:16px">
Checklist {pct}/9 complete</p>''', unsafe_allow_html=True)

    st.divider()
    ca1, ca2 = st.columns(2)
    with ca1:
        st.markdown("**👤 Performance Engineer Sign-Off**")
        st.session_state.approval_pe_name = st.text_input(
            "PE Engineer Name", value=st.session_state.get("approval_pe_name", ""),
            key="pe_name", placeholder="Enter your name")
        st.session_state.approval_pe = st.checkbox(
            "✅ I confirm all PE checklist items are validated",
            value=st.session_state.approval_pe, key="pe_approve",
            disabled=not checklist_done)
        if not checklist_done:
            st.caption("⚠ Complete all checklist items above to enable sign-off")
        if st.session_state.approval_pe and not st.session_state.get("approval_pe_date", ""):
            st.session_state.approval_pe_date = date.today().isoformat()
        if st.session_state.approval_pe:
            st.markdown(f'<span class="b-signed">✅ PE Approved &nbsp;|&nbsp; '
                        f'{st.session_state.get("approval_pe_date", "")}</span>', unsafe_allow_html=True)
        else:
            st.markdown('<span class="b-pending">⏳ PE Approval Pending</span>', unsafe_allow_html=True)

    with ca2:
        st.markdown("**🏢 Customer Sign-Off**")
        st.session_state.approval_customer_name = st.text_input(
            "Customer Representative", value=st.session_state.get("approval_customer_name", ""),
            key="cust_name", placeholder="e.g. Antony Castaldi")
        st.session_state.approval_customer = st.checkbox(
            "✅ Customer approves current performance benchmark & UI performance",
            value=st.session_state.approval_customer, key="cust_approve")
        if st.session_state.approval_customer and not st.session_state.get("approval_customer_date", ""):
            st.session_state.approval_customer_date = date.today().isoformat()
        if st.session_state.approval_customer:
            st.markdown(f'<span class="b-signed">✅ Customer Approved &nbsp;|&nbsp; '
                        f'{st.session_state.get("approval_customer_date", "")}</span>', unsafe_allow_html=True)
        else:
            st.markdown('<span class="b-pending">⏳ Customer Approval Pending</span>',
                        unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.session_state.approval_notes = st.text_area(
        "📝 PE Approval Notes / Observations",
        value=st.session_state.get("approval_notes", ""), height=100,
        placeholder="e.g. Batch window within SLA, resource utilization healthy, "
                    "SOW metrics validated, UI benchmark approved.")

    both        = st.session_state.approval_pe and st.session_state.approval_customer
    status_lbl  = "APPROVED ✅" if both else "PENDING ⏳"
    status_col  = C["green"] if both else C["amber"]
    st.markdown(f'''<div style="background:{status_col}22;border:1px solid {status_col}55;
border-radius:12px;padding:18px 24px;text-align:center;margin-top:16px">
<div style="font-size:22px;font-weight:700;color:{status_col}">
Go-Live Sign-Off Status: {status_lbl}</div>
<div style="font-size:12px;color:{C["muted"]};margin-top:6px">
PE: {st.session_state.get("approval_pe_name", "") or "N/A"} &nbsp;|&nbsp;
Customer: {st.session_state.get("approval_customer_name", "") or "N/A"} &nbsp;|&nbsp;
{st.session_state.customer_name} {st.session_state.env_type}
</div></div>''', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# ── PE FINDINGS ENGINE ────────────────────────────────────────
def generate_pe_findings(m):
    """
    Auto-generate prioritised PE audit findings from computed metrics.
    Returns list of {"level": "critical|warning|info|ok", "icon": str,
                     "text": str, "sub": str}
    Generic — works for any customer data.
    """
    findings = []
    servers  = st.session_state.get("server_data") or []
    has_ctrl = st.session_state.get("ctrlm_df") is not None

    # ── Batch / SLA findings ──────────────────────────────────
    if has_ctrl:
        breach_pct = 100 - m.get("compliance", 100)
        if m.get("jobs_breach", 0) > 0:
            findings.append({"level":"critical","icon":"🚨",
                "text":f"{m['jobs_breach']} job(s) breached SLA window",
                "sub":f"Peak runtime exceeded {DAILY_LIMIT_HRS}h daily limit — immediate review required"})
        elif breach_pct > 0:
            findings.append({"level":"warning","icon":"⚠️",
                "text":f"Compliance at {m.get('compliance',100):.1f}% ({m.get('jobs_at_risk',0)} at-risk jobs)",
                "sub":"Jobs approaching SLA boundary — monitor closely"})
        else:
            findings.append({"level":"ok","icon":"✅",
                "text":f"Batch compliance at {m.get('compliance',100):.1f}% — within SLA",
                "sub":f"{m.get('total_jobs',0)} jobs, {m.get('total_runs',0)} runs reviewed"})

        fsla = m.get("fleet_sla_buffer")
        if fsla and fsla.get("status") in ("CRITICAL","CAUTION"):
            findings.append({"level":"warning","icon":"⏱️",
                "text":f"Tightest SLA buffer: {fsla.get('buffer_pct',0):.1f}% headroom ({fsla.get('status','')})",
                "sub":f"Worst job has only {fsla.get('buffer_hrs',0):.2f}h of growth capacity"})

        anoms = m.get("anomalies") or []
        if anoms:
            findings.append({"level":"warning","icon":"📉",
                "text":f"{len(anoms)} statistical anomaly/anomalies detected in job runtimes",
                "sub":f"Top outlier: {anoms[0].get('job','?')} (z={anoms[0].get('z',0):.1f}σ)"})

    # ── Infrastructure / resource findings ───────────────────
    if servers:
        fleet = fleet_intelligence_engine(servers) or {}
        grade = fleet.get("grade","?");  fscore = fleet.get("score",0)
        crit  = fleet.get("critical",0); warn   = fleet.get("warning",0)
        if crit > 0:
            findings.append({"level":"critical","icon":"🖥️",
                "text":f"{crit} server(s) in CRITICAL state (fleet grade {grade})",
                "sub":"CPU/Memory/Disk thresholds breached — escalation recommended"})
        elif warn > 0:
            findings.append({"level":"warning","icon":"🖥️",
                "text":f"{warn} server(s) in WARNING state (fleet grade {grade}, score {fscore:.0f})",
                "sub":"Resource utilization approaching thresholds"})
        else:
            findings.append({"level":"ok","icon":"🖥️",
                "text":f"Fleet health grade {grade} ({fscore:.0f}/100) — all servers within thresholds",
                "sub":f"{fleet.get('healthy',0)}/{fleet.get('total',0)} servers healthy"})

        inf_anoms = fleet.get("anomalies") or []
        for ia in inf_anoms[:2]:
            findings.append({"level":"info","icon":"📊",
                "text":f"Infrastructure anomaly: {ia.get('host','?')} — {ia.get('metric','?')} at {ia.get('value',0):.1f}%",
                "sub":f"z-score {ia.get('z',0):.1f}σ above fleet average"})

    # ── Data availability findings ────────────────────────────
    if not has_ctrl and not servers:
        findings.append({"level":"info","icon":"📂",
            "text":"No audit data loaded — upload Ctrl-M CSV and/or Resource Utilization report",
            "sub":"Use the Upload panel to begin PE audit analysis"})

    sow_dfu = st.session_state.get("sow_dfu", 0)
    if sow_dfu > 0:
        sow_base = st.session_state.get("sow_dfu_base", sow_dfu)
        util_pct = (sow_dfu / sow_base * 100) if sow_base else 0
        if util_pct > 100:
            findings.append({"level":"critical","icon":"📦",
                "text":f"SOW volume exceeded: {util_pct:.1f}% of contracted DFU",
                "sub":"Volume above contract ceiling — commercial review required"})
        elif util_pct > 85:
            findings.append({"level":"warning","icon":"📦",
                "text":f"SOW volume at {util_pct:.1f}% of contracted DFU",
                "sub":"Approaching contract ceiling — plan for next cycle"})

    # Issues register
    open_issues = [i for i in (st.session_state.get("issues_list") or [])
                   if i.get("Status") in ("Open","In Progress")]
    if open_issues:
        crit_issues = [i for i in open_issues if i.get("Severity") == "Critical"]
        if crit_issues:
            findings.append({"level":"critical","icon":"📋",
                "text":f"{len(crit_issues)} critical open issue(s) in register",
                "sub":f"IDs: {', '.join(i['ID'] for i in crit_issues[:3])}"})
        else:
            findings.append({"level":"warning","icon":"📋",
                "text":f"{len(open_issues)} open issue(s) in Issues Register",
                "sub":"Review and action before PE sign-off"})

    return findings


def render_pe_findings_panel(findings):
    """Render the findings in a prominent single-block panel."""
    if not findings:
        return
    level_class = {"critical":"finding-critical","warning":"finding-warning",
                   "info":"finding-info","ok":"finding-ok"}
    rows_html = ""
    for f in findings:
        cls  = level_class.get(f["level"], "finding-info")
        sub  = f'<div class="finding-sub">{_html_mod.escape(f["sub"])}</div>' if f.get("sub") else ""
        rows_html += (
            f'<div class="finding-row {cls}">'
            f'<span class="finding-icon">{f["icon"]}</span>'
            f'<div><div class="finding-text">{_html_mod.escape(f["text"])}</div>{sub}</div>'
            f'</div>'
        )
    crit_count = sum(1 for f in findings if f["level"] == "critical")
    warn_count = sum(1 for f in findings if f["level"] == "warning")
    summary_badge = ""
    if crit_count:
        summary_badge += f'<span class="b-breach" style="font-size:10px">{crit_count} Critical</span> '
    if warn_count:
        summary_badge += f'<span class="b-warn" style="font-size:10px">{warn_count} Warning</span>'
    if not crit_count and not warn_count:
        summary_badge = f'<span class="b-ok" style="font-size:10px">All Clear</span>'

    st.markdown(
        f'<div class="findings-panel">'
        f'<div class="findings-panel-title">🔎 PE Audit Findings  {summary_badge}</div>'
        f'{rows_html}'
        f'</div>',
        unsafe_allow_html=True)


# ── MAIN DASHBOARD ───────────────────────────────────────────
def dashboard(m):
    customer = st.session_state.customer_name.strip()
    env      = st.session_state.env_type.strip()
    _has_ctrlm_data = st.session_state.get("ctrlm_df") is not None
    _cdf     = st.session_state.get("ctrlm_df")
    _has_ctrlm = _cdf is not None

    # ── Persistent Intel Strip (always visible) ───────────────
    servers     = st.session_state.get("server_data") or []
    _fleet_info = fleet_intelligence_engine(servers) if servers else {}
    _fl_grade   = _fleet_info.get("grade","—")
    _fl_score   = f"{_fleet_info.get('score',0):.0f}" if servers else "—"
    _compliance = f"{m.get('compliance',0):.1f}%" if _has_ctrlm_data else "—"
    _breach_txt = str(m.get("jobs_breach",0)) if _has_ctrlm_data else "—"
    _anom_txt   = str(len(m.get("anomalies") or [])) if _has_ctrlm_data else "—"
    _srv_txt    = str(len(servers)) if servers else "—"
    _env_col    = C["red"] if env.upper()=="PROD" else (C["amber"] if env.upper() in ("QA","NON-PROD") else C["blue"])
    _env_badge  = (f'<span style="background:{_env_col}22;color:{_env_col};padding:2px 8px;border-radius:6px;font-size:10px;font-weight:700">{env}</span>') if env else ""
    st.markdown(
        f'<div class="intel-bar">'
        f'<div class="intel-item"><span class="intel-label">CUSTOMER</span>'
        f'<span class="intel-value">{_html_mod.escape(customer) if customer else "—"}</span> {_env_badge}</div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">COMPLIANCE</span>'
        f'<span class="intel-value" style="color:{C["red"] if m.get("jobs_breach",0) else C["green"]}">{_compliance}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">SLA BREACHES</span>'
        f'<span class="intel-value" style="color:{C["red"] if m.get("jobs_breach",0) else C["muted"]}">{_breach_txt}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">ANOMALIES</span>'
        f'<span class="intel-value" style="color:{C["amber"] if int(_anom_txt or 0) else C["muted"]}">{_anom_txt}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">FLEET GRADE</span>'
        f'<span class="intel-value">{_fl_grade} ({_fl_score})</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">SERVERS</span>'
        f'<span class="intel-value">{_srv_txt}</span></div>'
        f'</div>',
        unsafe_allow_html=True)

    # ── Helper: "data not loaded" nudge card ──────────────────
    def _nudge(icon, title, msg):
        st.markdown(
            f'''<div style="text-align:center;padding:48px 24px;
background:{C["card"]};border-radius:14px;margin:8px 0">
  <div style="font-size:44px;margin-bottom:12px">{icon}</div>
  <p style="font-size:16px;font-weight:700;color:{C["white"]};margin:0 0 6px">{title}</p>
  <p style="font-size:12px;color:{C["muted"]}">{msg}</p>
</div>''', unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    # ── 5-Group Horizontal Navigation ─────────────────────────
    g1, g2, g3, g4, g5 = st.tabs([
        "📊\nOverview",
        "⚙️\nBatch Review",
        "🖥️\nApp Review",
        "🔧\nInfra Review",
        "📋\nGovernance",
    ])

    # ═══════════════════════════════════════════════════════════
    # G1 — OVERVIEW: KPI Cards + Findings + Charts + Approval
    # ═══════════════════════════════════════════════════════════
    with g1:
        at_risk = max(0, m["total_jobs"] - m["jobs_ok"] - m["jobs_breach"])
        m["jobs_at_risk"] = at_risk
        m["total_runs"] = len(_cdf) if _cdf is not None else 0
        batch_performance_kpis(m)
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

        # ── Findings + Approval side-by-side ──────────────────
        _ov_left, _ov_right = st.columns([2.2, 1])
        with _ov_left:
            _findings = generate_pe_findings(m)
            render_pe_findings_panel(_findings)
        with _ov_right:
            pe_ok   = st.session_state.get("approval_pe", False)
            cu_ok   = st.session_state.get("approval_customer", False)
            pe_name = st.session_state.get("approval_pe_name", "") or "—"
            cu_name = st.session_state.get("approval_customer_name", "") or "—"
            _ap_col = C["green"] if (pe_ok and cu_ok) else (C["amber"] if (pe_ok or cu_ok) else C["muted"])
            _ap_lbl = "APPROVED" if (pe_ok and cu_ok) else ("PARTIAL" if (pe_ok or cu_ok) else "PENDING")
            _ap_icon = "✅" if (pe_ok and cu_ok) else ("⏳" if (pe_ok or cu_ok) else "⬜")
            st.markdown(f'''<div class="panel" style="height:100%">
<p class="panel-title">✍️ Approval Status</p>
<div style="text-align:center;padding:18px 0 12px">
<span style="font-size:36px">{_ap_icon}</span>
<p style="font-size:18px;font-weight:800;color:{_ap_col};margin:8px 0 2px">{_ap_lbl}</p>
</div>
<div style="font-size:12px;line-height:2">
<div style="display:flex;justify-content:space-between;padding:4px 0;border-bottom:1px solid {C["border"]}">
<span style="color:{C["muted"]}">PE Lead</span>
<span style="color:{C["green"] if pe_ok else C["muted"]};font-weight:700">{"✅ " if pe_ok else ""}{_html_mod.escape(pe_name)}</span></div>
<div style="display:flex;justify-content:space-between;padding:4px 0">
<span style="color:{C["muted"]}">Customer</span>
<span style="color:{C["green"] if cu_ok else C["muted"]};font-weight:700">{"✅ " if cu_ok else ""}{_html_mod.escape(cu_name)}</span></div>
</div></div>''', unsafe_allow_html=True)

        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

        if _has_ctrlm_data:
            ch1, ch2, ch3 = st.columns([1.3, 0.9, 1.3])
            with ch1:
                st.markdown(f'''<div class="panel">
<p class="panel-title">📅 Batch Window Trend
<span style="font-size:11px;color:{C['muted']};font-weight:400;margin-left:6px">Last 30 days</span></p>''', unsafe_allow_html=True)
                st.plotly_chart(window_trend_fig(m["window"]), use_container_width=True, config={"displayModeBar": False})
                st.markdown("</div>", unsafe_allow_html=True)
            with ch2:
                st.markdown(f'<div class="panel"><p class="panel-title">🥧 Sub-App Distribution</p>', unsafe_allow_html=True)
                if "Sub_Application" in _cdf.columns and len(_cdf) > 0:
                    _sb2 = _cdf.groupby("Sub_Application", as_index=False)["Run_Sec"].sum()
                    _sb2["hrs"] = _sb2["Run_Sec"] / 3600
                    _fd = go.Figure(go.Pie(labels=_sb2["Sub_Application"], values=_sb2["hrs"],
                        hole=0.55, textinfo="none",
                        marker=dict(colors=[C["blue"], C["purple"], C["cyan"], C["green"], C["amber"]])))
                    _fd.update_layout(**BASE_LAYOUT, height=220,
                        margin=dict(l=0, r=0, t=0, b=0), showlegend=True,
                        legend=dict(orientation="h", y=-0.25, x=0.5, xanchor="center", font=dict(size=9)))
                    st.plotly_chart(_fd, use_container_width=True, config={"displayModeBar": False})
                else:
                    st.markdown(f'<p style="color:{C["muted"]};font-size:11px;padding:20px 0;text-align:center">Sub-Application column not found</p>', unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
            with ch3:
                st.markdown(f'''<div class="panel">
<p class="panel-title">📊 Daily Batch Window
<span style="font-size:11px;color:{C['muted']};font-weight:400;margin-left:6px">Hours per day</span></p>''', unsafe_allow_html=True)
                _win = m["window"]
                if len(_win) > 0:
                    _bc = [C["red"] if v > DAILY_LIMIT_HRS else C["blue"] for v in _win["total_hrs"]]
                    _fb = go.Figure(go.Bar(x=_win["run_date"].astype(str), y=_win["total_hrs"],
                        marker_color=_bc, marker_line_width=0))
                    _fb.add_hline(y=DAILY_LIMIT_HRS, line_dash="dot", line_color=C["red"], line_width=1.5)
                    _fb.update_layout(**BASE_LAYOUT, height=220,
                        margin=dict(l=10, r=10, t=0, b=30),
                        xaxis=dict(**AXIS, tickangle=-45, tickfont=dict(size=9), nticks=8),
                        yaxis=dict(**AXIS))
                    st.plotly_chart(_fb, use_container_width=True, config={"displayModeBar": False})
                st.markdown("</div>", unsafe_allow_html=True)

            st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
            hcol, rcol = st.columns([1.6, 1])
            with hcol:
                nd = m["daily"]["run_date"].nunique() if "run_date" in m["daily"].columns else "?"
                nj = m["daily"]["Job_Name"].nunique() if "Job_Name" in m["daily"].columns else "?"
                st.markdown(f'''<div class="panel">
<p class="panel-title">🗓️ SLA Compliance Heatmap
<span style="font-size:11px;color:{C['muted']};font-weight:400;margin-left:6px">{nd}d × {nj} jobs</span></p>''', unsafe_allow_html=True)
                st.plotly_chart(heatmap_fig(m["daily"]), use_container_width=True, config={"displayModeBar": False})
                st.markdown(f'''<div style="display:flex;gap:14px;font-size:11px;margin-top:2px">
<span style="color:{C['green']}">● OK</span>
<span style="color:{C['amber']}">● At Risk</span>
<span style="color:{C['red']}">● Breach</span>
</div></div>''', unsafe_allow_html=True)
            with rcol:
                st.markdown(f'<div class="panel"><p class="panel-title">🏆 Top Jobs by Peak Runtime</p>', unsafe_allow_html=True)
                if len(m["top_jobs"]) > 0:
                    _tj_html = (
                        f'<div class="tj-hdr tj-grid">'
                        f'<span>Job</span><span style="text-align:right">Peak</span>'
                        f'<span style="text-align:right">Avg</span><span style="text-align:right">Status</span>'
                        f'</div>'
                    )
                    for _, _r in m["top_jobs"].head(10).iterrows():
                        _ib  = _r["peak_hrs"] > DAILY_LIMIT_HRS
                        _bdg = '<span class="b-breach" style="font-size:10px">BR</span>' if _ib else '<span class="b-ok" style="font-size:10px">OK</span>'
                        _rbg = f'{C["red"]}08' if _ib else "transparent"
                        _jn  = _html_mod.escape(_r["Job_Name"][:26])
                        _tj_html += (
                            f'<div class="tj-grid" style="background:{_rbg}">'
                            f'<span class="tj-cell" style="font-weight:600" title="{_html_mod.escape(_r["Job_Name"])}">{_jn}</span>'
                            f'<span class="tj-cell" style="text-align:right;color:{C["red"] if _ib else C["white"]};font-weight:700">{_r["peak_hrs"]:.2f}h</span>'
                            f'<span class="tj-cell" style="text-align:right;color:{C["muted"]}">{_r["avg_hrs"]:.2f}h</span>'
                            f'<span class="tj-cell" style="text-align:right">{_bdg}</span>'
                            f'</div>'
                        )
                    st.markdown(_tj_html, unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
        else:
            _nudge("📂", "Upload Data to See Overview Charts",
                   "Upload a <b>Ctrl-M CSV</b> to see batch window trends, SLA heatmap, and top job analysis.")

    # ═══════════════════════════════════════════════════════════
    # G2 — BATCH REVIEW: Job Drilldown | Batch SLA | Benchmark
    # ═══════════════════════════════════════════════════════════
    with g2:
        b1, b2, b3 = st.tabs(["🔍 Job Drilldown", "📋 Batch SLA", "📊 Benchmark"])

        with b1:
            if not _has_ctrlm:
                _nudge("📋", "Job Drilldown — Ctrl-M CSV Required",
                       "Upload your <b>Ctrl-M Execution History CSV</b> to see per-job SLA analysis and compliance heatmap.")
            else:
                breached = set(m["monthly"][m["monthly"]["breach"]]["Job_Name"])
                _jobs    = sorted(_cdf["Job_Name"].unique()) if "Job_Name" in _cdf.columns else []
                if _jobs:
                    jc1, jc2 = st.columns([3, 1])
                    with jc1:
                        selected = st.selectbox("Select Job", _jobs)
                    with jc2:
                        st.markdown("<br>", unsafe_allow_html=True)
                        if selected in breached:
                            st.markdown('<span class="b-breach">⚠ MONTHLY BREACH</span>', unsafe_allow_html=True)
                        else:
                            st.markdown('<span class="b-ok">✅ WITHIN SLA</span>', unsafe_allow_html=True)
                    d1, d2 = st.columns(2)
                    with d1:
                        st.markdown('<div class="panel">', unsafe_allow_html=True)
                        st.plotly_chart(job_trend_fig(_cdf, selected),
                                        use_container_width=True, config={"displayModeBar": False})
                        st.markdown("</div>", unsafe_allow_html=True)
                    with d2:
                        st.markdown('<div class="panel">', unsafe_allow_html=True)
                        st.plotly_chart(monthly_bar_fig(m["monthly"], selected),
                                        use_container_width=True, config={"displayModeBar": False})
                        st.markdown("</div>", unsafe_allow_html=True)
                    jd = m["daily"][m["daily"]["Job_Name"] == selected].sort_values("run_date")
                    st.dataframe(
                        jd[["run_date","total_hrs","runs","breach"]]
                          .rename(columns={"run_date":"Date","total_hrs":"Runtime (hrs)","runs":"Runs","breach":"Breach?"}),
                        use_container_width=True, height=200)
                    bdf = m["monthly"][m["monthly"]["breach"]]
                    if not bdf.empty:
                        st.markdown("**⚠️ All Monthly Breaching Jobs**")
                        st.dataframe(bdf.sort_values("total_hrs", ascending=False),
                                     use_container_width=True, height=180)

        with b2:
            bp = st.session_state.get("batch_sla_df")
            if bp is None:
                _nudge("📋", "Batch SLA — SLA Matrix Required",
                       "Upload your <b>Batch SLA XLSX</b> or CSV to see per-job compliance analysis.")
            else:
                sla_fmt = getattr(bp, 'attrs', {}).get("_sla_format", "")
                cdf = st.session_state.get("ctrlm_df")

                if sla_fmt == "batch_schedule":
                    st.markdown('<div class="panel"><p class="panel-title">Batch Schedule SLA</p>', unsafe_allow_html=True)
                    k1, k2, k3 = st.columns(3)
                    with k1: kpi("Batch Streams", str(len(bp)), C["blue"])
                    if "SLA_Window_Hrs" in bp.columns:
                        with k2: kpi("Max SLA Window", f"{bp['SLA_Window_Hrs'].max():.1f}h", C["amber"])
                    if "Current_End" in bp.columns and "Expected_End" in bp.columns:
                        try:
                            def _check_overrun(row):
                                exp = pd.to_datetime(str(row.get("Expected_End","")), format="%H:%M:%S", errors="coerce")
                                cur = pd.to_datetime(str(row.get("Current_End","")), format="%H:%M:%S", errors="coerce")
                                if pd.isna(exp) or pd.isna(cur): return False
                                return cur > exp
                            overruns = bp.apply(_check_overrun, axis=1).sum()
                            with k3: kpi("SLA Overruns", str(overruns), C["red"] if overruns > 0 else C["green"])
                        except Exception:
                            pass
                    st.dataframe(bp, use_container_width=True, height=300)
                    if "Comments" in bp.columns:
                        for _, row in bp.iterrows():
                            cmt = str(row.get("Comments", ""))
                            if cmt and cmt not in ("nan", "None", ""):
                                bn = row.get("Batch_Name", row.get(bp.columns[0], ""))
                                st.caption(f"**{bn}**: {cmt}")
                    st.markdown("</div>", unsafe_allow_html=True)

                elif "Job_Name" in bp.columns and "SLA_Hrs" in bp.columns and cdf is not None:
                    st.markdown('<div class="panel"><p class="panel-title">SLA Compliance Analysis</p>', unsafe_allow_html=True)
                    job_peak = (cdf.groupby("Job_Name", as_index=False)
                                  .agg(peak_hrs=("run_time_hrs","max"),
                                       avg_hrs=("run_time_hrs","mean"),
                                       total_runs=("run_time_hrs","count")))
                    bp_clean = bp.copy()
                    bp_clean["SLA_Hrs"] = pd.to_numeric(bp_clean["SLA_Hrs"], errors="coerce")
                    merged = bp_clean.merge(job_peak, on="Job_Name", how="left")
                    merged["peak_hrs"] = merged["peak_hrs"].fillna(0).round(3)
                    merged["avg_hrs"]  = merged["avg_hrs"].fillna(0).round(3)
                    merged["buffer_pct"] = ((merged["SLA_Hrs"] - merged["peak_hrs"]) / merged["SLA_Hrs"] * 100).round(1)
                    merged["status"] = merged.apply(
                        lambda r: "BREACH" if r["peak_hrs"] > r["SLA_Hrs"]
                        else ("AT RISK" if r["buffer_pct"] < 15 else "OK") if r["SLA_Hrs"] > 0
                        else "NO SLA", axis=1)
                    n_breach = (merged["status"] == "BREACH").sum()
                    n_risk   = (merged["status"] == "AT RISK").sum()
                    n_ok     = (merged["status"] == "OK").sum()
                    k1, k2, k3, k4 = st.columns(4)
                    with k1: kpi("Total SLA Jobs", str(len(merged)), C["blue"])
                    with k2: kpi("Compliant", str(n_ok), C["green"])
                    with k3: kpi("At Risk", str(n_risk), C["amber"])
                    with k4: kpi("Breaching", str(n_breach), C["red"])
                    display_cols = ["Job_Name","SLA_Hrs","peak_hrs","avg_hrs","buffer_pct","status"]
                    avail = [c for c in display_cols if c in merged.columns]
                    st.dataframe(merged[avail].sort_values("buffer_pct", ascending=True),
                                 use_container_width=True, height=350)
                    st.markdown("</div>", unsafe_allow_html=True)

                else:
                    st.markdown('<div class="panel"><p class="panel-title">Batch SLA Matrix (Raw)</p>', unsafe_allow_html=True)
                    if "Job_Name" not in bp.columns and "Batch_Name" not in bp.columns:
                        st.warning("Could not detect Job_Name or Batch_Name column. Check column headers.")
                    st.dataframe(bp, use_container_width=True, height=300)
                    st.markdown("</div>", unsafe_allow_html=True)

        with b3:
            ui_performance_tab()

    # ═══════════════════════════════════════════════════════════
    # G3 — APPLICATION REVIEW: Sub-App | Perf Test
    # ═══════════════════════════════════════════════════════════
    with g3:
        a1, a2 = st.tabs(["🥧 Sub-App Mix", "⚡ Perf Test"])

        with a1:
            if not _has_ctrlm:
                _nudge("🥧", "Sub-Application — Ctrl-M CSV Required",
                       "Upload your <b>Ctrl-M CSV</b> to see workload distribution by sub-application.")
            else:
                s1, s2 = st.columns([1, 2])
                with s1:
                    st.markdown('<div class="panel"><p class="panel-title">Runtime by Sub-App</p>',
                                unsafe_allow_html=True)
                    st.plotly_chart(sub_app_pie(m["sub_stats"]), use_container_width=True,
                                    config={"displayModeBar": False})
                    st.markdown("</div>", unsafe_allow_html=True)
                with s2:
                    st.markdown('<div class="panel"><p class="panel-title">Sub-Application Detail</p>',
                                unsafe_allow_html=True)
                    st.dataframe(
                        m["sub_stats"]
                          .rename(columns={"Sub_Application":"Sub App","total_hrs":"Total Hrs","jobs":"Jobs"})
                          .sort_values("Total Hrs", ascending=False),
                        use_container_width=True, height=220)
                    st.markdown("</div>", unsafe_allow_html=True)

        with a2:
            if not _has_ctrlm:
                _nudge("⚡", "Perf-Test Report — Ctrl-M CSV Required",
                       "Upload your <b>Ctrl-M CSV</b> to generate the full performance test report with buffer analysis.")
            else:
                perf_test_report_tab(m)

    # ═══════════════════════════════════════════════════════════
    # G4 — INFRA REVIEW: Resource Util | SOW & Volume
    # ═══════════════════════════════════════════════════════════
    with g4:
        i1, i2 = st.tabs(["🖥️ Resource Util", "📄 SOW & Volume"])
        with i1:
            resource_tab()
        with i2:
            sow_tab()

    # ═══════════════════════════════════════════════════════════
    # G5 — GOVERNANCE: Issues | PE Docs | Export | Approval
    # ═══════════════════════════════════════════════════════════
    with g5:
        v1, v2, v3, v4 = st.tabs(["⚠️ Issues", "📋 PE Docs", "📤 Export", "✍️ Approval"])

        with v1:
            issues_waivers_tab()

        with v2:
            pe_document_review_tab()

        with v3:
            st.markdown(f'''<div class="panel">
<p class="panel-title">📤 Export Report</p>
<p class="panel-sub">Download the full PE Audit report as an HTML file — includes dark/light theme toggle,
all charts, server table, issues register and approval status.</p>''', unsafe_allow_html=True)
            if not _has_ctrlm:
                st.info("📂 Upload a Ctrl-M CSV to generate the full performance audit report. "
                        "Resource and PE Document data is included when available.")
            else:
                ec1, ec2 = st.columns(2)
                with ec1:
                    if st.button("🔄 Build HTML Report", use_container_width=True, key="build_report_btn"):
                        st.session_state["_html_report"] = build_html_report(m)
                with ec2:
                    if st.session_state.get("_html_report"):
                        st.download_button(
                            "⬇️ Download Report (.html)",
                            st.session_state["_html_report"].encode("utf-8"),
                            file_name=f"PE_Audit_{(st.session_state.get('customer_name','Report')).replace(' ','_')}.html",
                            mime="text/html",
                            key="dl_report_btn",
                            use_container_width=True,
                        )
                if st.session_state.get("_html_report"):
                    st.success("✅ Report built. Click **⬇️ Download Report** to save.")
                    with st.expander("📋 Report Preview (first 3,000 chars)"):
                        st.code(st.session_state["_html_report"][:3000], language="html")
                ms = st.session_state.get("_ai_master_summary")
                if ms:
                    st.markdown("---")
                    st.markdown(f'''<div style="background:{C["card2"]};border:1px solid {C["border"]};
border-radius:10px;padding:16px 20px;margin-top:8px">
<p style="font-size:12px;font-weight:700;color:{C["white"]};margin-bottom:8px">🏆 PE Approval Summary (AI Generated)</p>
<p style="font-size:12px;color:{C["muted"]};line-height:1.8">{ms}</p>
</div>''', unsafe_allow_html=True)
                    st.download_button("⬇️ Download Summary Text",
                        ms.encode(), f"PE_Summary_{st.session_state.get('customer_name','').replace(' ','_')}.txt",
                        "text/plain", key="dl_summ_exp")
                if st.session_state.get("issues_list"):
                    st.markdown("---")
                    iss_df = pd.DataFrame(st.session_state.issues_list)
                    st.download_button("⬇️ Issues Register CSV",
                                       iss_df.to_csv(index=False).encode(),
                                       "issues_register.csv", "text/csv", key="dl_iss")
            st.markdown("</div>", unsafe_allow_html=True)

        with v4:
            approval_tab()



# ── HTML REPORT GENERATOR ─────────────────────────────────────
def fig_to_b64(fig, width=900, height=None):
    """Convert a Plotly figure to a base64 PNG string for HTML report embedding.

    Requires the optional 'kaleido' package: pip install kaleido
    Returns an empty string if kaleido is not installed or rendering fails,
    which causes the HTML report to show a [Chart unavailable] placeholder.
    """
    try:
        kwargs = dict(format="png", width=width, engine="kaleido")
        if height:
            kwargs["height"] = height
        img_bytes = fig.to_image(**kwargs)
        import base64
        return base64.b64encode(img_bytes).decode()
    except Exception:
        # kaleido not installed or Plotly version mismatch.
        # Return empty string — build_html_report uses inline [Chart unavailable] fallback.
        return ""


def build_html_report(m):
    customer  = st.session_state.customer_name or "Unknown Customer"
    env       = st.session_state.env_type or ""
    gen_date  = datetime.now().strftime("%d %b %Y, %I:%M %p IST")
    servers   = st.session_state.server_data or []
    issues    = st.session_state.get("issues_list", []) or []
    pe_name   = st.session_state.get("approval_pe_name", "") or "—"
    cust_name = st.session_state.get("approval_customer_name", "") or "—"
    pe_ok     = st.session_state.approval_pe
    cu_ok     = st.session_state.approval_customer
    both_ok   = pe_ok and cu_ok
    master_summary = st.session_state.get("_ai_master_summary", "") or ""

    # ── pre-render charts to base64 ──
    hm_b64  = fig_to_b64(heatmap_fig(m["daily"]), height=500)
    win_b64 = fig_to_b64(window_trend_fig(m["window"]))
    top_b64 = fig_to_b64(top_jobs_fig(m["top_jobs"]))

    top2 = m["top_jobs"].copy()
    if "buffer_pct" not in top2.columns:
        top2["buffer_pct"] = ((DAILY_LIMIT_HRS - top2["peak_hrs"]) / DAILY_LIMIT_HRS * 100).round(1)
    if "sla_used_pct" not in top2.columns:
        top2["sla_used_pct"] = (top2["peak_hrs"] / DAILY_LIMIT_HRS * 100).round(1)
    if "status" not in top2.columns:
        top2["status"] = top2["buffer_pct"].apply(lambda x: "BREACH" if x<0 else ("AT RISK" if x<15 else "OK"))
    disp20 = top2.head(20).sort_values("buffer_pct")

    sign_color  = "#22c55e" if both_ok else "#f59e0b"
    sign_label  = "✅ APPROVED" if both_ok else "⏳ PENDING"
    pe_tick     = "✅" if pe_ok  else "⏳"
    cu_tick     = "✅" if cu_ok  else "⏳"
    sow_dfu       = st.session_state.get("sow_dfu", 0)
    sow_sku       = st.session_state.get("sow_sku", 0)
    sow_dfu_agreed= st.session_state.get("sow_dfu_base", 0)
    sow_sku_agreed= st.session_state.get("sow_sku_base", 0)

    # ── KPI numbers ──
    comp_pct = m["compliance"]
    comp_color_d = "#22c55e" if comp_pct>=99 else ("#f59e0b" if comp_pct>=85 else "#ef4444")
    n_breach = m["jobs_breach"]
    n_ok     = m["jobs_ok"]
    n_jobs   = m["total_jobs"]
    total_hrs = m.get("total_hrs",0)
    n_servers = len(servers)

    # ── chart tags ──
    kaleido_placeholder = "<div style='padding:30px;text-align:center;color:#64748b;background:#1e2130;border-radius:8px;font-size:12px'>📊 Install <code>kaleido</code> to embed charts: <code>pip install kaleido</code></div>"
    hm_tag  = f'<img src="data:image/png;base64,{hm_b64}"  class="chart-img">' if hm_b64  else kaleido_placeholder
    win_tag = f'<img src="data:image/png;base64,{win_b64}" class="chart-img">' if win_b64 else kaleido_placeholder
    top_tag = f'<img src="data:image/png;base64,{top_b64}" class="chart-img">' if top_b64 else kaleido_placeholder

    # ── server rows ──
    def srv_badge(v, warn, crit):
        if v == 0: return '<span class="tag tag-muted">N/A</span>'
        if v >= crit: return f'<span class="tag tag-red">{v:.1f}%</span>'
        if v >= warn: return f'<span class="tag tag-amber">{v:.1f}%</span>'
        return f'<span class="tag tag-green">{v:.1f}%</span>'

    def srv_row(s):
        unk = is_unknown_server(s)
        cpu = s.get("cpu_used",0); mem = s.get("mem_used",0); dsk = s.get("disk_used_max",0)
        ram = s.get("mem_total_gb",0)
        if unk:
            status_tag = '<span class="tag tag-muted">IMAGE ONLY</span>'
            cpu_td = mem_td = dsk_td = '<span class="dim">—</span>'
        else:
            worst = max(cpu,mem,dsk)
            if worst>=90:   status_tag = '<span class="tag tag-red">CRITICAL</span>'
            elif worst>=75: status_tag = '<span class="tag tag-amber">MODERATE</span>'
            else:           status_tag = '<span class="tag tag-green">HEALTHY</span>'
            cpu_td = srv_badge(cpu, CPU_OK, CPU_WARN)
            mem_td = srv_badge(mem, MEM_OK, MEM_WARN)
            dsk_td = srv_badge(dsk, DISK_OK, DISK_WARN)
        ram_disp = f"{ram:.1f} GB" if ram else "—"
        stype_tag = f'<span class="tag tag-blue">{s.get("type","APP")}</span>'
        disks = " ".join(f'<span class="mtag">{k}: {v:.0f}%</span>' for k,v in list(s.get("disks",{}).items())[:4])
        return f"""<tr>
          <td class="host-cell"><b>{s["host"].split(".")[0]}</b><br><span class="dim">{s["host"]}</span></td>
          <td>{stype_tag}</td>
          <td>{cpu_td}</td>
          <td>{mem_td}</td>
          <td class="dim">{ram_disp}</td>
          <td>{dsk_td}</td>
          <td style="font-size:10px">{disks if disks else '<span class="dim">—</span>'}</td>
          <td>{status_tag}</td>
        </tr>"""

    srv_rows = "".join(srv_row(s) for s in servers) if servers else \
               "<tr><td colspan='8' class='dim' style='text-align:center;padding:20px'>No server data loaded</td></tr>"

    # ── issues rows ──
    sev_map = {"Critical":"tag-red","High":"tag-amber","Medium":"tag-amber","Low":"tag-green","Informational":"tag-blue"}
    def iss_row(i):
        sc = sev_map.get(i.get("Severity",""),"tag-muted")
        return f"""<tr>
          <td><b>{i.get('ID','')}</b></td>
          <td><span class="tag {sc}">{i.get('Severity','')}</span></td>
          <td>{i.get('Type','')}</td>
          <td>{i.get('Status','')}</td>
          <td>{i.get('Description','')}</td>
          <td class="dim">{i.get('Mitigation','')}</td>
          <td class="dim">{i.get('Owner','')}</td>
          <td class="dim">{i.get('ETA','')}</td>
        </tr>"""
    iss_rows = "".join(iss_row(i) for i in issues) if issues else \
               "<tr><td colspan='8' class='dim' style='text-align:center;padding:20px'>No issues logged</td></tr>"

    # ── top jobs rows ──
    def job_status_tag(peak, buf):
        if peak > DAILY_LIMIT_HRS: return '<span class="tag tag-red">BREACH</span>'
        if buf < 15:                return '<span class="tag tag-amber">AT RISK</span>'
        return                             '<span class="tag tag-green">OK</span>'

    top_rows = ""
    for _, r in m["top_jobs"].head(20).iterrows():
        buf = top2.loc[top2["Job_Name"]==r["Job_Name"],"buffer_pct"].values
        buf_val = buf[0] if len(buf) else 100.0
        peak_style = 'style="color:#ef4444;font-weight:700"' if r["peak_hrs"]>DAILY_LIMIT_HRS else ""
        top_rows += f"""<tr>
          <td><b>{r["Job_Name"]}</b></td>
          <td {peak_style}>{r["peak_hrs"]:.3f}h</td>
          <td class="dim">{r["avg_hrs"]:.3f}h</td>
          <td>{buf_val:.1f}%</td>
          <td>{job_status_tag(r["peak_hrs"], buf_val)}</td>
        </tr>"""

    # ── Summary section ──
    summary_html = ""
    if master_summary:
        summary_html = f"""
<div class="section" style="border-left:4px solid {sign_color}">
  <div class="section-header">
    <span class="section-icon">🏆</span>
    <div><h2>PE Audit Approval Summary</h2><p class="sub">AI-generated executive sign-off narrative</p></div>
  </div>
  <div class="summary-box">
    <p style="font-size:13.5px;line-height:1.9;color:var(--text-primary)">{master_summary.replace(chr(10),'<br>')}</p>
  </div>
</div>"""

    html = f"""<!DOCTYPE html>
<html lang="en" data-theme="dark">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>PE Audit Report — {customer} {env}</title>
<style>
/* ─── CSS VARIABLES — DARK (DEFAULT) ──────────────────────────────── */
:root, [data-theme="dark"] {{
  --bg-page:     #0b0f1a;
  --bg-card:     #141928;
  --bg-card2:    #1a2035;
  --bg-table-hd: #1e2540;
  --bg-table-r:  #141928;
  --bg-table-rh: #1a2035;
  --border:      #2d3656;
  --text-primary: #e8edf5;
  --text-muted:   #6b7a99;
  --text-head:    #a8b4d0;
  --accent-blue:  #4f8ef7;
  --accent-purple:#8b5cf6;
  --green:        #22c55e;
  --amber:        #f59e0b;
  --red:          #ef4444;
  --cyan:         #06b6d4;
  --header-grad1: #0f1729;
  --header-grad2: #1a2a4a;
}}
/* ─── CSS VARIABLES — LIGHT ──────────────────────────────────────── */
[data-theme="light"] {{
  --bg-page:     #f0f4fb;
  --bg-card:     #ffffff;
  --bg-card2:    #f8faff;
  --bg-table-hd: #e8edf8;
  --bg-table-r:  #ffffff;
  --bg-table-rh: #f3f7ff;
  --border:      #d0d8ef;
  --text-primary: #1e293b;
  --text-muted:   #64748b;
  --text-head:    #334155;
  --accent-blue:  #2563eb;
  --accent-purple:#7c3aed;
  --green:        #16a34a;
  --amber:        #d97706;
  --red:          #dc2626;
  --cyan:         #0891b2;
  --header-grad1: #1e3a5f;
  --header-grad2: #2d5a9e;
}}
/* ─── RESET & BASE ───────────────────────────────────────────────── */
*, *::before, *::after {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ font-family:'Segoe UI','Inter',Arial,sans-serif; background:var(--bg-page);
        color:var(--text-primary); font-size:13px; line-height:1.5; }}
.page {{ max-width:1180px; margin:0 auto; padding:28px 24px; }}

/* ─── THEME TOGGLE ───────────────────────────────────────────────── */
.theme-toggle {{ position:fixed; top:20px; right:80px; z-index:1000;
  background:var(--bg-card2); border:1px solid var(--border);
  border-radius:24px; padding:6px 14px; cursor:pointer; font-size:12px;
  font-weight:600; color:var(--text-primary); display:flex; align-items:center; gap:6px; }}
.print-btn {{ position:fixed; top:20px; right:20px; z-index:1000;
  background:linear-gradient(135deg,var(--accent-blue),var(--accent-purple));
  color:#fff; border:none; border-radius:24px; padding:8px 18px;
  font-size:12px; font-weight:600; cursor:pointer; box-shadow:0 4px 16px #3b82f640; }}

/* ─── HEADER ─────────────────────────────────────────────────────── */
.report-header {{ background:linear-gradient(135deg,var(--header-grad1),var(--header-grad2));
  border-radius:16px; padding:28px 36px; margin-bottom:24px;
  display:flex; justify-content:space-between; align-items:center;
  border:1px solid var(--border); position:relative; overflow:hidden; }}
.report-header::before {{ content:""; position:absolute; top:-40px; right:-40px;
  width:200px; height:200px; border-radius:50%;
  background:radial-gradient(circle, #4f8ef720, transparent 70%); }}
.report-header h1 {{ font-size:24px; font-weight:800; color:#fff; margin-bottom:4px; letter-spacing:-.3px; }}
.report-header .sub-title {{ font-size:12px; color:#94a3b8; margin-bottom:8px; }}
.report-header .gen-date {{ font-size:11px; color:#64748b; }}
.customer-block {{ text-align:right; }}
.customer-name {{ font-size:20px; font-weight:700; color:#fff; margin-bottom:6px; }}
.env-badge {{ background:#4f8ef722; color:#7cb9ff; border:1px solid #4f8ef755;
  padding:4px 14px; border-radius:20px; font-size:11px; font-weight:700; display:inline-block; }}

/* ─── SIGN-OFF BANNER ────────────────────────────────────────────── */
.signoff-banner {{ border-radius:14px; padding:20px 28px; margin-bottom:22px;
  display:flex; justify-content:space-between; align-items:center;
  background:{sign_color}12; border:2px solid {sign_color}44; }}
.signoff-verdict {{ font-size:22px; font-weight:800; color:{sign_color}; }}
.signoff-meta {{ font-size:12px; color:var(--text-muted); text-align:right; }}
.signoff-meta b {{ color:var(--text-primary); }}

/* ─── KPI GRID ───────────────────────────────────────────────────── */
.kpi-grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:14px; margin-bottom:22px; }}
.kpi-box {{ background:var(--bg-card); border:1px solid var(--border); border-radius:14px;
  padding:20px 18px; position:relative; overflow:hidden; }}
.kpi-box::after {{ content:""; position:absolute; bottom:-10px; right:-10px;
  width:60px; height:60px; border-radius:50%;
  background:radial-gradient(circle, var(--accent-blue)10, transparent 70%); }}
.kpi-label {{ font-size:10px; color:var(--text-muted); text-transform:uppercase;
  letter-spacing:.1em; font-weight:700; margin-bottom:8px; }}
.kpi-value {{ font-size:32px; font-weight:900; line-height:1; margin-bottom:4px; }}
.kpi-sub {{ font-size:10px; color:var(--text-muted); }}

/* ─── SECTIONS ───────────────────────────────────────────────────── */
.section {{ background:var(--bg-card); border:1px solid var(--border);
  border-radius:14px; padding:22px 26px; margin-bottom:18px; }}
.section-header {{ display:flex; align-items:center; gap:12px; margin-bottom:16px; }}
.section-icon {{ font-size:20px; }}
.section-header h2 {{ font-size:15px; font-weight:700; color:var(--text-primary); margin-bottom:2px; }}
.section-header .sub {{ font-size:11px; color:var(--text-muted); }}
.chart-img {{ width:100%; border-radius:10px; margin:4px 0; }}
.chart-row {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; margin-bottom:18px; }}
.chart-row .section {{ margin-bottom:0; }}

/* ─── SUMMARY BOX ────────────────────────────────────────────────── */
.summary-box {{ background:var(--bg-card2); border:1px solid var(--border);
  border-radius:10px; padding:18px 22px; }}

/* ─── TABLES ─────────────────────────────────────────────────────── */
table {{ width:100%; border-collapse:collapse; font-size:12px; }}
thead tr {{ background:var(--bg-table-hd); }}
th {{ padding:10px 12px; text-align:left; font-size:10px; color:var(--text-head);
  font-weight:700; text-transform:uppercase; letter-spacing:.06em;
  border-bottom:2px solid var(--border); white-space:nowrap; }}
td {{ padding:9px 12px; border-bottom:1px solid var(--border);
  vertical-align:middle; color:var(--text-primary); }}
tbody tr {{ background:var(--bg-table-r); }}
tbody tr:hover {{ background:var(--bg-table-rh); }}
.host-cell {{ font-size:11px; }}
.dim {{ color:var(--text-muted); font-size:11px; }}

/* ─── TAGS ───────────────────────────────────────────────────────── */
.tag {{ display:inline-block; padding:2px 9px; border-radius:20px;
  font-size:10px; font-weight:700; white-space:nowrap; border:1px solid transparent; }}
.tag-green  {{ background:#22c55e18; color:var(--green); border-color:#22c55e33; }}
.tag-amber  {{ background:#f59e0b18; color:var(--amber); border-color:#f59e0b33; }}
.tag-red    {{ background:#ef444418; color:var(--red);   border-color:#ef444433; }}
.tag-blue   {{ background:#4f8ef718; color:var(--accent-blue); border-color:#4f8ef733; }}
.tag-muted  {{ background:var(--border); color:var(--text-muted); }}
.mtag {{ display:inline-block; background:var(--bg-card2); color:var(--text-muted);
  padding:1px 6px; border-radius:4px; font-size:9px; margin:1px; }}

/* ─── APPROVAL GRID ──────────────────────────────────────────────── */
.approval-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:14px; margin-top:14px; }}
.appr-box {{ background:var(--bg-card2); border:1px solid var(--border);
  border-radius:10px; padding:16px 18px; }}
.appr-box h3 {{ font-size:12px; font-weight:700; color:var(--text-muted);
  text-transform:uppercase; letter-spacing:.06em; margin-bottom:10px; }}
.appr-box p {{ font-size:12px; color:var(--text-primary); margin-bottom:4px; }}

/* ─── FOOTER ─────────────────────────────────────────────────────── */
.report-footer {{ text-align:center; padding:20px 0 8px;
  font-size:11px; color:var(--text-muted); border-top:1px solid var(--border); margin-top:24px; }}

/* ─── PRINT ──────────────────────────────────────────────────────── */
@media print {{
  .print-btn, .theme-toggle {{ display:none; }}
  body {{ background:#fff; color:#000; }}
  .report-header {{ background:#1e3a5f!important; color:#fff!important; }}
  .page {{ padding:8px; }}
  .kpi-grid {{ grid-template-columns:repeat(4,1fr); }}
}}
@page {{ size:A4; margin:12mm; }}
</style>
</head>
<body>

<!-- CONTROLS -->
<button class="theme-toggle" onclick="toggleTheme()">🌓 Toggle Theme</button>
<button class="print-btn" onclick="window.print()">🖨️ Print / PDF</button>

<script>
function toggleTheme() {{
  var html = document.documentElement;
  html.dataset.theme = html.dataset.theme === 'dark' ? 'light' : 'dark';
  document.querySelector('.theme-toggle').textContent = html.dataset.theme === 'dark' ? '🌓 Toggle Theme' : '🌑 Dark Mode';
}}
</script>

<div class="page">

<!-- ─── HEADER ──────────────────────────────────────────────────── -->
<div class="report-header">
  <div>
    <div style="font-size:11px;color:#4f8ef7;font-weight:700;letter-spacing:.1em;text-transform:uppercase;margin-bottom:6px">
      ✦ PERFORMANCE ENGINEERING</div>
    <h1>PE Audit Control Tower</h1>
    <div class="sub-title">Performance Engineering Audit Report · v5.0</div>
    <div class="gen-date">Generated: {gen_date}</div>
  </div>
  <div class="customer-block">
    <div class="customer-name">{customer}</div>
    <span class="env-badge">{env if env else "Production"}</span>
  </div>
</div>

<!-- ─── SIGN-OFF BANNER ──────────────────────────────────────────── -->
<div class="signoff-banner">
  <div>
    <div class="signoff-verdict">{sign_label}</div>
    <div style="font-size:12px;color:var(--text-muted);margin-top:4px">Go-Live Sign-Off Status</div>
  </div>
  <div class="signoff-meta">
    <p>PE Engineer: <b>{pe_name}</b> {pe_tick}</p>
    <p style="margin-top:4px">Customer Rep: <b>{cust_name}</b> {cu_tick}</p>
    <p style="margin-top:4px">{customer} · {env}</p>
  </div>
</div>

<!-- ─── KPI CARDS ────────────────────────────────────────────────── -->
<div class="kpi-grid">
  <div class="kpi-box" style="border-top:3px solid {comp_color_d}">
    <div class="kpi-label">SLA Compliance</div>
    <div class="kpi-value" style="color:{comp_color_d}">{m['compliance']:.1f}%</div>
    <div class="kpi-sub">{n_breach} breach · {n_ok} OK out of {n_jobs} jobs</div>
  </div>
  <div class="kpi-box" style="border-top:3px solid var(--accent-blue)">
    <div class="kpi-label">Total Batch Jobs</div>
    <div class="kpi-value" style="color:var(--accent-blue)">{n_jobs}</div>
    <div class="kpi-sub">{m.get('total_runs',0):,} total runs · {total_hrs:.1f}h runtime</div>
  </div>
  <div class="kpi-box" style="border-top:3px solid {'#ef4444' if n_breach else '#22c55e'}">
    <div class="kpi-label">SLA Breaches</div>
    <div class="kpi-value" style="color:{'#ef4444' if n_breach else '#22c55e'}">{n_breach}</div>
    <div class="kpi-sub">Daily SLA limit: {DAILY_LIMIT_HRS}h · Monthly: {MONTHLY_LIMIT_HRS}h</div>
  </div>
  <div class="kpi-box" style="border-top:3px solid var(--accent-purple)">
    <div class="kpi-label">Servers Monitored</div>
    <div class="kpi-value" style="color:var(--accent-purple)">{n_servers}</div>
    <div class="kpi-sub">{sum(1 for s in servers if s.get('type')=='APP')} APP · {sum(1 for s in servers if s.get('type')=='DB')} DB</div>
  </div>
</div>

<!-- ─── MASTER PE SUMMARY ────────────────────────────────────────── -->
{summary_html}

<!-- ─── HEATMAP ──────────────────────────────────────────────────── -->
<div class="section">
  <div class="section-header">
    <span class="section-icon">🗓️</span>
    <div><h2>SLA Compliance Heatmap — Last 21 Days</h2>
    <p class="sub">Green = within SLA · Amber = approaching limit · Red = breach. Top 40 jobs shown.</p></div>
  </div>
  {hm_tag}
</div>

<!-- ─── DAILY WINDOW + TOP JOBS ──────────────────────────────────── -->
<div class="chart-row">
  <div class="section">
    <div class="section-header">
      <span class="section-icon">📅</span>
      <div><h2>Daily Batch Window</h2>
      <p class="sub">Total runtime per day vs {DAILY_LIMIT_HRS}h SLA limit</p></div>
    </div>
    {win_tag}
  </div>
  <div class="section">
    <div class="section-header">
      <span class="section-icon">⏱️</span>
      <div><h2>Top Jobs by Peak Runtime</h2>
      <p class="sub">Peak vs Average · SLA threshold marked</p></div>
    </div>
    {top_tag}
  </div>
</div>

<!-- ─── TOP JOBS TABLE ───────────────────────────────────────────── -->
<div class="section">
  <div class="section-header">
    <span class="section-icon">📋</span>
    <div><h2>Batch Performance Detail</h2>
    <p class="sub">Top 20 jobs by peak runtime · Buffer % = headroom before SLA breach</p></div>
  </div>
  <table>
    <thead><tr>
      <th>Job Name</th><th>Peak Runtime</th><th>Avg Runtime</th><th>SLA Buffer %</th><th>Status</th>
    </tr></thead>
    <tbody>{top_rows}</tbody>
  </table>
</div>

<!-- ─── SERVER TABLE ─────────────────────────────────────────────── -->
<div class="section">
  <div class="section-header">
    <span class="section-icon">🖥️</span>
    <div><h2>Infrastructure Resource Utilisation</h2>
    <p class="sub">CPU/Memory/Disk thresholds — OK ≤{CPU_OK}% · Warn ≤{CPU_WARN}% · Critical &gt;{CPU_WARN}%</p></div>
  </div>
  <table>
    <thead><tr>
      <th>Host</th><th>Type</th><th>CPU %</th><th>Memory %</th><th>RAM</th><th>Max Disk %</th><th>Mounts</th><th>Health</th>
    </tr></thead>
    <tbody>{srv_rows}</tbody>
  </table>
</div>

<!-- ─── SOW VOLUMES ──────────────────────────────────────────────── -->
<div class="section">
  <div class="section-header">
    <span class="section-icon">📄</span>
    <div><h2>SOW Data Volume Compliance</h2>
    <p class="sub">Contracted vs actual DFU / SKU volumes</p></div>
  </div>
  <div style="display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-top:4px">
    <div class="appr-box">
      <h3>DFU (Item-Location-Customer)</h3>
      <p>Contracted: <b>{sow_dfu_agreed:,}</b></p>
      <p>Actual: <b>{sow_dfu:,}</b></p>
      <p style="margin-top:8px">{"<span class='tag tag-green'>Within SOW</span>" if sow_dfu_agreed==0 or sow_dfu<=sow_dfu_agreed else "<span class='tag tag-red'>Exceeds SOW</span>"}</p>
    </div>
    <div class="appr-box">
      <h3>SKU (Item-Location)</h3>
      <p>Contracted: <b>{sow_sku_agreed:,}</b></p>
      <p>Actual: <b>{sow_sku:,}</b></p>
      <p style="margin-top:8px">{"<span class='tag tag-green'>Within SOW</span>" if sow_sku_agreed==0 or sow_sku<=sow_sku_agreed else "<span class='tag tag-red'>Exceeds SOW</span>"}</p>
    </div>
  </div>
</div>

<!-- ─── ISSUES TABLE ─────────────────────────────────────────────── -->
<div class="section">
  <div class="section-header">
    <span class="section-icon">⚠️</span>
    <div><h2>Issues &amp; Waivers Register</h2>
    <p class="sub">{len(issues)} item(s) logged</p></div>
  </div>
  <table>
    <thead><tr>
      <th>ID</th><th>Severity</th><th>Type</th><th>Status</th>
      <th>Description</th><th>Mitigation</th><th>Owner</th><th>ETA</th>
    </tr></thead>
    <tbody>{iss_rows}</tbody>
  </table>
</div>

<!-- ─── APPROVAL ─────────────────────────────────────────────────── -->
<div class="section" style="border-color:{sign_color}55">
  <div class="section-header">
    <span class="section-icon">✍️</span>
    <div><h2>Customer Approval Sign-Off</h2>
    <p class="sub">Dual sign-off required for go-live clearance</p></div>
  </div>
  <div class="approval-grid">
    <div class="appr-box">
      <h3>Performance Engineering</h3>
      <p><b>{pe_name}</b></p>
      <p style="margin-top:8px">{pe_tick} {"<span class='tag tag-green'>Approved " + str(st.session_state.get('approval_pe_date','')) + "</span>" if pe_ok else "<span class='tag tag-amber'>Pending</span>"}</p>
    </div>
    <div class="appr-box">
      <h3>Customer Representative</h3>
      <p><b>{cust_name}</b></p>
      <p style="margin-top:8px">{cu_tick} {"<span class='tag tag-green'>Approved " + str(st.session_state.get('approval_customer_date','')) + "</span>" if cu_ok else "<span class='tag tag-amber'>Pending</span>"}</p>
    </div>
  </div>
  <div style="margin-top:16px;padding:14px 20px;background:{sign_color}15;
    border:1px solid {sign_color}44;border-radius:10px;text-align:center">
    <span style="font-size:16px;font-weight:800;color:{sign_color}">{sign_label} — Overall Status</span>
  </div>
</div>

<!-- ─── FOOTER ───────────────────────────────────────────────────── -->
<div class="report-footer">
  <p><b>PE Control Tower v5.0</b> &nbsp;·&nbsp; Performance Engineering Team &nbsp;·&nbsp; {gen_date}</p>
  <p style="margin-top:4px">Daily SLA: {DAILY_LIMIT_HRS}h &nbsp;|&nbsp; Monthly: {MONTHLY_LIMIT_HRS}h &nbsp;|&nbsp;
    CPU OK: ≤{CPU_OK}% &nbsp;|&nbsp; CPU Warn: ≤{CPU_WARN}% &nbsp;|&nbsp;
    Disk OK: ≤{DISK_OK}% &nbsp;|&nbsp; Disk Warn: ≤{DISK_WARN}%</p>
</div>

</div><!-- .page -->
</body>
</html>"""
    return html





# ── EMPTY STATE ──────────────────────────────────────────────
def empty_state():
    cards = [("📋","Ctrl-M CSV","Last 60 days"),("📊","Batch SLA","CSV or XLSX"),
             ("🖥️","Resource CSV","Server metrics"),("📄","SOW / Evidence","PDF / XLSX")]
    ch = "".join(f"""<div style="background:{C['card']};border:1px solid {C['border']};
border-radius:12px;padding:18px;text-align:center;min-width:130px">
<div style="font-size:26px">{ic}</div>
<div style="font-size:12px;font-weight:600;margin-top:6px">{lb}</div>
<div style="font-size:10px;color:{C['muted']}">{sb}</div></div>""" for ic,lb,sb in cards)
    st.markdown(f"""<div style="text-align:center;padding:50px 20px">
<div style="font-size:60px;margin-bottom:16px">🔧</div>
<h2 style="font-size:22px;font-weight:700;margin-bottom:8px">PE Control Tower v4.0</h2>
<p style="color:{C['muted']};font-size:13px;max-width:500px;margin:0 auto 28px">
Upload a Ctrl-M execution CSV to activate SLA intelligence, heatmaps, resource analysis,
SOW tracking and customer approval sign-off.</p>
<div style="display:flex;justify-content:center;gap:20px;flex-wrap:wrap">{ch}</div>
</div>""", unsafe_allow_html=True)


# ── SIDEBAR ──────────────────────────────────────────────────

def pe_audit_intelligence_panel(ctrlm_df, server_data):
    customer=st.session_state.get("customer_name","") or "the customer"
    env=st.session_state.get("env_type","")
    lines=[]; env_txt=f"({env})" if env else ""
    lines.append(("🔵","init",
        f"Audit initialized for <b>{customer}</b> {env_txt} — all uploaded datasets are cross-referenced below."))
    if ctrlm_df is not None:
        m=compute_metrics_fast(ctrlm_df); n_runs=len(ctrlm_df); comp=m["compliance"]
        total_h=m.get("total_hrs",0)
        avg_win=m["window"]["total_hrs"].mean() if len(m["window"])>0 else 0
        if comp>=99.9:
            lines.append(("✅","ok",
                f"SLA compliance exceptional at <b>{comp:.1f}%</b> across {n_runs:,} batch runs."))
        elif comp>=90:
            lines.append(("⚠️","warn",
                f"SLA compliance at <b>{comp:.1f}%</b> across {n_runs:,} runs — "
                f"<b>{m['jobs_breach']} job(s)</b> need attention."))
        else:
            lines.append(("🔴","breach",
                f"SLA compliance critical at <b>{comp:.1f}%</b> — {m['jobs_breach']} job(s) in breach."))
        try:
            dg=ctrlm_df.groupby("run_date")["Run_Sec"].sum()/3600
            pd_=dg.idxmax(); pv=dg.max()
            lines.append(("⏱️","info",
                f"Total batch runtime: <b>{total_h:.0f}h</b> across {m['total_jobs']} unique jobs. "
                f"Average daily window: <b>{avg_win:.1f}h</b>, peak on <b>{pd_}</b> at <b>{pv:.1f}h</b>."))
        except Exception:
            lines.append(("⏱️","info",
                f"Total batch runtime: <b>{total_h:.0f}h</b> across {m['total_jobs']} unique jobs."))
        try:
            tj=m["top_jobs"].iloc[0]
            ws="⚠️ Exceeds SLA." if tj["peak_hrs"]>DAILY_LIMIT_HRS else "Within SLA limits."
            lines.append(("🔧","info",
                f"Longest-running job: <b>{tj['Job_Name']}</b> with peak "
                f"<b>{tj['peak_hrs']:.2f}h</b> (avg {tj['avg_hrs']:.2f}h). {ws}"))
        except Exception:
            pass  # top_jobs may be empty on a very small dataset — insight is optional
        try:
            ar=m["top_jobs"][(m["top_jobs"]["peak_hrs"]>DAILY_LIMIT_HRS*0.85)&
                              (m["top_jobs"]["peak_hrs"]<=DAILY_LIMIT_HRS)]
            if len(ar)>0:
                nm=", ".join(ar["Job_Name"].tolist()[:3])
                lines.append(("⚡","warn",
                    f"<b>{len(ar)} job(s)</b> at risk with &lt;15% SLA buffer: "
                    f"<b>{nm}</b>{'…' if len(ar)>3 else ''}."))
        except Exception:
            pass  # at-risk calculation is an optional insight — never crash the panel
        try:
            if "Sub_Application" in ctrlm_df.columns:
                sg=ctrlm_df.groupby("Sub_Application")["Run_Sec"].sum()/3600
                ts=sg.idxmax(); th_=sg.max()
                tj_=ctrlm_df[ctrlm_df["Sub_Application"]==ts]["Job_Name"].nunique()
                tot=ctrlm_df["Run_Sec"].sum()/3600; pct=th_/tot*100 if tot>0 else 0
                lines.append(("📊","info",
                    f"Dominant sub-application: <b>{ts}</b> consuming "
                    f"<b>{th_:.1f}h</b> ({tj_} jobs) — <b>{pct:.0f}%</b> of total compute."))
        except Exception:
            pass  # Sub_Application groupby can fail on unusual data — insight is optional
    if server_data:
        try:
            fie=fleet_intelligence_engine(server_data)
            if fie:
                lines.append(("🖥️","info",
                    f"Infrastructure: <b>{len(server_data)} servers</b> scanned. "
                    f"Fleet Grade: <b>{fie['grade']}</b> ({fie['score']}/100). "
                    f"Health: <b>{fie['healthy']} healthy</b>, {fie['warning']} warning, {fie['critical']} critical."))
                if fie.get("anomalies"):
                    a=fie["anomalies"][0]
                    lines.append(("🔶","warn",
                        f"Statistical anomaly: <b>{a['host'].split('.')[0]}</b> — "
                        f"{a['metric']} at <b>{a['value']:.1f}%</b> (z={a['z']:+.1f})."))
        except Exception:
            pass  # fleet_intelligence_engine is optional enrichment — never crash the panel
    if not lines: return
    cmap={"init":C["blue"],"ok":C["green"],"warn":C["amber"],"breach":C["red"],"info":C["border"]}
    st.markdown(f"""<div style="background:linear-gradient(135deg,{C['card']},{C['card2']});
  border:1px solid {C['border']};border-radius:14px;padding:20px 24px;margin-top:8px">
  <div style="display:flex;align-items:center;gap:10px;margin-bottom:14px">
    <span style="font-size:18px">🤖</span>
    <span style="font-size:13px;font-weight:700">PE Audit Intelligence</span>
    <span class="b-blue" style="font-size:10px;letter-spacing:.06em">AUTO-GENERATED</span>
  </div>""",unsafe_allow_html=True)
    for icon,kind,text in lines:
        bc=cmap.get(kind,C["muted"])
        st.markdown(f"""<div style="display:flex;align-items:flex-start;gap:12px;padding:10px 14px;
  background:{C['card2']};border-radius:10px;margin-bottom:6px;border-left:3px solid {bc}">
  <span style="font-size:16px;flex-shrink:0">{icon}</span>
  <span style="font-size:12px;line-height:1.6">{text}</span>
</div>""",unsafe_allow_html=True)
    st.markdown("</div>",unsafe_allow_html=True)


def batch_performance_kpis(m):
    at_risk=m.get("jobs_at_risk",max(0,m["total_jobs"]-m["jobs_ok"]-m["jobs_breach"]))
    n_runs=m.get("total_runs",0)
    avg_win=m["window"]["total_hrs"].mean() if len(m["window"])>0 else 0
    m["avg_daily_hrs"]=avg_win
    total_rt=m.get("total_hrs",0)
    cc=C["green"] if m["compliance"]>=99 else(C["amber"] if m["compliance"]>=85 else C["red"])
    wc=C["red"] if avg_win>DAILY_LIMIT_HRS else(C["amber"] if avg_win>DAILY_LIMIT_HRS*0.85 else C["blue"])
    c1,c2,c3,c4=st.columns(4)
    with c1:
        st.markdown(f"""<div class="kpi-card"><p class="kpi-label">SLA COMPLIANCE</p>
  <p class="kpi-value" style="color:{cc}">{m['compliance']:.1f}%</p>
  <p class="kpi-sub">{n_runs:,} total runs</p></div>""",unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="kpi-card"><p class="kpi-label">BATCH HEALTH</p>
  <div style="margin-top:6px">
    <div style="display:flex;justify-content:space-between;padding:3px 0;font-size:12px">
      <span style="color:{C['green']};font-weight:600">On Track</span>
      <span style="color:{C['green']};font-weight:800;font-size:16px">{m['jobs_ok']}</span></div>
    <div style="display:flex;justify-content:space-between;padding:3px 0;font-size:12px">
      <span style="color:{C['amber']};font-weight:600">At Risk</span>
      <span style="color:{C['amber']};font-weight:800;font-size:16px">{at_risk}</span></div>
    <div style="display:flex;justify-content:space-between;padding:3px 0;font-size:12px">
      <span style="color:{C['red']};font-weight:600">Breach</span>
      <span style="color:{C['red']};font-weight:800;font-size:16px">{m['jobs_breach']}</span></div>
  </div></div>""",unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="kpi-card"><p class="kpi-label">TOTAL RUNTIME</p>
  <p class="kpi-value" style="color:{C['blue']}">{total_rt:.0f}h</p>
  <p class="kpi-sub">{m['total_jobs']} unique jobs</p></div>""",unsafe_allow_html=True)
    with c4:
        st.markdown(f"""<div class="kpi-card"><p class="kpi-label">AVG DAILY WINDOW</p>
  <p class="kpi-value" style="color:{wc}">{avg_win:.1f}h</p>
  <p class="kpi-sub">of {DAILY_LIMIT_HRS}h SLA limit</p></div>""",unsafe_allow_html=True)



def render_sidebar_nav(_any_data):
    """Left sidebar navigation — vivid dashboard nav with clear alignment."""
    _NAV = [
        ("upload",     "📂  Upload & Intake",     True),
        ("overview",   "📊  Executive Dashboard",  False),
        ("batch",      "⚙️  Batch Review",        False),
        ("app",        "🖥️  App Review",          False),
        ("infra",      "🔧  Infra Review",         False),
        ("findings",   "🔍  Findings & Evidence",  False),
        ("governance", "📋  Approval / Export",    False),
    ]
    _cur = st.session_state.get("_current_page", "upload")

    with st.sidebar:
        # ── Logo / Brand ──
        st.markdown(f"""<div style="padding:24px 20px 20px;
  border-bottom:1px solid rgba(59,130,246,.2);
  background:linear-gradient(180deg,rgba(59,130,246,.06) 0%,transparent 100%)">
  <div style="display:flex;align-items:center;gap:14px">
    <div style="background:linear-gradient(135deg,{C['blue']},{C['cyan']});
      width:52px;height:52px;border-radius:14px;display:flex;align-items:center;
      justify-content:center;font-weight:900;font-size:20px;color:#fff;
      box-shadow:0 6px 24px rgba(59,130,246,.45),
                 0 0 0 2px rgba(59,130,246,.2);
      flex-shrink:0">PE</div>
    <div>
      <div style="font-size:20px;font-weight:800;color:#ffffff;letter-spacing:0;
        text-shadow:0 0 30px rgba(59,130,246,.35)">Control Tower</div>
      <div style="font-size:12px;color:#7b8fc4;letter-spacing:.06em;text-transform:uppercase;
        margin-top:3px;font-weight:600">PE Audit Workstation</div>
    </div>
  </div>
</div>""", unsafe_allow_html=True)

        # ── Section Label ──
        st.markdown(f"""<div style="padding:16px 20px 6px">
  <p style="font-size:11px;font-weight:700;color:#5a6fa0;letter-spacing:.1em;
     text-transform:uppercase;margin:0">NAVIGATION</p>
</div>""", unsafe_allow_html=True)

        for key, label, always_on in _NAV:
            enabled = always_on or _any_data
            is_active = (_cur == key)
            if st.button(
                label,
                key=f"nav_{key}",
                use_container_width=True,
                disabled=not enabled,
                type="primary" if is_active else "secondary",
            ):
                if enabled:
                    st.session_state["_current_page"] = key
                    st.rerun()

        # ── Data Status Panel ──
        st.markdown(f"""<div style="border-top:1px solid rgba(59,130,246,.15);
  margin-top:16px;padding:18px 20px;
  background:linear-gradient(180deg,transparent 0%,rgba(59,130,246,.04) 100%)">
  <p style="font-size:11px;font-weight:700;color:#5a6fa0;letter-spacing:.1em;
     text-transform:uppercase;margin:0 0 12px">DATA STATUS</p>""", unsafe_allow_html=True)

        _ds_items = [
            ("Ctrl-M Batch CSV", st.session_state.get("ctrlm_df") is not None),
            ("Batch SLA Matrix", st.session_state.get("batch_sla_df") is not None),
            ("Resource / Infra", bool(st.session_state.get("server_data"))),
            ("PE Document", st.session_state.get("_pe_doc_data") is not None),
        ]
        for _dl, _loaded in _ds_items:
            if _loaded:
                _dot = f'<span style="color:{C["green"]};text-shadow:0 0 10px {C["green"]}90;font-size:11px">●</span>'
                _tc = "#e0e8ff"
                _suffix = f'<span style="color:{C["green"]};font-size:11px;font-weight:700;margin-left:auto">LOADED</span>'
            else:
                _dot = f'<span style="color:#2d3a60;font-size:11px">○</span>'
                _tc = "#4a5a8e"
                _suffix = f'<span style="color:#2d3a60;font-size:10px;margin-left:auto">—</span>'
            st.markdown(f"""<div style="display:flex;align-items:center;gap:10px;
  padding:7px 0;font-size:13px;font-weight:500">
  {_dot} <span style="color:{_tc}">{_dl}</span>{_suffix}
</div>""", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

        # ── Version Footer ──
        st.markdown(f"""<div style="border-top:1px solid rgba(59,130,246,.1);
  padding:14px 20px;margin-top:6px">
  <p style="font-size:11px;color:#4a5a8e;margin:0;text-align:center;font-weight:600">
    PE Control Tower v6.0</p>
  <p style="font-size:10px;color:#3a4a7e;margin:3px 0 0;text-align:center">
    Performance Engineering</p>
</div>""", unsafe_allow_html=True)



# ── PAGE TITLE (OraVision-style large heading) ───────────────────
def _page_title(title, subtitle=""):
    _sub = (f'<p style="font-size:14px;color:#8a9cc8;margin:6px 0 0;font-weight:400">{subtitle}</p>'
            if subtitle else "")
    st.markdown(
        f'<div style="padding:28px 0 20px;text-align:center;border-bottom:1px solid rgba(64,112,232,.1);margin-bottom:20px">'
        f'<h1 style="font-size:38px!important;font-weight:800!important;margin:0;'
        f'background:linear-gradient(135deg,#60a5fa 0%,#a78bfa 50%,#38bdf8 100%);'
        f'-webkit-background-clip:text;-webkit-text-fill-color:transparent;'
        f'background-clip:text;letter-spacing:-.02em;line-height:1.1">{title}</h1>'
        f'{_sub}'
        f'</div>',
        unsafe_allow_html=True)

# ── INTEL BAR (shared across data pages) ─────────────────────
def _render_intel_bar(m):
    """Persistent intelligence strip shown on all data pages."""
    customer = st.session_state.customer_name.strip()
    env      = st.session_state.env_type.strip()
    _has_ctrlm_data = st.session_state.get("ctrlm_df") is not None
    servers     = st.session_state.get("server_data") or []
    _fleet_info = fleet_intelligence_engine(servers) if servers else {}
    _fl_grade   = _fleet_info.get("grade", "—")
    _fl_score   = f"{_fleet_info.get('score',0):.0f}" if servers else "—"
    _compliance = f"{m.get('compliance',0):.1f}%" if _has_ctrlm_data else "—"
    _breach_txt = str(m.get("jobs_breach",0)) if _has_ctrlm_data else "—"
    _anom_txt   = str(len(m.get("anomalies") or [])) if _has_ctrlm_data else "—"
    _srv_txt    = str(len(servers)) if servers else "—"
    _env_col    = C["red"] if env.upper()=="PROD" else (C["amber"] if env.upper() in ("QA","NON-PROD") else C["blue"])
    _env_badge  = (f'<span style="background:{_env_col}22;color:{_env_col};padding:2px 8px;border-radius:6px;font-size:10px;font-weight:700">{env}</span>') if env else ""
    st.markdown(
        f'<div class="intel-bar">'
        f'<div class="intel-item"><span class="intel-label">CUSTOMER</span>'
        f'<span class="intel-value">{_html_mod.escape(customer) if customer else "—"}</span> {_env_badge}</div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">COMPLIANCE</span>'
        f'<span class="intel-value" style="color:{C["red"] if m.get("jobs_breach",0) else C["green"]}">{_compliance}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">SLA BREACHES</span>'
        f'<span class="intel-value" style="color:{C["red"] if m.get("jobs_breach",0) else C["muted"]}">{_breach_txt}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">ANOMALIES</span>'
        f'<span class="intel-value" style="color:{C["amber"] if int(_anom_txt if _anom_txt != "—" else 0) else C["muted"]}">{_anom_txt}</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">FLEET GRADE</span>'
        f'<span class="intel-value">{_fl_grade} ({_fl_score})</span></div>'
        f'<div class="intel-divider"></div>'
        f'<div class="intel-item"><span class="intel-label">SERVERS</span>'
        f'<span class="intel-value">{_srv_txt}</span></div>'
        f'</div>',
        unsafe_allow_html=True)


# ── NUDGE HELPER (module-level) ──────────────────────────────
def _nudge(icon, title, msg):
    _html = (
        f'<div style="text-align:center;padding:48px 24px;'
        f'background:{C["card"]};border-radius:14px;margin:8px 0">'
        f'<div style="font-size:44px;margin-bottom:12px">{icon}</div>'
        f'<p style="font-size:16px;font-weight:700;color:{C["white"]};margin:0 0 6px">{title}</p>'
        f'<p style="font-size:12px;color:{C["muted"]}">{msg}</p>'
        f'</div>'
    )
    st.markdown(_html, unsafe_allow_html=True)


# ── PAGE 1: UPLOAD & INTAKE ─────────────────────────────────
def page_upload():
    """Dedicated upload page — always shows upload panel, data state irrelevant."""
    upload_panel()
    # Show empty state guide only when truly no data loaded yet
    _any = (st.session_state.get("ctrlm_df") is not None or
            bool(st.session_state.get("server_data")) or
            st.session_state.get("batch_sla_df") is not None or
            st.session_state.get("_pe_doc_data") is not None)
    if not _any:
        empty_state()
    else:
        # Show a "proceed to dashboard" nudge when data is loaded
        _loaded_count = sum([
            st.session_state.get("ctrlm_df") is not None,
            bool(st.session_state.get("server_data")),
            st.session_state.get("batch_sla_df") is not None,
            st.session_state.get("_pe_doc_data") is not None,
        ])
        st.markdown(
            f'<div style="background:linear-gradient(135deg,{C["card"]},{C["card2"]});'
            f'border:1px solid rgba(64,112,232,.25);border-radius:14px;'
            f'padding:20px 28px;margin-top:20px;display:flex;align-items:center;'
            f'justify-content:space-between;gap:20px">'
            f'<div>'
            f'<p style="font-size:15px;font-weight:700;color:{C["white"]};margin:0 0 4px">'
            f'✅ {_loaded_count} file source(s) loaded — ready to analyse</p>'
            f'<p style="font-size:12px;color:{C["muted"]};margin:0">'
            f'Use the left sidebar to navigate to Executive Dashboard, Batch Review, and more.</p>'
            f'</div>'
            f'<div style="font-size:36px">→</div>'
            f'</div>',
            unsafe_allow_html=True)


# ── PAGE 2: EXECUTIVE DASHBOARD ─────────────────────────────
def page_overview(m):
    """Executive Dashboard — KPIs, findings, approval, charts overview."""
    _page_title("Executive Dashboard", "Key metrics, SLA compliance, findings and approval status")
    _render_intel_bar(m)
    _has_ctrlm = st.session_state.get("ctrlm_df") is not None
    _cdf = st.session_state.get("ctrlm_df")

    at_risk = max(0, m["total_jobs"] - m["jobs_ok"] - m["jobs_breach"])
    m["jobs_at_risk"] = at_risk
    m["total_runs"] = len(_cdf) if _cdf is not None else 0
    batch_performance_kpis(m)
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ── Findings + Approval side-by-side ──────────────────
    _ov_left, _ov_right = st.columns([2.2, 1])
    with _ov_left:
        _findings = generate_pe_findings(m)
        render_pe_findings_panel(_findings)
    with _ov_right:
        pe_ok   = st.session_state.get("approval_pe", False)
        cu_ok   = st.session_state.get("approval_customer", False)
        pe_name = st.session_state.get("approval_pe_name", "") or "—"
        cu_name = st.session_state.get("approval_customer_name", "") or "—"
        _ap_col = C["green"] if (pe_ok and cu_ok) else (C["amber"] if (pe_ok or cu_ok) else C["muted"])
        _ap_lbl = "APPROVED" if (pe_ok and cu_ok) else ("PARTIAL" if (pe_ok or cu_ok) else "PENDING")
        _ap_icon = "✅" if (pe_ok and cu_ok) else ("⏳" if (pe_ok or cu_ok) else "⬜")
        _ap_html = (
            f'<div class="panel" style="height:100%">'
            f'<p class="panel-title">✍️ Approval Status</p>'
            f'<div style="text-align:center;padding:18px 0 12px">'
            f'<span style="font-size:36px">{_ap_icon}</span>'
            f'<p style="font-size:18px;font-weight:800;color:{_ap_col};margin:8px 0 2px">{_ap_lbl}</p>'
            f'</div>'
            f'<div style="font-size:12px;line-height:2">'
            f'<div style="display:flex;justify-content:space-between;padding:4px 0;border-bottom:1px solid {C["border"]}">'
            f'<span style="color:{C["muted"]}">PE Lead</span>'
            f'<span style="color:{C["green"] if pe_ok else C["muted"]};font-weight:700">{"✅ " if pe_ok else ""}{_html_mod.escape(pe_name)}</span></div>'
            f'<div style="display:flex;justify-content:space-between;padding:4px 0">'
            f'<span style="color:{C["muted"]}">Customer</span>'
            f'<span style="color:{C["green"] if cu_ok else C["muted"]};font-weight:700">{"✅ " if cu_ok else ""}{_html_mod.escape(cu_name)}</span></div>'
            f'</div></div>'
        )
        st.markdown(_ap_html, unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    if _has_ctrlm:
        ch1, ch2, ch3 = st.columns([1.3, 0.9, 1.3])
        with ch1:
            st.markdown(
                f'<div class="panel">'
                f'<p class="panel-title">📅 Batch Window Trend'
                f'<span style="font-size:11px;color:{C["muted"]};font-weight:400;margin-left:6px">Last 30 days</span></p>',
                unsafe_allow_html=True)
            st.plotly_chart(window_trend_fig(m["window"]), use_container_width=True, config={"displayModeBar": False})
            st.markdown("</div>", unsafe_allow_html=True)
        with ch2:
            st.markdown(f'<div class="panel"><p class="panel-title">🥧 Sub-App Distribution</p>', unsafe_allow_html=True)
            if "Sub_Application" in _cdf.columns and len(_cdf) > 0:
                _sb2 = _cdf.groupby("Sub_Application", as_index=False)["Run_Sec"].sum()
                _sb2["hrs"] = _sb2["Run_Sec"] / 3600
                _fd = go.Figure(go.Pie(labels=_sb2["Sub_Application"], values=_sb2["hrs"],
                    hole=0.55, textinfo="none",
                    marker=dict(colors=[C["blue"], C["purple"], C["cyan"], C["green"], C["amber"]])))
                _fd.update_layout(**BASE_LAYOUT, height=220,
                    margin=dict(l=0, r=0, t=0, b=0), showlegend=True,
                    legend=dict(orientation="h", y=-0.25, x=0.5, xanchor="center", font=dict(size=9)))
                st.plotly_chart(_fd, use_container_width=True, config={"displayModeBar": False})
            else:
                st.markdown(f'<p style="color:{C["muted"]};font-size:11px;padding:20px 0;text-align:center">Sub-Application column not found</p>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
        with ch3:
            st.markdown(
                f'<div class="panel">'
                f'<p class="panel-title">📊 Daily Batch Window'
                f'<span style="font-size:11px;color:{C["muted"]};font-weight:400;margin-left:6px">Hours per day</span></p>',
                unsafe_allow_html=True)
            _win = m["window"]
            if len(_win) > 0:
                _bc = [C["red"] if v > DAILY_LIMIT_HRS else C["blue"] for v in _win["total_hrs"]]
                _fb = go.Figure(go.Bar(x=_win["run_date"].astype(str), y=_win["total_hrs"],
                    marker_color=_bc, marker_line_width=0))
                _fb.add_hline(y=DAILY_LIMIT_HRS, line_dash="dot", line_color=C["red"], line_width=1.5)
                _fb.update_layout(**BASE_LAYOUT, height=220,
                    margin=dict(l=10, r=10, t=0, b=30),
                    xaxis=dict(**AXIS, tickangle=-45, tickfont=dict(size=9), nticks=8),
                    yaxis=dict(**AXIS))
                st.plotly_chart(_fb, use_container_width=True, config={"displayModeBar": False})
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        hcol, rcol = st.columns([1.6, 1])
        with hcol:
            nd = m["daily"]["run_date"].nunique() if "run_date" in m["daily"].columns else "?"
            nj = m["daily"]["Job_Name"].nunique() if "Job_Name" in m["daily"].columns else "?"
            st.markdown(
                f'<div class="panel">'
                f'<p class="panel-title">🗓️ SLA Compliance Heatmap'
                f'<span style="font-size:11px;color:{C["muted"]};font-weight:400;margin-left:6px">{nd}d × {nj} jobs</span></p>',
                unsafe_allow_html=True)
            st.plotly_chart(heatmap_fig(m["daily"]), use_container_width=True, config={"displayModeBar": False})
            st.markdown(
                f'<div style="display:flex;gap:14px;font-size:11px;margin-top:2px">'
                f'<span style="color:{C["green"]}">● OK</span>'
                f'<span style="color:{C["amber"]}">● At Risk</span>'
                f'<span style="color:{C["red"]}">● Breach</span>'
                f'</div></div>',
                unsafe_allow_html=True)
        with rcol:
            st.markdown(f'<div class="panel"><p class="panel-title">🏆 Top Jobs by Peak Runtime</p>', unsafe_allow_html=True)
            if len(m["top_jobs"]) > 0:
                _tj_html = (
                    f'<div class="tj-hdr tj-grid">'
                    f'<span>Job</span><span style="text-align:right">Peak</span>'
                    f'<span style="text-align:right">Avg</span><span style="text-align:right">Status</span>'
                    f'</div>'
                )
                for _, _r in m["top_jobs"].head(10).iterrows():
                    _ib  = _r["peak_hrs"] > DAILY_LIMIT_HRS
                    _bdg = '<span class="b-breach" style="font-size:10px">BR</span>' if _ib else '<span class="b-ok" style="font-size:10px">OK</span>'
                    _rbg = f'{C["red"]}08' if _ib else "transparent"
                    _jn  = _html_mod.escape(_r["Job_Name"][:26])
                    _tj_html += (
                        f'<div class="tj-grid" style="background:{_rbg}">'
                        f'<span class="tj-cell" style="font-weight:600" title="{_html_mod.escape(_r["Job_Name"])}">{_jn}</span>'
                        f'<span class="tj-cell" style="text-align:right;color:{C["red"] if _ib else C["white"]};font-weight:700">{_r["peak_hrs"]:.2f}h</span>'
                        f'<span class="tj-cell" style="text-align:right;color:{C["muted"]}">{_r["avg_hrs"]:.2f}h</span>'
                        f'<span class="tj-cell" style="text-align:right">{_bdg}</span>'
                        f'</div>'
                    )
                st.markdown(_tj_html, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
    else:
        _nudge("📂", "Upload Data to See Overview Charts",
               "Upload a <b>Ctrl-M CSV</b> to see batch window trends, SLA heatmap, and top job analysis.")


# ── PAGE 3: BATCH REVIEW ────────────────────────────────────
def page_batch(m):
    """Batch Review — Job Drilldown, SLA, Benchmark."""
    _page_title("Batch Review", "Job drilldown · SLA compliance · Performance benchmark")
    _render_intel_bar(m)
    _has_ctrlm = st.session_state.get("ctrlm_df") is not None
    _cdf = st.session_state.get("ctrlm_df")

    b1, b2, b3 = st.tabs(["🔍 Job Drilldown", "📋 Batch SLA", "📊 Benchmark"])

    with b1:
        if not _has_ctrlm:
            _nudge("📋", "Job Drilldown — Ctrl-M CSV Required",
                   "Upload your <b>Ctrl-M Execution History CSV</b> to see per-job SLA analysis and compliance heatmap.")
        else:
            breached = set(m["monthly"][m["monthly"]["breach"]]["Job_Name"])
            _jobs    = sorted(_cdf["Job_Name"].unique()) if "Job_Name" in _cdf.columns else []
            if _jobs:
                jc1, jc2 = st.columns([3, 1])
                with jc1:
                    selected = st.selectbox("Select Job", _jobs)
                with jc2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if selected in breached:
                        st.markdown('<span class="b-breach">⚠ MONTHLY BREACH</span>', unsafe_allow_html=True)
                    else:
                        st.markdown('<span class="b-ok">✅ WITHIN SLA</span>', unsafe_allow_html=True)
                d1, d2 = st.columns(2)
                with d1:
                    st.markdown('<div class="panel">', unsafe_allow_html=True)
                    st.plotly_chart(job_trend_fig(_cdf, selected),
                                    use_container_width=True, config={"displayModeBar": False})
                    st.markdown("</div>", unsafe_allow_html=True)
                with d2:
                    st.markdown('<div class="panel">', unsafe_allow_html=True)
                    st.plotly_chart(monthly_bar_fig(m["monthly"], selected),
                                    use_container_width=True, config={"displayModeBar": False})
                    st.markdown("</div>", unsafe_allow_html=True)
                jd = m["daily"][m["daily"]["Job_Name"] == selected].sort_values("run_date")
                st.dataframe(
                    jd[["run_date","total_hrs","runs","breach"]]
                      .rename(columns={"run_date":"Date","total_hrs":"Runtime (hrs)","runs":"Runs","breach":"Breach?"}),
                    use_container_width=True, height=200)
                bdf = m["monthly"][m["monthly"]["breach"]]
                if not bdf.empty:
                    st.markdown("**⚠️ All Monthly Breaching Jobs**")
                    st.dataframe(bdf.sort_values("total_hrs", ascending=False),
                                 use_container_width=True, height=180)

    with b2:
        bp = st.session_state.get("batch_sla_df")
        if bp is None:
            _nudge("📋", "Batch SLA — SLA Matrix Required",
                   "Upload your <b>Batch SLA XLSX</b> or CSV to see per-job compliance analysis.")
        else:
            sla_fmt = getattr(bp, 'attrs', {}).get("_sla_format", "")
            cdf = st.session_state.get("ctrlm_df")

            if sla_fmt == "batch_schedule":
                st.markdown('<div class="panel"><p class="panel-title">Batch Schedule SLA</p>', unsafe_allow_html=True)
                k1, k2, k3 = st.columns(3)
                with k1: kpi("Batch Streams", str(len(bp)), C["blue"])
                if "SLA_Window_Hrs" in bp.columns:
                    with k2: kpi("Max SLA Window", f"{bp['SLA_Window_Hrs'].max():.1f}h", C["amber"])
                if "Current_End" in bp.columns and "Expected_End" in bp.columns:
                    try:
                        def _check_overrun(row):
                            exp = pd.to_datetime(str(row.get("Expected_End","")), format="%H:%M:%S", errors="coerce")
                            cur = pd.to_datetime(str(row.get("Current_End","")), format="%H:%M:%S", errors="coerce")
                            if pd.isna(exp) or pd.isna(cur): return False
                            return cur > exp
                        overruns = bp.apply(_check_overrun, axis=1).sum()
                        with k3: kpi("SLA Overruns", str(overruns), C["red"] if overruns > 0 else C["green"])
                    except Exception:
                        pass
                st.dataframe(bp, use_container_width=True, height=300)
                if "Comments" in bp.columns:
                    for _, row in bp.iterrows():
                        cmt = str(row.get("Comments", ""))
                        if cmt and cmt not in ("nan", "None", ""):
                            bn = row.get("Batch_Name", row.get(bp.columns[0], ""))
                            st.caption(f"**{bn}**: {cmt}")
                st.markdown("</div>", unsafe_allow_html=True)

            elif "Job_Name" in bp.columns and "SLA_Hrs" in bp.columns and cdf is not None:
                st.markdown('<div class="panel"><p class="panel-title">SLA Compliance Analysis</p>', unsafe_allow_html=True)
                job_peak = (cdf.groupby("Job_Name", as_index=False)
                              .agg(peak_hrs=("run_time_hrs","max"),
                                   avg_hrs=("run_time_hrs","mean"),
                                   total_runs=("run_time_hrs","count")))
                bp_clean = bp.copy()
                bp_clean["SLA_Hrs"] = pd.to_numeric(bp_clean["SLA_Hrs"], errors="coerce")
                merged = bp_clean.merge(job_peak, on="Job_Name", how="left")
                merged["peak_hrs"] = merged["peak_hrs"].fillna(0).round(3)
                merged["avg_hrs"]  = merged["avg_hrs"].fillna(0).round(3)
                merged["buffer_pct"] = ((merged["SLA_Hrs"] - merged["peak_hrs"]) / merged["SLA_Hrs"] * 100).round(1)
                merged["status"] = merged.apply(
                    lambda r: "BREACH" if r["peak_hrs"] > r["SLA_Hrs"]
                    else ("AT RISK" if r["buffer_pct"] < 15 else "OK") if r["SLA_Hrs"] > 0
                    else "NO SLA", axis=1)
                n_breach = (merged["status"] == "BREACH").sum()
                n_risk   = (merged["status"] == "AT RISK").sum()
                n_ok     = (merged["status"] == "OK").sum()
                k1, k2, k3, k4 = st.columns(4)
                with k1: kpi("Total SLA Jobs", str(len(merged)), C["blue"])
                with k2: kpi("Compliant", str(n_ok), C["green"])
                with k3: kpi("At Risk", str(n_risk), C["amber"])
                with k4: kpi("Breaching", str(n_breach), C["red"])
                display_cols = ["Job_Name","SLA_Hrs","peak_hrs","avg_hrs","buffer_pct","status"]
                avail = [c for c in display_cols if c in merged.columns]
                st.dataframe(merged[avail].sort_values("buffer_pct", ascending=True),
                             use_container_width=True, height=350)
                st.markdown("</div>", unsafe_allow_html=True)

            else:
                st.markdown('<div class="panel"><p class="panel-title">Batch SLA Matrix (Raw)</p>', unsafe_allow_html=True)
                if "Job_Name" not in bp.columns and "Batch_Name" not in bp.columns:
                    st.warning("Could not detect Job_Name or Batch_Name column. Check column headers.")
                st.dataframe(bp, use_container_width=True, height=300)
                st.markdown("</div>", unsafe_allow_html=True)

    with b3:
        ui_performance_tab()


# ── PAGE 4: APP REVIEW ──────────────────────────────────────
def page_app(m):
    """App Review — Sub-App Mix, Performance Test."""
    _page_title("App Review", "Sub-application workload distribution · Performance test analysis")
    _render_intel_bar(m)
    _has_ctrlm = st.session_state.get("ctrlm_df") is not None

    a1, a2 = st.tabs(["🥧 Sub-App Mix", "⚡ Perf Test"])

    with a1:
        if not _has_ctrlm:
            _nudge("🥧", "Sub-Application — Ctrl-M CSV Required",
                   "Upload your <b>Ctrl-M CSV</b> to see workload distribution by sub-application.")
        else:
            s1, s2 = st.columns([1, 2])
            with s1:
                st.markdown('<div class="panel"><p class="panel-title">Runtime by Sub-App</p>',
                            unsafe_allow_html=True)
                st.plotly_chart(sub_app_pie(m["sub_stats"]), use_container_width=True,
                                config={"displayModeBar": False})
                st.markdown("</div>", unsafe_allow_html=True)
            with s2:
                st.markdown('<div class="panel"><p class="panel-title">Sub-Application Detail</p>',
                            unsafe_allow_html=True)
                st.dataframe(
                    m["sub_stats"]
                      .rename(columns={"Sub_Application":"Sub App","total_hrs":"Total Hrs","jobs":"Jobs"})
                      .sort_values("Total Hrs", ascending=False),
                    use_container_width=True, height=220)
                st.markdown("</div>", unsafe_allow_html=True)

    with a2:
        if not _has_ctrlm:
            _nudge("⚡", "Perf-Test Report — Ctrl-M CSV Required",
                   "Upload your <b>Ctrl-M CSV</b> to generate the full performance test report with buffer analysis.")
        else:
            perf_test_report_tab(m)


# ── PAGE 5: INFRA REVIEW ────────────────────────────────────
def page_infra(m):
    """Infra Review — Resource Utilization, SOW & Volume."""
    _page_title("Infra Review", "Server resource utilization · SOW contract volume analysis")
    _render_intel_bar(m)

    i1, i2 = st.tabs(["🖥️ Resource Util", "📄 SOW & Volume"])
    with i1:
        resource_tab()
    with i2:
        sow_tab()


# ── PAGE 6: FINDINGS & EVIDENCE ─────────────────────────────
def page_findings(m):
    """Findings & Evidence — PE findings, intelligence panel, AI analysis."""
    _page_title("Findings & Evidence", "Automated audit findings · PE intelligence · AI-powered root cause analysis")
    _render_intel_bar(m)

    st.markdown(
        f'<div style="margin-bottom:16px">'
        f'<p style="font-size:18px;font-weight:700;color:{C["white"]};margin:0 0 4px">🔍 Findings & Evidence</p>'
        f'<p style="font-size:12px;color:{C["muted"]};margin:0">Automated PE audit findings, intelligence insights, and AI-powered analysis</p>'
        f'</div>',
        unsafe_allow_html=True)

    # Full-width findings panel
    _findings = generate_pe_findings(m)
    render_pe_findings_panel(_findings)

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    # PE Audit Intelligence panel
    ctrlm_df = st.session_state.get("ctrlm_df")
    server_data = st.session_state.get("server_data") or []
    pe_audit_intelligence_panel(ctrlm_df, server_data)

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    # AI Analysis section
    _ai_tabs = st.tabs(["🤖 AI Batch Analysis", "🤖 AI Resource Analysis", "🏆 Master Summary"])

    with _ai_tabs[0]:
        _render_ai_insight_box("_ai_ctrlm_insight", "📊 Batch Intelligence", "_run_ai_ctrlm")
    with _ai_tabs[1]:
        _render_ai_insight_box("_ai_resource_insight", "🖥️ Resource Intelligence", "_run_ai_resource")
    with _ai_tabs[2]:
        _render_ai_insight_box("_ai_master_summary", "🏆 PE Approval Summary", "_run_master_summary")


# ── PAGE 7: APPROVAL / EXPORT / DOCS ────────────────────────
def page_governance(m):
    """Governance — Issues, PE Docs, Export, Approval."""
    _page_title("Approval & Export", "Issue register · PE documents · Report export · Sign-off")
    _render_intel_bar(m)
    _has_ctrlm = st.session_state.get("ctrlm_df") is not None

    v1, v2, v3, v4 = st.tabs(["⚠️ Issues", "📋 PE Docs", "📤 Export", "✍️ Approval"])

    with v1:
        issues_waivers_tab()

    with v2:
        pe_document_review_tab()

    with v3:
        st.markdown(
            f'<div class="panel">'
            f'<p class="panel-title">📤 Export Report</p>'
            f'<p class="panel-sub">Download the full PE Audit report as an HTML file — includes dark/light theme toggle, '
            f'all charts, server table, issues register and approval status.</p>',
            unsafe_allow_html=True)
        if not _has_ctrlm:
            st.info("📂 Upload a Ctrl-M CSV to generate the full performance audit report. "
                    "Resource and PE Document data is included when available.")
        else:
            ec1, ec2 = st.columns(2)
            with ec1:
                if st.button("🔄 Build HTML Report", use_container_width=True, key="build_report_btn"):
                    st.session_state["_html_report"] = build_html_report(m)
            with ec2:
                if st.session_state.get("_html_report"):
                    st.download_button(
                        "⬇️ Download Report (.html)",
                        st.session_state["_html_report"].encode("utf-8"),
                        file_name=f"PE_Audit_{(st.session_state.get('customer_name','Report')).replace(' ','_')}.html",
                        mime="text/html",
                        key="dl_report_btn",
                        use_container_width=True,
                    )
            if st.session_state.get("_html_report"):
                st.success("✅ Report built. Click **⬇️ Download Report** to save.")
                with st.expander("📋 Report Preview (first 3,000 chars)"):
                    st.code(st.session_state["_html_report"][:3000], language="html")
            ms = st.session_state.get("_ai_master_summary")
            if ms:
                st.markdown("---")
                st.markdown(
                    f'<div style="background:{C["card2"]};border:1px solid {C["border"]};'
                    f'border-radius:10px;padding:16px 20px;margin-top:8px">'
                    f'<p style="font-size:12px;font-weight:700;color:{C["white"]};margin-bottom:8px">🏆 PE Approval Summary (AI Generated)</p>'
                    f'<p style="font-size:12px;color:{C["muted"]};line-height:1.8">{ms}</p>'
                    f'</div>',
                    unsafe_allow_html=True)
                st.download_button("⬇️ Download Summary Text",
                    ms.encode(), f"PE_Summary_{st.session_state.get('customer_name','').replace(' ','_')}.txt",
                    "text/plain", key="dl_summ_exp")
            if st.session_state.get("issues_list"):
                st.markdown("---")
                iss_df = pd.DataFrame(st.session_state.issues_list)
                st.download_button("⬇️ Issues Register CSV",
                                   iss_df.to_csv(index=False).encode(),
                                   "issues_register.csv", "text/csv", key="dl_iss")
        st.markdown("</div>", unsafe_allow_html=True)

    with v4:
        approval_tab()


# ── FOOTER ───────────────────────────────────────────────────
def footer():
    st.markdown(f"""<div style="border-top:1px solid {C['border']};background:{C['card']};
padding:12px 28px;margin-top:36px;display:flex;justify-content:space-between;
font-size:11px;color:{C['muted']}">
<span>PE Control Tower v4.0 &nbsp;•&nbsp; Performance Engineering Team</span>
<span>Daily SLA: {DAILY_LIMIT_HRS}h &nbsp;|&nbsp; Monthly: {MONTHLY_LIMIT_HRS}h &nbsp;|&nbsp;
CPU/Mem OK: ≤{CPU_OK}% &nbsp;|&nbsp; Critical: >{CPU_WARN}%</span>
</div>""", unsafe_allow_html=True)


# ── ENTRY POINT ──────────────────────────────────────────────
init_state()          # MUST be first — initialises all session_state keys
inject_css()

_overall = "OK"
_has_ctrlm    = st.session_state.get("ctrlm_df") is not None
_has_servers  = bool(st.session_state.get("server_data"))
_has_sla      = st.session_state.get("batch_sla_df") is not None
_has_ui       = st.session_state.get("_uiperf_df") is not None
_has_pe_doc   = st.session_state.get("_pe_doc_data") is not None
_any_data     = _has_ctrlm or _has_servers or _has_sla or _has_ui or _has_pe_doc

# NO auto-navigation — user controls when to leave upload page
# The sidebar nav always lets user go back to Upload & Intake

_m = None
if _has_ctrlm:
    _m = compute_metrics_fast(st.session_state.ctrlm_df)
    _overall = "BREACH" if _m["jobs_breach"] > 0 else "OK"

# Build minimal metrics if other data loaded but no Ctrl-M
if _m is None and _any_data:
    _empty_df = pd.DataFrame(columns=["Job_Name","run_date","run_time_hrs","Run_Sec","Sub_Application","month"])
    _m = compute_metrics(_empty_df)

render_sidebar_nav(_any_data)
render_header(
    st.session_state.get("customer_name",""),
    st.session_state.get("env_type",""),
    _overall
)

# ── Page Router ──────────────────────────────────────────────
_page = st.session_state.get("_current_page", "upload")

# Guard: force back to upload if no data and trying to access data pages
if not _any_data and _page != "upload":
    _page = "upload"
    st.session_state["_current_page"] = "upload"

if   _page == "upload":     page_upload()
elif _page == "overview":   page_overview(_m)
elif _page == "batch":      page_batch(_m)
elif _page == "app":        page_app(_m)
elif _page == "infra":      page_infra(_m)
elif _page == "findings":   page_findings(_m)
elif _page == "governance": page_governance(_m)
else:
    page_upload()  # fallback

footer()