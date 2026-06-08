"""
Generic resource utilization parser — single entry point.

parse_resource_file(file_bytes: bytes, filename: str) -> list[dict]

Pipeline (exact per spec):
  1. detect_file_mode()      → TEXT_PDF | TEXT_DOCX | IMAGE_PDF | IMAGE_DOCX | CSV | XLSX
  2. extract_text()          → plain text from PDF or DOCX
  3. parse_text_to_servers() → CPU_RE.findall() + MEM_RE.findall() (paired, per block)
       - Available/Free Memory found  → mem_used = round(100 - available, 2)
       - Memory Used / utilization    → use directly
  4. _is_valid_hostname()    → filters English words / metric labels, keeps real hostnames
  5. save_resource_session() → clears stale _SESSION cache, returns fresh list

Replace usage:
    OLD:  parse_resource_docx(buf), parse_zabbix_pdf_text(text)
    NEW:  parse_resource_file(file_bytes, filename)
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# ── Module-level session cache (cleared on every new upload) ─────
# Prevents stale DB hostnames from a previous customer's file from
# bleeding into the next upload. Single-user PE Dashboard pattern.
_SESSION: Dict[str, Any] = {}


# ─────────────────────────────────────────────────────────────────
# English word / metric-label blocklist
# "server", "cpu", "total" etc. must never be treated as hostnames
# Real server names ALWAYS contain at least one digit (e.g. tsbb191525041)
# ─────────────────────────────────────────────────────────────────
_BLOCKED: frozenset = frozenset({
    # Articles, conjunctions, pronouns
    "the", "and", "for", "are", "but", "not", "you", "all", "any", "can",
    "her", "was", "one", "our", "out", "day", "get", "has", "him", "his",
    "how", "man", "new", "now", "old", "see", "two", "way", "who", "did",
    "its", "let", "put", "say", "she", "too", "use", "via", "per", "set",
    # IT metric and report terms
    "cpu", "mem", "disk", "idle", "time", "sla", "max", "min", "avg",
    "trend", "total", "used", "free", "value", "latest", "status", "graphs",
    "system", "report", "performance", "utilization", "resource", "monitoring",
    "available", "percent", "memory", "storage", "network", "interface",
    "ram", "bandwidth", "uncached", "process", "chronyd",
    # Generic doc structure terms
    "server", "servers", "node", "host", "hosts", "cluster", "group",
    "section", "table", "page", "date", "name", "type", "title",
    "summary", "details", "overview", "results", "information",
})

# Regex for real server hostnames: letter-prefix + digits pattern
# Matches: tsbb191525041, prbg241530001, appsvr01, db01, web-node-3
_HOSTNAME_RE = re.compile(
    r'^[a-z]'                     # must start with a letter
    r'[a-z0-9._-]{1,}'            # letters/digits/dots/dashes
    r'\d'                         # must contain at least one digit (by end)
    r'[a-z0-9._-]{0,}$',          # optional suffix
    re.IGNORECASE,
)


def _is_valid_hostname(name: str, strict: bool = True) -> bool:
    """
    Return True only for real server hostnames.

    strict=True  (Strategies A-C): enforce digit + letter-prefix rules.
    strict=False (Strategy D):     accept label-style names like
                                   "ApplicationServer1" from DOCX headings.
    """
    if not name or len(name) < 3:
        return False
    # Strip FQDN suffix — work with the short hostname only
    base = name.split(".")[0].lower().strip()
    if not base or len(base) < 2:
        return False
    # Hard-block known English / metric words
    if base in _BLOCKED:
        return False
    # Reject underscore-joined metric labels: RAM_utilization, CPU_last_15days, VM_Uncached_Bandwidth
    parts = base.replace("_", " ").split()
    if any(p in _BLOCKED for p in parts):
        return False
    if strict:
        # Must contain at least one digit
        if not re.search(r"\d", base):
            return False
        # Must start with a letter
        if not re.match(r"^[a-z]", base, re.I):
            return False
        # Reject OS version strings like "3.10.0-1160"
        if re.match(r'^\d', base):
            return False
    return True


# ─────────────────────────────────────────────────────────────────
# Step 1 — detect_file_mode
# ─────────────────────────────────────────────────────────────────
def detect_file_mode(file_bytes: bytes, filename: str) -> str:
    """Detect resource file type.  Delegates to sla_parser.detect_resource_mode().
    Returns one of: TEXT_PDF | TEXT_DOCX | IMAGE_PDF | IMAGE_DOCX | CSV | XLSX | UNKNOWN
    """
    try:
        from services.sla_parser import detect_resource_mode
        mode = detect_resource_mode(file_bytes, filename)
        # Upgrade UNKNOWN for known spreadsheet extensions
        if mode == "UNKNOWN":
            ext = (filename or "").rsplit(".", 1)[-1].lower()
            if ext in ("xlsx", "xls"):
                return "XLSX"
        return mode
    except Exception:
        ext = (filename or "").rsplit(".", 1)[-1].lower()
        if ext == "pdf":                return "TEXT_PDF"
        if ext == "docx":              return "TEXT_DOCX"
        if ext == "csv":               return "CSV"
        if ext in ("xlsx", "xls"):     return "XLSX"
        return "UNKNOWN"


# ─────────────────────────────────────────────────────────────────
# Step 2 — extract_text
# ─────────────────────────────────────────────────────────────────
def extract_text(file_bytes: bytes, mode: str) -> str:
    """Extract plain text from a PDF or DOCX.  Returns '' for image-only modes."""
    if mode in ("IMAGE_DOCX", "IMAGE_PDF"):
        return ""

    if mode == "TEXT_PDF":
        # Strategy 1: PyMuPDF (fitz) — preserves layout / column order far better
        #             than pypdf which concatenates words in stream order.
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            parts: List[str] = []
            for page in doc:
                # sort=True guarantees reading-order (top→bottom, left→right)
                txt = page.get_text("text", sort=True)
                if txt:
                    parts.append(txt)
            doc.close()
            joined = "\n".join(parts)
            if joined.strip():
                logger.debug("extract_text: PDF via PyMuPDF (%d chars)", len(joined))
                return joined
        except ImportError:
            pass
        except Exception as exc:
            logger.warning("extract_text PyMuPDF failed: %s", exc)

        # Strategy 2: pypdf / PyPDF2 fallback
        try:
            from services.resource_parser import extract_pdf_text
            text = extract_pdf_text(io.BytesIO(file_bytes))
            logger.debug("extract_text: PDF via pypdf (%d chars)", len(text))
            return text
        except Exception as exc:
            logger.warning("extract_text pypdf failed: %s", exc)
            return ""

    if mode == "TEXT_DOCX":
        return _extract_docx_structured(file_bytes)

    return ""


def _extract_docx_structured(file_bytes: bytes) -> str:
    """
    Extract DOCX text in document order (paragraphs interleaved with tables),
    emitting ``<<SECTION: heading_text >>`` markers whenever a heading paragraph
    or bold+short paragraph is encountered.

    This preserves the heading→metrics association needed for Strategy D
    (Leonardo / Distell / CEAT-style DOCXs).
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn as _qn

        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []

        def _is_heading_para(para) -> bool:
            """True when the paragraph looks like a section heading."""
            style = (para.style.name or "").lower()
            if "heading" in style or style.startswith("title"):
                return True
            t = para.text.strip()
            if not t or len(t) > 120:
                return False
            words = t.split()
            # Short paragraph ending with ":" → treat as server section label
            # e.g. "Application Server 1 :" / "Database Server :"
            if t.endswith(":") and len(words) <= 6:
                return True
            # Bold majority → treat as heading if the line is short
            if para.runs:
                bold_chars = sum(len(r.text) for r in para.runs
                                 if r.bold and r.text.strip())
                total = max(len(t), 1)
                if bold_chars / total > 0.5 and len(words) <= 8:
                    return True
            return False

        # Walk body XML in document order to interleave paragraphs and tables
        body = doc.element.body
        for child in body.iterchildren():
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if local == "p":                          # paragraph
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                t = para.text.strip()
                if not t:
                    continue
                if _is_heading_para(para):
                    parts.append(f"\n<<SECTION: {t} >>")
                else:
                    parts.append(t)

            elif local == "tbl":                      # table
                from docx.table import Table
                tbl = Table(child, doc)
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

        result = "\n".join(parts)
        logger.debug("_extract_docx_structured: %d chars", len(result))
        return result

    except Exception as exc:
        logger.warning("_extract_docx_structured failed (%s) — using simple fallback", exc)
        # Minimal fallback
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))
            return "\n".join(lines)
        except Exception:
            return ""


# ─────────────────────────────────────────────────────────────────
# Compiled regex patterns (module-level — compiled once at import)
# All patterns use findall() so hostname and value are always PAIRED
# ─────────────────────────────────────────────────────────────────

# CPU idle → used = 100 - idle  (anchored: no [^:] greed beyond \n)
_CPU_IDLE_RE = re.compile(
    r'CPU\s+idle\s+time[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# CPU idle trend/SLA (for cpu_avg)
_CPU_IDLE_SLA_RE = re.compile(
    r'CPU\s+idle\s+time\s+(?:trend|SLA)[^:\n]{0,40}:\s*(\d[\d.]+)',
    re.IGNORECASE)

# CPU utilization/usage/load direct (also: "CPU % avg", "CPU used")
_CPU_USED_RE = re.compile(
    r'CPU\s*(?:%\s*)?(?:utilization|usage|load|used|busy|percent|average|avg)[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# Azure Monitor / Grafana: "Avg CPU Percentage" or "cpu_usage_percent"
_CPU_AZURE_RE = re.compile(
    r'(?:Avg(?:erage)?\s+CPU|cpu_usage|processorcpuutilization)[^:\n]{0,40}[:\s]+(\d[\d.]+)\s*%?',
    re.IGNORECASE)

# Available / Free memory % → mem_used = 100 - value
# Covers: "Available Memory [%]", "Available memory:", "Free memory %",
#         "Available memory percent:", "Free Memory (%):"
_MEM_AVAIL_RE = re.compile(
    r'(?:Available|Free)\s+[Mm]emory'
    r'(?:\s*[\[(%]\s*%?\s*[)\]]?|\s*percent|\s*%)?'
    r'[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# Memory used / utilization / percent used → direct value
# Also handles "Memory Usage", "Mem Used", "Used Memory %", "memoryutilization"
_MEM_USED_RE = re.compile(
    r'(?:[Mm]emory\s*(?:utilization|used|usage|consumed|percent\s*used|%)|'
    r'Used\s+[Mm]emory\s*%?|'
    r'Mem(?:ory)?\s+Used|'
    r'memoryutilization|'
    r'memory_usage_percent)'
    r'[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# Total memory → GB figure
_MEM_TOTAL_RE = re.compile(
    r'Total\s+[Mm]emory[^:\n]{0,30}:\s*(\d[\d.]+)\s*GB',
    re.IGNORECASE)

# Free disk space on /mount (%) → used = 100 - free
_DISK_FREE_RE = re.compile(
    r'Free\s+disk\s+space\s+on\s+(\S+)\s*\(percentage\)[^:\n]{0,40}:'
    r'\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# Available disk space on /mount → used = 100 - available
_DISK_AVAIL_RE = re.compile(
    r'[Aa]vailable\s+[Dd]isk\s+space\s+(?:in\s+%?\s+)?on\s+(\S+)'
    r'[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# Used disk space on /mount → direct
_DISK_USED_RE = re.compile(
    r'(?:[Uu]sed\s+disk\s+space\s+on|[Dd]isk\s+(?:usage|utilization)\s+(?:on|for)?|'
    r'[Ff]ilesystem\s+(?:usage\s+)?on)\s+(\S+)'
    r'[^:\n]{0,60}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)

# ASM disk groups (Oracle) — used% is direct (no inversion)
_DISK_ASM_RE = re.compile(
    r'ASM\s+(?:disk\s+group\s+)?(\w+)\s*[:\s]+(?:used|utilization)[:\s]+(\d[\d.]+)\s*%?',
    re.IGNORECASE)

# Tabular format: "CPU | 16.4% | ..." or "CPU: 16.4" (simple inline)
_CPU_INLINE_RE  = re.compile(r'^\s*[Cc][Pp][Uu]\s*[:|]\s*(\d[\d.]+)\s*%?\s*$', re.MULTILINE)
_MEM_INLINE_RE  = re.compile(r'^\s*[Mm]em(?:ory)?\s*[:|]\s*(\d[\d.]+)\s*%?\s*$', re.MULTILINE)
_DISK_INLINE_RE = re.compile(r'^\s*[Dd]isk\s*[:|]\s*(\d[\d.]+)\s*%?\s*$', re.MULTILINE)

# Pipe-delimited table row: "CPU % | 16.4 | ..."
_TABLE_CPU_RE  = re.compile(r'CPU\s*%?\s*\|\s*(\d[\d.]+)', re.IGNORECASE)
_TABLE_MEM_RE  = re.compile(r'Mem(?:ory)?\s*%?\s*\|\s*(\d[\d.]+)', re.IGNORECASE)
_TABLE_DISK_RE = re.compile(r'Disk\s*%?\s*\|\s*(\d[\d.]+)', re.IGNORECASE)

# Paired CPU pattern for Strategy C fallback
# Captures (hostname, cpu_value) in a single findall() — no fragile zip()
_CPU_PAIR_RE = re.compile(
    r'CPU[^:\n]{0,80}(?:for|on)\s+([a-z]{2,8}\d{3,}[a-z0-9.]*)'
    r'[^:\n]{0,40}:\s*(?:latest\s+value:\s*)?(\d[\d.]+)',
    re.IGNORECASE)


# ─────────────────────────────────────────────────────────────────
# Block-level paired metric extractor
# ─────────────────────────────────────────────────────────────────
def _extract_metrics_from_block(block: str) -> Dict[str, Any]:
    """
    Extract CPU / Memory / Disk from a single server's text block.

    Using findall() inside a block ensures the hostname (block header) is
    always PAIRED with the values found — no cross-server contamination.

    Available Memory % → mem_used = round(100 - available, 2)  [INVERTED]
    Memory Used %      → mem_used = direct value
    """
    rec: Dict[str, Any] = {
        "cpu_used": 0.0, "cpu_avg": 0.0,
        "mem_used": 0.0, "mem_total_gb": 0.0,
        "disk_used_max": 0.0, "disks": {},
    }

    # ── CPU ──────────────────────────────────────────────────────
    # Prefer idle→invert (most accurate Zabbix metric)
    m = _CPU_IDLE_RE.search(block)
    if m:
        rec["cpu_used"] = round(100.0 - float(m.group(1)), 2)

    m = _CPU_IDLE_SLA_RE.search(block)
    if m:
        rec["cpu_avg"] = round(100.0 - float(m.group(1)), 2)

    # Fallback: direct CPU utilization
    if rec["cpu_used"] == 0.0:
        m = _CPU_USED_RE.search(block)
        if m:
            rec["cpu_used"] = round(float(m.group(1)), 2)

    # Fallback: Azure Monitor / Grafana style
    if rec["cpu_used"] == 0.0:
        m = _CPU_AZURE_RE.search(block)
        if m:
            rec["cpu_used"] = round(float(m.group(1)), 2)

    # Fallback: simple inline "CPU: 16.4%" or table "CPU % | 16.4"
    if rec["cpu_used"] == 0.0:
        for pat in (_CPU_INLINE_RE, _TABLE_CPU_RE):
            m = pat.search(block)
            if m:
                rec["cpu_used"] = round(float(m.group(1)), 2)
                break

    # Clamp to valid range
    rec["cpu_used"] = min(100.0, max(0.0, rec["cpu_used"]))
    rec["cpu_avg"]  = min(100.0, max(0.0, rec["cpu_avg"]))

    # ── Memory ───────────────────────────────────────────────────
    # RULE: Available/Free Memory % → mem_used = 100 - available
    m = _MEM_AVAIL_RE.search(block)
    if m:
        val = float(m.group(1))
        rec["mem_used"] = round(100.0 - val, 2) if 0.0 < val <= 100.0 else 0.0
        logger.debug("  mem from Available: %.1f%% avail → %.1f%% used", val, rec["mem_used"])
    else:
        # Direct Memory Used %
        m = _MEM_USED_RE.search(block)
        if m:
            rec["mem_used"] = round(float(m.group(1)), 2)

    # Fallback: simple inline "Mem: 45%" or table "Memory % | 45"
    if rec["mem_used"] == 0.0:
        for pat in (_MEM_INLINE_RE, _TABLE_MEM_RE):
            m = pat.search(block)
            if m:
                rec["mem_used"] = round(float(m.group(1)), 2)
                break

    # Memory total GB (informational)
    m = _MEM_TOTAL_RE.search(block)
    if m:
        rec["mem_total_gb"] = float(m.group(1))

    # Clamp
    rec["mem_used"] = min(100.0, max(0.0, rec["mem_used"]))

    # ── Disk ─────────────────────────────────────────────────────
    # Free disk → invert (findall returns [(mount, pct), ...] — always PAIRED)
    for mount, pct in _DISK_FREE_RE.findall(block):
        rec["disks"][mount] = round(100.0 - float(pct), 2)

    # Available disk → invert (PAIRED)
    for mount, pct in _DISK_AVAIL_RE.findall(block):
        if mount not in rec["disks"]:
            v = float(pct)
            rec["disks"][mount] = round(100.0 - v, 2) if 0.0 < v <= 100.0 else 0.0

    # Used disk → direct (PAIRED)
    for mount, pct in _DISK_USED_RE.findall(block):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # ASM disk groups (Oracle) — used% direct, no inversion
    for label, pct in _DISK_ASM_RE.findall(block):
        key = f"ASM:{label.upper()}"
        if key not in rec["disks"]:
            rec["disks"][key] = round(float(pct), 2)

    # Fallback: simple inline "Disk: 45%" or table row
    if not rec["disks"]:
        for pat in (_DISK_INLINE_RE, _TABLE_DISK_RE):
            m = pat.search(block)
            if m:
                rec["disks"]["/"] = round(float(m.group(1)), 2)
                break

    # Normalize bad mount keys (empty string / parens → '/')
    cleaned: Dict[str, float] = {}
    for mnt, v in rec["disks"].items():
        key = mnt if (mnt and not mnt.startswith("(")) else "/"
        cleaned[key] = max(cleaned.get(key, 0.0), v)
    rec["disks"] = cleaned
    if rec["disks"]:
        rec["disk_used_max"] = max(rec["disks"].values())

    return rec


# ─────────────────────────────────────────────────────────────────
# Step 3 — parse_text_to_servers
# ─────────────────────────────────────────────────────────────────
def parse_text_to_servers(text: str) -> List[Dict[str, Any]]:
    """
    Parse plain text → list of server metric dicts.

    Uses paired CPU_RE.findall() / MEM_RE.findall() inside each block so
    hostname and values are always bound together — no cross-server mis-pairing.

    Strategy A: Zabbix "System Status for HOSTNAME" blocks (richest data)
    Strategy B: "Graphs for HOSTNAME" ToC entries → image_only stubs
    Strategy C: Paired CPU hostname+value regex (fallback for other formats)
    """
    if not text or not text.strip():
        return []

    try:
        from services.resource_parser import _infer_server_type
    except Exception:
        def _infer_server_type(h, ctx="", hint=""):  # type: ignore[override]
            return "DB" if any(k in h.lower() for k in ["db", "sql", "ora"]) else "APP"

    servers: Dict[str, Dict[str, Any]] = {}  # host_lower → record

    def _key(h: str) -> str:
        return h.split(".")[0].lower()

    def _upsert(host: str, rec: Dict[str, Any]) -> None:
        k = _key(host)
        new_has_data = rec.get("cpu_used", 0) > 0 or rec.get("mem_used", 0) > 0 or rec.get("disk_used_max", 0) > 0
        if k not in servers:
            servers[k] = rec
            return
        existing = servers[k]
        old_has_data = existing.get("cpu_used", 0) > 0 or existing.get("mem_used", 0) > 0 or existing.get("disk_used_max", 0) > 0
        if new_has_data and not old_has_data:
            servers[k] = rec
        elif new_has_data and old_has_data:
            for fld in ("cpu_used", "cpu_avg", "mem_used", "mem_total_gb", "disk_used_max"):
                existing[fld] = max(float(existing.get(fld, 0)), float(rec.get(fld, 0)))
            existing["disks"].update(rec.get("disks", {}))
            if existing["disks"]:
                existing["disk_used_max"] = max(existing["disks"].values())

    # ── Strategy A: "System Status for HOSTNAME" blocks ──────────
    blocks = re.split(r'System Status for\s+', text, flags=re.IGNORECASE)
    for block in blocks[1:]:
        lines = block.strip().splitlines()
        if not lines:
            continue
        host_raw = lines[0].strip().split()[0].rstrip(".")
        if not _is_valid_hostname(host_raw, strict=True):
            logger.debug("Strategy A: skipped invalid hostname %r", host_raw)
            continue

        metrics = _extract_metrics_from_block(block)
        rec: Dict[str, Any] = {
            "host": host_raw,
            "type": _infer_server_type(host_raw, block[:500]),
            "_image_only": False,
            **metrics,
        }
        _upsert(host_raw, rec)
        logger.debug(
            "Strategy A: %-30s  cpu=%.1f%%  mem=%.1f%%  disk=%.1f%%",
            host_raw, metrics["cpu_used"], metrics["mem_used"], metrics["disk_used_max"])

    # ── Strategy B: "Graphs for HOSTNAME" ToC entries ────────────
    for host in re.findall(r'Graphs for\s+(\S+)', text, re.IGNORECASE):
        host = host.rstrip(".")
        if not _is_valid_hostname(host, strict=True):
            continue
        if _key(host) in servers:
            continue
        servers[_key(host)] = {
            "host": host,
            "type": _infer_server_type(host),
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }
        logger.debug("Strategy B stub: %s", host)

    # ── Strategy C: Paired hostname+value regex (fallback) ───────
    for host, cpu_val in _CPU_PAIR_RE.findall(text):
        host = host.rstrip(".")
        if not _is_valid_hostname(host, strict=True):
            continue
        k = _key(host)
        if k in servers and servers[k].get("cpu_used", 0) > 0:
            continue
        servers.setdefault(k, {
            "host": host,
            "type": _infer_server_type(host),
            "cpu_used": round(float(cpu_val), 2),
            "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": False,
        })
        logger.debug("Strategy C: %s  cpu=%.1f%%", host, float(cpu_val))

    # ── Strategy D: <<SECTION: label >> blocks (Leonardo/CEAT/Distell DOCX) ──
    # _extract_docx_structured() emits these markers when it encounters heading
    # paragraphs.  Each block contains the section label + body text (metrics
    # as text tables, and optionally a real FQDN hostname).
    _SEC_RE = re.compile(r'<<SECTION:\s*(.+?)\s*>>(.+?)(?=<<SECTION:|$)',
                         re.DOTALL | re.IGNORECASE)
    _fqdn_re = re.compile(
        r'\b([a-z]{2,8}\d{3,}[a-z0-9]*(?:\.[a-z0-9-]+)*)\b', re.I)
    for sec_label, sec_body in _SEC_RE.findall(text):
        label = sec_label.strip().rstrip(":")
        body  = sec_body.strip()
        if not label:
            continue

        # Try to find a real FQDN / server hostname — search LABEL first, then body
        fqdn_m = _fqdn_re.search(label) or _fqdn_re.search(body)
        if fqdn_m:
            real_host = fqdn_m.group(1).rstrip(". ")
        else:
            # No FQDN found — use label as the display name
            # Sanitise: remove trailing colon/spaces, replace runs of spaces with _
            real_host = re.sub(r'\s+', '_', label)

        # Skip metric-label sections that don't identify a server
        if not _is_valid_hostname(real_host, strict=False):
            continue

        full_ctx = label + "\n" + body
        metrics = _extract_metrics_from_block(full_ctx)

        rec = {
            "host":         real_host,
            "label":        label,          # human-readable display name
            "type":         _infer_server_type(real_host, full_ctx[:600], label),
            "_image_only":  not (
                metrics["cpu_used"] > 0
                or metrics["mem_used"] > 0
                or metrics["disk_used_max"] > 0
            ),
            **metrics,
        }
        k = _key(real_host)
        # Only add if not already discovered by Strategy A/B
        if k not in servers:
            servers[k] = rec
            logger.debug(
                "Strategy D: label=%r host=%r  cpu=%.1f%%  mem=%.1f%%  disk=%.1f%%",
                label, real_host,
                metrics["cpu_used"], metrics["mem_used"], metrics["disk_used_max"],
            )
        else:
            # Merge: update any zero fields
            _upsert(real_host, rec)

    # ── Strategy E: "ServerLabel: hostname" / "ServerLabel hostname" lines ──
    # Handles patterns like:
    #   --DB Server: prbg901403001
    #   APP1 Server tsbb081402011 (Available Memory)
    #   Server1: tsbb191525041.jdadelivers.com
    #   SRE UI Server: prbg901425001
    #   SRE BATCH1 Server: prbg901430001
    # These are common in DOCXs that list servers without <<SECTION:>> markers.
    _HOST_PAT = r'([a-z]{2,8}\d{3,}[a-z0-9]*(?:\.[a-z0-9.-]+)?)'
    _LABEL_HOST_RE = re.compile(
        r'(?:^|[\n])[\-\u2013\u2014]*\s*'                       # optional leading dashes
        r'(?:'
            r'(?:App(?:lication)?|DB|Database|SRE|Batch|Collab|Cognos|Utility|ACT|DMD|Web|ETL|Server)'
            r'[\w\s#&]*?'                                        # anything after the type keyword
        r')'
        r'[:\s]+\s*'                                             # separator (colon or whitespace)
        + _HOST_PAT,                                             # hostname capture
        re.IGNORECASE | re.MULTILINE,
    )
    # Also: "APP1 Server hostname (some metric)"
    _LABEL_HOST2_RE = re.compile(
        r'(?:^|[\n])[\-\u2013\u2014]*\s*'
        r'(\w[\w\s#&]*?)\s+'                                    # label capture
        + _HOST_PAT +                                            # hostname
        r'\s*(?:\(|$)',                                           # followed by '(' or EOL
        re.IGNORECASE | re.MULTILINE,
    )
    for m in _LABEL_HOST_RE.finditer(text):
        host = m.group(1).rstrip(". ")
        if not _is_valid_hostname(host, strict=True):
            continue
        k = _key(host)
        if k in servers:
            continue
        ctx = text[max(0, m.start()-100):m.end()+200]
        label_ctx = m.group(0).strip().lstrip("-–— ")
        servers[k] = {
            "host": host, "label": label_ctx,
            "type": _infer_server_type(host, ctx, label_ctx),
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }
        logger.debug("Strategy E: %s (from label line)", host)

    for m in _LABEL_HOST2_RE.finditer(text):
        label_part = m.group(1).strip()
        host = m.group(2).rstrip(". ")
        if not _is_valid_hostname(host, strict=True):
            continue
        k = _key(host)
        if k in servers:
            continue
        ctx = text[max(0, m.start()-100):m.end()+200]
        servers[k] = {
            "host": host, "label": label_part,
            "type": _infer_server_type(host, ctx, label_part),
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }
        logger.debug("Strategy E2: %s (label-host pattern)", host)

    # ── Strategy F: Zabbix PDF URL extraction ──
    # "Graphs for CUSTOMER" pages contain URLs like https://host.jdadelivers.com
    # Extract unique hostnames from URLs when no server hostnames were found by
    # earlier strategies (common in Public URL / Zabbix exports).
    _URL_HOST_RE = re.compile(
        r'https?://([a-z0-9][-a-z0-9]*\.jdadelivers\.com)', re.IGNORECASE)
    for m in _URL_HOST_RE.finditer(text):
        url_host = m.group(1).rstrip(".")
        short = url_host.split(".")[0]
        k = _key(url_host)
        if k in servers:
            continue
        servers[k] = {
            "host": url_host, "label": short,
            "type": "APP",
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }
        logger.debug("Strategy F: %s (from URL)", url_host)

    all_recs = list(servers.values())
    with_data = [s for s in all_recs
                 if s.get("cpu_used", 0) > 0 or s.get("mem_used", 0) > 0 or s.get("disk_used_max", 0) > 0]
    logger.info("parse_text_to_servers: %d total, %d with data", len(all_recs), len(with_data))
    return with_data if with_data else all_recs


# ─────────────────────────────────────────────────────────────────
# fleet_grade (lightweight, independent of resource_calculator)
# ─────────────────────────────────────────────────────────────────
def fleet_grade(servers: List[Dict[str, Any]]) -> str:
    """Quick fleet health grade: CRITICAL | WARNING | MODERATE | HEALTHY | UNKNOWN."""
    if not servers:
        return "UNKNOWN"
    scores: List[float] = []
    for s in servers:
        cpu  = float(s.get("cpu_used", 0) or 0)
        mem  = float(s.get("mem_used", 0) or 0)
        disk = float(s.get("disk_used_max", 0) or 0)
        if cpu == 0 and mem == 0 and disk == 0:
            continue
        scores.append((100 - cpu) * 0.30 + (100 - mem) * 0.40 + (100 - disk) * 0.30)
    if not scores:
        return "UNKNOWN"
    avg = sum(scores) / len(scores)
    if avg >= 80: return "HEALTHY"
    if avg >= 60: return "MODERATE"
    if avg >= 40: return "WARNING"
    return "CRITICAL"


# ─────────────────────────────────────────────────────────────────
# Step 5 — save_resource_session
# ─────────────────────────────────────────────────────────────────
def save_resource_session(
    servers: List[Dict[str, Any]],
    customer_name: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    Clear ALL stale server data from the previous upload, store fresh results.

    Call on every new resource file upload to prevent hostname contamination
    from a previous customer's report bleeding into the current analysis.

    In a stateless FastAPI app this guards against module-level variable reuse
    across rapid sequential uploads (e.g., CI test harness, concurrent sessions).
    """
    global _SESSION
    _SESSION.clear()
    _SESSION["servers"]       = servers
    _SESSION["customer_name"] = customer_name
    logger.info(
        "save_resource_session: cleared stale data → %d new server(s) (customer=%s)",
        len(servers), customer_name or "unknown",
    )
    return servers


# ─────────────────────────────────────────────────────────────────
# Tabular resource parser (CSV / XLSX)
# ─────────────────────────────────────────────────────────────────
def _parse_tabular_resource(
    file_bytes: bytes, filename: str, mode: str,
) -> List[Dict[str, Any]]:
    """Parse CSV or XLSX file containing server host/metric data.

    Handles multiple formats:
      - Server fill metrics CSV (comment header lines starting with #)
      - Server inventory XLSX with columns like Server Name, Category, etc.
      - Generic CSV/XLSX with host + cpu/mem/disk columns

    Returns list of server dicts, or empty list if the file isn't a resource table.
    """
    try:
        import pandas as pd
        from services.resource_parser import _infer_server_type
    except ImportError:
        return []

    df = None
    fn = (filename or "").lower()
    try:
        if fn.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(file_bytes), comment="#",
                             skipinitialspace=True, on_bad_lines="skip")
        elif fn.endswith((".xlsx", ".xls")):
            df = pd.read_excel(io.BytesIO(file_bytes))
        else:
            # Try CSV first, then XLSX
            try:
                df = pd.read_csv(io.BytesIO(file_bytes), comment="#",
                                 skipinitialspace=True, on_bad_lines="skip")
            except Exception:
                try:
                    df = pd.read_excel(io.BytesIO(file_bytes))
                except Exception:
                    return []
    except Exception as exc:
        logger.warning("_parse_tabular_resource: read failed: %s", exc)
        return []

    if df is None or df.empty:
        return []

    # Normalise column names: strip, lower, replace spaces/special chars
    df.columns = [str(c).strip() for c in df.columns]
    col_map = {c.lower().replace(" ", "_").replace("%", "pct"): c for c in df.columns}

    # ── Detect host column ────────────────────────────────────
    host_col = None
    for candidate in ["host", "hostname", "server_name", "server", "server_details",
                       "servername", "fqdn", "node", "instance"]:
        if candidate in col_map:
            host_col = col_map[candidate]
            break
    if host_col is None:
        # Check for columns containing "server" or "host" in name
        for orig in df.columns:
            low = orig.lower()
            if "server" in low or "host" in low or "node" in low:
                host_col = orig
                break

    if host_col is None:
        logger.debug("_parse_tabular_resource: no host column found in %s", list(df.columns))
        return []

    # If the first row looks like a sub-header, use it and drop
    first_val = str(df.iloc[0][host_col]).strip().lower() if len(df) > 0 else ""
    if first_val in ("", "nan", "server name", "hostname", "server", "host", "instance"):
        # Re-read with the second row of actual data as the header
        # Try using first data row as column names
        new_cols = [str(v).strip() if str(v).strip() and str(v) != "nan" else df.columns[i]
                    for i, v in enumerate(df.iloc[0])]
        df = df.iloc[1:].reset_index(drop=True)
        df.columns = new_cols
        col_map = {c.lower().replace(" ", "_").replace("%", "pct"): c for c in df.columns}
        # Re-detect host column
        host_col = None
        for candidate in ["host", "hostname", "server_name", "server", "servername",
                           "fqdn", "node", "instance"]:
            if candidate in col_map:
                host_col = col_map[candidate]
                break
        if host_col is None:
            for orig in df.columns:
                low = orig.lower()
                if "server" in low or "host" in low:
                    host_col = orig
                    break
        if host_col is None:
            return []

    # ── Detect metric columns ─────────────────────────────────
    def _find_col(*candidates):
        for c in candidates:
            if c in col_map:
                return col_map[c]
        return None

    cpu_col  = _find_col("cpu_used", "cpu_pct", "cpu", "cpu_utilization", "cpu_usage",
                         "avg_cpu", "cpu_avg", "processorcpuutilization")
    mem_col  = _find_col("mem_used", "memory_pct", "mem_pct", "memory", "mem",
                         "memory_used", "memory_usage", "memoryutilization", "ram_pct")
    disk_col = _find_col("disk_used_max", "disk_pct", "disk", "disk_used",
                         "disk_usage", "storage_pct", "filesystem_pct")
    type_col = _find_col("type", "server_type", "category", "role")
    mem_gb_col = _find_col("mem_total_gb", "memory_gb", "total_memory", "ram_gb", "ram")

    servers: List[Dict[str, Any]] = []
    for _, row in df.iterrows():
        host_val = str(row.get(host_col, "")).strip()
        if not host_val or host_val.lower() in ("nan", "", "none"):
            continue

        # Determine server type from column or hostname
        stype = "APP"
        if type_col:
            raw_type = str(row.get(type_col, "")).strip().lower()
            if any(k in raw_type for k in ["db", "database", "sql", "oracle"]):
                stype = "DB"
            elif any(k in raw_type for k in ["sre", "batch"]):
                stype = "SRE"
            elif any(k in raw_type for k in ["app", "application", "web"]):
                stype = "APP"
            else:
                stype = _infer_server_type(host_val, "", raw_type)
        else:
            stype = _infer_server_type(host_val)

        def _float(col):
            if col is None:
                return 0.0
            v = row.get(col)
            try:
                f = float(v)
                return round(f, 2) if f == f else 0.0  # NaN check: NaN != NaN
            except (ValueError, TypeError):
                return 0.0

        cpu  = _float(cpu_col)
        mem  = _float(mem_col)
        disk = _float(disk_col)
        mem_gb = _float(mem_gb_col)
        has_data = cpu > 0 or mem > 0 or disk > 0

        servers.append({
            "host":         host_val,
            "type":         stype,
            "cpu_used":     cpu,
            "cpu_avg":      0.0,
            "mem_used":     mem,
            "mem_total_gb": mem_gb,
            "disk_used_max": disk,
            "disks":        {"/" : disk} if disk > 0 else {},
            "_image_only":  not has_data,
        })

    logger.info("_parse_tabular_resource: %s → %d server(s)", filename, len(servers))
    return servers


# ─────────────────────────────────────────────────────────────────
# Master entry point — parse_resource_file
# ─────────────────────────────────────────────────────────────────
def parse_resource_file(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Single entry point for all resource file types.

    Pipeline:
      1. detect_file_mode()      → TEXT_PDF | TEXT_DOCX | IMAGE_PDF | IMAGE_DOCX | CSV | XLSX
      2. extract_text()          → plain text (empty for image-only modes)
      3. parse_text_to_servers() → paired CPU/MEM/Disk extraction with hostname filter
      4. DOCX fallback           → python-docx structured parser if text extraction yields 0
      5. save_resource_session() → clears stale _SESSION, returns fresh list

    Returns a list of server dicts.  Empty list means image-only mode —
    caller should trigger Vision enrichment via _run_vision_enrichment().

    Each server dict has keys:
        host, type, cpu_used, cpu_avg, mem_used, mem_total_gb,
        disk_used_max, disks, _image_only, [_customer_name]
    """
    # Step 1 — detect file mode (RULE 4)
    mode = detect_file_mode(file_bytes, filename)
    logger.info("parse_resource_file: %s  mode=%s", filename, mode)

    # ── CSV / XLSX: tabular server data ──────────────────────────
    if mode in ("CSV", "XLSX", "UNKNOWN"):
        tabular = _parse_tabular_resource(file_bytes, filename, mode)
        if tabular:
            return save_resource_session(tabular)

    # Step 2 — extract text
    text = extract_text(file_bytes, mode)

    # Step 3 — parse text → server records
    servers: List[Dict[str, Any]] = []
    if text.strip():
        servers = parse_text_to_servers(text)

    # Step 4 — DOCX structured fallback
    # If TEXT_DOCX but parse_text_to_servers() found nothing, fall back to the
    # full python-docx heading/table parser which handles Leonardo/Distell layouts.
    if not servers and mode == "TEXT_DOCX":
        try:
            from services.resource_parser import parse_resource_docx
            fallback, _ = parse_resource_docx(io.BytesIO(file_bytes))
            if fallback:
                servers = fallback
                logger.info("parse_resource_file: DOCX fallback parser yielded %d servers", len(servers))
        except Exception as exc:
            logger.warning("DOCX fallback parser failed: %s", exc)

    # Mark image_only flag on each record
    for s in servers:
        has_data = s.get("cpu_used", 0) > 0 or s.get("mem_used", 0) > 0 or s.get("disk_used_max", 0) > 0
        s.setdefault("_image_only", not has_data)

    # ── NVIDIA LLM fallback ──────────────────────────────────────
    # When the regex parser produced nothing useful from a text-bearing
    # Zabbix / Azure Monitor / Grafana report, ask the NVIDIA NIM LLM to
    # read the raw text and return a structured server list.  Triggered
    # only when we *have* text but no usable server records.
    def _all_zero(srvs):
        if not srvs:
            return True
        for s in srvs:
            if (s.get("cpu_used", 0) > 0
                or s.get("mem_used", 0) > 0
                or s.get("disk_used_max", 0) > 0):
                return False
        return True

    if text.strip() and _all_zero(servers):
        try:
            from services.nvidia_llm import extract_servers_from_text
            from services.config_store import get_nvidia_key
            nv_key = get_nvidia_key()
            if nv_key:
                llm_servers = extract_servers_from_text(text, nv_key)
                if llm_servers:
                    logger.info(
                        "parse_resource_file: NVIDIA LLM fallback yielded %d server(s)",
                        len(llm_servers),
                    )
                    servers = llm_servers
        except Exception as exc:
            logger.warning("parse_resource_file: NVIDIA LLM fallback failed: %s", exc)

    if not servers:
        logger.info("parse_resource_file: no text metrics extracted — image_only, Vision required")

    # Extract customer name if tagged by the DOCX parser
    customer_name: Optional[str] = None
    for s in servers:
        if s.get("_customer_name"):
            customer_name = s["_customer_name"]
            break

    # Step 5 — clear stale session, store fresh results
    return save_resource_session(servers, customer_name)
