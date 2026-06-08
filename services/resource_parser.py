"""
Resource utilization parsers — extracted verbatim from app_v2.py
(Streamlit references stripped). Pure functions, no global state.

Public API:
    extract_pdf_text(file_obj, max_pages=300) -> str
    parse_resource_docx(file_obj)             -> tuple[list[dict], bool]
    parse_zabbix_pdf_text(text)               -> list[dict]
    get_health_score(cpu, mem, disk, server_type="APP") -> float

Private helpers (kept private with underscore prefix):
    _infer_server_type(host, context="", doc_section_hint="")
    _parse_zabbix_block(blk)
    _calculate_host_health(cpu_pct, mem_pct, disk_pct)

The regex patterns and parsing rules are unchanged from the original
Streamlit monolith — only `st.*` calls have been removed.
"""
from __future__ import annotations

import io
import re


# ─────────────────────────────────────────────────────────────────
# PDF text extraction (helper used by the upload pipeline)
# ─────────────────────────────────────────────────────────────────
def extract_pdf_text(file_obj, max_pages: int = 300) -> str:
    """Extract text from PDF. Handles large (100+ page) Zabbix reports efficiently.
    Uses list join instead of string concatenation for performance."""
    parts = []
    try:
        from pypdf import PdfReader
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        total = min(len(reader.pages), max_pages)
        for i in range(total):
            t = reader.pages[i].extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts)
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        total = min(len(reader.pages), max_pages)
        for i in range(total):
            t = reader.pages[i].extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts)
    except ImportError:
        pass
    return ""


# ─────────────────────────────────────────────────────────────────
# Health score (F7 — host health)
# ─────────────────────────────────────────────────────────────────
def _calculate_host_health(cpu_pct, mem_pct, disk_pct):
    """F7 — Host Health Score: continuous weighted formula (not zone-based).
    Returns health_score 0-100: higher = healthier."""
    def _clip(v):
        v = float(v)
        if v < 0:
            return 0.0
        if v > 100:
            return 100.0
        return v

    cpu = _clip(cpu_pct)
    mem = _clip(mem_pct)
    disk = _clip(disk_pct)
    # Weights: CPU 30%, Memory 40%, Disk 30%
    health_score = (100 - cpu) * 0.30 + (100 - mem) * 0.40 + (100 - disk) * 0.30
    if health_score >= 80:
        status = "HEALTHY"
    elif health_score >= 60:
        status = "WARNING"
    else:
        status = "CRITICAL"
    return {
        "health_score": round(health_score, 1),
        "status": status,
        "cpu_pct": round(cpu, 1),
        "mem_pct": round(mem, 1),
        "disk_pct": round(disk, 1),
    }


def get_health_score(cpu, mem, disk, server_type="APP"):
    """Delegates to F7 calculate_host_health with server-type-aware thresholds.
    DB servers: memory 80-92% is expected (SGA/PGA allocation) — reduced penalty.
    DB servers: 40%+ CPU triggers increasing score reduction (CPU contention sensitive).
    Returns -1 for unknown servers (all-zero metrics)."""
    if cpu == 0 and disk == 0 and mem == 0:
        return -1   # UNKNOWN — not healthy, no data
    score = _calculate_host_health(cpu, mem, disk)["health_score"]
    # DB-specific CPU penalty: 40%+ CPU on a DB server progressively degrades score
    if server_type == "DB" and cpu >= 40:
        penalty = min(25, (cpu - 40) * 0.625)  # up to -25 pts at 80% CPU
        score = max(0, score - penalty)
    # DB-specific memory leniency: 80-92% memory is expected (SGA/PGA)
    # Restore up to 20 health points that the base formula unfairly deducted
    if server_type == "DB" and 75 <= mem <= 92:
        # The base formula penalizes mem at 40% weight. For a DB at 88%:
        # base penalty = 88 * 0.40 = 35.2 pts lost. But 88% is normal for DB.
        # Restore proportional credit: at 88% DB gets back ~16 pts
        restore = min(20.0, (mem - 75) * 1.2)  # scale 0-20 over 75-92% range
        score = min(100, score + restore)
    return score


# ─────────────────────────────────────────────────────────────────
# Environment inference (PROD / TEST / DEV)
# ─────────────────────────────────────────────────────────────────
def _infer_environment(host: str) -> str:
    """Infer server environment from hostname prefix/patterns.

    Convention:
      p* = PROD, t* = TEST, d* = DEV
    Also checks for explicit keywords: prod, prd, test, tst, uat, dev, stg, staging, qa.
    Returns 'PROD', 'TEST', 'DEV', or '' (unknown).
    """
    h = (host or "").lower().strip()
    if not h:
        return ""

    # Explicit keywords anywhere in hostname (highest priority)
    if re.search(r'\b(prod|prd)\b|[-_](prod|prd)[-_\d.]', h):
        return "PROD"
    if re.search(r'\b(test|tst|uat|qa)\b|[-_](test|tst|uat|qa)[-_\d.]', h):
        return "TEST"
    if re.search(r'\b(dev|stg|staging)\b|[-_](dev|stg|staging)[-_\d.]', h):
        return "DEV"

    # First-character prefix convention: p=PROD, t=TEST, d=DEV
    if h[0] == 'p':
        return "PROD"
    if h[0] == 't':
        return "TEST"
    if h[0] == 'd':
        return "DEV"

    return ""


# ─────────────────────────────────────────────────────────────────
# Server-type inference
# ─────────────────────────────────────────────────────────────────
def _infer_server_type(host, context="", doc_section_hint=""):
    """Detect APP / DB / SRE from hostname prefix, naming convention, section hint, or context.

    Works across all customer naming conventions:
    - JDA/BY: tsXXNNNNNN — classify by numeric range (SRE 1525-1535, ACT 1540-1549)
    - Generic: hostnames containing 'db', 'sql', 'ora', 'data', 'mongo', 'redis', 'pg'
    - SRE: hostnames or headings containing 'sre'
    - doc_section_hint: heading text from DOCX section (highest priority after JDA range)
    - Context: surrounding text mentions database-related keywords
    """
    h = host.lower()
    # JDA/BY convention: tsXXNNNNNN — classify by numeric suffix range
    m = re.match(r'ts[a-z]{2}(\d+)', h)
    if m:
        n = int(m.group(1))
        if 1525 <= n <= 1535: return "SRE"    # SRE servers
        if 1540 <= n <= 1549: return "APP"    # ACT servers
        return "DB"                            # all other ts* = DB

    # doc_section_hint from DOCX heading takes priority over hostname guesses
    if doc_section_hint:
        dh = doc_section_hint.lower()
        if re.search(r'\bsre\b', dh):
            return "SRE"
        if re.search(r'\bapp(?:lication)?\b|\bweb\b|\bbatch\b|\betl\b|\bact\b', dh):
            return "APP"
        if re.search(r'\bdb\b|\bdatabase\b|\boracle\b|\bsql\b|\bdata\s*server\b', dh):
            return "DB"

    # SRE indicators in hostname
    if "sre" in h:
        return "SRE"

    # Generic DB indicators in hostname
    if any(k in h for k in ["db","oracle","sql","data","mongo","redis","postgres",
                              "pg","mysql","mssql","mariadb","cassandra","elastic"]): return "DB"
    # Common DB hostname patterns: *-db-*, *db01, etc.
    if re.search(r'[-_]db[-_\d]|db\d{1,3}$', h): return "DB"
    if context:
        cl = context.lower()
        if any(k in cl for k in ["oracle","database"," db ","db server","db instance",
                                   "sql server","mysql","postgres","mongodb","tablespace",
                                   "datafile","redo log","archive log"]): return "DB"
    return "APP"


# ─────────────────────────────────────────────────────────────────
# Zabbix text-block parser
# ─────────────────────────────────────────────────────────────────
def _parse_zabbix_block(blk):
    """Extract CPU/Mem/Disk from a single Zabbix data block."""
    rec = {"cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
           "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {}}

    # Memory total
    m = re.search(r'Total\s+[Mm]emory\s*:\s*([\d.]+)\s*GB', blk)
    if m: rec["mem_total_gb"] = float(m.group(1))

    # CPU idle → used
    m = re.search(r'CPU idle time\s*:\s*latest value:\s*([\d.]+)', blk)
    if m: rec["cpu_used"] = round(100 - float(m.group(1)), 2)
    m = re.search(r'CPU idle time trend/SLA:\s*([\d.]+)', blk)
    if m: rec["cpu_avg"] = round(100 - float(m.group(1)), 2)

    # CPU utilization direct
    if rec["cpu_used"] == 0.0:
        for cpat in [r'CPU\s+utilization\s*:\s*latest value:\s*([\d.]+)',
                     r'CPU\s+usage\s*:\s*latest value:\s*([\d.]+)',
                     r'CPU\s+load\s*:\s*latest value:\s*([\d.]+)']:
            m = re.search(cpat, blk, re.I)
            if m: rec["cpu_used"] = round(float(m.group(1)), 2); break

    # Memory used %
    for mem_pat in [
        r'(?:Available|Free)\s+[Mm]emory\s*(?:%|percent)[^:]*:\s*latest value:\s*([\d.]+)',
        r'(?:Available|Free)\s+[Mm]emory\s*:\s*latest value:\s*([\d.]+)\s*%',
        r'[Mm]emory\s+utilization\s*:\s*latest value:\s*([\d.]+)',
        r'[Mm]emory\s+used\s*:\s*latest value:\s*([\d.]+)',
        r'Used\s+[Mm]emory\s*%[^:]*:\s*([\d.]+)',
    ]:
        m = re.search(mem_pat, blk, re.I)
        if m:
            val = float(m.group(1))
            if re.match(r'(?:Available|Free)', mem_pat, re.I):
                rec["mem_used"] = round(100 - val, 2) if val <= 100 else 0.0
            else:
                rec["mem_used"] = round(val, 2)
            break

    # Disk — multiple Zabbix/monitoring text formats
    # Pattern 1: "Free disk space on /mount (percentage) : 98.95 %"
    for mount, pct in re.findall(
            r'Free disk space on (\S+)\s*\(percentage\)\s*:\s*([\d.]+)', blk):
        rec["disks"][mount] = round(100 - float(pct), 2)

    # Pattern 2: "Disk space on /mount : used 45.2%" or "Disk utilization /mount : 45%"
    for mount, pct in re.findall(
            r'[Dd]isk\s+(?:space|utilization)\s+(?:on\s+)?(\S+)\s*:\s*(?:used\s+)?([\d.]+)\s*%', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 3: "Used disk space on /mount (percentage) : latest value: 45.2"
    for mount, pct in re.findall(
            r'[Uu]sed\s+disk\s+space\s+on\s+(\S+)\s*\(percentage\)\s*:\s*(?:latest value:\s*)?([\d.]+)', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 4: "Filesystem /mount : used 45%" or "/mount used: 45.2%"
    for mount, pct in re.findall(
            r'(?:Filesystem|Volume)\s+(\S+)\s*:\s*used\s+([\d.]+)\s*%', blk, re.I):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(float(pct), 2)

    # Pattern 5: "Available disk space on /mount (percentage) : latest value: 85"
    for mount, pct in re.findall(
            r'[Aa]vailable\s+[Dd]isk\s+space\s+(?:in\s+%?\s+)?on\s+(\S+)\s*(?:\(percentage\))?\s*:\s*(?:latest value:\s*)?([\d.]+)', blk):
        if mount not in rec["disks"]:
            rec["disks"][mount] = round(100 - float(pct), 2) if float(pct) <= 100 else 0.0

    # BUG-2: Oracle ASM disk groups — Zabbix reports used% directly (no inversion)
    # Format: "ASM disk group DATA: used 78.5%" or "ASM: DATA used: 78%"
    for label, pct in re.findall(
            r'ASM\s+(?:disk\s+group\s+)?(\w+)\s*[:\s]+(?:used|utilization)[:\s]+([\d.]+)\s*%?', blk, re.I):
        key = f"ASM:{label.upper()}"
        if key not in rec["disks"]:
            rec["disks"][key] = round(float(pct), 2)   # used% directly, no 100- inversion
    # Also catch "diskgroup DATA used 78%" format
    for label, pct in re.findall(
            r'[Dd]iskgroup\s+(\w+)\s+used\s+([\d.]+)\s*%?', blk):
        key = f"ASM:{label.upper()}"
        if key not in rec["disks"]:
            rec["disks"][key] = round(float(pct), 2)

    # BUG-1: Normalize empty/invalid mount keys → root '/'
    # Happens when Zabbix text omits mount point for the root partition
    cleaned_disks = {}
    for mnt, v in rec["disks"].items():
        # Reject mounts that are obviously wrong captures (parens, empty)
        norm = mnt if (mnt and not mnt.startswith("(")) else "/"
        cleaned_disks[norm] = max(cleaned_disks.get(norm, 0.0), v)
    rec["disks"] = cleaned_disks

    # Track known disk mount points from graph titles (even without values)
    # "Disk space usage /mount" or "Available Disk space in % on /mount"
    _disk_mounts_seen = set(rec["disks"].keys())
    for mount in re.findall(r'[Dd]isk\s+space\s+usage\s+(/\S*)', blk):
        _disk_mounts_seen.add(mount)
    for mount in re.findall(r'Available\s+[Dd]isk\s+space\s+in\s+%\s+on\s+(/\S*)', blk):
        _disk_mounts_seen.add(mount)
    for label in re.findall(r'ASM\s+(?:disk\s+group\s+)?(\w+)', blk, re.I):
        _disk_mounts_seen.add(f"ASM:{label.upper()}")
    rec["_disk_mounts_known"] = list(_disk_mounts_seen)

    if rec["disks"]:
        rec["disk_used_max"] = max(rec["disks"].values())

    # BUG-3: Oracle DB metrics — uptime, active sessions, instance status
    oracle = {}
    _m = re.search(r'(?:Oracle\s+)?(?:DB\s+)?[Uu]ptime[:\s]+([\d.]+)\s*(?:days?|d\b)', blk)
    if _m: oracle["uptime_days"] = float(_m.group(1))
    _m = re.search(r'(?:Active\s+)?(?:sessions?|connections?)\s*:\s*(?:latest value:\s*)?([\d]+)', blk, re.I)
    if _m: oracle["sessions"] = int(_m.group(1))
    _m = re.search(r'(?:Oracle\s+)?[Ii]nstance\s+(?:status|state)\s*[:\s]+(\w+)', blk)
    if _m: oracle["instance_status"] = _m.group(1).upper()
    if oracle:
        rec["oracle"] = oracle

    return rec


# ─────────────────────────────────────────────────────────────────
# Zabbix PDF text parser
# ─────────────────────────────────────────────────────────────────
def parse_zabbix_pdf_text(text):
    """Universal Zabbix/resource utilization PDF text parser.

    Handles ALL known formats:
      A) Header-block layout: "System Status for <host>" sections with text metrics
      B) UTZ/Monitor-style: "Graphs for <host>" only (data in images, text is ToC)
      C) Image-only PDFs: title + screenshots, extract server names from any pattern
      D) Mixed: some servers have text metrics, others are image-only

    Returns list of server dicts. Image-only servers have zeroed metrics
    and are flagged for Gemini Vision OCR enrichment.
    """
    servers_map = {}  # short_host_lower → rec (case-insensitive dedup)
    def _dedup_key(h):
        """Normalize hostname to lowercase short name for dedup."""
        return h.split(".")[0].lower()

    # ── Strategy A: "System Status for <host>" header-block layout ──
    status_blocks = re.split(r'System Status for\s+', text)
    for blk in status_blocks[1:]:
        lines = blk.strip().splitlines()
        host = lines[0].strip().split()[0] if lines else "unknown"
        host = host.rstrip(".")
        if not host or host == "unknown": continue

        rec = _parse_zabbix_block(blk)
        rec["host"] = host
        rec["type"] = _infer_server_type(host, blk)
        rec["_image_only"] = False

        has_data = (rec["cpu_used"] > 0 or rec["mem_total_gb"] > 0
                    or rec["disk_used_max"] > 0 or rec["mem_used"] > 0)

        key = _dedup_key(host)
        if key in servers_map:
            existing = servers_map[key]
            e_has = (existing["cpu_used"] > 0 or existing["mem_total_gb"] > 0
                     or existing["disk_used_max"] > 0 or existing["mem_used"] > 0)
            if has_data and not e_has:
                servers_map[key] = rec
            elif has_data and e_has:
                for k in ["cpu_used","cpu_avg","mem_used","mem_total_gb","disk_used_max"]:
                    existing[k] = max(existing[k], rec[k])
                existing["disks"].update(rec["disks"])
                if existing["disks"]:
                    existing["disk_used_max"] = max(existing["disks"].values())
        else:
            servers_map[key] = rec

    # ── Strategy B: "Graphs for <host>" (UTZ/Monitor format — ToC-only PDFs) ──
    # These PDFs have NO "System Status for" but list servers via "Graphs for"
    graph_hosts = list(dict.fromkeys(re.findall(
        r'Graphs for\s+(\S+(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)', text)))
    for host in graph_hosts:
        host = host.rstrip(".")
        key = _dedup_key(host)
        if key in servers_map:
            continue  # already have from Strategy A
        # Find context around this host for DB detection
        idx = text.find(host)
        context = text[idx:idx+500] if idx >= 0 else ""
        stype = _infer_server_type(host, context)
        # Check if Oracle references exist for this host anywhere
        host_upper = host.split(".")[0].upper()
        if re.search(rf'Oracle.*{re.escape(host_upper)}|{re.escape(host_upper)}.*Oracle', text, re.I):
            stype = "DB"
        servers_map[key] = {
            "host": host, "type": stype,
            "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
            "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
            "_image_only": True,
        }

    # ── Strategy C: Fallback hostname extraction from any pattern ──
    # Catches "Trends and metrics for <host>", standalone FQDNs, etc.
    FQDN_RE = re.compile(r'\b([a-z]{2,6}\d{6,}[0-9a-z]*(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)\b', re.I)
    for m in FQDN_RE.finditer(text):
        host = m.group(1)
        key = _dedup_key(host)
        if key not in servers_map:
            servers_map[key] = {
                "host": host, "type": _infer_server_type(host),
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
                "_image_only": True,
            }
    # BUG-6: Short-hostname fallback — catches tsXXNNNN patterns with fewer digits
    # than FQDN_RE's 6-digit minimum (e.g. tsbc1234 = 4 digits)
    SHORTHOST_RE = re.compile(r'\b([a-z]{2,4}\d{4,5}[0-9a-z]*)\b', re.I)
    for m in SHORTHOST_RE.finditer(text):
        host = m.group(1)
        key = _dedup_key(host)
        if key not in servers_map:
            servers_map[key] = {
                "host": host, "type": _infer_server_type(host),
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0, "disks": {},
                "_image_only": True,
            }

    # ── Dedup: prefer entries with data, remove pure ToC stubs ──
    servers_all = list(servers_map.values())
    servers_with_data = [s for s in servers_all
                         if s["cpu_used"] > 0 or s["mem_total_gb"] > 0 or s["disk_used_max"] > 0]
    return servers_with_data if servers_with_data else servers_all


# ─────────────────────────────────────────────────────────────────
# DOCX resource utilization parser
# ─────────────────────────────────────────────────────────────────
def parse_resource_docx(file_obj):
    """
    Parse resource utilisation DOCX — handles ALL real PE report formats:
    A) Distell:   "Server1: tsbb191525041.jdadelivers.com" / bare FQDN
    B) Leonardo:  "Application Server 1 :", "SRE UI IO :", "Utility :"
    C) Generic:   "SRE : hostname.domain.com"
    D) CEAt:      "DB Server 1", "DB Server 2", "App Server" (no colon/bold)
    E) Generic:   any paragraph whose text contains a known server keyword
                  and is short enough to be a heading
    Returns (servers_list, image_only: bool)
    """
    try:
        from docx import Document
    except ImportError:
        return [], True

    file_obj.seek(0)
    try:
        doc = Document(file_obj)
    except Exception as _docx_err:
        # File is corrupted or not a valid DOCX — return empty list gracefully.
        # The caller is expected to surface a friendly error to the user.
        _ = _docx_err
        return [], True

    # ── Try structured tables first (text-based metrics) ─────────
    servers_with_data = []
    for tbl in doc.tables:
        if len(tbl.rows) < 2: continue
        hdrs = [c.text.strip().lower() for c in tbl.rows[0].cells]
        has_host   = any(k in h for h in hdrs for k in ["host","server","name"])
        has_metric = any(k in " ".join(hdrs) for k in ["cpu","disk","mem","memory"])
        if not (has_host and has_metric): continue
        for row in tbl.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells[0]: continue
            rec = {"host":cells[0],"type":_infer_server_type(cells[0]),
                   "cpu_used":0.0,"cpu_avg":0.0,
                   "mem_used":0.0,"mem_total_gb":0.0,"disk_used_max":0.0,"disks":{}}
            for hi,h in enumerate(hdrs):
                if hi>=len(cells): break
                cell_text = cells[hi].strip()
                # Skip columns that are clearly not numeric metrics
                # (OS version strings like "3.10.0-1160", hostnames, category labels)
                if any(skip in h for skip in ["o/s","os","oper","version","categor","name","server","host","region","module","enterprise","product"]):
                    continue
                # Extract the FIRST plain integer or simple decimal (not dotted version strings)
                # A valid metric is: digits optionally followed by ONE decimal point + digits
                vm = re.search(r"\b(\d{1,6}(?:\.\d{1,3})?)\b", cell_text)
                if not vm:
                    continue
                candidate = vm.group(1)
                # Reject version-like strings: more than one dot means it's a version, not a metric
                if cell_text.count(".") > 1:
                    continue
                try:
                    val = float(candidate)
                except ValueError:
                    continue
                if "cpu" in h:                              rec["cpu_used"]=val
                elif "mem" in h and "gb" in h:              rec["mem_total_gb"]=val
                elif "mem" in h and ("%" in h or "used" in h):  rec["mem_used"]=val
                elif "disk" in h or "storage" in h:         rec["disk_used_max"]=val
            servers_with_data.append(rec)
    if servers_with_data:
        return servers_with_data, False

    # ── Regex helpers ─────────────────────────────────────────────
    # Generic server hostnames: short alpha prefix + digits, optionally followed by FQDN
    # e.g. tsbb191525041.domain.com, prbg241530001, dvbb941426001.local
    FQDN_RE    = re.compile(
        r"\b([a-z]{2,6}[0-9]{3,}[0-9a-z]*"
        r"(?:\.\w+\.(?:com|net|local|org|internal|lan|corp|int|cloud))?)\b",
        re.IGNORECASE)
    SHORT_HOST = re.compile(
        r"\b([a-z]{2,6}[0-9]{4,}[0-9a-z]*)\b", re.IGNORECASE)

    # Server role / type keywords
    ROLE_RE = re.compile(
        r"\b(application|app|database|db|sre|batch|utility|act|cognos|"
        r"ui|io|server|node|worker|integration|etl|web|gateway|report|"
        r"scheduler|mq|middleware|proxy|balancer|cache|search|analytics)\b",
        re.IGNORECASE)

    SERVER_HEADING_RE = re.compile(
        r"\b(app(?:lication)?\s*server\s*\d*|"
        r"db\s*server\s*\d*|"
        r"database\s*server\s*\d*|"
        r"sre\s*(?:ui|io|batch|app|server)?\s*\d*|"
        r"utility\s*(?:server)?\s*\d*|"
        r"act\s*server\s*\d*|"
        r"cognos\s*(?:server)?\s*\d*|"
        r"batch\s*server\s*\d*|"
        r"web\s*server\s*\d*|"
        r"etl\s*server\s*\d*|"
        r"server\s*\d+)",
        re.IGNORECASE)

    def infer_type(text):
        t = text.lower()
        if any(k in t for k in ["db ","database"," db","oracle","sql","data base","dbserver"]): return "DB"
        return "APP"

    def extract_hostname(text):
        m = FQDN_RE.search(text)
        if m: return m.group(1)
        m = SHORT_HOST.search(text)
        if m: return m.group(1)
        return None

    def para_is_bold(para):
        if not para.runs: return False
        bold_chars = sum(len(r.text) for r in para.runs if r.bold and r.text.strip())
        total_chars = len(para.text.strip())
        return total_chars > 0 and bold_chars / total_chars > 0.4

    # Company / org name suffixes — lines with these are titles, never servers
    COMPANY_RE = re.compile(
        r"\b(ltd|limited|inc|corp|corporation|pty|gmbh|b\.?v|s\.?a|"
        r"ag|llc|plc|co\.|group|holdings|international|industries)\b",
        re.IGNORECASE)

    def is_server_heading(para, excluded=None):
        t = para.text.strip()
        if not t or len(t) > 150: return False

        # Never a server if it's in the excluded titles set
        if excluded and t in excluded: return False

        # Never a server if it contains a company name suffix (Ltd, Pty, Corp, Inc)
        # UNLESS it also has a real hostname
        has_host = bool(FQDN_RE.search(t) or SHORT_HOST.search(t))
        has_company = bool(COMPANY_RE.search(t))
        if has_company and not has_host: return False

        # 1) Contains a real JDA/Zabbix hostname → always a server entry
        if has_host: return True

        # 2) Matches known server heading pattern (DB Server 1, App Server 2, etc.)
        if SERVER_HEADING_RE.search(t): return True

        # 3) Bold/underline + role keyword — but NOT if it looks like a document title
        is_styled = para_is_bold(para) or any(
            r.underline for r in para.runs if r.text.strip())
        if is_styled and ROLE_RE.search(t):
            # Extra guard: colon-split left side must not be a company/title
            parts = re.split(r"\s*:\s*", t, maxsplit=1)
            if len(parts) == 2:
                left = parts[0].strip()
                # If left has no hostname and has company suffix → it's a title
                if COMPANY_RE.search(left) and not FQDN_RE.search(left): return False
                # If left looks like "Document Name" (Title Case multi-word, no digits)
                # and right looks like a description → skip
                if (not re.search(r"\d", left) and
                    len(left.split()) >= 2 and
                    not SERVER_HEADING_RE.search(left) and
                    not FQDN_RE.search(left)):
                    # Right side must have a hostname to count
                    if not FQDN_RE.search(parts[1]) and not SHORT_HOST.search(parts[1]):
                        return False
            return True

        # 4) Short line + role keyword (Leonardo: "SRE UI IO :", "Utility :")
        stripped = t.rstrip(": ")
        if len(stripped) < 60 and ROLE_RE.search(stripped):
            words = stripped.split()
            if len(words) <= 6:
                # Guard: must not have company suffix on left of colon
                parts = re.split(r"\s*:\s*", stripped, maxsplit=1)
                left = parts[0].strip()
                if COMPANY_RE.search(left): return False
                return True

        return False

    # ── Collect all text from paragraphs + single-cell tables ────
    all_paras = list(doc.paragraphs)
    # Also add single-cell table entries (some DOCXs use 1-cell tables as headings)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c for c in row.cells if c.text.strip()]
            if len(cells) == 1:
                for para in cells[0].paragraphs:
                    if para.text.strip():
                        all_paras.append(para)

    # ── Walk all paragraphs ───────────────────────────────────────
    seen, servers = set(), []
    customer_name = None
    excluded_titles = set()   # lines used as doc title/customer → never server headings

    # ── PASS 1: extract customer name + build excluded set ────────
    for para in all_paras:
        t = para.text.strip()
        if not t or len(t) > 80: continue
        if para_is_bold(para) and not FQDN_RE.search(t) and not SHORT_HOST.search(t):
            # Has company suffix → definitely a title line
            if COMPANY_RE.search(t):
                excluded_titles.add(t)
                if not customer_name:
                    cname = re.sub(
                        r"[:\s]*(sre|servers?|resource|utilization|report|"
                        r"performance|metrics|consumption|test|prod|uat|dev|"
                        r"asia|2022|2023|2024|2025|2026).*$",
                        "", t, flags=re.IGNORECASE).strip().rstrip(":,.")
                    if 3 < len(cname) < 60 and not SERVER_HEADING_RE.match(cname):
                        customer_name = cname
            # Bold non-hostname short line that looks like a document title
            elif not customer_name and len(t) > 5:
                cname = re.sub(
                    r"[:\s]*(sre|servers?|resource|utilization|report|"
                    r"performance|metrics|consumption|test|prod|uat|dev|"
                    r"asia|2022|2023|2024|2025|2026).*$",
                    "", t, flags=re.IGNORECASE).strip().rstrip(":,.")
                if 3 < len(cname) < 60 and not SERVER_HEADING_RE.match(cname):
                    customer_name = cname
                    excluded_titles.add(t)

    # ── PASS 2: extract server headings ──────────────────────────
    # Build a list of (label, hostname, type) tuples
    # Handles patterns:
    #   "Server1: tsbb191525041.jdadelivers.com" (label:host on same line)
    #   "Application Server 1" then "tsbb911502021" (label then host on next line)
    #   "Database Server:" (label only, no hostname)
    #   Skip: "------- Application Server-1 end------" (separator lines)
    pending_label = None  # label from previous paragraph, waiting for hostname

    for para in all_paras:
        t = para.text.strip()
        if not t: continue

        # Skip separator lines (--- ... ---)
        if re.match(r'^[-─—=_]{5,}', t): continue

        # Check if this line is just a bare hostname (follow-up to a label)
        bare_host = extract_hostname(t)
        if bare_host and len(t.replace(bare_host, "").strip(":. ")) < 5:
            # This is a hostname-only line
            if pending_label:
                # Merge with previous label
                display = bare_host
                label = pending_label
                stype = infer_type(pending_label)
                # Also check hostname for DB
                if _infer_server_type(bare_host) == "DB": stype = "DB"
                pending_label = None
                if display not in seen:
                    seen.add(display)
                    servers.append({
                        "host": display, "label": label, "type": stype,
                        "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                        "mem_total_gb": 0.0, "disk_used_max": 0.0,
                        "disks": {}, "_image_only": True,
                    })
                continue
            else:
                # Standalone hostname
                display = bare_host
                stype = _infer_server_type(bare_host)
                if display not in seen:
                    seen.add(display)
                    servers.append({
                        "host": display, "label": display, "type": stype,
                        "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                        "mem_total_gb": 0.0, "disk_used_max": 0.0,
                        "disks": {}, "_image_only": True,
                    })
                continue

        # Clear pending label if this is not a hostname
        if pending_label:
            # Previous label had no hostname — store as label-only server
            display = pending_label.rstrip(": ").strip()
            stype = infer_type(pending_label)
            if display not in seen and len(display) > 2:
                seen.add(display)
                servers.append({
                    "host": display, "label": display, "type": stype,
                    "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                    "mem_total_gb": 0.0, "disk_used_max": 0.0,
                    "disks": {}, "_image_only": True,
                })
            pending_label = None

        if not is_server_heading(para, excluded=excluded_titles): continue

        # ── Parse heading into host + label ──────────────────────
        colon_split = re.split(r"\s*:\s*", t, maxsplit=1)
        hostname = None
        label    = t

        if len(colon_split) == 2:
            left, right = colon_split[0].strip(), colon_split[1].strip()
            h_right = extract_hostname(right)
            h_left  = extract_hostname(left)
            if h_right:
                hostname = h_right
                label    = left if left else right
            elif h_left:
                hostname = h_left
                label    = right if right else left
            else:
                label = left or right
        else:
            hostname = extract_hostname(t)
            label    = t

        if hostname:
            display = hostname
            stype = infer_type(t)
            if _infer_server_type(hostname) == "DB": stype = "DB"
            if display not in seen:
                seen.add(display)
                servers.append({
                    "host": display, "label": label.rstrip(": ").strip(),
                    "type": stype,
                    "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                    "mem_total_gb": 0.0, "disk_used_max": 0.0,
                    "disks": {}, "_image_only": True,
                })
        else:
            # No hostname on this heading line — set as pending label
            # so the next paragraph (if it's a bare hostname) can merge
            pending_label = label.rstrip(": ").strip()

    # Flush any remaining pending_label at end of loop
    if pending_label:
        display = pending_label.rstrip(": ").strip()
        stype = infer_type(pending_label)
        if display not in seen and len(display) > 2:
            seen.add(display)
            servers.append({
                "host": display, "label": display, "type": stype,
                "cpu_used": 0.0, "cpu_avg": 0.0, "mem_used": 0.0,
                "mem_total_gb": 0.0, "disk_used_max": 0.0,
                "disks": {}, "_image_only": True,
            })

    if servers:
        # Tag each server record with the detected customer name so callers
        # can surface it in the UI without re-parsing.
        if customer_name:
            for s in servers:
                s.setdefault("_customer_name", customer_name)
        return servers, True

    return [], True
